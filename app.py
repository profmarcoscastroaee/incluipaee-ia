
# INCLUISRM V53 - Perfil docente, modo maker inclusivo e projetos norteadores no AEE
# Atualização: integra perfil pedagógico/tecnológico do professor AEE e docente regular, modo maker inclusivo e projetos interdisciplinares sem caracterizar reforço escolar.

import os
import re
import json
import sqlite3
import calendar
from datetime import datetime, date, time, timezone, timedelta
try:
    from zoneinfo import ZoneInfo
except Exception:
    ZoneInfo = None
from pathlib import Path
from html import escape

import streamlit as st
import pandas as pd
import altair as alt

if "form_reset" not in st.session_state:
    st.session_state.form_reset = {}


def get_form_key(nome):
    if nome not in st.session_state.form_reset:
        st.session_state.form_reset[nome] = 0
    return f"{nome}_{st.session_state.form_reset[nome]}"


def resetar_form(nome):
    if nome not in st.session_state.form_reset:
        st.session_state.form_reset[nome] = 0
    st.session_state.form_reset[nome] += 1


try:
    from openai import OpenAI
except Exception:
    OpenAI = None

# ======================================================
# CONFIGURAÇÕES GERAIS
# ======================================================
st.set_page_config(
    page_title="INCLUISRM",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
)

DB_PATH = Path("inclui_paee.db")
LOGO_PATH = "logosrm.png"

# Pastas da Base de Conhecimento IA
BASE_CONHECIMENTO_DIR = Path("base_conhecimento")
PASTA_CIENTIFICA = BASE_CONHECIMENTO_DIR / "cientifica"
PASTA_PEDAGOGICA = BASE_CONHECIMENTO_DIR / "pedagogica"
CHROMA_DIR = Path("chroma_db_incluisrm")
DOCUMENTOS_AVALIACOES_DIR = Path("documentos_avaliacoes")
RELATORIOS_VISUAIS_DOCENTE_DIR = Path("relatorios_visuais_docente")

PASTA_BNCC = BASE_CONHECIMENTO_DIR / "bncc"
PASTA_BNCC_COMPUTACAO = BASE_CONHECIMENTO_DIR / "bncc_computacao"
PASTA_ESTRATEGIAS = BASE_CONHECIMENTO_DIR / "estrategias"
PASTA_TECNOLOGIA_ASSISTIVA = BASE_CONHECIMENTO_DIR / "tecnologia_assistiva"
# Pastas criadas manualmente pelo usuário no servidor/projeto.
# O sistema passa a reconhecer estes nomes exatamente como aparecem na estrutura de arquivos.
PASTA_TECNOLOGIA_USUARIO = BASE_CONHECIMENTO_DIR / "Tecnologia"
PASTA_LEI_USUARIO = BASE_CONHECIMENTO_DIR / "Lei"
PASTA_MAKER_INCLUSIVO = BASE_CONHECIMENTO_DIR / "maker_inclusivo"
PASTA_LEGISLACAO = BASE_CONHECIMENTO_DIR / "legislacao"
PASTA_PROJETOS_INTERDISCIPLINARES = BASE_CONHECIMENTO_DIR / "projetos_interdisciplinares"
PASTA_PERFIS_PEDAGOGICOS = BASE_CONHECIMENTO_DIR / "perfis_pedagogicos"

PASTAS_BASE_IA = {
    "cientifica": PASTA_CIENTIFICA,
    "pedagogica": PASTA_PEDAGOGICA,
    "bncc": PASTA_BNCC,
    "bncc_computacao": PASTA_BNCC_COMPUTACAO,
    "estrategias": PASTA_ESTRATEGIAS,
    "tecnologia_assistiva": PASTA_TECNOLOGIA_ASSISTIVA,
    "tecnologia": PASTA_TECNOLOGIA_USUARIO,
    "maker_inclusivo": PASTA_MAKER_INCLUSIVO,
    "legislacao": PASTA_LEGISLACAO,
    "lei": PASTA_LEI_USUARIO,
    "projetos_interdisciplinares": PASTA_PROJETOS_INTERDISCIPLINARES,
    "perfis_pedagogicos": PASTA_PERFIS_PEDAGOGICOS,
}

BASE_CONHECIMENTO_DIR.mkdir(parents=True, exist_ok=True)
for _pasta_base in PASTAS_BASE_IA.values():
    _pasta_base.mkdir(parents=True, exist_ok=True)
CHROMA_DIR.mkdir(parents=True, exist_ok=True)
DOCUMENTOS_AVALIACOES_DIR.mkdir(parents=True, exist_ok=True)
RELATORIOS_VISUAIS_DOCENTE_DIR.mkdir(parents=True, exist_ok=True)

APP_NAME = "INCLUISRM"
APP_SUBTITLE = "Sistema Inteligente de Articulação Pedagógica Inclusiva"
APP_VERSION = "V53"
APP_VERSION_LABEL = "Infográfico Docente • Layout compacto em painel • Evidências + Sugestões"
# Fuso fixo UTC-3 usado por Recife/Pernambuco.
# Usar timezone/timedelta evita erro em ambientes Render sem base tzdata completa.
FUSO_LOCAL = timezone(timedelta(hours=-3), name="America/Recife")
os.environ["TZ"] = "America/Recife"
try:
    import time as _time_module
    if hasattr(_time_module, "tzset"):
        _time_module.tzset()
except Exception:
    pass


def agora_local():
    """Retorna data/hora no fuso local America/Recife, independente do UTC do servidor."""
    return datetime.now(timezone.utc).astimezone(FUSO_LOCAL)


# ======================================================
# CSS / LAYOUT PROFISSIONAL
# ======================================================
st.markdown(
    """
<style>
/* =============================
   INCLUISRM - Layout profissional
   ============================= */

.stApp {
    background: linear-gradient(180deg, #f7f9fc 0%, #eef3f9 100%);
}

.block-container {
    padding-top: 2.2rem;
    padding-bottom: 2.5rem;
    padding-left: 2.4rem;
    padding-right: 2.4rem;
    max-width: 1400px;
}

/* Sidebar */
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #0f172a 0%, #111827 100%);
    border-right: 1px solid rgba(255,255,255,0.08);
}

section[data-testid="stSidebar"] * {
    color: #ffffff !important;
}

section[data-testid="stSidebar"] hr {
    border-color: rgba(255,255,255,0.14);
}

.sidebar-logo-card {
    background: transparent;
    border-radius: 0;
    padding: 0;
    margin: 4px 0 14px 0;
    box-shadow: none;
}

.sidebar-logo-card img {
    border-radius: 14px;
}

.sidebar-title {
    font-size: 20px;
    font-weight: 900;
    letter-spacing: 0.5px;
    margin-top: 8px;
}

.sidebar-subtitle {
    font-size: 12px;
    color: #cbd5e1 !important;
    line-height: 1.35;
    margin-bottom: 12px;
}

/* Hero */
.app-hero {
    background: linear-gradient(135deg, #ffffff 0%, #edf7ff 100%);
    border: 1px solid #dbeafe;
    border-radius: 24px;
    padding: 26px 30px;
    margin-bottom: 22px;
    box-shadow: 0 14px 35px rgba(15, 23, 42, 0.08);
}

.app-title {
    font-size: 38px;
    font-weight: 900;
    color: #0f172a;
    margin: 0;
    line-height: 1.05;
    letter-spacing: -0.5px;
}

.app-subtitle {
    font-size: 17px;
    color: #475569;
    margin-top: 8px;
    margin-bottom: 0;
}

.app-badge {
    display: inline-block;
    background: #dcfce7;
    color: #166534 !important;
    border: 1px solid #bbf7d0;
    border-radius: 999px;
    padding: 6px 12px;
    font-size: 13px;
    font-weight: 700;
    margin-bottom: 10px;
}

.subtitulo {
    font-size: 23px;
    font-weight: 850;
    color: #0f172a;
    margin-bottom: 12px;
}

.descricao {
    color: #64748b;
    font-size: 16px;
    margin-bottom: 20px;
}

/* Containers com borda do Streamlit */
div[data-testid="stVerticalBlock"] div[data-testid="stVerticalBlockBorderWrapper"] {
    border-radius: 20px;
    border: 1px solid #e2e8f0;
    box-shadow: 0 10px 25px rgba(15, 23, 42, 0.05);
    background: #ffffff;
}

/* Métricas */
[data-testid="stMetric"] {
    background: #ffffff;
    border: 1px solid #e2e8f0;
    padding: 18px;
    border-radius: 18px;
    box-shadow: 0 8px 24px rgba(15, 23, 42, 0.05);
}

[data-testid="stMetricValue"] {
    font-size: 34px;
    font-weight: 900;
    color: #0f172a;
}

/* Botões */
.stButton > button,
.stDownloadButton > button,
div[data-testid="stLinkButton"] a {
    border-radius: 12px !important;
    font-weight: 700 !important;
    border: 1px solid #cbd5e1 !important;
    box-shadow: 0 4px 12px rgba(15, 23, 42, 0.06);
}

/* Inputs */
div[data-baseweb="input"],
div[data-baseweb="select"],
div[data-baseweb="textarea"] {
    border-radius: 12px;
}

/* Dataframes */
[data-testid="stDataFrame"] {
    border-radius: 16px;
    overflow: hidden;
    border: 1px solid #e2e8f0;
}

div[data-testid="stAlert"] {
    border-radius: 14px;
}

.streamlit-expanderHeader {
    border-radius: 12px;
    font-weight: 700;
}

hr {
    margin-top: 1.1rem;
    margin-bottom: 1.2rem;
    border-color: #e2e8f0;
}



/* Ajuste de largura da barra lateral para nomes longos no menu */
section[data-testid="stSidebar"] {
    min-width: 360px !important;
    max-width: 360px !important;
    width: 360px !important;
}

section[data-testid="stSidebar"] > div {
    min-width: 360px !important;
    max-width: 360px !important;
    width: 360px !important;
}

section[data-testid="stSidebar"] label,
section[data-testid="stSidebar"] p,
section[data-testid="stSidebar"] span {
    white-space: normal !important;
    overflow-wrap: normal !important;
    word-break: normal !important;
}

section[data-testid="stSidebar"] [role="radiogroup"] label {
    padding-top: 4px !important;
    padding-bottom: 4px !important;
}

#MainMenu {visibility: hidden;}
footer {visibility: hidden;}
header {visibility: hidden;}
</style>
""",
    unsafe_allow_html=True,
)


# ======================================================
# BANCO DE DADOS
# ======================================================
# ======================================================
# CONEXÃO COM BANCO DE DADOS
# - No Render: usa PostgreSQL pela variável DATABASE_URL
# - Localmente: se não houver DATABASE_URL, usa SQLite para testes
# ======================================================
try:
    import psycopg2
except Exception:
    psycopg2 = None

DATABASE_URL = os.getenv("DATABASE_URL")
USAR_POSTGRES = bool(DATABASE_URL)

if psycopg2 is not None:
    INTEGRITY_ERRORS = (sqlite3.IntegrityError, psycopg2.IntegrityError)
else:
    INTEGRITY_ERRORS = (sqlite3.IntegrityError,)


class CursorPostgresCompat:
    """Pequena camada de compatibilidade para reaproveitar o código escrito para SQLite.
    Converte placeholders ? para %s e AUTOINCREMENT para SERIAL.
    """
    def __init__(self, cursor):
        self.cursor = cursor

    def execute(self, query, params=None):
        query = query.replace("SERIAL PRIMARY KEY", "SERIAL PRIMARY KEY")
        query = query.replace("?", "%s")
        if params is None:
            return self.cursor.execute(query)
        return self.cursor.execute(query, params)

    def fetchone(self):
        return self.cursor.fetchone()

    def fetchall(self):
        return self.cursor.fetchall()

    def close(self):
        return self.cursor.close()


class ConexaoPostgresCompat:
    def __init__(self, conn):
        self.conn = conn

    def cursor(self):
        return CursorPostgresCompat(self.conn.cursor())

    def commit(self):
        return self.conn.commit()

    def rollback(self):
        return self.conn.rollback()

    def close(self):
        return self.conn.close()


def conectar():
    if USAR_POSTGRES:
        if psycopg2 is None:
            raise RuntimeError("psycopg2-binary não está instalado. Adicione psycopg2-binary ao requirements.txt")
        return ConexaoPostgresCompat(psycopg2.connect(DATABASE_URL))
    return sqlite3.connect(DB_PATH)


def obter_conexao_pandas(conn):
    """Retorna a conexão real para uso com pandas.
    No PostgreSQL, a conexão do app é envolvida por ConexaoPostgresCompat;
    o pandas precisa receber a conexão psycopg2 original.
    """
    return conn.conn if hasattr(conn, "conn") else conn


def listar_tabelas_banco():
    """Lista tabelas do banco atual com compatibilidade SQLite e PostgreSQL.
    Evita o uso de sqlite_master quando o sistema está no Render/PostgreSQL.
    """
    conn = conectar()
    conn_pandas = obter_conexao_pandas(conn)
    if USAR_POSTGRES:
        query = """
        SELECT table_name AS name
        FROM information_schema.tables
        WHERE table_schema = 'public'
          AND table_type = 'BASE TABLE'
        ORDER BY table_name
        """
    else:
        query = "SELECT name FROM sqlite_master WHERE type='table' ORDER BY name"
    df = pd.read_sql_query(query, conn_pandas)
    conn.close()
    return df["name"].tolist()


def carregar_tabela_dataframe(tabela):
    """Carrega uma tabela em DataFrame com conexão compatível com SQLite/PostgreSQL."""
    conn = conectar()
    conn_pandas = obter_conexao_pandas(conn)
    df = pd.read_sql_query(f'SELECT * FROM {tabela}', conn_pandas)
    conn.close()
    return df


def limpar_cache_dados():
    """Limpa os caches de consulta após qualquer gravação/edição/exclusão.
    Isso evita dados desatualizados depois de cadastrar, editar ou excluir registros.
    """
    try:
        st.cache_data.clear()
    except Exception:
        pass


def coluna_existe(cursor, tabela, coluna):
    if USAR_POSTGRES:
        cursor.execute(
            """
            SELECT column_name
            FROM information_schema.columns
            WHERE table_name = %s AND column_name = %s
            """,
            (tabela, coluna),
        )
        return cursor.fetchone() is not None
    cursor.execute(f"PRAGMA table_info({tabela})")
    return coluna in [c[1] for c in cursor.fetchall()]


def adicionar_coluna_se_nao_existe(cursor, tabela, coluna, definicao):
    if not coluna_existe(cursor, tabela, coluna):
        cursor.execute(f"ALTER TABLE {tabela} ADD COLUMN {coluna} {definicao}")


def criar_tabelas():
    conn = conectar()
    cursor = conn.cursor()

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS estudantes (
            id SERIAL PRIMARY KEY,
            codigo TEXT UNIQUE NOT NULL,
            ano_serie TEXT,
            turma TEXT,
            perfil TEXT,
            observacoes TEXT,
            criado_em TEXT
        )
        """
    )
    adicionar_coluna_se_nao_existe(cursor, "estudantes", "turno", "TEXT")
    adicionar_coluna_se_nao_existe(cursor, "estudantes", "dias_atendimento", "TEXT")
    adicionar_coluna_se_nao_existe(cursor, "estudantes", "horario_preferencial", "TEXT")

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS professores (
            id SERIAL PRIMARY KEY,
            nome_referencia TEXT,
            escola TEXT,
            regional TEXT,
            formacao TEXT,
            carga_horaria TEXT,
            turno_atuacao TEXT,
            observacoes TEXT,
            criado_em TEXT
        )
        """
    )

    # Perfil pedagógico/tecnológico do professor AEE para personalizar as sugestões da IA.
    for coluna, definicao in [
        ("areas_interesse", "TEXT"),
        ("nivel_tecnologico", "TEXT"),
        ("modo_maker", "TEXT"),
        ("interesse_formacao_maker", "TEXT"),
        ("projetos_interesse", "TEXT"),
        ("preferencias_metodologicas", "TEXT"),
        ("recursos_professor", "TEXT"),
        ("recursos_professor_uso", "TEXT"),
        ("recursos_professor_observacoes", "TEXT"),
    ]:
        adicionar_coluna_se_nao_existe(cursor, "professores", coluna, definicao)


    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS estudante_professor (
            id SERIAL PRIMARY KEY,
            estudante_id INTEGER NOT NULL,
            professor_id INTEGER NOT NULL,
            criado_em TEXT,
            FOREIGN KEY(estudante_id) REFERENCES estudantes(id),
            FOREIGN KEY(professor_id) REFERENCES professores(id),
            UNIQUE(estudante_id, professor_id)
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS entrevistas_familia (
            id SERIAL PRIMARY KEY,
            estudante_id INTEGER NOT NULL,
            data_registro TEXT,
            rotina TEXT,
            saude TEXT,
            comunicacao TEXT,
            autonomia TEXT,
            socializacao TEXT,
            interesses TEXT,
            observacoes TEXT,
            FOREIGN KEY(estudante_id) REFERENCES estudantes(id)
        )
        """
    )

    # Campos ampliados da Entrevista com a Família
    # Mantém compatibilidade com bancos antigos e adiciona as novas colunas automaticamente.
    for coluna, definicao in [
        # Controle histórico da entrevista familiar por ano letivo
        ("ano_letivo", "TEXT"),
        ("tipo_registro", "TEXT"),
        ("auxilio_governamental", "TEXT"),
        ("auxilio_quais", "TEXT"),
        ("historico_familiar", "TEXT"),
        ("historico_quais", "TEXT"),
        ("repetiu_ano", "TEXT"),
        ("repetiu_qtd", "TEXT"),
        ("trocou_escola", "TEXT"),
        ("trocou_qtd", "TEXT"),
        ("motivo_troca", "TEXT"),
        ("situacao_escolar", "TEXT"),
        ("interesse_escola", "TEXT"),
        ("organiza_materiais", "TEXT"),
        ("resistencia_escola", "TEXT"),
        ("relacao_colegas", "TEXT"),
        ("relacao_professores", "TEXT"),
        ("leva_alimentacao", "TEXT"),
        ("merenda_escolar", "TEXT"),
        ("alergia_alimentar", "TEXT"),
        ("alergia_quais", "TEXT"),
        ("obs_diversas", "TEXT"),
        ("motivo_escolha", "TEXT"),
        ("outros_motivos", "TEXT"),
        ("conhecimento_aee", "TEXT"),
        ("doenca_preexistente", "TEXT"),
        ("convulsoes", "TEXT"),
        ("acompanhamentos", "TEXT"),
        ("acompanhamento_outro", "TEXT"),
        ("frequencia_acompanhamento", "TEXT"),
        ("frequencia_outro", "TEXT"),
        ("alimentacao_saudavel", "TEXT"),
        ("seletividade_alimentar", "TEXT"),
        ("dieta_sensorial", "TEXT"),
        ("suplemento_alimentar", "TEXT"),
        ("suplemento_qual", "TEXT"),
        ("alimenta_sonda", "TEXT"),
        ("dorme_bem", "TEXT"),
        ("medicacao", "TEXT"),
        ("medicacao_qual", "TEXT"),
        ("tempo_medicacao_tratamentos", "TEXT"),
        ("obs_saude", "TEXT"),
        ("lateralidade", "TEXT"),
        ("estereotipias", "TEXT"),
        ("estereotipias_quais", "TEXT"),
        ("segura_objetos_duas_maos", "TEXT"),
        ("tamanho_objetos", "TEXT"),
        ("pega_lapis", "TEXT"),
        ("engatinhou", "TEXT"),
        ("idade_andou", "TEXT"),
        ("usa_fraldas", "TEXT"),
        ("usa_sonda_alivio", "TEXT"),
        ("autonomia_atividades", "TEXT"),
        ("autonomia_outros", "TEXT"),
        ("atende_comandos", "TEXT"),
        ("gosta_toque", "TEXT"),
        ("obs_psicomotor", "TEXT"),
        ("verbal", "TEXT"),
        ("consegue_comunicar", "TEXT"),
        ("problemas_fala", "TEXT"),
        ("ecolalia", "TEXT"),
        ("da_recado", "TEXT"),
        ("comunicacao_alternativa", "TEXT"),
        ("comunicacao_alternativa_qual", "TEXT"),
        ("relacao_pai", "TEXT"),
        ("relacao_mae", "TEXT"),
        ("relacao_parentes", "TEXT"),
        ("relacao_irmaos", "TEXT"),
        ("relacao_estudantes", "TEXT"),
        ("tem_melhor_amigo", "TEXT"),
        ("tipo_melhor_amigo", "TEXT"),
        ("adapta_ambiente", "TEXT"),
        ("flexivel_rotina", "TEXT"),
        ("respeita_regras", "TEXT"),
        ("chora_facilidade", "TEXT"),
        ("brinca_como", "TEXT"),
        ("interesses_lazer", "TEXT"),
        ("familia_gosta", "TEXT"),
        ("familia_nao_gosta", "TEXT"),
        ("ambiente_estudo_casa", "TEXT"),
        ("habilidades", "TEXT"),
        ("oportunidades_melhoria", "TEXT"),
        ("outras_info_familia", "TEXT"),
    ]:
        adicionar_coluna_se_nao_existe(cursor, "entrevistas_familia", coluna, definicao)

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS avaliacoes (
            id SERIAL PRIMARY KEY,
            estudante_id INTEGER NOT NULL,
            data_registro TEXT,
            barreiras TEXT,
            potencialidades TEXT,
            comunicacao TEXT,
            interacao TEXT,
            autonomia TEXT,
            aprendizagem TEXT,
            resumo_laudo TEXT,
            FOREIGN KEY(estudante_id) REFERENCES estudantes(id)
        )
        """
    )

    # Controle histórico por ano letivo e apoio da IA para Avaliação Pedagógica.
    # Mantém compatibilidade com bancos antigos e adiciona as novas colunas automaticamente.
    for coluna, definicao in [
        ("ano_letivo", "TEXT"),
        ("tipo_registro", "TEXT"),
        ("avaliacao_anterior_id", "INTEGER"),
        ("analise_comparativa_ia", "TEXT"),
        ("sugestao_nova_avaliacao_ia", "TEXT"),
        ("origem_documento", "TEXT"),
        ("texto_documento_extra", "TEXT"),
    ]:
        adicionar_coluna_se_nao_existe(cursor, "avaliacoes", coluna, definicao)

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS avaliacoes_documentos (
            id SERIAL PRIMARY KEY,
            avaliacao_id INTEGER,
            estudante_id INTEGER NOT NULL,
            ano_letivo TEXT,
            tipo_documento TEXT,
            nome_original TEXT,
            nome_arquivo_salvo TEXT,
            caminho_arquivo TEXT,
            observacao TEXT,
            data_upload TEXT,
            FOREIGN KEY(estudante_id) REFERENCES estudantes(id),
            FOREIGN KEY(avaliacao_id) REFERENCES avaliacoes(id)
        )
        """
    )


    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS escutas_docentes (
            id SERIAL PRIMARY KEY,
            estudante_id INTEGER NOT NULL,
            data_registro TEXT,
            ano_letivo TEXT,
            professor_nome TEXT,
            codigo_docente TEXT,
            componente_curricular TEXT,
            outro_componente TEXT,
            turma TEXT,
            tempo_acompanhamento TEXT,
            participacao_sala TEXT,
            comunicacao TEXT,
            interacao_social TEXT,
            aprendizagem TEXT,
            barreiras_percebidas TEXT,
            potencialidades_observadas TEXT,
            estrategias_funcionam TEXT,
            adaptacoes_utilizadas TEXT,
            recomendacoes_docente TEXT,
            nivel_participacao INTEGER DEFAULT 5,
            nivel_autonomia INTEGER DEFAULT 5,
            nivel_engajamento INTEGER DEFAULT 5,
            observacoes TEXT,
            FOREIGN KEY(estudante_id) REFERENCES estudantes(id)
        )
        """
    )

    adicionar_coluna_se_nao_existe(cursor, "escutas_docentes", "codigo_docente", "TEXT")

    # Perfil de atuação do docente da sala regular para orientar relatórios mais realistas.
    for coluna, definicao in [
        ("areas_interesse_docente", "TEXT"),
        ("nivel_tecnologico_docente", "TEXT"),
        ("modo_maker_docente", "TEXT"),
        ("interesse_formacao_maker_docente", "TEXT"),
        ("interesse_projetos_interdisciplinares", "TEXT"),
        ("projeto_interesse_docente", "TEXT"),
        ("preferencias_metodologicas_docente", "TEXT"),
    ]:
        adicionar_coluna_se_nao_existe(cursor, "escutas_docentes", coluna, definicao)


    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS relatorios_docente (
            id SERIAL PRIMARY KEY,
            estudante_id INTEGER NOT NULL,
            data_geracao TEXT,
            ano_letivo TEXT,
            componente_destino TEXT,
            professor_destino TEXT,
            titulo TEXT,
            conteudo TEXT,
            fontes_utilizadas TEXT,
            observacoes TEXT,
            FOREIGN KEY(estudante_id) REFERENCES estudantes(id)
        )
        """
    )


    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS relatorios_visuais_docente (
            id SERIAL PRIMARY KEY,
            estudante_id INTEGER NOT NULL,
            data_geracao TEXT,
            ano_letivo TEXT,
            componente_destino TEXT,
            tipo_relatorio TEXT,
            titulo TEXT,
            nome_arquivo TEXT,
            caminho_arquivo TEXT,
            fontes_utilizadas TEXT,
            observacoes TEXT,
            FOREIGN KEY(estudante_id) REFERENCES estudantes(id)
        )
        """
    )

    for coluna, definicao in [
        ("ano_letivo", "TEXT"),
        ("componente_destino", "TEXT"),
        ("tipo_relatorio", "TEXT"),
        ("titulo", "TEXT"),
        ("nome_arquivo", "TEXT"),
        ("caminho_arquivo", "TEXT"),
        ("fontes_utilizadas", "TEXT"),
        ("observacoes", "TEXT"),
    ]:
        adicionar_coluna_se_nao_existe(cursor, "relatorios_visuais_docente", coluna, definicao)

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS estudos_caso (
            id SERIAL PRIMARY KEY,
            estudante_id INTEGER NOT NULL,
            data_registro TEXT,
            contextualizacao TEXT,
            queixa_principal TEXT,
            potencialidades TEXT,
            dificuldades TEXT,
            estrategias TEXT,
            intervencoes TEXT,
            avaliacao TEXT,
            consideracoes TEXT,
            FOREIGN KEY(estudante_id) REFERENCES estudantes(id)
        )
        """
    )

    # Campos ampliados do Estudo de Caso / Plano AEE conforme modelo GRE.
    # Não armazenamos CPF, RG, endereço ou laudo médico completo; esses campos ficam em branco nos documentos finais.
    # Armazenamos apenas CID e síntese pedagógica/funcional do laudo, quando pertinente ao AEE.
    for coluna, definicao in [
        ("etapa_modalidade", "TEXT"),
        ("ano_etapa", "TEXT"),
        ("laudo", "TEXT"),
        ("deficiencia", "TEXT"),
        ("cid", "TEXT"),
        ("observacoes_laudo", "TEXT"),
        ("altas_habilidades", "TEXT"),
        ("bpc", "TEXT"),
        ("escola_nome", "TEXT"),
        ("unidade_aee", "TEXT"),
        ("gestor_nome", "TEXT"),
        ("gestor_contato", "TEXT"),
        ("professor_nome", "TEXT"),
        ("professor_contato", "TEXT"),
        ("matricula_professor", "TEXT"),
        ("especialidade_professor", "TEXT"),
        ("periodo_inicio", "TEXT"),
        ("periodo_fim", "TEXT"),
        ("frequencia_atendimento", "TEXT"),
        ("tempo_atendimento_semana", "TEXT"),
        ("formato_atendimento", "TEXT"),
        ("percurso_educacional", "TEXT"),
        ("motivo_encaminhamento_aee", "TEXT"),
        ("precisa_transporte_inclusivo", "TEXT"),
        ("recebe_transporte_inclusivo", "TEXT"),
        ("precisa_profissional_apoio", "TEXT"),
        ("justificativa_apoio", "TEXT"),
        ("acompanhado_profissional_apoio", "TEXT"),
        ("nome_profissional_apoio", "TEXT"),
        ("recursos_tecnologia_assistiva", "TEXT"),
        ("observacoes_ambiente_educacional", "TEXT"),
        ("habilidades_observadas", "TEXT"),
        ("habilidades_a_desenvolver", "TEXT"),
        ("indicadores_altas_habilidades", "TEXT"),
        ("recursos_surdez", "TEXT"),
        ("observacoes_surdez", "TEXT"),
        # Controle histórico por ano letivo e apoio da IA para novo estudo de caso
        ("ano_letivo", "TEXT"),
        ("tipo_registro", "TEXT"),
        ("estudo_anterior_id", "INTEGER"),
        ("analise_comparativa_ia", "TEXT"),
        ("sugestao_novo_estudo_ia", "TEXT"),
    ]:
        adicionar_coluna_se_nao_existe(cursor, "estudos_caso", coluna, definicao)

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS planos_aee (
            id SERIAL PRIMARY KEY,
            estudante_id INTEGER NOT NULL,
            data_registro TEXT,
            habilidades_prioritarias TEXT,
            recursos_acessibilidade TEXT,
            objetivos_gerais TEXT,
            objetivos_especificos TEXT,
            metodologia TEXT,
            estrategias TEXT,
            prazo TEXT,
            acoes_escola TEXT,
            barreiras_identificadas TEXT,
            parcerias TEXT,
            avaliacao_acompanhamento TEXT,
            observacoes TEXT,
            FOREIGN KEY(estudante_id) REFERENCES estudantes(id)
        )
        """
    )

    # Campos completos da Parte 2 - Plano de Atendimento Educacional Especializado (modelo GRE).
    # Mantém compatibilidade com bancos antigos, incluindo instalações já publicadas no Render/PostgreSQL.
    for coluna, definicao in [
        ("habilidades_prioritarias", "TEXT"),
        ("recursos_acessibilidade", "TEXT"),
        ("objetivos_gerais", "TEXT"),
        ("objetivos_especificos", "TEXT"),
        ("metodologia", "TEXT"),
        ("estrategias", "TEXT"),
        ("prazo", "TEXT"),
        ("acoes_escola", "TEXT"),
        ("barreiras_identificadas", "TEXT"),
        ("parcerias", "TEXT"),
        ("avaliacao_acompanhamento", "TEXT"),
        ("observacoes", "TEXT"),
    ]:
        adicionar_coluna_se_nao_existe(cursor, "planos_aee", coluna, definicao)

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS paees (
            id SERIAL PRIMARY KEY,
            estudante_id INTEGER NOT NULL,
            data_geracao TEXT,
            conteudo TEXT,
            FOREIGN KEY(estudante_id) REFERENCES estudantes(id)
        )
        """
    )


    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS plano_aee_ia (
            id SERIAL PRIMARY KEY,
            estudante_id INTEGER NOT NULL,
            data_geracao TEXT,
            mes_referencia TEXT,
            ano_referencia TEXT,
            qtd_atendimentos_semana INTEGER DEFAULT 1,
            tipo_geracao TEXT,
            diagnostico_ia TEXT,
            sugestao_geral TEXT,
            objetivos_prioritarios TEXT,
            recursos_sugeridos TEXT,
            estrategias_recomendadas TEXT,
            plano_mensal TEXT,
            sugestoes_semanais TEXT,
            observacoes TEXT,
            FOREIGN KEY(estudante_id) REFERENCES estudantes(id)
        )
        """
    )

    # Histórico do módulo Plano AEE - IA: sugestão geral, plano mensal e evolução.
    for coluna, definicao in [
        ("mes_referencia", "TEXT"),
        ("ano_referencia", "TEXT"),
        ("qtd_atendimentos_semana", "INTEGER DEFAULT 1"),
        ("tipo_geracao", "TEXT"),
        ("diagnostico_ia", "TEXT"),
        ("sugestao_geral", "TEXT"),
        ("objetivos_prioritarios", "TEXT"),
        ("recursos_sugeridos", "TEXT"),
        ("estrategias_recomendadas", "TEXT"),
        ("plano_mensal", "TEXT"),
        ("sugestoes_semanais", "TEXT"),
        ("observacoes", "TEXT"),
    ]:
        adicionar_coluna_se_nao_existe(cursor, "plano_aee_ia", coluna, definicao)

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS relatorios (
            id SERIAL PRIMARY KEY,
            estudante_id INTEGER NOT NULL,
            data_geracao TEXT,
            titulo TEXT,
            conteudo TEXT,
            FOREIGN KEY(estudante_id) REFERENCES estudantes(id)
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS atendimentos (
            id SERIAL PRIMARY KEY,
            estudante_id INTEGER NOT NULL,
            data_atendimento TEXT,
            objetivo TEXT,
            atividade TEXT,
            resposta_estudante TEXT,
            avancos TEXT,
            dificuldades TEXT,
            evolucao TEXT,
            qtd_atividades INTEGER DEFAULT 1,
            nivel_resposta INTEGER DEFAULT 5,
            nivel_avanco INTEGER DEFAULT 5,
            nivel_dificuldade INTEGER DEFAULT 5,
            nivel_engajamento INTEGER DEFAULT 5,
            nivel_evolucao REAL DEFAULT 5,
            encaminhamentos TEXT,
            FOREIGN KEY(estudante_id) REFERENCES estudantes(id)
        )
        """
    )

    for coluna, definicao in [
        ("evolucao", "TEXT"),
        ("qtd_atividades", "INTEGER DEFAULT 1"),
        ("nivel_resposta", "INTEGER DEFAULT 5"),
        ("nivel_avanco", "INTEGER DEFAULT 5"),
        ("nivel_dificuldade", "INTEGER DEFAULT 5"),
        ("nivel_engajamento", "INTEGER DEFAULT 5"),
        ("nivel_evolucao", "REAL DEFAULT 5"),
        ("encaminhamentos", "TEXT"),
    ]:
        adicionar_coluna_se_nao_existe(cursor, "atendimentos", coluna, definicao)

    # Recursos pedagógicos vinculados a cada atendimento.
    # Armazena apenas nomes, categorias, links e observações do material usado,
    # sem upload de imagens ou registro visual do estudante.
    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS recursos_atendimento (
            id SERIAL PRIMARY KEY,
            atendimento_id INTEGER NOT NULL,
            estudante_id INTEGER NOT NULL,
            nome_recurso TEXT,
            categoria_recurso TEXT,
            link_recurso TEXT,
            observacao_uso TEXT,
            criado_em TEXT,
            FOREIGN KEY(estudante_id) REFERENCES estudantes(id),
            FOREIGN KEY(atendimento_id) REFERENCES atendimentos(id)
        )
        """
    )

    for coluna, definicao in [
        ("atendimento_id", "INTEGER"),
        ("estudante_id", "INTEGER"),
        ("nome_recurso", "TEXT"),
        ("categoria_recurso", "TEXT"),
        ("link_recurso", "TEXT"),
        ("observacao_uso", "TEXT"),
        ("criado_em", "TEXT"),
    ]:
        adicionar_coluna_se_nao_existe(cursor, "recursos_atendimento", coluna, definicao)

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS agenda (
            id SERIAL PRIMARY KEY,
            estudante_id INTEGER NOT NULL,
            data_agendamento TEXT,
            dia_semana TEXT,
            hora_inicio TEXT,
            hora_fim TEXT,
            tipo_atendimento TEXT,
            observacoes TEXT,
            criado_em TEXT,
            FOREIGN KEY(estudante_id) REFERENCES estudantes(id)
        )
        """
    )

    # Campos ampliados da Agenda para controle de presença e integração com o Quadro Semanal GRE.
    # Mantém compatibilidade com bancos antigos.
    for coluna, definicao in [
        ("status_presenca", "TEXT DEFAULT 'Agendado'"),
        ("atendimento_id", "INTEGER"),
        ("preenchimento_ia", "TEXT"),
    ]:
        adicionar_coluna_se_nao_existe(cursor, "agenda", coluna, definicao)


    # Banco de Recursos Pedagógicos e Tecnologia Assistiva da Escola.
    # Permite cadastrar ou importar recursos reais disponíveis na unidade escolar
    # para que a IA gere sugestões contextualizadas ao que a escola possui.
    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS recursos_escola (
            id SERIAL PRIMARY KEY,
            escola_nome TEXT,
            nome_recurso TEXT NOT NULL,
            categoria TEXT,
            descricao TEXT,
            quantidade INTEGER DEFAULT 1,
            localizacao TEXT,
            publico_indicado TEXT,
            objetivo_pedagogico TEXT,
            status TEXT,
            origem TEXT,
            link_referencia TEXT,
            observacoes TEXT,
            criado_em TEXT
        )
        """
    )

    for coluna, definicao in [
        ("escola_nome", "TEXT"),
        ("nome_recurso", "TEXT"),
        ("categoria", "TEXT"),
        ("descricao", "TEXT"),
        ("quantidade", "INTEGER DEFAULT 1"),
        ("localizacao", "TEXT"),
        ("publico_indicado", "TEXT"),
        ("objetivo_pedagogico", "TEXT"),
        ("status", "TEXT"),
        ("origem", "TEXT"),
        ("link_referencia", "TEXT"),
        ("observacoes", "TEXT"),
        ("criado_em", "TEXT"),
    ]:
        adicionar_coluna_se_nao_existe(cursor, "recursos_escola", coluna, definicao)


    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS documentos_gre_gerados (
            id SERIAL PRIMARY KEY,
            estudante_id INTEGER,
            tipo_documento TEXT,
            nome_arquivo TEXT,
            data_geracao TEXT,
            observacao TEXT,
            FOREIGN KEY(estudante_id) REFERENCES estudantes(id)
        )
        """
    )

    conn.commit()
    conn.close()
    limpar_cache_dados()


# Cria/atualiza as tabelas apenas uma vez por sessão do Streamlit.
# Isso reduz o tempo de troca entre telas, principalmente usando PostgreSQL.
if "tabelas_criadas" not in st.session_state:
    criar_tabelas()
    st.session_state.tabelas_criadas = True


# ======================================================
# FUNÇÕES UTILITÁRIAS
# ======================================================
PERFIS = [
    "Não informado",
    "Deficiência intelectual",
    "Deficiência visual",
    "Deficiência auditiva/surdez",
    "Deficiência física",
    "TEA",
    "TEA - Nível I",
    "TEA - Nível II",
    "TEA - Nível III",
    "Altas habilidades/superdotação",
    "Deficiência múltipla",
    "Outro",
]

DIAS_SEMANA = [
    "Segunda-feira",
    "Terça-feira",
    "Quarta-feira",
    "Quinta-feira",
    "Sexta-feira",
]

MAPA_MESES = {
    "Janeiro": 1,
    "Fevereiro": 2,
    "Março": 3,
    "Abril": 4,
    "Maio": 5,
    "Junho": 6,
    "Julho": 7,
    "Agosto": 8,
    "Setembro": 9,
    "Outubro": 10,
    "Novembro": 11,
    "Dezembro": 12,
}

MAPA_DIAS_SEMANA = {
    "Segunda-feira": 0,
    "Terça-feira": 1,
    "Quarta-feira": 2,
    "Quinta-feira": 3,
    "Sexta-feira": 4,
}

def gerar_datas_atendimentos_mes(ano, mes_nome, dias_atendimento):
    """Calcula automaticamente as datas reais de atendimento do mês.

    Usa o mês, ano e os dias de atendimento cadastrados/selecionados para gerar
    a quantidade correta de encontros. Ex.: segunda e sexta em um mês podem gerar
    8, 9 ou 10 atendimentos, conforme o calendário.
    """
    try:
        ano_int = int(ano)
    except Exception:
        ano_int = agora_local().year

    mes_num = MAPA_MESES.get(str(mes_nome), agora_local().month)
    dias_codigo = [MAPA_DIAS_SEMANA[d] for d in dias_atendimento if d in MAPA_DIAS_SEMANA]

    if not dias_codigo:
        return []

    total_dias = calendar.monthrange(ano_int, mes_num)[1]
    datas = []
    for dia in range(1, total_dias + 1):
        data_atual = date(ano_int, mes_num, dia)
        if data_atual.weekday() in dias_codigo:
            datas.append(data_atual)
    return datas

def datas_atendimentos_para_texto(datas):
    if not datas:
        return "Nenhuma data calculada. Verifique mês, ano e dias de atendimento."
    linhas = []
    for idx, data_atual in enumerate(datas, start=1):
        nome_dia = DIAS_SEMANA[data_atual.weekday()] if data_atual.weekday() < len(DIAS_SEMANA) else "Dia"
        linhas.append(f"Atendimento {idx}: {data_atual.strftime('%d/%m/%Y')} ({nome_dia})")
    return "\n".join(linhas)


CAMPOS_ENTREVISTA_FAMILIA = [
    "data_registro",
    "ano_letivo", "tipo_registro",
    "auxilio_governamental", "auxilio_quais", "historico_familiar", "historico_quais",
    "repetiu_ano", "repetiu_qtd", "trocou_escola", "trocou_qtd", "motivo_troca",
    "situacao_escolar", "interesse_escola", "organiza_materiais", "resistencia_escola",
    "relacao_colegas", "relacao_professores", "leva_alimentacao", "merenda_escolar",
    "alergia_alimentar", "alergia_quais", "obs_diversas",
    "motivo_escolha", "outros_motivos", "conhecimento_aee",
    "doenca_preexistente", "convulsoes", "acompanhamentos", "acompanhamento_outro",
    "frequencia_acompanhamento", "frequencia_outro", "alimentacao_saudavel",
    "seletividade_alimentar", "dieta_sensorial", "suplemento_alimentar", "suplemento_qual",
    "alimenta_sonda", "dorme_bem", "medicacao", "medicacao_qual",
    "tempo_medicacao_tratamentos", "obs_saude",
    "lateralidade", "estereotipias", "estereotipias_quais", "segura_objetos_duas_maos",
    "tamanho_objetos", "pega_lapis", "engatinhou", "idade_andou",
    "usa_fraldas", "usa_sonda_alivio", "autonomia_atividades", "autonomia_outros",
    "atende_comandos", "gosta_toque", "obs_psicomotor",
    "verbal", "consegue_comunicar", "problemas_fala", "ecolalia", "da_recado",
    "comunicacao_alternativa", "comunicacao_alternativa_qual",
    "relacao_pai", "relacao_mae", "relacao_parentes", "relacao_irmaos", "relacao_estudantes",
    "tem_melhor_amigo", "tipo_melhor_amigo", "adapta_ambiente", "flexivel_rotina",
    "respeita_regras", "chora_facilidade", "brinca_como", "interesses_lazer",
    "familia_gosta", "familia_nao_gosta", "ambiente_estudo_casa",
    "habilidades", "oportunidades_melhoria", "outras_info_familia",
]

OPCOES_TIPO_ENTREVISTA_FAMILIA = [
    "Registro histórico",
    "Entrevista familiar atual",
]

CAMPOS_AVALIACAO = [
    "data_registro",
    "ano_letivo",
    "tipo_registro",
    "avaliacao_anterior_id",
    "analise_comparativa_ia",
    "sugestao_nova_avaliacao_ia",
    "origem_documento",
    "texto_documento_extra",
    "barreiras",
    "potencialidades",
    "comunicacao",
    "interacao",
    "autonomia",
    "aprendizagem",
    "resumo_laudo",
]

OPCOES_TIPO_AVALIACAO = [
    "Registro histórico",
    "Avaliação pedagógica atual",
    "Nova avaliação com base em avaliação anterior",
    "Avaliação Pedagógica - Extra / Documento Livre",
]


OPCOES_AREAS_CONHECIMENTO = [
    "Geral – todas as áreas",
    "Matemática",
    "Português",
    "História",
    "Geografia",
    "Ciências",
    "Biologia",
    "Física",
    "Química",
    "Artes",
    "Educação Física",
    "Inglês",
    "Espanhol",
    "Filosofia",
    "Sociologia",
    "Tecnologia / Computação",
    "Projeto de Vida",
    "Sala de Leitura",
    "Outras",
]

OPCOES_BASES_CONHECIMENTO_IA = {
    "cientifica": "Base Científica Inclusiva",
    "pedagogica": "Base Pedagógica AEE",
    "bncc": "BNCC / Currículo por área",
    "bncc_computacao": "BNCC Computação",
    "estrategias": "Estratégias Pedagógicas Inclusivas",
    "tecnologia_assistiva": "Tecnologia Assistiva",
    "tecnologia": "Tecnologia",
    "maker_inclusivo": "Maker Inclusivo / STEAM",
    "legislacao": "Legislação e Normativas",
    "lei": "Lei",
    "projetos_interdisciplinares": "Projetos Interdisciplinares",
    "perfis_pedagogicos": "Perfis Pedagógicos e Barreiras",
}


OPCOES_ESTRATEGIAS_INCLUSIVAS = [
    "Apoio visual",
    "Rotina estruturada",
    "Instruções curtas e objetivas",
    "Tempo ampliado",
    "Atividades concretas/manipuláveis",
    "Mediação gradual",
    "Atividades em dupla ou pequenos grupos",
    "Comunicação alternativa/aumentativa",
    "Tecnologia assistiva",
    "Gamificação",
    "Recursos sensoriais",
    "Recursos de impressão 3D / maker",
    "Leitura compartilhada",
    "Avaliação oral ou por demonstração prática",
    "Redução de cópia extensa",
    "Outro",
]

CAMPOS_ESCUTA_DOCENTE = [
    "data_registro",
    "ano_letivo",
    "codigo_docente",
    "componente_curricular",
    "outro_componente",
    "turma",
    "tempo_acompanhamento",
    "participacao_sala",
    "comunicacao",
    "interacao_social",
    "aprendizagem",
    "barreiras_percebidas",
    "potencialidades_observadas",
    "estrategias_funcionam",
    "adaptacoes_utilizadas",
    "recomendacoes_docente",
    "nivel_participacao",
    "nivel_autonomia",
    "nivel_engajamento",
    "observacoes",
    "areas_interesse_docente",
    "nivel_tecnologico_docente",
    "modo_maker_docente",
    "interesse_formacao_maker_docente",
    "interesse_projetos_interdisciplinares",
    "projeto_interesse_docente",
    "preferencias_metodologicas_docente",
]

CAMPOS_RELATORIO_DOCENTE = [
    "data_geracao",
    "ano_letivo",
    "componente_destino",
    "professor_destino",
    "titulo",
    "conteudo",
    "fontes_utilizadas",
    "observacoes",
]

CAMPOS_RELATORIO_VISUAL_DOCENTE = [
    "data_geracao",
    "ano_letivo",
    "componente_destino",
    "tipo_relatorio",
    "titulo",
    "nome_arquivo",
    "caminho_arquivo",
    "fontes_utilizadas",
    "observacoes",
]

CAMPOS_ESTUDO_CASO = [
    "data_registro",
    "ano_letivo", "tipo_registro", "estudo_anterior_id", "analise_comparativa_ia", "sugestao_novo_estudo_ia",
    "contextualizacao", "queixa_principal", "potencialidades", "dificuldades", "estrategias", "intervencoes", "avaliacao", "consideracoes",
    "etapa_modalidade", "ano_etapa", "laudo", "deficiencia", "cid", "observacoes_laudo", "altas_habilidades", "bpc",
    "escola_nome", "unidade_aee", "gestor_nome", "gestor_contato",
    "professor_nome", "professor_contato", "matricula_professor", "especialidade_professor",
    "periodo_inicio", "periodo_fim", "frequencia_atendimento", "tempo_atendimento_semana", "formato_atendimento",
    "percurso_educacional", "motivo_encaminhamento_aee", "precisa_transporte_inclusivo", "recebe_transporte_inclusivo",
    "precisa_profissional_apoio", "justificativa_apoio", "acompanhado_profissional_apoio", "nome_profissional_apoio",
    "recursos_tecnologia_assistiva", "observacoes_ambiente_educacional",
    "habilidades_observadas", "habilidades_a_desenvolver", "indicadores_altas_habilidades",
    "recursos_surdez", "observacoes_surdez",
]

CAMPOS_PLANO_AEE = [
    "data_registro",
    "habilidades_prioritarias",
    "recursos_acessibilidade",
    "objetivos_gerais",
    "objetivos_especificos",
    "metodologia",
    "estrategias",
    "prazo",
    "acoes_escola",
    "barreiras_identificadas",
    "parcerias",
    "avaliacao_acompanhamento",
    "observacoes",
]

CAMPOS_PLANO_AEE_IA = [
    "data_geracao",
    "mes_referencia",
    "ano_referencia",
    "qtd_atendimentos_semana",
    "tipo_geracao",
    "diagnostico_ia",
    "sugestao_geral",
    "objetivos_prioritarios",
    "recursos_sugeridos",
    "estrategias_recomendadas",
    "plano_mensal",
    "sugestoes_semanais",
    "observacoes",
]

MESES_REFERENCIA = [
    "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
    "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"
]


OPCOES_ANO_ETAPA = [
    # Ensino Fundamental
    "1º ano do EF", "2º ano do EF", "3º ano do EF", "4º ano do EF", "5º ano do EF",
    "6º ano do EF", "7º ano do EF", "8º ano do EF", "9º ano do EF",
    # Ensino Médio
    "1º ano do EM", "2º ano do EM", "3º ano do EM",
    # EJA - Ensino Fundamental
    "EJA Ensino Fundamental - Módulo 5", "EJA Ensino Fundamental - Módulo 6",
    "EJA Ensino Fundamental - Módulo 7", "EJA Ensino Fundamental - Módulo 8",
    # EJA - Ensino Médio
    "EJA Ensino Médio - Módulo 1", "EJA Ensino Médio - Módulo 2", "EJA Ensino Médio - Módulo 3",
    "Outro / não informado",
]

OPCOES_ESPECIALIDADE_AEE = [
    "Professor(a) do AEE",
    "Professor(a) Brailista",
    "Professor(a) Instrutor(a) de LIBRAS",
    "Professor(a) Intérprete de LIBRAS",
]

OPCOES_FORMATO_ATENDIMENTO = [
    "Individual",
    "Coletivo",
    "Em sala de aula",
    "Domiciliar",
    "Hospitalar",
    "Outro",
]

OPCOES_HABILIDADES_PEDAGOGICAS = [
    "Inteligência Interpessoal (interesse pelo sentimento do outro)",
    "Interesse e habilidade em atividades sensoriais (tátil, auditivo ou outros sentidos)",
    "Raciocínio lógico-matemático (interesse por cálculos e números)",
    "Realiza as quatro operações",
    "Resolve situações-problema com autonomia e de maneira diversificada",
    "Interesse em temas científicos",
    "Compreende explicações e cumpre comandos",
    "Demonstra criatividade",
    "Acompanha as atividades propostas do grupo/classe",
    "Usa recursos tecnológicos com autonomia (tablet, celular, computador)",
    "Fluência na leitura",
    "Escreve textos com autonomia",
    "Apropriação do sistema de escrita alfabética adequada à idade",
    "Habilidades artísticas",
    "Comunica seus desejos e necessidades",
    "Apresenta atenção compartilhada",
    "Interesse por leitura",
    "Domina Libras",
    "Domina Braille",
    "Possui identidade surda",
    "Usa comunicação aumentativa e alternativa com autonomia",
    "Outro",
]

OPCOES_HABILIDADES_A_DESENVOLVER = [
    "Inteligência Interpessoal",
    "Atividades sensoriais",
    "Raciocínio lógico-matemático",
    "Resolução de situações-problema com autonomia",
    "Compreensão de explicações e comandos",
    "Criatividade",
    "Utilização de recursos tecnológicos com autonomia",
    "Fluência na leitura",
    "Escrita de textos com autonomia",
    "Apropriação do sistema de escrita alfabética",
    "Expressão de seus desejos e necessidades",
    "Atenção compartilhada",
    "Interesse por leitura",
    "Domínio de Libras",
    "Domínio de Braille",
    "Utilização da comunicação aumentativa e alternativa com autonomia",
    "Outro",
]

OPCOES_INDICADORES_AHSD = [
    "Aprende fácil e rapidamente",
    "Original, imaginativa, criativa, não convencional",
    "Pensa de forma incomum para resolver problemas",
    "Persistente, independente, autodirecionada",
    "Persuasiva, capaz de influenciar os outros",
    "Inquisitiva, cética, curiosa",
    "Adapta-se a uma variedade de situações e novos ambientes",
    "Criativa ao construir com materiais incomuns",
    "Habilidade nas artes (música, desenho, dança etc.)",
    "Entende a importância da natureza (tempo, lua, sol, estrelas, solo etc.)",
    "Vocabulário excepcional, verbalmente fluente",
    "Aprende facilmente novas línguas",
    "Trabalha independente, mostra iniciativa",
    "Bom julgamento, lógica",
    "Usa recursos tecnológicos com autonomia",
    "Versátil, muitos interesses, interesse além da idade cronológica",
    "Mostra insights e percepções incomuns",
    "Demonstra alto nível de sensibilidade e empatia com relação aos outros",
    "Apresenta excelente senso de humor",
    "Expressa ideias e reações, frequentemente de forma argumentativa",
    "Outro",
]

OPCOES_RECURSOS_TA = [
    "Recursos de comunicação alternativa/aumentativa",
    "Pranchas de CAA",
    "Tablet/celular/computador com recurso acessível",
    "Chromebook",
    "Materiais concretos/manipuláveis",
    "Recursos táteis/sensoriais",
    "Jogos pedagógicos acessíveis",
    "Wordwall",
    "Impressão 3D/recurso maker inclusivo",
    "Robótica educacional",
    "Recursos em Libras",
    "Recursos em Braille",
    "Leitor de tela/ampliação",
    "Órteses/adaptações de acesso",
    "Outro",
]

OPCOES_RECURSOS_SURDEZ = [
    "Aparelho auditivo",
    "Implante coclear",
    "Libras",
    "Leitura labial",
    "Intérprete de Libras",
    "Não se aplica",
]

OPCOES_HABILIDADES_PRIORITARIAS_SRM = [
    "Desenvolvimento de vida autônoma",
    "Enriquecimento curricular (para aluno com altas habilidades e superdotação)",
    "Ensino da informática acessível",
    "Ensino da Língua Brasileira de Sinais",
    "Ensino da Língua Portuguesa como segunda língua",
    "Ensino das técnicas de cálculo no soroban",
    "Ensino do sistema Braille",
    "Ensino de técnicas de orientação e mobilidade",
    "Ensino do uso da comunicação aumentativa e alternativa",
    "Ensino do uso de recursos ópticos e não ópticos",
]

OPCOES_RECURSOS_ACESSIBILIDADE = [
    "Recurso de auxílio à escrita",
    "Colmeia para teclado",
    "Teclado expandido",
    "Suporte de punho ajustável",
    "Mouses especiais",
    "Acionadores",
    "Leitor automático",
    "Ampliador de texto",
    "Réguas braille",
    "Multiplano",
    "Reglete",
    "Máquina braille",
    "Pranchas de CAA",
    "Texto apoiado",
    "Textos e livros em Braille",
    "Audiodescrição",
    "Descrição de imagens",
    "Hand Talk",
    "Wordwall",
    "Gemini",
    "Copilot",
    "Tablet",
    "Chromebook",
    "Impressão 3D",
    "Robótica educacional",
    "Outros",
]

OPCOES_ACOES_ESCOLA = [
    "Implemento de tecnologia assistiva para uso do(a) estudante",
    "Formação continuada de professores, profissionais de apoio, equipe gestora, demais funcionários da escola e famílias com a temática da educação inclusiva",
    "Parcerias com profissionais de saúde, UBS, USF, Conselho Tutelar, CREAS e CRAS",
    "Articulação sobre o PDDE Equidade",
    "Articulação de horários para o diálogo entre professor do AEE, professor da sala comum e profissional de apoio",
    "Solicitação de transporte escolar inclusivo",
    "Outros",
]

OPCOES_BARREIRAS = [
    "Barreiras atitudinais",
    "Barreiras arquitetônicas",
    "Barreiras comunicacionais",
    "Barreiras curriculares",
    "Outros",
]

OPCOES_CATEGORIAS_RECURSOS_ESCOLA = [
    "Tecnologia Assistiva",
    "Recurso pedagógico",
    "Comunicação Aumentativa e Alternativa (CAA)",
    "Recurso visual",
    "Recurso tátil/sensorial",
    "Material manipulável",
    "Impressão 3D / Cultura maker",
    "Robótica educacional",
    "Tecnologia digital",
    "Acessibilidade física/motora",
    "Baixa visão / ampliação",
    "Libras / Surdez",
    "Braille",
    "Jogos pedagógicos",
    "Música / expressão artística",
    "Outro",
]

OPCOES_STATUS_RECURSOS_ESCOLA = [
    "Disponível",
    "Disponível com agendamento",
    "Em manutenção",
    "Emprestado",
    "Indisponível",
]

OPCOES_AREAS_INTERESSE_DOCENTE = [
    "Alfabetização",
    "Comunicação Alternativa (CAA)",
    "Recursos visuais",
    "Artes",
    "Música",
    "Jogos pedagógicos",
    "Tecnologia educacional",
    "Cultura maker",
    "Robótica educacional",
    "Impressão 3D",
    "Programação",
    "Ciências",
    "Matemática manipulável",
    "Leitura e escrita",
    "Coordenação motora",
    "Autonomia e vida diária",
    "Ensino médio / projetos interdisciplinares",
    "Atividades desplugadas",
    "Projetos com materiais de baixo custo",
    "Outro",
]

OPCOES_NIVEL_TECNOLOGICO_DOCENTE = [
    "Básico",
    "Intermediário",
    "Avançado",
]

OPCOES_NIVEL_PROJETO_NORTEADOR = [
    "Desplugado / baixa tecnologia",
    "Maker básico",
    "Maker intermediário",
    "Maker avançado",
]

OPCOES_PROJETOS_NORTEADORES = [
    "Carrinho seguidor de linha",
    "Carrinho movido à energia solar",
    "Estação meteorológica simples",
    "Horta inteligente",
    "Robótica simples",
    "Podcast escolar",
    "Impressão 3D / protótipo tátil",
    "Automação residencial simples",
    "Jogo educativo",
    "Ebook interativo",
    "Experimento científico acessível",
    "Projeto livre / personalizado",
]

OPCOES_RECURSOS_PROFESSOR_AEE = [
    "Tablet",
    "Notebook",
    "Chromebook",
    "Celular",
    "Recursos de Comunicação Aumentativa e Alternativa (CAA)",
    "Pranchas de comunicação",
    "Pictogramas / cartões visuais",
    "Materiais impressos em papel",
    "Jogos pedagógicos",
    "Materiais concretos/manipuláveis",
    "Materiais táteis/sensoriais",
    "Caneta 3D",
    "Impressora 3D / peças 3D",
    "Kit de robótica educacional",
    "Arduino / sensores / motores",
    "Recursos de música / som",
    "Outros",
]


def hoje_str():
    """Data/hora local para salvar histórico e documentos."""
    return agora_local().strftime("%d/%m/%Y %H:%M")


def normalizar_data_historico(data_texto):
    """Mantém o texto salvo. Registros antigos em UTC não são alterados automaticamente.
    Novos registros usam hoje_str() em America/Recife.
    """
    return data_texto or "Data não informada"


def formatar_data(data_obj):
    if isinstance(data_obj, (datetime, date)):
        return data_obj.strftime("%d/%m/%Y")
    return str(data_obj)


def data_para_date(data_texto):
    try:
        return datetime.strptime(data_texto, "%d/%m/%Y").date()
    except Exception:
        return agora_local().date()


def hora_para_time(hora_texto, padrao="08:00"):
    try:
        return datetime.strptime(str(hora_texto)[:5], "%H:%M").time()
    except Exception:
        return datetime.strptime(padrao, "%H:%M").time()


def limitar_escala(valor, padrao=5):
    try:
        valor = int(valor)
    except Exception:
        valor = padrao
    return max(1, min(10, valor))


def calcular_indice_geral(nivel_resposta, nivel_avanco, nivel_dificuldade, nivel_engajamento):
    resposta = limitar_escala(nivel_resposta)
    avanco = limitar_escala(nivel_avanco)
    dificuldade = limitar_escala(nivel_dificuldade)
    engajamento = limitar_escala(nivel_engajamento)
    barreira_invertida = 11 - dificuldade
    indice = (resposta * 0.30) + (avanco * 0.35) + (engajamento * 0.20) + (barreira_invertida * 0.15)
    return round(max(1, min(10, indice)), 1)


def interpretar_indice(indice):
    if indice <= 3:
        return "Baixa evolução / necessidade de maior apoio"
    if indice <= 6:
        return "Evolução parcial / acompanhamento em desenvolvimento"
    if indice <= 8:
        return "Boa evolução / resposta positiva"
    return "Evolução elevada / maior autonomia e participação"


def opcoes_estudantes_por_id(estudantes):
    ids = [e[0] for e in estudantes]
    mapa = {e[0]: f"{e[1]} - {e[2] or 'Ano/Série não informado'} - {e[4] or 'Perfil não informado'}" for e in estudantes}
    return ids, mapa


def render_app_header():
    st.markdown(
        f"""
        <div class="app-hero">
            <span class="app-badge">AEE • Memória Pedagógica • Articulação Docente • IA • {APP_VERSION} {APP_VERSION_LABEL}</span>
            <h1 class="app-title">INCLUISRM</h1>
            <p class="app-subtitle">Sistema Inteligente de Articulação Pedagógica Inclusiva</p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def gerar_docx_documento(conteudo, nome_base, tipo="documento"):
    """Gera documento Word editável com os mesmos textos do PDF.
    Observação: dados sensíveis permanecem como linhas em branco para preenchimento manual.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    nome_arquivo = f"{tipo}_{nome_base}.docx".replace("/", "-").replace("\\", "-")
    doc = Document()

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("INCLUISRM\nSistema Inteligente de Articulação Pedagógica Inclusiva")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph("")

    for linha in conteudo.split("\n"):
        linha = linha.rstrip()
        if not linha:
            doc.add_paragraph("")
            continue

        par = doc.add_paragraph()
        texto_limpo = linha.strip()
        run = par.add_run(texto_limpo)

        if texto_limpo.isupper() or (len(texto_limpo) > 2 and texto_limpo[:2].isdigit() and "." in texto_limpo[:4]):
            run.bold = True
        run.font.size = Pt(11)

    doc.add_paragraph("")
    rodape = doc.add_paragraph(f"Gerado em {agora_local().strftime('%d/%m/%Y %H:%M')} pelo INCLUISRM.")
    rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in rodape.runs:
        r.font.size = Pt(9)

    doc.save(nome_arquivo)
    return nome_arquivo



# ======================================================
# RELATÓRIOS VISUAIS - PLANO AEE IA
# ======================================================
def limpar_marcadores_relatorio(texto):
    """Remove marcações simples de Markdown usadas pela IA para melhorar Word/PDF."""
    texto = str(texto or "")
    texto = texto.replace("**", "")
    texto = re.sub(r"^#{1,6}\s*", "", texto, flags=re.MULTILINE)
    texto = texto.replace("---", "")
    return texto.strip()


def eh_titulo_relatorio(linha):
    linha = linha.strip()
    if not linha:
        return False
    if linha.upper() == linha and len(linha) > 8:
        return True
    if re.match(r"^(\d+\.|SEMANA\s+\d+|ATENDIMENTO\s+\d+|Semana\s+\d+|Atendimento\s+\d+)", linha, flags=re.I):
        return True
    palavras_chave = [
        "Identificação", "Objetivo", "Habilidades", "Recursos", "Estratégias", "Avaliação",
        "Indicadores", "Ajustes", "Registro", "Organização", "Diagnóstico", "Sugestão"
    ]
    return any(linha.lower().startswith(k.lower()) for k in palavras_chave)


def extrair_linhas_chave_valor(texto):
    pares = []
    for raw in str(texto or "").splitlines():
        linha = limpar_marcadores_relatorio(raw).strip(" -•\t")
        if ":" in linha and len(linha) < 140:
            chave, valor = linha.split(":", 1)
            if chave and valor and len(chave) <= 45:
                pares.append((chave.strip(), valor.strip()))
    return pares


def obter_tema_documento(tipo="documento", titulo_doc="", conteudo=""):
    """Define cor e explicação do papel de cada documento pedagógico."""
    tipo = str(tipo or "documento").lower()
    titulo = str(titulo_doc or "").lower()
    conteudo_txt = str(conteudo or "").lower()[:1500]

    temas = {
        "perfil": {
            "cor": "#1d4ed8",
            "cor_clara": "#dbeafe",
            "titulo": "Perfil Pedagógico Inteligente",
            "papel": "Este documento apoia o professor do AEE na leitura pedagógica inicial do estudante, organizando potencialidades, barreiras, necessidades educacionais e recursos de acessibilidade que podem orientar o planejamento das intervenções. A IA não substitui o professor: ela organiza informações e sugere caminhos pedagógicos para análise profissional.",
        },
        "mensal": {
            "cor": "#15803d",
            "cor_clara": "#dcfce7",
            "titulo": "Plano Mensal AEE - IA",
            "papel": "Este documento organiza as ações mensais do Atendimento Educacional Especializado, indicando objetivos, atividades, recursos, estratégias de mediação e formas de acompanhamento. Sua finalidade é apoiar o registro contínuo da evolução pedagógica do estudante e orientar os atendimentos na SRM.",
        },
        "paee": {
            "cor": "#7c3aed",
            "cor_clara": "#ede9fe",
            "titulo": "Plano AEE / PAEE",
            "papel": "Este documento consolida o planejamento educacional especializado do estudante, registrando objetivos, habilidades prioritárias, recursos de acessibilidade, metodologia, estratégias, barreiras, parcerias e avaliação. Ele serve como referência para o acompanhamento pedagógico e para a articulação entre AEE, sala regular, gestão e família.",
        },
        "docente": {
            "cor": "#ea580c",
            "cor_clara": "#ffedd5",
            "titulo": "Relatório Pedagógico de Apoio ao Docente",
            "papel": "Este documento orienta o professor da sala regular com estratégias práticas, adaptações, recursos e recomendações pedagógicas objetivas, favorecendo a participação do estudante nas atividades curriculares e fortalecendo a articulação com o AEE.",
        },
        "gre": {
            "cor": "#334155",
            "cor_clara": "#e2e8f0",
            "titulo": "Relatório GRE",
            "papel": "Este documento reúne informações pedagógicas e registros relevantes para fins de acompanhamento institucional, articulação com a gestão escolar e encaminhamentos à GRE quando necessário.",
        },
        "avaliacao": {
            "cor": "#0f766e",
            "cor_clara": "#ccfbf1",
            "titulo": "Avaliação Pedagógica",
            "papel": "Este documento registra observações pedagógicas iniciais, barreiras, potencialidades, autonomia, comunicação, interação e aprendizagem, subsidiando a elaboração do estudo de caso e do Plano AEE.",
        },
        "estudo": {
            "cor": "#2563eb",
            "cor_clara": "#dbeafe",
            "titulo": "Estudo de Caso",
            "papel": "Este documento sistematiza o percurso educacional, necessidades, barreiras, potencialidades e estratégias observadas, servindo de base para a elaboração do Plano AEE e para a tomada de decisões pedagógicas.",
        },
        "documento": {
            "cor": "#0f172a",
            "cor_clara": "#f1f5f9",
            "titulo": "Documento Pedagógico",
            "papel": "Este documento organiza informações pedagógicas para apoiar o registro, o acompanhamento e a tomada de decisão no Atendimento Educacional Especializado.",
        },
    }

    if "perfil pedagógico" in titulo or "perfil pedagógico" in conteudo_txt or "síntese pedagógica do estudante" in conteudo_txt:
        return temas["perfil"]
    if "mensal" in titulo or "plano mensal" in conteudo_txt:
        return temas["mensal"]
    if tipo == "plano" or "plano aee" in titulo or "paee" in titulo:
        return temas["paee"]
    if "apoio ao docente" in titulo or "apoio ao docente" in conteudo_txt:
        return temas["docente"]
    if tipo == "relatorio" or "relatório gre" in titulo or "relatorio gre" in titulo:
        return temas["gre"]
    if tipo == "avaliacao" or "avaliação pedagógica" in titulo or "avaliacao pedagogica" in titulo:
        return temas["avaliacao"]
    if tipo == "estudo" or "estudo de caso" in titulo:
        return temas["estudo"]
    return temas["documento"]


def gerar_docx_plano_aee_ia_visual(conteudo, nome_base, titulo_doc="PLANO MENSAL AEE - IA"):
    """Gera Word com layout visual para planos IA: capa, cards, seções e tabela de atendimentos."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def set_cell_bg(cell, fill):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fill)
        tc_pr.append(shd)

    def set_cell_text(cell, text, bold=False, color="111827", size=9):
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(str(text or ""))
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor.from_string(color)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    tema = obter_tema_documento("plano_ia_visual", titulo_doc, conteudo)
    cor_principal = tema["cor"].replace("#", "")
    cor_clara = tema["cor_clara"].replace("#", "")

    nome_arquivo = f"plano_aee_ia_visual_{nome_base}.docx".replace("/", "-").replace("\\", "-")
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.45)
    sec.bottom_margin = Inches(0.45)
    sec.left_margin = Inches(0.55)
    sec.right_margin = Inches(0.55)

    # Capa / cabeçalho
    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = True
    set_cell_bg(header.cell(0,0), cor_principal)
    set_cell_bg(header.cell(0,1), cor_clara)
    set_cell_text(header.cell(0,0), "INCLUISRM\nSistema de Gestão do Atendimento Educacional Especializado", True, "FFFFFF", 11)
    set_cell_text(header.cell(0,1), f"{titulo_doc}\nDocumento pedagógico gerado com apoio de IA", True, "0F172A", 12)

    doc.add_paragraph("")

    papel = doc.add_table(rows=1, cols=1)
    set_cell_bg(papel.cell(0,0), cor_clara)
    set_cell_text(papel.cell(0,0), "Papel deste documento: " + tema["papel"], False, "0F172A", 9)

    doc.add_paragraph("")

    # Cards de identificação
    pares = extrair_linhas_chave_valor(conteudo)[:8]
    if pares:
        doc.add_heading("Identificação segura", level=2)
        t = doc.add_table(rows=0, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for chave, valor in pares:
            row = t.add_row().cells
            set_cell_bg(row[0], "DBEAFE")
            set_cell_bg(row[1], "F8FAFC")
            set_cell_text(row[0], chave, True, "1E3A8A", 9)
            set_cell_text(row[1], valor, False, "111827", 9)

    doc.add_paragraph("")
    doc.add_heading("Roteiro pedagógico do atendimento", level=2)

    texto = limpar_marcadores_relatorio(conteudo)
    for raw in texto.splitlines():
        linha = raw.strip()
        if not linha or linha == "---":
            continue
        if eh_titulo_relatorio(linha):
            p = doc.add_paragraph()
            r = p.add_run(linha)
            r.bold = True
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor.from_string(cor_principal)
            continue
        if linha.startswith("-") or linha.startswith("•"):
            p = doc.add_paragraph(style=None)
            p.paragraph_format.left_indent = Inches(0.18)
            r = p.add_run("• " + linha.lstrip("-• "))
            r.font.size = Pt(10)
            continue
        p = doc.add_paragraph()
        r = p.add_run(linha)
        r.font.size = Pt(10)

    doc.add_paragraph("")
    foot = doc.add_table(rows=1, cols=1)
    set_cell_bg(foot.cell(0,0), "F1F5F9")
    set_cell_text(
        foot.cell(0,0),
        f"Documento gerado em {agora_local().strftime('%d/%m/%Y %H:%M')} pelo INCLUISRM • LabTec3DI/UFRPE • Uso pedagógico no AEE",
        False,
        "475569",
        8,
    )

    doc.save(nome_arquivo)
    return nome_arquivo


def gerar_pdf_plano_aee_ia_visual(conteudo, nome_base, titulo_doc="PLANO MENSAL AEE - IA"):
    """Gera PDF com visual mais profissional usando cards, cores e seções."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image

    tema = obter_tema_documento("plano_ia_visual", titulo_doc, conteudo)
    cor_principal = tema["cor"]
    cor_clara = tema["cor_clara"]

    nome_arquivo = f"plano_aee_ia_visual_{nome_base}.pdf".replace("/", "-").replace("\\", "-")
    doc = SimpleDocTemplate(nome_arquivo, pagesize=A4, rightMargin=1.35*cm, leftMargin=1.35*cm, topMargin=1.15*cm, bottomMargin=1.15*cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitlePlanoIA", parent=styles["Title"], alignment=TA_CENTER, fontSize=16, leading=20, textColor=colors.HexColor(cor_principal), spaceAfter=8)
    subtitle_style = ParagraphStyle("SubtitlePlanoIA", parent=styles["Normal"], alignment=TA_CENTER, fontSize=9, textColor=colors.HexColor("#475569"), spaceAfter=12)
    sec_style = ParagraphStyle("SecPlanoIA", parent=styles["Heading2"], fontSize=12, leading=15, textColor=colors.HexColor(cor_principal), spaceBefore=8, spaceAfter=5)
    normal_style = ParagraphStyle("NormalPlanoIA", parent=styles["Normal"], fontSize=9.5, leading=13, textColor=colors.HexColor("#111827"), spaceAfter=4)
    bullet_style = ParagraphStyle("BulletPlanoIA", parent=normal_style, leftIndent=12, firstLineIndent=-8)
    small_style = ParagraphStyle("SmallPlanoIA", parent=styles["Normal"], alignment=TA_CENTER, fontSize=8, textColor=colors.HexColor("#64748b"))

    elementos = []
    try:
        logo = Image(LOGO_PATH, width=4.0*cm, height=1.9*cm)
        logo.hAlign = "CENTER"
        elementos.append(logo)
        elementos.append(Spacer(1, 6))
    except Exception:
        pass

    capa = Table(
        [[Paragraph("<b>INCLUISRM</b><br/>Sistema de Gestão do Atendimento Educacional Especializado", normal_style),
          Paragraph(f"<b>{escape(titulo_doc)}</b><br/>Relatório visual de planejamento pedagógico", normal_style)]],
        colWidths=[8.4*cm, 9.1*cm],
    )
    capa.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (0,0), colors.HexColor(cor_principal)),
        ("TEXTCOLOR", (0,0), (0,0), colors.white),
        ("BACKGROUND", (1,0), (1,0), colors.HexColor(cor_clara)),
        ("BOX", (0,0), (-1,-1), 0.6, colors.HexColor("#cbd5e1")),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("LEFTPADDING", (0,0), (-1,-1), 10),
        ("RIGHTPADDING", (0,0), (-1,-1), 10),
        ("TOPPADDING", (0,0), (-1,-1), 10),
        ("BOTTOMPADDING", (0,0), (-1,-1), 10),
    ]))
    elementos.append(capa)
    elementos.append(Spacer(1, 12))
    elementos.append(Paragraph(titulo_doc, title_style))
    elementos.append(Paragraph("Planejamento organizado em linguagem pedagógica, com foco em execução, registro e acompanhamento evolutivo.", subtitle_style))

    papel_box = Table([[Paragraph("<b>Papel deste documento:</b> " + escape(tema["papel"]), normal_style)]], colWidths=[17.5*cm])
    papel_box.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,-1), colors.HexColor(cor_clara)),
        ("BOX", (0,0), (-1,-1), 0.45, colors.HexColor(cor_principal)),
        ("LEFTPADDING", (0,0), (-1,-1), 8),
        ("RIGHTPADDING", (0,0), (-1,-1), 8),
        ("TOPPADDING", (0,0), (-1,-1), 7),
        ("BOTTOMPADDING", (0,0), (-1,-1), 7),
    ]))
    elementos.append(papel_box)
    elementos.append(Spacer(1, 10))

    pares = extrair_linhas_chave_valor(conteudo)[:8]
    if pares:
        dados = [[Paragraph("<b>Campo</b>", normal_style), Paragraph("<b>Informação</b>", normal_style)]]
        for chave, valor in pares:
            dados.append([Paragraph(escape(chave), normal_style), Paragraph(escape(valor), normal_style)])
        tabela = Table(dados, colWidths=[5.2*cm, 12.3*cm])
        tabela.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor(cor_principal)),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("BACKGROUND", (0,1), (0,-1), colors.HexColor(cor_clara)),
            ("BACKGROUND", (1,1), (1,-1), colors.HexColor("#f8fafc")),
            ("GRID", (0,0), (-1,-1), 0.35, colors.HexColor("#cbd5e1")),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("LEFTPADDING", (0,0), (-1,-1), 7),
            ("RIGHTPADDING", (0,0), (-1,-1), 7),
            ("TOPPADDING", (0,0), (-1,-1), 5),
            ("BOTTOMPADDING", (0,0), (-1,-1), 5),
        ]))
        elementos.append(tabela)
        elementos.append(Spacer(1, 10))

    texto = limpar_marcadores_relatorio(conteudo)
    for raw in texto.splitlines():
        linha = raw.strip()
        if not linha:
            continue
        if eh_titulo_relatorio(linha):
            elementos.append(Spacer(1, 4))
            elementos.append(Paragraph(escape(linha), sec_style))
        elif linha.startswith("-") or linha.startswith("•"):
            elementos.append(Paragraph("• " + escape(linha.lstrip("-• ")), bullet_style))
        else:
            elementos.append(Paragraph(escape(linha), normal_style))

    elementos.append(Spacer(1, 14))
    rodape = Table([[Paragraph(f"Gerado em {agora_local().strftime('%d/%m/%Y %H:%M')} pelo INCLUISRM • LabTec3DI/UFRPE • Documento pedagógico de apoio ao AEE", small_style)]], colWidths=[17.5*cm])
    rodape.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,-1), colors.HexColor("#f1f5f9")), ("BOX", (0,0), (-1,-1), 0.3, colors.HexColor("#cbd5e1")), ("TOPPADDING", (0,0), (-1,-1), 6), ("BOTTOMPADDING", (0,0), (-1,-1), 6)]))
    elementos.append(rodape)
    doc.build(elementos)
    return nome_arquivo


def export_buttons(texto, nome_base, tipo_pdf="documento"):
    """Botões de exportação organizados em uma única linha.

    Fluxo:
    - TXT baixa diretamente;
    - PDF/Word primeiro geram o arquivo e, em seguida, habilitam o botão de download.

    Essa organização evita que os botões fiquem empilhados/desalinhados na tela.
    """
    chave_pdf = f"pdf_{nome_base}_{tipo_pdf}"
    chave_docx = f"docx_{nome_base}_{tipo_pdf}"

    col_txt, col_pdf_gerar, col_pdf_baixar, col_word_gerar, col_word_baixar = st.columns([1, 1, 1, 1, 1])

    with col_txt:
        st.download_button(
            "📄 Baixar TXT",
            data=texto,
            file_name=f"{nome_base}.txt",
            mime="text/plain",
            key=f"txt_{nome_base}_{tipo_pdf}",
            use_container_width=True,
        )

    with col_pdf_gerar:
        if st.button("📕 Gerar PDF", key=f"gerar_pdf_{nome_base}_{tipo_pdf}", use_container_width=True):
            if tipo_pdf == "plano_ia_visual":
                arquivo = gerar_pdf_plano_aee_ia_visual(texto, nome_base)
            else:
                arquivo = gerar_pdf_documento(texto, nome_base, tipo=tipo_pdf)
            st.session_state[chave_pdf] = arquivo

    with col_pdf_baixar:
        if chave_pdf in st.session_state:
            with open(st.session_state[chave_pdf], "rb") as f:
                st.download_button(
                    "⬇️ Baixar PDF",
                    data=f,
                    file_name=f"{nome_base}.pdf",
                    mime="application/pdf",
                    key=f"download_pdf_{nome_base}_{tipo_pdf}",
                    use_container_width=True,
                )
        else:
            st.button("⬇️ Baixar PDF", key=f"download_pdf_disabled_{nome_base}_{tipo_pdf}", disabled=True, use_container_width=True)

    with col_word_gerar:
        if st.button("📘 Gerar Word", key=f"gerar_docx_{nome_base}_{tipo_pdf}", use_container_width=True):
            try:
                if tipo_pdf == "plano_ia_visual":
                    arquivo = gerar_docx_plano_aee_ia_visual(texto, nome_base)
                else:
                    arquivo = gerar_docx_documento(texto, nome_base, tipo=tipo_pdf)
                st.session_state[chave_docx] = arquivo
            except ModuleNotFoundError:
                st.error("Biblioteca python-docx não instalada. Rode: pip install python-docx")

    with col_word_baixar:
        if chave_docx in st.session_state:
            with open(st.session_state[chave_docx], "rb") as f:
                st.download_button(
                    "⬇️ Baixar Word",
                    data=f,
                    file_name=f"{nome_base}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_docx_{nome_base}_{tipo_pdf}",
                    use_container_width=True,
                )
        else:
            st.button("⬇️ Baixar Word", key=f"download_docx_disabled_{nome_base}_{tipo_pdf}", disabled=True, use_container_width=True)



# ======================================================
# CRUD - RECURSOS PEDAGÓGICOS E TECNOLOGIA ASSISTIVA DA ESCOLA
# ======================================================
def cadastrar_recurso_escola(
    escola_nome,
    nome_recurso,
    categoria,
    descricao,
    quantidade,
    localizacao,
    publico_indicado,
    objetivo_pedagogico,
    status,
    origem,
    link_referencia,
    observacoes,
):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        """
        INSERT INTO recursos_escola (
            escola_nome, nome_recurso, categoria, descricao, quantidade,
            localizacao, publico_indicado, objetivo_pedagogico, status,
            origem, link_referencia, observacoes, criado_em
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            escola_nome,
            nome_recurso,
            categoria,
            descricao,
            int(quantidade or 1),
            localizacao,
            publico_indicado,
            objetivo_pedagogico,
            status,
            origem,
            link_referencia,
            observacoes,
            hoje_str(),
        ),
    )
    conn.commit()
    conn.close()
    limpar_cache_dados()


@st.cache_data(ttl=30, show_spinner=False)
def listar_recursos_escola(escola_nome=None, apenas_disponiveis=False):
    conn = conectar()
    cursor = conn.cursor()

    query = """
        SELECT id, escola_nome, nome_recurso, categoria, descricao, quantidade,
               localizacao, publico_indicado, objetivo_pedagogico, status,
               origem, link_referencia, observacoes, criado_em
        FROM recursos_escola
    """
    filtros = []
    params = []

    if escola_nome and str(escola_nome).strip():
        filtros.append("(escola_nome = ? OR escola_nome IS NULL OR escola_nome = '')")
        params.append(escola_nome)

    if apenas_disponiveis:
        filtros.append("(status IS NULL OR status = '' OR status IN ('Disponível', 'Disponível com agendamento'))")

    if filtros:
        query += " WHERE " + " AND ".join(filtros)

    query += " ORDER BY escola_nome, categoria, nome_recurso"

    cursor.execute(query, tuple(params))
    dados = cursor.fetchall()
    conn.close()
    return dados


def excluir_recurso_escola(recurso_id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM recursos_escola WHERE id = ?", (recurso_id,))
    conn.commit()
    conn.close()
    limpar_cache_dados()


def importar_recursos_escola_dataframe(df, escola_padrao=""):
    """Importa recursos por planilha CSV/XLSX.

    Colunas aceitas:
    escola_nome, nome_recurso, categoria, descricao, quantidade, localizacao,
    publico_indicado, objetivo_pedagogico, status, origem, link_referencia, observacoes.

    Também aceita variações simples de nomes como recurso, nome, qtd, local.
    """
    if df is None or df.empty:
        return 0, []

    mapa_colunas = {
        "escola": "escola_nome",
        "unidade": "escola_nome",
        "escola_nome": "escola_nome",
        "recurso": "nome_recurso",
        "nome": "nome_recurso",
        "nome_recurso": "nome_recurso",
        "nome do recurso": "nome_recurso",
        "categoria": "categoria",
        "descrição": "descricao",
        "descricao": "descricao",
        "quantidade": "quantidade",
        "qtd": "quantidade",
        "local": "localizacao",
        "localização": "localizacao",
        "localizacao": "localizacao",
        "público indicado": "publico_indicado",
        "publico indicado": "publico_indicado",
        "publico_indicado": "publico_indicado",
        "objetivo": "objetivo_pedagogico",
        "objetivo pedagógico": "objetivo_pedagogico",
        "objetivo pedagogico": "objetivo_pedagogico",
        "objetivo_pedagogico": "objetivo_pedagogico",
        "status": "status",
        "situação": "status",
        "situacao": "status",
        "origem": "origem",
        "link": "link_referencia",
        "link_referencia": "link_referencia",
        "observação": "observacoes",
        "observacao": "observacoes",
        "observações": "observacoes",
        "observacoes": "observacoes",
    }

    df_trabalho = df.copy()
    df_trabalho.columns = [
        mapa_colunas.get(str(c).strip().lower(), str(c).strip().lower())
        for c in df_trabalho.columns
    ]

    importados = 0
    erros = []

    for idx, row in df_trabalho.iterrows():
        nome = str(row.get("nome_recurso", "") or "").strip()
        if not nome:
            erros.append(f"Linha {idx + 2}: recurso sem nome.")
            continue

        def campo(nome_col, padrao=""):
            valor = row.get(nome_col, padrao)
            if pd.isna(valor):
                return padrao
            return str(valor).strip()

        qtd_raw = row.get("quantidade", 1)
        try:
            qtd = int(float(qtd_raw)) if not pd.isna(qtd_raw) else 1
        except Exception:
            qtd = 1

        cadastrar_recurso_escola(
            escola_nome=campo("escola_nome", escola_padrao),
            nome_recurso=nome,
            categoria=campo("categoria", "Recurso pedagógico"),
            descricao=campo("descricao", ""),
            quantidade=qtd,
            localizacao=campo("localizacao", ""),
            publico_indicado=campo("publico_indicado", ""),
            objetivo_pedagogico=campo("objetivo_pedagogico", ""),
            status=campo("status", "Disponível"),
            origem=campo("origem", "Importação por planilha"),
            link_referencia=campo("link_referencia", ""),
            observacoes=campo("observacoes", ""),
        )
        importados += 1

    return importados, erros


def listar_recursos_escola_texto(escola_nome=None, limite=40):
    """Monta o contexto dos recursos institucionais da escola para uso nos prompts da IA.

    Regras pedagógicas:
    - Se não houver recursos cadastrados, a IA deve registrar explicitamente a ausência no banco institucional da escola.
    - Se houver recursos cadastrados, a IA deve utilizá-los apenas quando forem coerentes
      com o perfil pedagógico, as barreiras, os objetivos e as necessidades do estudante.
    - Os recursos do professor AEE devem ser considerados em conjunto com os recursos da escola,
      não apenas quando a escola não possuir recursos cadastrados.
    - A IA não deve forçar o uso de recurso incompatível apenas porque ele está cadastrado ou porque o professor o possui.
    """
    recursos = listar_recursos_escola(escola_nome=escola_nome, apenas_disponiveis=True)

    if not recursos:
        return (
            "Não há recursos pedagógicos ou tecnologias assistivas cadastradas no banco institucional da escola no momento. "
            "Quando a seção tratar dos recursos da escola, registre essa informação de forma explícita. "
            "Ainda assim, considere também os recursos pedagógicos/tecnologias assistivas informados no perfil do professor AEE, "
            "quando houver, como recursos de apoio disponibilizados pelo professor para uso mediado nos atendimentos. "
            "Esses recursos do professor não devem ser apresentados como patrimônio institucional da escola. "
            "Não invente recursos como se já estivessem disponíveis no banco institucional da unidade escolar."
        )

    linhas = []
    linhas.append("RECURSOS PEDAGÓGICOS E TECNOLOGIAS ASSISTIVAS CADASTRADOS NA ESCOLA:")
    for r in recursos[:limite]:
        (
            _id, escola, nome, categoria, descricao, quantidade, localizacao,
            publico, objetivo, status, origem, link, obs, criado_em
        ) = r
        linhas.append(
            f"- {nome} | Categoria: {categoria or 'Não informada'} | Quantidade: {quantidade or 1} | "
            f"Local: {localizacao or 'Não informado'} | Status: {status or 'Não informado'} | "
            f"Público indicado: {publico or 'Não informado'} | "
            f"Objetivo pedagógico: {objetivo or descricao or 'Não informado'}"
        )

    linhas.append("")
    linhas.append(
        "ORIENTAÇÃO PARA A IA: considere conjuntamente os recursos cadastrados pela escola e os recursos informados no perfil do professor AEE. "
        "Priorize os recursos institucionais da escola quando estiverem disponíveis e forem compatíveis, mas complemente as sugestões com os recursos do professor quando eles forem coerentes com o perfil do estudante, os objetivos pedagógicos e a estratégia proposta. "
        "Não force o uso de recursos incompatíveis e não invente recursos não cadastrados. "
        "Quando utilizar recursos do professor, descreva-os como recursos de apoio disponibilizados pelo professor do AEE, e não como patrimônio institucional da escola. "
        "Quando houver recursos cadastrados na escola, mas nenhum for compatível com determinada necessidade, registre isso de forma pedagógica e sugira alternativas possíveis, inclusive recursos do professor se estiverem informados e forem adequados."
    )
    return "\n".join(linhas)


def renderizar_pagina_recursos_escola():
    st.markdown('<div class="subtitulo">🏫 Recursos Pedagógicos e Tecnologia Assistiva da Escola</div>', unsafe_allow_html=True)
    st.markdown(
        """
        <div class="descricao">
        Cadastre ou importe os recursos reais disponíveis na escola para que o Plano AEE - IA e o Plano Mensal
        possam sugerir atividades pedagógicas contextualizadas com a estrutura existente.
        </div>
        """,
        unsafe_allow_html=True,
    )

    aba1, aba2, aba3, aba4 = st.tabs([
        "➕ Cadastro manual",
        "📥 Importar planilha",
        "📦 Banco de recursos",
        "📊 Indicadores",
    ])

    with aba1:
        st.markdown("#### Cadastrar recurso da escola")
        with st.form("form_recurso_escola"):
            col1, col2 = st.columns(2)
            with col1:
                escola_nome = st.text_input("Escola/unidade", placeholder="Ex.: EREM Ernesto Silva")
                nome_recurso = st.text_input("Nome do recurso *", placeholder="Ex.: Tablet, prancha CAA, impressora 3D")
                categoria = st.selectbox("Categoria", OPCOES_CATEGORIAS_RECURSOS_ESCOLA)
                quantidade = st.number_input("Quantidade", min_value=1, value=1, step=1)
                status = st.selectbox("Status", OPCOES_STATUS_RECURSOS_ESCOLA)
            with col2:
                localizacao = st.text_input("Localização", placeholder="Ex.: Sala de Recursos, laboratório, biblioteca")
                publico_indicado = st.text_input("Público indicado", placeholder="Ex.: TEA, baixa visão, deficiência física, todos")
                origem = st.text_input("Origem/aquisição", placeholder="Ex.: PDDE, doação, escola, projeto")
                link_referencia = st.text_input("Link de referência", placeholder="Opcional")

            descricao = st.text_area("Descrição do recurso", height=90)
            objetivo_pedagogico = st.text_area("Objetivo pedagógico / possibilidades de uso", height=90)
            observacoes = st.text_area("Observações", height=80)

            salvar = st.form_submit_button("💾 Salvar recurso", use_container_width=True)
            if salvar:
                if not nome_recurso.strip():
                    st.error("Informe o nome do recurso.")
                else:
                    cadastrar_recurso_escola(
                        escola_nome, nome_recurso, categoria, descricao, quantidade,
                        localizacao, publico_indicado, objetivo_pedagogico, status,
                        origem, link_referencia, observacoes
                    )
                    st.success("Recurso cadastrado com sucesso.")

    with aba2:
        st.markdown("#### Importar recursos por planilha")
        st.info(
            "A planilha pode ter colunas como: escola_nome, nome_recurso, categoria, descricao, quantidade, "
            "localizacao, publico_indicado, objetivo_pedagogico, status, origem, link_referencia e observacoes."
        )
        escola_padrao = st.text_input("Escola padrão para linhas sem escola informada", key="escola_padrao_importacao")
        arquivo = st.file_uploader("Enviar planilha CSV ou Excel", type=["csv", "xlsx", "xls"])

        if arquivo is not None:
            try:
                if arquivo.name.lower().endswith(".csv"):
                    df_import = pd.read_csv(arquivo)
                else:
                    df_import = pd.read_excel(arquivo)

                st.dataframe(df_import.head(20), use_container_width=True)

                if st.button("📥 Importar recursos da planilha", use_container_width=True):
                    total, erros = importar_recursos_escola_dataframe(df_import, escola_padrao=escola_padrao)
                    st.success(f"{total} recurso(s) importado(s) com sucesso.")
                    if erros:
                        st.warning("Algumas linhas não foram importadas:")
                        st.write(erros[:20])
            except Exception as e:
                st.error(f"Não foi possível ler a planilha: {e}")

    with aba3:
        st.markdown("#### Banco de recursos cadastrados")
        recursos = listar_recursos_escola()
        if not recursos:
            st.info("Nenhum recurso cadastrado até o momento.")
        else:
            df = pd.DataFrame(
                recursos,
                columns=[
                    "ID", "Escola", "Recurso", "Categoria", "Descrição", "Quantidade",
                    "Localização", "Público indicado", "Objetivo pedagógico", "Status",
                    "Origem", "Link", "Observações", "Criado em"
                ],
            )
            st.dataframe(df, use_container_width=True, hide_index=True)

            st.markdown("#### Excluir recurso")
            ids = df["ID"].tolist()
            mapa_ids = {row["ID"]: f"{row['Recurso']} - {row['Escola'] or 'Sem escola informada'}" for _, row in df.iterrows()}
            recurso_id = st.selectbox("Selecione o recurso para excluir", ids, format_func=lambda x: mapa_ids.get(x, str(x)))
            if st.button("🗑️ Excluir recurso selecionado"):
                excluir_recurso_escola(recurso_id)
                st.success("Recurso excluído.")
                st.rerun()

    with aba4:
        recursos = listar_recursos_escola()
        if not recursos:
            st.info("Cadastre recursos para visualizar indicadores.")
        else:
            df = pd.DataFrame(
                recursos,
                columns=[
                    "ID", "Escola", "Recurso", "Categoria", "Descrição", "Quantidade",
                    "Localização", "Público indicado", "Objetivo pedagógico", "Status",
                    "Origem", "Link", "Observações", "Criado em"
                ],
            )
            total_itens = len(df)
            total_unidades = int(pd.to_numeric(df["Quantidade"], errors="coerce").fillna(1).sum())
            categorias = df["Categoria"].nunique()
            disponiveis = df[df["Status"].isin(["Disponível", "Disponível com agendamento"])].shape[0]

            col1, col2, col3, col4 = st.columns(4)
            col1.metric("Tipos de recursos", total_itens)
            col2.metric("Unidades cadastradas", total_unidades)
            col3.metric("Categorias", categorias)
            col4.metric("Disponíveis", disponiveis)

            graf = df.groupby("Categoria", dropna=False)["Quantidade"].sum().reset_index()
            graf["Categoria"] = graf["Categoria"].fillna("Não informada")
            chart = alt.Chart(graf).mark_bar().encode(
                x=alt.X("Quantidade:Q", title="Quantidade"),
                y=alt.Y("Categoria:N", sort="-x", title="Categoria"),
                tooltip=["Categoria", "Quantidade"],
            )
            st.altair_chart(chart, use_container_width=True)



# ======================================================
# CRUD - ESTUDANTES
# ======================================================
def cadastrar_estudante(codigo, ano_serie, turma, turno, perfil, observacoes, dias_atendimento, horario_preferencial):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        """
        INSERT INTO estudantes (
            codigo, ano_serie, turma, turno, perfil, observacoes,
            dias_atendimento, horario_preferencial, criado_em
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (codigo, ano_serie, turma, turno, perfil, observacoes, dias_atendimento, horario_preferencial, hoje_str()),
    )
    conn.commit()
    conn.close()
    limpar_cache_dados()


@st.cache_data(ttl=30, show_spinner=False)
def listar_estudantes():
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT id, codigo, ano_serie, turma, perfil, observacoes, turno, dias_atendimento, horario_preferencial
        FROM estudantes
        ORDER BY codigo
        """
    )
    dados = cursor.fetchall()
    conn.close()
    return dados


@st.cache_data(ttl=30, show_spinner=False)
def buscar_estudante(estudante_id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT id, codigo, ano_serie, turma, perfil, observacoes, turno, dias_atendimento, horario_preferencial
        FROM estudantes
        WHERE id = ?
        """,
        (estudante_id,),
    )
    dado = cursor.fetchone()
    conn.close()
    return dado


def atualizar_estudante(estudante_id, codigo, ano_serie, turma, turno, perfil, observacoes, dias_atendimento, horario_preferencial):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        """
        UPDATE estudantes
        SET codigo=?, ano_serie=?, turma=?, turno=?, perfil=?, observacoes=?,
            dias_atendimento=?, horario_preferencial=?
        WHERE id=?
        """,
        (codigo, ano_serie, turma, turno, perfil, observacoes, dias_atendimento, horario_preferencial, estudante_id),
    )
    conn.commit()
    conn.close()
    limpar_cache_dados()


def excluir_estudante(estudante_id):
    conn = conectar()
    cursor = conn.cursor()
    for tabela in [
        "agenda", "atendimentos", "relatorios", "paees", "plano_aee_ia", "planos_aee",
        "estudos_caso", "avaliacoes", "entrevistas_familia", "estudante_professor"
    ]:
        cursor.execute(f"DELETE FROM {tabela} WHERE estudante_id=?", (estudante_id,))
    cursor.execute("DELETE FROM estudantes WHERE id=?", (estudante_id,))
    conn.commit()
    conn.close()
    limpar_cache_dados()


# ======================================================
# CRUD - PROFESSOR
# ======================================================
def salvar_professor(
    nome_referencia,
    escola,
    regional,
    formacao,
    carga_horaria,
    turno_atuacao,
    observacoes,
    areas_interesse="",
    nivel_tecnologico="Básico",
    modo_maker="Não",
    interesse_formacao_maker="Não",
    projetos_interesse="",
    preferencias_metodologicas="",
    recursos_professor="",
    recursos_professor_uso="",
    recursos_professor_observacoes="",
):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        """
        INSERT INTO professores (
            nome_referencia, escola, regional, formacao, carga_horaria,
            turno_atuacao, observacoes, criado_em,
            areas_interesse, nivel_tecnologico, modo_maker,
            interesse_formacao_maker, projetos_interesse, preferencias_metodologicas,
            recursos_professor, recursos_professor_uso, recursos_professor_observacoes
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            nome_referencia, escola, regional, formacao, carga_horaria,
            turno_atuacao, observacoes, hoje_str(),
            areas_interesse, nivel_tecnologico, modo_maker,
            interesse_formacao_maker, projetos_interesse, preferencias_metodologicas,
            recursos_professor, recursos_professor_uso, recursos_professor_observacoes,
        ),
    )
    conn.commit()
    conn.close()
    limpar_cache_dados()


@st.cache_data(ttl=30, show_spinner=False)
def listar_professores():
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT id, nome_referencia, escola, regional, formacao, carga_horaria, turno_atuacao, observacoes, criado_em,
               areas_interesse, nivel_tecnologico, modo_maker, interesse_formacao_maker, projetos_interesse, preferencias_metodologicas,
               recursos_professor, recursos_professor_uso, recursos_professor_observacoes
        FROM professores
        ORDER BY id DESC
        """
    )
    dados = cursor.fetchall()
    conn.close()
    return dados


def atualizar_professor(
    professor_id,
    nome_referencia,
    escola,
    regional,
    formacao,
    carga_horaria,
    turno_atuacao,
    observacoes,
    areas_interesse="",
    nivel_tecnologico="Básico",
    modo_maker="Não",
    interesse_formacao_maker="Não",
    projetos_interesse="",
    preferencias_metodologicas="",
    recursos_professor="",
    recursos_professor_uso="",
    recursos_professor_observacoes="",
):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        """
        UPDATE professores
        SET nome_referencia=?, escola=?, regional=?, formacao=?,
            carga_horaria=?, turno_atuacao=?, observacoes=?,
            areas_interesse=?, nivel_tecnologico=?, modo_maker=?,
            interesse_formacao_maker=?, projetos_interesse=?, preferencias_metodologicas=?,
            recursos_professor=?, recursos_professor_uso=?, recursos_professor_observacoes=?
        WHERE id=?
        """,
        (
            nome_referencia, escola, regional, formacao, carga_horaria,
            turno_atuacao, observacoes, areas_interesse, nivel_tecnologico,
            modo_maker, interesse_formacao_maker, projetos_interesse,
            preferencias_metodologicas, recursos_professor, recursos_professor_uso,
            recursos_professor_observacoes, professor_id,
        ),
    )
    conn.commit()
    conn.close()
    limpar_cache_dados()


def excluir_professor(professor_id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM estudante_professor WHERE professor_id=?", (professor_id,))
    cursor.execute("DELETE FROM professores WHERE id=?", (professor_id,))
    conn.commit()
    conn.close()
    limpar_cache_dados()


@st.cache_data(ttl=30, show_spinner=False)
def contar_alunos_do_professor(professor_id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT COUNT(*) FROM estudante_professor WHERE professor_id=?",
        (professor_id,),
    )
    total = cursor.fetchone()[0]
    conn.close()
    return total


def vincular_professor_estudante(estudante_id, professor_id):
    """Vincula professor ao estudante, limitando cada professor a até 14 estudantes."""
    conn = conectar()
    cursor = conn.cursor()

    cursor.execute(
        "SELECT COUNT(*) FROM estudante_professor WHERE professor_id=?",
        (professor_id,),
    )
    total = cursor.fetchone()[0]

    cursor.execute(
        """
        SELECT id FROM estudante_professor
        WHERE estudante_id=? AND professor_id=?
        """,
        (estudante_id, professor_id),
    )
    ja_existe = cursor.fetchone()

    if ja_existe:
        conn.close()
        return False, "Este estudante já está vinculado a este professor."

    if total >= 14:
        conn.close()
        return False, "Este professor já possui 14 estudantes vinculados."

    cursor.execute(
        """
        INSERT INTO estudante_professor (estudante_id, professor_id, criado_em)
        VALUES (?, ?, ?)
        """,
        (estudante_id, professor_id, hoje_str()),
    )
    conn.commit()
    conn.close()
    limpar_cache_dados()
    return True, "Vínculo realizado com sucesso."


def remover_vinculo_professor_estudante(vinculo_id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM estudante_professor WHERE id=?", (vinculo_id,))
    conn.commit()
    conn.close()
    limpar_cache_dados()


@st.cache_data(ttl=30, show_spinner=False)
def listar_estudantes_do_professor(professor_id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT ep.id, e.id, e.codigo, e.ano_serie, e.turma, e.perfil
        FROM estudante_professor ep
        JOIN estudantes e ON e.id = ep.estudante_id
        WHERE ep.professor_id=?
        ORDER BY e.codigo
        """,
        (professor_id,),
    )
    dados = cursor.fetchall()
    conn.close()
    return dados


@st.cache_data(ttl=30, show_spinner=False)
def listar_professores_do_estudante(estudante_id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT ep.id, p.id, p.nome_referencia, p.escola, p.regional, p.carga_horaria, p.turno_atuacao
        FROM estudante_professor ep
        JOIN professores p ON p.id = ep.professor_id
        WHERE ep.estudante_id=?
        ORDER BY p.nome_referencia
        """,
        (estudante_id,),
    )
    dados = cursor.fetchall()
    conn.close()
    return dados


@st.cache_data(ttl=30, show_spinner=False)
def buscar_professor_responsavel():
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT id, nome_referencia, escola, regional, formacao,
               carga_horaria, turno_atuacao, observacoes, criado_em
        FROM professores
        ORDER BY id DESC
        LIMIT 1
        """
    )
    dado = cursor.fetchone()
    conn.close()
    return dado


def texto_professor_responsavel():
    professor = buscar_professor_responsavel()
    if not professor:
        return "Nenhum professor AEE cadastrado."
    return (
        f"{professor[1] or 'Não informado.'} | "
        f"Escola: {professor[2] or 'Não informado.'} | "
        f"Regional/GRE: {professor[3] or 'Não informado.'} | "
        f"Carga horária: {professor[5] or 'Não informado.'} | "
        f"Turno: {professor[6] or 'Não informado.'}"
    )


def texto_professores_vinculados(estudante_id=None):
    """Protótipo personalizado: usa o professor responsável mais recente."""
    return texto_professor_responsavel()


# ======================================================
# CRUD - ENTREVISTA, AVALIAÇÃO, ESTUDO, PLANO
# ======================================================
def inserir_registro(tabela, campos, valores):
    conn = conectar()
    cursor = conn.cursor()
    placeholders = ", ".join(["?"] * len(valores))
    cursor.execute(
        f"INSERT INTO {tabela} ({', '.join(campos)}) VALUES ({placeholders})",
        valores,
    )
    conn.commit()
    conn.close()
    limpar_cache_dados()


@st.cache_data(ttl=30, show_spinner=False)
def listar_por_estudante(tabela, campos, estudante_id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        f"SELECT id, {', '.join(campos)} FROM {tabela} WHERE estudante_id=? ORDER BY id DESC",
        (estudante_id,),
    )
    dados = cursor.fetchall()
    conn.close()
    return dados


def excluir_registro(tabela, registro_id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(f"DELETE FROM {tabela} WHERE id=?", (registro_id,))
    conn.commit()
    conn.close()
    limpar_cache_dados()


def atualizar_registro(tabela, campos, valores, registro_id):
    conn = conectar()
    cursor = conn.cursor()
    sets = ", ".join([f"{campo}=?" for campo in campos])
    cursor.execute(f"UPDATE {tabela} SET {sets} WHERE id=?", list(valores) + [registro_id])
    conn.commit()
    conn.close()
    limpar_cache_dados()


@st.cache_data(ttl=30, show_spinner=False)
def ultima_linha(tabela, campos, estudante_id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        f"SELECT {', '.join(campos)} FROM {tabela} WHERE estudante_id=? ORDER BY id DESC LIMIT 1",
        (estudante_id,),
    )
    dado = cursor.fetchone()
    conn.close()
    return dado


# ======================================================
# CRUD - AVALIAÇÃO
# ======================================================
def salvar_avaliacao(
    estudante_id,
    barreiras="",
    potencialidades="",
    comunicacao="",
    interacao="",
    autonomia="",
    aprendizagem="",
    resumo_laudo="",
    ano_letivo="",
    tipo_registro="Avaliação pedagógica atual",
    avaliacao_anterior_id=None,
    analise_comparativa_ia="",
    sugestao_nova_avaliacao_ia="",
    origem_documento="",
    texto_documento_extra="",
):
    inserir_registro(
        "avaliacoes",
        [
            "estudante_id", "data_registro",
            "ano_letivo", "tipo_registro", "avaliacao_anterior_id", "analise_comparativa_ia", "sugestao_nova_avaliacao_ia",
            "origem_documento", "texto_documento_extra",
            "barreiras", "potencialidades", "comunicacao", "interacao", "autonomia", "aprendizagem", "resumo_laudo",
        ],
        [
            estudante_id, hoje_str(),
            ano_letivo, tipo_registro, avaliacao_anterior_id, analise_comparativa_ia, sugestao_nova_avaliacao_ia,
            origem_documento, texto_documento_extra,
            barreiras, potencialidades, comunicacao, interacao, autonomia, aprendizagem, resumo_laudo,
        ],
    )


@st.cache_data(ttl=30, show_spinner=False)
def listar_avaliacoes(estudante_id):
    return listar_por_estudante(
        "avaliacoes",
        CAMPOS_AVALIACAO,
        estudante_id,
    )


@st.cache_data(ttl=30, show_spinner=False)
def ultima_avaliacao(estudante_id):
    return ultima_linha(
        "avaliacoes",
        CAMPOS_AVALIACAO,
        estudante_id,
    )


@st.cache_data(ttl=30, show_spinner=False)
def buscar_avaliacao_por_id(avaliacao_id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        f"SELECT id, {', '.join(CAMPOS_AVALIACAO)} FROM avaliacoes WHERE id=?",
        (avaliacao_id,),
    )
    dado = cursor.fetchone()
    conn.close()
    return dado


def salvar_documento_avaliacao(
    estudante_id,
    avaliacao_id,
    ano_letivo,
    tipo_documento,
    arquivo,
    observacao="",
):
    """Salva documento anexado à avaliação pedagógica.

    Observação importante:
    - No Render, arquivos salvos em pasta local podem ser perdidos em novo deploy.
    - Para produção, o ideal é usar armazenamento persistente/externo.
    """
    if arquivo is None:
        return None

    pasta_estudante = DOCUMENTOS_AVALIACOES_DIR / f"estudante_{estudante_id}"
    pasta_estudante.mkdir(parents=True, exist_ok=True)

    timestamp = agora_local().strftime("%Y%m%d_%H%M%S")
    nome_seguro = arquivo.name.replace("/", "_").replace("\\", "_")
    nome_salvo = f"{timestamp}_{nome_seguro}"
    caminho = pasta_estudante / nome_salvo

    with open(caminho, "wb") as f:
        f.write(arquivo.getbuffer())

    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        """
        INSERT INTO avaliacoes_documentos (
            avaliacao_id, estudante_id, ano_letivo, tipo_documento,
            nome_original, nome_arquivo_salvo, caminho_arquivo,
            observacao, data_upload
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            avaliacao_id,
            estudante_id,
            ano_letivo,
            tipo_documento,
            arquivo.name,
            nome_salvo,
            str(caminho),
            observacao,
            hoje_str(),
        ),
    )
    conn.commit()
    conn.close()
    limpar_cache_dados()
    return str(caminho)


@st.cache_data(ttl=30, show_spinner=False)
def listar_documentos_avaliacao(estudante_id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT id, avaliacao_id, ano_letivo, tipo_documento,
               nome_original, caminho_arquivo, observacao, data_upload
        FROM avaliacoes_documentos
        WHERE estudante_id=?
        ORDER BY id DESC
        """,
        (estudante_id,),
    )
    dados = cursor.fetchall()
    conn.close()
    return dados


def excluir_documento_avaliacao(documento_id, caminho_arquivo):
    try:
        if caminho_arquivo and Path(caminho_arquivo).exists():
            Path(caminho_arquivo).unlink()
    except Exception:
        pass

    conn = conectar()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM avaliacoes_documentos WHERE id=?", (documento_id,))
    conn.commit()
    conn.close()
    limpar_cache_dados()


# ======================================================
# CRUD - ARTICULAÇÃO PEDAGÓGICA INCLUSIVA / ESCUTA DOCENTE
# ======================================================
def salvar_escuta_docente(
    estudante_id,
    ano_letivo,
    codigo_docente,
    componente_curricular,
    outro_componente,
    turma,
    tempo_acompanhamento,
    participacao_sala,
    comunicacao,
    interacao_social,
    aprendizagem,
    barreiras_percebidas,
    potencialidades_observadas,
    estrategias_funcionam,
    adaptacoes_utilizadas,
    recomendacoes_docente,
    nivel_participacao,
    nivel_autonomia,
    nivel_engajamento,
    observacoes,
    areas_interesse_docente="",
    nivel_tecnologico_docente="Básico",
    modo_maker_docente="Não",
    interesse_formacao_maker_docente="Não",
    interesse_projetos_interdisciplinares="Não",
    projeto_interesse_docente="",
    preferencias_metodologicas_docente="",
):
    inserir_registro(
        "escutas_docentes",
        [
            "estudante_id",
            "data_registro",
            "ano_letivo",
            "codigo_docente",
            "componente_curricular",
            "outro_componente",
            "turma",
            "tempo_acompanhamento",
            "participacao_sala",
            "comunicacao",
            "interacao_social",
            "aprendizagem",
            "barreiras_percebidas",
            "potencialidades_observadas",
            "estrategias_funcionam",
            "adaptacoes_utilizadas",
            "recomendacoes_docente",
            "nivel_participacao",
            "nivel_autonomia",
            "nivel_engajamento",
            "observacoes",
            "areas_interesse_docente",
            "nivel_tecnologico_docente",
            "modo_maker_docente",
            "interesse_formacao_maker_docente",
            "interesse_projetos_interdisciplinares",
            "projeto_interesse_docente",
            "preferencias_metodologicas_docente",
        ],
        [
            estudante_id,
            hoje_str(),
            ano_letivo,
            codigo_docente,
            componente_curricular,
            outro_componente,
            turma,
            tempo_acompanhamento,
            participacao_sala,
            comunicacao,
            interacao_social,
            aprendizagem,
            barreiras_percebidas,
            potencialidades_observadas,
            estrategias_funcionam,
            adaptacoes_utilizadas,
            recomendacoes_docente,
            nivel_participacao,
            nivel_autonomia,
            nivel_engajamento,
            observacoes,
            areas_interesse_docente,
            nivel_tecnologico_docente,
            modo_maker_docente,
            interesse_formacao_maker_docente,
            interesse_projetos_interdisciplinares,
            projeto_interesse_docente,
            preferencias_metodologicas_docente,
        ],
    )


@st.cache_data(ttl=30, show_spinner=False)
def listar_escutas_docentes(estudante_id):
    """Lista escutas docentes sem expor nome pessoal.
    Usa codigo_docente como campo principal e mantém fallback técnico para registros antigos.
    """
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT id,
               data_registro,
               ano_letivo,
               COALESCE(NULLIF(codigo_docente, ''), NULLIF(professor_nome, ''), 'Não informado') AS codigo_docente,
               componente_curricular,
               outro_componente,
               turma,
               tempo_acompanhamento,
               participacao_sala,
               comunicacao,
               interacao_social,
               aprendizagem,
               barreiras_percebidas,
               potencialidades_observadas,
               estrategias_funcionam,
               adaptacoes_utilizadas,
               recomendacoes_docente,
               nivel_participacao,
               nivel_autonomia,
               nivel_engajamento,
               observacoes,
               areas_interesse_docente,
               nivel_tecnologico_docente,
               modo_maker_docente,
               interesse_formacao_maker_docente,
               interesse_projetos_interdisciplinares,
               projeto_interesse_docente,
               preferencias_metodologicas_docente
        FROM escutas_docentes
        WHERE estudante_id=?
        ORDER BY id DESC
        """,
        (estudante_id,),
    )
    dados = cursor.fetchall()
    conn.close()
    return dados


def excluir_escuta_docente(escuta_id):
    excluir_registro("escutas_docentes", escuta_id)


def texto_escuta_docente(estudante, escuta):
    dados = dict(zip(["id"] + CAMPOS_ESCUTA_DOCENTE, escuta))

    def v(campo):
        valor = dados.get(campo)
        return valor if valor not in (None, "") else "Não informado."

    componente = v("componente_curricular")
    if componente == "Outras" and dados.get("outro_componente"):
        componente = dados.get("outro_componente")

    return f"""
ARTICULAÇÃO PEDAGÓGICA INCLUSIVA - ESCUTA DOCENTE

Código interno do estudante: {estudante[1]}
Ano/Série: {estudante[2] or 'Não informado.'}
Turma: {estudante[3] or 'Não informado.'}
Data do registro: {v('data_registro')}
Ano letivo: {v('ano_letivo')}

Código interno do docente / referência pedagógica: {v('codigo_docente')}
Componente curricular / área: {componente}
Turma observada: {v('turma')}
Tempo de acompanhamento: {v('tempo_acompanhamento')}

PARTICIPAÇÃO EM SALA:
{v('participacao_sala')}

COMUNICAÇÃO:
{v('comunicacao')}

INTERAÇÃO SOCIAL:
{v('interacao_social')}

APRENDIZAGEM:
{v('aprendizagem')}

BARREIRAS PERCEBIDAS:
{v('barreiras_percebidas')}

POTENCIALIDADES OBSERVADAS:
{v('potencialidades_observadas')}

ESTRATÉGIAS QUE FUNCIONAM:
{v('estrategias_funcionam')}

ADAPTAÇÕES UTILIZADAS:
{v('adaptacoes_utilizadas')}

RECOMENDAÇÕES DO DOCENTE:
{v('recomendacoes_docente')}

INDICADORES:
Participação: {v('nivel_participacao')}/10
Autonomia: {v('nivel_autonomia')}/10
Engajamento: {v('nivel_engajamento')}/10

PERFIL PEDAGÓGICO/TECNOLÓGICO DO DOCENTE DA SALA REGULAR:
Áreas de interesse/atuação: {v('areas_interesse_docente')}
Nível tecnológico: {v('nivel_tecnologico_docente')}
Modo maker inclusivo ativado: {v('modo_maker_docente')}
Interesse em formação maker/tecnologias: {v('interesse_formacao_maker_docente')}
Interesse em projetos interdisciplinares: {v('interesse_projetos_interdisciplinares')}
Projeto/tema que gostaria de trabalhar: {v('projeto_interesse_docente')}
Preferências metodológicas: {v('preferencias_metodologicas_docente')}

OBSERVAÇÕES:
{v('observacoes')}
""".strip()


def salvar_relatorio_docente(
    estudante_id,
    ano_letivo,
    componente_destino,
    professor_destino,
    titulo,
    conteudo,
    fontes_utilizadas,
    observacoes="",
):
    inserir_registro(
        "relatorios_docente",
        [
            "estudante_id",
            "data_geracao",
            "ano_letivo",
            "componente_destino",
            "professor_destino",
            "titulo",
            "conteudo",
            "fontes_utilizadas",
            "observacoes",
        ],
        [
            estudante_id,
            hoje_str(),
            ano_letivo,
            componente_destino,
            professor_destino,
            titulo,
            conteudo,
            fontes_utilizadas,
            observacoes,
        ],
    )


@st.cache_data(ttl=30, show_spinner=False)
def listar_relatorios_docente(estudante_id):
    return listar_por_estudante("relatorios_docente", CAMPOS_RELATORIO_DOCENTE, estudante_id)


def excluir_relatorio_docente(relatorio_id):
    excluir_registro("relatorios_docente", relatorio_id)



def salvar_relatorio_visual_docente(
    estudante_id,
    ano_letivo,
    componente_destino,
    tipo_relatorio,
    titulo,
    caminho_arquivo,
    fontes_utilizadas="",
    observacoes="",
):
    """Salva no histórico o PDF visual gerado para consulta futura."""
    caminho = str(caminho_arquivo or "")
    nome_arquivo = Path(caminho).name if caminho else ""
    inserir_registro(
        "relatorios_visuais_docente",
        [
            "estudante_id",
            "data_geracao",
            "ano_letivo",
            "componente_destino",
            "tipo_relatorio",
            "titulo",
            "nome_arquivo",
            "caminho_arquivo",
            "fontes_utilizadas",
            "observacoes",
        ],
        [
            estudante_id,
            hoje_str(),
            ano_letivo,
            componente_destino,
            tipo_relatorio,
            titulo,
            nome_arquivo,
            caminho,
            fontes_utilizadas,
            observacoes,
        ],
    )


@st.cache_data(ttl=30, show_spinner=False)
def listar_relatorios_visuais_docente(estudante_id):
    return listar_por_estudante("relatorios_visuais_docente", CAMPOS_RELATORIO_VISUAL_DOCENTE, estudante_id)


def excluir_relatorio_visual_docente(relatorio_visual_id, caminho_arquivo=""):
    """Exclui o registro do histórico e tenta remover o PDF salvo no servidor."""
    try:
        if caminho_arquivo and Path(caminho_arquivo).exists():
            Path(caminho_arquivo).unlink()
    except Exception:
        pass
    excluir_registro("relatorios_visuais_docente", relatorio_visual_id)


def caminho_relatorio_visual_docente(nome_arquivo):
    """Retorna caminho seguro para salvar PDFs visuais docentes."""
    RELATORIOS_VISUAIS_DOCENTE_DIR.mkdir(parents=True, exist_ok=True)
    nome_seguro = re.sub(r"[^A-Za-z0-9_\-.]", "_", str(nome_arquivo or "relatorio_visual.pdf"))
    return RELATORIOS_VISUAIS_DOCENTE_DIR / nome_seguro


def texto_relatorio_docente(estudante, relatorio):
    dados = dict(zip(["id"] + CAMPOS_RELATORIO_DOCENTE, relatorio))

    def v(campo):
        valor = dados.get(campo)
        return valor if valor not in (None, "") else "Não informado."

    return f"""
{v('titulo')}

Código interno do estudante: {estudante[1]}
Ano/Série: {estudante[2] or 'Não informado.'}
Turma: {estudante[3] or 'Não informado.'}
Ano letivo: {v('ano_letivo')}
Componente/área de destino: {v('componente_destino')}
Destinatário: ________________________________________
Data de geração: {v('data_geracao')}

{v('conteudo')}

FONTES UTILIZADAS PELO SISTEMA:
{v('fontes_utilizadas')}

OBSERVAÇÕES:
{v('observacoes')}

Observação institucional:
Este relatório possui finalidade exclusivamente pedagógica. Seu objetivo é apoiar o planejamento inclusivo do docente, sem substituir a avaliação profissional do professor, da equipe pedagógica ou do AEE.
""".strip()


def montar_contexto_relatorio_docente(estudante_id, limite_atendimentos=12):
    estudante = buscar_estudante(estudante_id)

    avaliacoes = listar_avaliacoes(estudante_id)[:5]
    estudos = listar_por_estudante("estudos_caso", CAMPOS_ESTUDO_CASO, estudante_id)[:3]
    entrevistas = listar_por_estudante("entrevistas_familia", CAMPOS_ENTREVISTA_FAMILIA, estudante_id)[:2]
    escutas = listar_escutas_docentes(estudante_id)[:10]
    atendimentos = listar_atendimentos(estudante_id)[:limite_atendimentos]
    documentos = listar_documentos_avaliacao(estudante_id)

    textos_avaliacoes = []
    for av in avaliacoes:
        try:
            textos_avaliacoes.append(texto_avaliacao(estudante, av))
        except Exception:
            textos_avaliacoes.append(str(av))

    textos_estudos = []
    for est in estudos:
        try:
            textos_estudos.append(texto_estudo_caso(estudante, est))
        except Exception:
            textos_estudos.append(str(est))

    textos_escutas = []
    for esc in escutas:
        try:
            textos_escutas.append(texto_escuta_docente(estudante, esc))
        except Exception:
            textos_escutas.append(str(esc))

    texto_entrevistas = "\n\n".join([str(e) for e in entrevistas]) or "Nenhuma entrevista familiar localizada."
    texto_atendimentos = "\n".join([str(a) for a in atendimentos]) or "Nenhum atendimento registrado."
    texto_documentos = "\n".join([
        f"- {d[4]} | Ano: {d[2] or 'não informado'} | Tipo: {d[3] or 'não informado'} | Obs.: {d[6] or ''}"
        for d in documentos
    ]) or "Nenhum documento histórico anexado."

    fontes = []
    if avaliacoes:
        fontes.append(f"{len(avaliacoes)} avaliação(ões) pedagógica(s)")
    if estudos:
        fontes.append(f"{len(estudos)} estudo(s) de caso")
    if escutas:
        fontes.append(f"{len(escutas)} escuta(s) docente(s)")
    if atendimentos:
        fontes.append(f"{len(atendimentos)} atendimento(s)")
    if documentos:
        fontes.append(f"{len(documentos)} documento(s) histórico(s) anexado(s)")
    if entrevistas:
        fontes.append(f"{len(entrevistas)} entrevista(s) familiar(es) usada(s) apenas como contexto pedagógico")

    contexto = f"""
DADOS DO ESTUDANTE
Código interno: {estudante[1]}
Ano/Série: {estudante[2]}
Turma: {estudante[3]}
Perfil educacional cadastrado: {estudante[4]}
Observações gerais do cadastro: {estudante[5] or ''}

AVALIAÇÕES PEDAGÓGICAS / DOCUMENTOS LIVRES:
{chr(10).join(textos_avaliacoes) if textos_avaliacoes else 'Nenhuma avaliação pedagógica localizada.'}

ESTUDOS DE CASO:
{chr(10).join(textos_estudos) if textos_estudos else 'Nenhum estudo de caso localizado.'}

ESCUTAS DOCENTES DA SALA REGULAR:
{chr(10).join(textos_escutas) if textos_escutas else 'Nenhuma escuta docente registrada.'}

ATENDIMENTOS DO AEE:
{texto_atendimentos}

DOCUMENTOS HISTÓRICOS ANEXADOS:
{texto_documentos}

ENTREVISTAS FAMILIARES:
{texto_entrevistas}
""".strip()

    return contexto, "; ".join(fontes) if fontes else "Nenhuma fonte registrada."


def gerar_relatorio_apoio_docente_ia(
    estudante_id,
    ano_letivo,
    componente_destino,
    professor_destino="",
    foco_docente="",
):
    estudante = buscar_estudante(estudante_id)
    contexto, fontes = montar_contexto_relatorio_docente(estudante_id)

    componente_normalizado = str(componente_destino or "")
    modo_geral = "geral" in componente_normalizado.lower()
    bases_disciplina = bases_para_componente(componente_normalizado)
    contexto_base_ia = ""
    fontes_base_ia = ""
    if not modo_geral and bases_disciplina:
        pergunta_base = (
            f"BNCC, BNCC Computação e estratégias inclusivas para {componente_normalizado} "
            f"considerando AEE, acessibilidade, adaptações pedagógicas e participação do estudante."
        )
        resultados_base = buscar_na_base_conhecimento(pergunta_base, bases=bases_disciplina, limite=4)
        contexto_base_ia = montar_contexto_base(resultados_base)
        fontes_base_ia = arquivos_consultados_texto(resultados_base)
        if fontes_base_ia and "Nenhum" not in fontes_base_ia:
            fontes = f"{fontes}; Base de Conhecimento IA: {fontes_base_ia}"
    else:
        contexto_base_ia = (
            "Modo geral selecionado: não consultar BNCC por disciplina. "
            "Gerar orientações amplas para todas as áreas do conhecimento, sem recomendações específicas de um componente curricular."
        )

    client = obter_cliente_openai()
    if client is None:
        conteudo_fallback = f"""
1. SÍNTESE PEDAGÓGICA FUNCIONAL
IA não configurada. Configure OPENAI_API_KEY para gerar a síntese automática. Ainda assim, o sistema reuniu as fontes pedagógicas disponíveis para revisão manual.

2. POTENCIALIDADES OBSERVADAS
Preencher com base nas avaliações pedagógicas, estudos de caso, escutas docentes e atendimentos registrados.

3. BARREIRAS PEDAGÓGICAS OBSERVADAS
Preencher com base nos registros do AEE e nas observações dos docentes da sala regular.

4. ESTRATÉGIAS PEDAGÓGICAS RECOMENDADAS
Preencher com orientações objetivas, como apoio visual, instruções curtas, mediação gradual, tempo ampliado e diferentes formas de participação.

5. RECOMENDAÇÕES AVALIATIVAS
Considerar participação, evolução individual, comunicação funcional, engajamento e diferentes formas de expressão do conhecimento.

6. ORIENTAÇÕES PARA O DOCENTE
Usar este relatório como apoio pedagógico, sem expor dados familiares ou informações sensíveis desnecessárias.

Fontes reunidas: {fontes}
""".strip()
        return conteudo_fallback, fontes

    prompt = f"""
Você é um especialista em Atendimento Educacional Especializado, educação inclusiva, desenho universal para aprendizagem e articulação pedagógica entre AEE e sala regular.

TAREFA:
Gerar um RELATÓRIO PEDAGÓGICO DE APOIO AO DOCENTE para orientar o professor da área indicada, cruzando as informações do sistema.

NOME DO RELATÓRIO:
Relatório Pedagógico de Apoio ao Docente

DESTINATÁRIO:
Componente/área: {componente_destino}
Destinatário: campo em branco para preenchimento manual no documento final
Ano letivo: {ano_letivo}
Foco solicitado pelo AEE: {foco_docente or 'Não informado'}

REGRAS OBRIGATÓRIAS:
- Não usar linguagem médica ou clínica.
- Não expor dados familiares.
- Não citar CPF, endereço, telefone, nomes de responsáveis ou dados sensíveis.
- Não inventar informações.
- Trabalhar somente com os dados fornecidos.
- Usar linguagem pedagógica, objetiva, acolhedora e institucional.
- O documento deve orientar o professor da sala regular, sem rotular o estudante.
- Focar em participação, aprendizagem, barreiras, potencialidades, estratégias e avaliação inclusiva.
- Se o componente for "Geral – todas as áreas", gerar orientações amplas, sem adaptações específicas de Matemática, História, Português ou outra disciplina.
- Se o componente for uma disciplina específica, utilizar a Base de Conhecimento IA consultada para relacionar, quando pertinente, BNCC, BNCC Computação e estratégias inclusivas da área.
- Não copiar trechos longos da base; sintetizar pedagogicamente em linguagem acessível ao docente.
- Considerar o perfil pedagógico/tecnológico do docente da sala regular, quando registrado na Escuta Docente.
- Não sugerir robótica, impressão 3D, programação ou tecnologias complexas quando o docente indicar baixa familiaridade tecnológica ou não ativar modo maker.
- Quando houver interesse em projetos interdisciplinares, propor caminhos pedagógicos acessíveis e progressivos, deixando claro que o AEE não é reforço escolar.
- Quando os dados forem insuficientes, escrever que a informação não foi localizada nos registros.

DADOS DISPONÍVEIS NO SISTEMA:
{contexto}

BASE DE CONHECIMENTO IA CONSULTADA:
{contexto_base_ia}

Produza o relatório com as seções abaixo:

1. FINALIDADE DO DOCUMENTO
Explique que o relatório apoia o planejamento pedagógico inclusivo do docente.

2. SÍNTESE PEDAGÓGICA FUNCIONAL
Síntese objetiva de como o estudante aprende, participa e responde às mediações.

3. POTENCIALIDADES OBSERVADAS
Listar potencialidades pedagógicas observadas nos registros.

4. BARREIRAS PEDAGÓGICAS E DESAFIOS EM SALA
Listar barreiras educacionais sem linguagem medicalizante.

5. ESTRATÉGIAS QUE TENDEM A FAVORECER A PARTICIPAÇÃO
Recomendações práticas para o docente da área.

6. RECOMENDAÇÕES PARA ADAPTAÇÃO DAS ATIVIDADES
Orientações de adaptação curricular, acessibilidade pedagógica e mediação.

7. RECOMENDAÇÕES AVALIATIVAS
Orientar formas de avaliação inclusiva considerando participação, evolução individual, comunicação, engajamento e diferentes formas de expressão.

8. PONTOS DE ATENÇÃO PARA ARTICULAÇÃO COM O AEE
Indicar como o professor pode dialogar com o AEE, sem expor dados sensíveis.

9. FECHAMENTO INSTITUCIONAL
Finalizar com uma mensagem pedagógica, acolhedora e profissional.
"""
    try:
        resposta = client.responses.create(model="gpt-4.1-mini", input=prompt)
        return (resposta.output_text or "").strip(), fontes
    except Exception as e:
        return f"Não foi possível gerar o relatório com IA agora. Erro: {e}", fontes


# ======================================================
# INFOGRÁFICO DOCENTE COM IA (JSON ESTRUTURADO)
# ======================================================
def normalizar_lista_infografico(valor, limite=6, tamanho_item=95):
    """Garante lista curta para caber no painel visual sem cortar texto."""
    if valor is None:
        itens = []
    elif isinstance(valor, list):
        itens = valor
    else:
        texto = limpar_marcadores_relatorio(str(valor))
        texto = texto.replace("\r", "\n")
        itens = []
        for linha in texto.splitlines():
            partes = re.split(r"\s+-\s+|\s*;\s*", linha)
            for parte in partes:
                parte = parte.strip().lstrip("•-–—0123456789. ").strip()
                if len(parte) > 3:
                    itens.append(parte)
        if not itens:
            itens = [p.strip(" .;:-") for p in re.split(r"(?<=[.;])\s+", texto) if len(p.strip()) > 6]

    saida = []
    vistos = set()
    for item in itens:
        item = re.sub(r"\s+", " ", limpar_marcadores_relatorio(str(item))).strip(" .;:-")
        if not item:
            continue
        if len(item) > tamanho_item:
            item = item[:tamanho_item].rsplit(" ", 1)[0] + "..."
        chave = item.lower()
        if chave not in vistos:
            vistos.add(chave)
            saida.append(item)
        if len(saida) >= limite:
            break
    return saida


def extrair_itens_secao_docente(texto, numero_secao, max_itens=6, tamanho_item=95):
    """Extrai itens de uma seção numerada do relatório docente."""
    texto = str(texto or "")
    padrao = rf"(?is)(?:^|\n)\s*{numero_secao}\.\s*.*?(?=\n\s*\d+\.\s|\Z)"
    achado = re.search(padrao, texto)
    if not achado:
        return []
    trecho = achado.group(0)
    trecho = re.sub(rf"(?is)^\s*{numero_secao}\.\s*[^\n]*", "", trecho).strip()
    return normalizar_lista_infografico(trecho, limite=max_itens, tamanho_item=tamanho_item)


def extrair_json_da_resposta(texto):
    """Extrai JSON mesmo quando a IA devolve texto com ```json ...```."""
    texto = str(texto or "").strip()
    texto = re.sub(r"^```(?:json)?", "", texto.strip(), flags=re.I).strip()
    texto = re.sub(r"```$", "", texto.strip()).strip()

    try:
        return json.loads(texto)
    except Exception:
        pass

    inicio = texto.find("{")
    fim = texto.rfind("}")
    if inicio >= 0 and fim > inicio:
        return json.loads(texto[inicio:fim + 1])
    raise ValueError("A resposta da IA não contém JSON válido.")


def normalizar_dados_infografico_docente(dados, relatorio_docente_txt="", estudante=None):
    """Normaliza o JSON do infográfico com regra ética:
    - campos observacionais só usam o que está documentado;
    - campos de ação pedagógica podem trazer sugestões, desde que não pareçam diagnóstico;
    - quando não houver informação, explicita ausência de registro.
    """
    dados = dados if isinstance(dados, dict) else {}
    perfil = estudante[4] if estudante and len(estudante) > 4 else "Não informado"

    NAO_INFO = "Não informado nos registros analisados."

    # Campos observacionais: não podem ser inventados nem deduzidos pelo CID/condição.
    fallback_observado = {
        "quem_e_estudante": extrair_itens_secao_docente(relatorio_docente_txt, 2, 6, 90) or [NAO_INFO],
        "como_aprende_melhor": extrair_itens_secao_docente(relatorio_docente_txt, 5, 6, 80) or [NAO_INFO],
        "o_que_dificulta": extrair_itens_secao_docente(relatorio_docente_txt, 4, 6, 80) or [NAO_INFO],
        "potencialidades": extrair_itens_secao_docente(relatorio_docente_txt, 3, 8, 70) or [NAO_INFO],
        "barreiras_pedagogicas": extrair_itens_secao_docente(relatorio_docente_txt, 4, 6, 75) or [NAO_INFO],
    }

    # Campo contextual: a condição/CID não é evidência individual; é referência para cuidado pedagógico.
    fallback_condicao = [
        f"Condição informada: {perfil}.",
        "A condição é referência contextual, não diagnóstico produzido pelo sistema.",
        "Cada estudante aprende, comunica e participa de forma própria.",
    ] if perfil and perfil != "Não informado" else [NAO_INFO]

    # Campos sugestivos: podem orientar o docente, mas sempre como possibilidade pedagógica.
    fallback_sugestivo = {
        "o_que_funciona_em_sala": extrair_itens_secao_docente(relatorio_docente_txt, 5, 8, 70) or [
            "Sugestão: usar instruções curtas e objetivas.",
            "Sugestão: organizar atividades em etapas.",
            "Sugestão: oferecer apoio visual quando necessário.",
        ],
        "atencao_docente": extrair_itens_secao_docente(relatorio_docente_txt, 8, 6, 80) or [
            "Sugestão: observar respostas do estudante antes de ampliar a demanda.",
            "Sugestão: evitar generalizações a partir da condição informada.",
        ],
        "avaliacao_flexivel": extrair_itens_secao_docente(relatorio_docente_txt, 7, 6, 80) or [
            "Sugestão: considerar diferentes formas de demonstrar aprendizagem.",
            "Sugestão: valorizar processo, participação e evolução registrada.",
        ],
        "articulacao_aee": extrair_itens_secao_docente(relatorio_docente_txt, 8, 5, 85) or [
            "Sugestão: dialogar com o AEE para ajustar estratégias e recursos.",
            "Sugestão: registrar o que funcionou na sala regular.",
        ],
        "recursos_sugeridos": [
            "Sugestão: apoio visual estruturado.",
            "Sugestão: materiais concretos, quando fizer sentido pedagógico.",
            "Sugestão: checklist de etapas.",
            "Sugestão: tecnologia com finalidade pedagógica.",
        ],
        "indicadores_de_avanco": [
            "Sugestão: observar maior participação nas atividades.",
            "Sugestão: observar permanência na tarefa com menor apoio.",
            "Sugestão: observar autonomia gradual e comunicação funcional.",
        ],
    }

    campos_observados = [
        "quem_e_estudante", "como_aprende_melhor", "o_que_dificulta",
        "potencialidades", "barreiras_pedagogicas"
    ]
    campos_sugestivos = [
        "o_que_funciona_em_sala", "atencao_docente", "avaliacao_flexivel",
        "articulacao_aee", "recursos_sugeridos", "indicadores_de_avanco"
    ]

    normalizado = {}

    for campo in campos_observados:
        limite = 8 if campo == "potencialidades" else 6
        itens = normalizar_lista_infografico(dados.get(campo), limite=limite, tamanho_item=85)
        if not itens:
            itens = normalizar_lista_infografico(fallback_observado.get(campo), limite=limite, tamanho_item=85)
        normalizado[campo] = itens or [NAO_INFO]

    normalizado["condicao_em_linguagem_pedagogica"] = normalizar_lista_infografico(
        dados.get("condicao_em_linguagem_pedagogica") or fallback_condicao,
        limite=4,
        tamanho_item=88,
    ) or [NAO_INFO]

    for campo in campos_sugestivos:
        limite = 8 if campo in ["o_que_funciona_em_sala", "recursos_sugeridos"] else 6
        itens = normalizar_lista_infografico(dados.get(campo), limite=limite, tamanho_item=85)
        if not itens:
            itens = normalizar_lista_infografico(fallback_sugestivo.get(campo), limite=limite, tamanho_item=85)
        # Garante que ações geradas por fallback fiquem como sugestões, não como fatos.
        if campo in campos_sugestivos:
            itens = [i if i.lower().startswith(("sugestão", "pode", "considerar", "avaliar")) else f"Sugestão: {i}" for i in itens]
        normalizado[campo] = itens

    foco = str(dados.get("foco_pedagogico") or "").strip()
    if not foco:
        foco = "Favorecer participação, aprendizagem e acessibilidade, valorizando evidências registradas e usando estratégias pedagógicas como possibilidades flexíveis."
    normalizado["foco_pedagogico"] = re.sub(r"\s+", " ", foco).strip()[:220]

    observacao = str(dados.get("observacao_etica") or "").strip()
    if not observacao:
        observacao = (
            "As informações observadas derivam dos registros educacionais analisados. "
            "As estratégias são sugestões pedagógicas e não representam diagnóstico, laudo ou definição fixa sobre o estudante."
        )
    normalizado["observacao_etica"] = re.sub(r"\s+", " ", observacao).strip()[:280]
    return normalizado


def gerar_conteudo_infografico_docente_ia(relatorio_docente_txt, estudante, ano_letivo, componente):
    """IA transforma o relatório textual em JSON curto e pronto para o painel infográfico."""
    client = obter_cliente_openai()

    if client is None:
        return normalizar_dados_infografico_docente({}, relatorio_docente_txt, estudante)

    codigo = estudante[1] if estudante and len(estudante) > 1 else "Não informado"
    ano_serie = estudante[2] if estudante and len(estudante) > 2 else "Não informado"
    turma = estudante[3] if estudante and len(estudante) > 3 else "Não informado"
    perfil = estudante[4] if estudante and len(estudante) > 4 else "Não informado"

    prompt = f"""
Você é especialista em AEE, educação inclusiva, DUA e design de informação pedagógica.

TAREFA:
Leia o RELATÓRIO PEDAGÓGICO DE APOIO AO DOCENTE e transforme o conteúdo em JSON curto para um PAINEL INFOGRÁFICO DOCENTE.

REGRA CENTRAL:
O sistema NÃO produz diagnóstico, NÃO confirma laudo e NÃO define o estudante pela condição informada.
O painel deve apoiar o professor da sala regular com informações pedagógicas e sugestões flexíveis.

FONTES PERMITIDAS:
Use EXCLUSIVAMENTE as informações presentes nos registros enviados ao prompt:
- relatório pedagógico docente;
- avaliação pedagógica, se citada no relatório;
- estudo de caso, se citado no relatório;
- registros do AEE, se citados no relatório;
- escuta docente, se citada no relatório.

REGRAS DE EVIDÊNCIA:
- Não invente informações.
- Não crie habilidades, interesses, comportamentos ou dificuldades sem registro.
- Não deduza características individuais apenas pela condição, CID ou laudo.
- Se uma informação não estiver nos registros, escreva: "Não informado nos registros analisados."
- Dados observados do estudante devem vir dos documentos, não da hipótese da IA.

USO DA CONDIÇÃO/CID:
A condição informada no cadastro pode ser usada APENAS como referência contextual.
Ela pode orientar POSSIBILIDADES pedagógicas, mas nunca deve ser tratada como evidência individual.
Quando usar a condição para orientar o professor, escreva com linguagem sugestiva e não determinista:
- "pode necessitar..."
- "pode se beneficiar..."
- "em determinados contextos, pode..."
- "alguns estudantes com perfil semelhante podem..."
Sempre deixe claro que cada estudante aprende, comunica e participa de forma própria.

LINGUAGEM:
- Não use linguagem clínica ou medicalizante.
- Não exponha dados familiares ou sensíveis.
- Use frases curtas, objetivas e úteis para o professor da sala regular.
- Cada item deve ter no máximo 90 caracteres.
- Priorize potencialidades e formas de participação antes das barreiras.
- Estratégias, recursos, avaliação e atenção docente devem aparecer como SUGESTÕES.

NUNCA escreva:
- "O estudante tem dificuldade de..." se isso não estiver documentado.
- "O estudante não consegue..." sem evidência nos registros.
- "Por ser TEA/CID..., ele apresenta...".

PODE escrever:
- "Nos registros, observa-se..." quando houver evidência.
- "Sugestão: ..." para ação pedagógica.
- "Pode se beneficiar de..." quando for orientação pedagógica não determinista.

IDENTIFICAÇÃO PEDAGÓGICA:
Código: {codigo}
Ano/Série: {ano_serie}
Turma: {turma}
Ano letivo: {ano_letivo}
Componente/área: {componente}
Condição informada no cadastro: {perfil}

JSON obrigatório, exatamente com estas chaves:
{{
  "quem_e_estudante": ["somente fatos observados nos registros"],
  "condicao_em_linguagem_pedagogica": ["condição como referência contextual, sem diagnóstico"],
  "como_aprende_melhor": ["somente formas de aprendizagem documentadas"],
  "o_que_dificulta": ["somente dificuldades registradas"],
  "potencialidades": ["habilidades e interesses observados nos registros"],
  "barreiras_pedagogicas": ["barreiras documentadas"],
  "o_que_funciona_em_sala": ["sugestões pedagógicas ou estratégias registradas"],
  "atencao_docente": ["sugestões de atenção pedagógica"],
  "avaliacao_flexivel": ["sugestões avaliativas flexíveis"],
  "articulacao_aee": ["ações de articulação AEE-sala regular"],
  "recursos_sugeridos": ["recursos possíveis, em linguagem sugestiva"],
  "indicadores_de_avanco": ["indicadores observáveis, não diagnósticos"],
  "foco_pedagogico": "frase central curta",
  "observacao_etica": "observação dizendo que não é diagnóstico e que sugestões são pedagógicas"
}}

RELATÓRIO PEDAGÓGICO:
{relatorio_docente_txt}
"""
    try:
        resposta = client.responses.create(model="gpt-4.1-mini", input=prompt)
        dados = extrair_json_da_resposta(resposta.output_text or "")
        return normalizar_dados_infografico_docente(dados, relatorio_docente_txt, estudante)
    except Exception as e:
        st.warning(f"Não foi possível estruturar o infográfico com IA. Usando síntese automática local. Erro: {e}")
        return normalizar_dados_infografico_docente({}, relatorio_docente_txt, estudante)


# ======================================================
# PAINEL VISUAL DE APOIO AO DOCENTE
# ======================================================
def extrair_secao_relatorio_docente(texto, numero_secao, limite=700):
    """Extrai uma seção numerada do relatório docente gerado pela IA.
    Caso a seção não exista, retorna texto vazio.
    """
    texto = str(texto or "")
    padrao = rf"(?is)(?:^|\n)\s*{numero_secao}\.\s*.*?(?=\n\s*\d+\.\s|\Z)"
    achado = re.search(padrao, texto)
    if not achado:
        return ""
    trecho = achado.group(0)
    trecho = re.sub(rf"(?is)^\s*{numero_secao}\.\s*[^\n]*", "", trecho).strip()
    trecho = limpar_marcadores_relatorio(trecho)
    trecho = re.sub(r"\s+", " ", trecho).strip()
    if len(trecho) > limite:
        trecho = trecho[:limite].rsplit(" ", 1)[0] + "..."
    return trecho


def texto_padrao_painel_docente(valor, padrao):
    valor = str(valor or "").strip()
    return valor if valor else padrao



def gerar_dados_painel_visual_docente(estudante, conteudo_relatorio):
    """Monta os cards do relatório visual docente usando o relatório textual e o cadastro do estudante.

    O foco é pedagógico: ajudar o professor a perceber quem é o estudante,
    o que a condição informada pode significar no contexto escolar e, principalmente,
    quais habilidades, interesses e potencialidades podem orientar a aprendizagem.
    """
    perfil_cadastrado = estudante[4] if estudante and len(estudante) > 4 else "Não informado"

    quem_estudante = texto_padrao_painel_docente(
        extrair_secao_relatorio_docente(conteudo_relatorio, 2),
        "Estudante com trajetória singular, formas próprias de comunicação, participação e aprendizagem. Observe seus sinais, interesses, respostas e formas de demonstrar compreensão."
    )

    significado_condicao = texto_padrao_painel_docente(
        f"Condição informada: {perfil_cadastrado}. Esta informação deve ser compreendida pedagogicamente, sem reduzir o estudante ao laudo ou ao CID. No contexto escolar, ela ajuda a planejar apoios, acessibilidade, comunicação, organização da rotina e formas de participação.",
        "Condição não informada. Registrar apenas informações educacionais necessárias para favorecer participação, acessibilidade e aprendizagem."
    )

    como_aprende = texto_padrao_painel_docente(
        extrair_secao_relatorio_docente(conteudo_relatorio, 5) or extrair_secao_relatorio_docente(conteudo_relatorio, 6),
        "Aprende melhor quando a atividade respeita seu ritmo, sua forma de comunicação, seus interesses e seu alcance, com instruções claras, apoio visual, mediação gradual e recursos concretos."
    )

    potencialidades = texto_padrao_painel_docente(
        extrair_secao_relatorio_docente(conteudo_relatorio, 3),
        "Registrar habilidades, interesses, preferências, formas de participação, autonomia possível, comunicação já existente e recursos que favorecem engajamento."
    )

    barreiras = texto_padrao_painel_docente(
        extrair_secao_relatorio_docente(conteudo_relatorio, 4),
        "Observar barreiras pedagógicas, comunicacionais, sensoriais, curriculares ou atitudinais que possam dificultar a participação."
    )

    estrategias = texto_padrao_painel_docente(
        extrair_secao_relatorio_docente(conteudo_relatorio, 5) or extrair_secao_relatorio_docente(conteudo_relatorio, 6),
        "Utilizar apoio visual, CAA quando necessário, instruções curtas, divisão da tarefa em etapas, tempo ampliado, demonstração prática e valorização das tentativas."
    )

    pontos_atencao = texto_padrao_painel_docente(
        extrair_secao_relatorio_docente(conteudo_relatorio, 8),
        "Evitar pressa, excesso de informações simultâneas e avaliação baseada em uma única forma de resposta. Registrar avanços e estratégias que funcionaram."
    )

    return {
        "quem_estudante": quem_estudante,
        "significado_condicao": significado_condicao,
        "como_aprende": como_aprende,
        "barreiras_observadas": barreiras,
        "potencialidades_interesses": potencialidades,
        "pontos_atencao": pontos_atencao,
        "estrategias_rapidas": estrategias,
        # Compatibilidade com versões anteriores do código
        "perfil_pedagogico": quem_estudante,
    }


def gerar_relatorio_visual_docente_html(
    estudante,
    ano_letivo,
    componente_destino,
    perfil_pedagogico,
    como_aprende,
    barreiras_observadas,
    potencialidades_interesses,
    pontos_atencao,
    estrategias_rapidas,
    conteudo_relatorio,
    fontes_geradas="",
):
    """Gera HTML visual do Painel de Apoio ao Docente.
    Pode ser exibido no Streamlit e baixado como arquivo HTML.
    """
    codigo = estudante[1] if estudante and len(estudante) > 1 else "Não informado"
    ano_serie = estudante[2] if estudante and len(estudante) > 2 else "Não informado"
    turma = estudante[3] if estudante and len(estudante) > 3 else "Não informado"

    perfil_pedagogico = escape(str(perfil_pedagogico or ""))
    como_aprende = escape(str(como_aprende or ""))
    barreiras_observadas = escape(str(barreiras_observadas or ""))
    potencialidades_interesses = escape(str(potencialidades_interesses or ""))
    pontos_atencao = escape(str(pontos_atencao or ""))
    estrategias_rapidas = escape(str(estrategias_rapidas or ""))
    conteudo_html = escape(str(conteudo_relatorio or "")).replace("\n", "<br>")
    fontes_html = escape(str(fontes_geradas or "Não informado"))

    html_relatorio = f"""
<style>
.visual-docente-wrapper {{
    background: linear-gradient(180deg, #f8fafc 0%, #eef6ff 100%);
    border-radius: 30px;
    padding: 28px;
    border: 1px solid #dbeafe;
}}
.visual-docente-header {{
    background: linear-gradient(135deg, #0f172a 0%, #1e3a8a 100%);
    color: #ffffff;
    border-radius: 28px;
    padding: 36px 40px;
    margin-bottom: 30px;
    box-shadow: 0 14px 36px rgba(15,23,42,0.20);
}}
.visual-docente-header h1 {{
    font-size: 42px;
    line-height: 1.15;
    margin: 0 0 14px 0;
    font-weight: 950;
    color: #ffffff;
}}
.visual-docente-header p {{
    font-size: 21px;
    line-height: 1.65;
    color: #e0f2fe;
    margin: 0;
}}
.visual-docente-meta {{
    display: flex;
    flex-wrap: wrap;
    gap: 12px;
    margin-top: 22px;
}}
.visual-docente-chip {{
    background: rgba(255,255,255,0.14);
    color: #ffffff;
    border: 1px solid rgba(255,255,255,0.24);
    border-radius: 999px;
    padding: 9px 15px;
    font-size: 17px;
    font-weight: 800;
}}
.card-neuro {{
    background: #ffffff;
    border-radius: 28px;
    padding: 42px;
    margin-top: 20px;
    margin-bottom: 32px;
    border: 1px solid #dbeafe;
    box-shadow: 0 12px 35px rgba(15,23,42,0.08);
}}
.card-neuro h2 {{
    font-size: 40px;
    font-weight: 950;
    color: #0f172a;
    margin-bottom: 14px;
    line-height: 1.2;
}}
.card-neuro .descricao-painel {{
    font-size: 22px;
    line-height: 1.8;
    color: #475569;
    margin-bottom: 28px;
    font-weight: 500;
}}
.grid-visual {{
    display: grid;
    grid-template-columns: repeat(auto-fit, minmax(340px, 1fr));
    gap: 24px;
    margin-top: 30px;
}}
.mini-card {{
    border-radius: 24px;
    padding: 32px;
    box-shadow: 0 8px 24px rgba(15,23,42,0.06);
    border-left: 10px solid #2563eb;
}}
.mini-card h3 {{
    font-size: 30px;
    font-weight: 950;
    margin-bottom: 18px;
    color: #0f172a;
    line-height: 1.3;
}}
.mini-card p {{
    font-size: 22px;
    line-height: 1.9;
    color: #334155;
    font-weight: 500;
    margin: 0;
}}
.azul {{ background: #dbeafe; border-left-color: #2563eb; }}
.roxo {{ background: #ede9fe; border-left-color: #7c3aed; }}
.laranja {{ background: #ffedd5; border-left-color: #ea580c; }}
.verde {{ background: #dcfce7; border-left-color: #16a34a; }}
.vermelho {{ background: #fee2e2; border-left-color: #dc2626; }}
.ciano {{ background: #cffafe; border-left-color: #0891b2; }}
.relatorio-texto {{
    background: #ffffff;
    border-radius: 24px;
    padding: 38px;
    margin-top: 30px;
    border: 1px solid #e2e8f0;
    box-shadow: 0 10px 28px rgba(15,23,42,0.06);
    font-size: 22px;
    line-height: 1.9;
    color: #334155;
}}
.relatorio-texto h2, .relatorio-texto h3 {{
    font-size: 30px;
    font-weight: 950;
    color: #0f172a;
    margin-top: 18px;
    margin-bottom: 16px;
}}
.fontes-painel {{
    background: #f8fafc;
    border: 1px solid #e2e8f0;
    border-radius: 18px;
    padding: 22px;
    margin-top: 22px;
    color: #475569;
    font-size: 17px;
    line-height: 1.6;
}}
@media print {{
    .visual-docente-wrapper {{ padding: 0; background: #ffffff; border: none; }}
    .mini-card {{ break-inside: avoid; }}
    .card-neuro {{ break-inside: avoid; }}
}}
</style>

<div class="visual-docente-wrapper">
    <div class="visual-docente-header">
        <h1>🧩 Painel Visual de Apoio ao Docente</h1>
        <p>Resumo pedagógico rápido para apoiar o professor da sala regular no planejamento inclusivo, na adaptação das atividades e na articulação com o AEE.</p>
        <div class="visual-docente-meta">
            <span class="visual-docente-chip">Código: {escape(str(codigo or 'Não informado'))}</span>
            <span class="visual-docente-chip">Ano/Série: {escape(str(ano_serie or 'Não informado'))}</span>
            <span class="visual-docente-chip">Turma: {escape(str(turma or 'Não informado'))}</span>
            <span class="visual-docente-chip">Ano letivo: {escape(str(ano_letivo or 'Não informado'))}</span>
            <span class="visual-docente-chip">Área: {escape(str(componente_destino or 'Geral'))}</span>
        </div>
    </div>

    <div class="card-neuro">
        <h2>🧭 Apoio pedagógico rápido</h2>
        <p class="descricao-painel">Este painel não substitui o relatório completo. Ele funciona como um guia visual de consulta rápida para o professor, destacando informações pedagógicas essenciais para favorecer participação, aprendizagem e acessibilidade.</p>

        <div class="grid-visual">
            <div class="mini-card azul">
                <h3>📘 Perfil pedagógico</h3>
                <p>{perfil_pedagogico}</p>
            </div>

            <div class="mini-card roxo">
                <h3>🧠 Como o estudante aprende</h3>
                <p>{como_aprende}</p>
            </div>

            <div class="mini-card laranja">
                <h3>⚠️ Barreiras observadas</h3>
                <p>{barreiras_observadas}</p>
            </div>

            <div class="mini-card verde">
                <h3>🌱 Potencialidades e interesses</h3>
                <p>{potencialidades_interesses}</p>
            </div>

            <div class="mini-card vermelho">
                <h3>❤️ O que merece atenção</h3>
                <p>{pontos_atencao}</p>
            </div>

            <div class="mini-card ciano">
                <h3>🎯 Estratégias rápidas</h3>
                <p>{estrategias_rapidas}</p>
            </div>
        </div>
    </div>

    <div class="relatorio-texto">
        <h2>📄 Relatório completo de apoio ao docente</h2>
        {conteudo_html}
    </div>

    <div class="fontes-painel">
        <strong>Fontes utilizadas pelo sistema:</strong><br>
        {fontes_html}
    </div>
</div>
"""
    return html_relatorio



def gerar_pdf_relatorio_visual_docente(
    estudante,
    ano_letivo,
    componente_destino,
    quem_estudante,
    significado_condicao,
    como_aprende,
    barreiras_observadas,
    potencialidades_interesses,
    pontos_atencao,
    estrategias_rapidas,
    conteudo_relatorio,
    fontes_geradas="",
    nome_base=None,
):
    """Gera o PDF pedagógico limpo em layout corrigido.

    V49: substitui tabelas longas por desenho controlado em canvas, com cards reais,
    quebras previsíveis e textos resumidos para evitar sobreposição, corte e aparência
    de tabela defeituosa.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase.pdfmetrics import stringWidth

    codigo = estudante[1] if estudante and len(estudante) > 1 else "Não informado"
    ano_serie = estudante[2] if estudante and len(estudante) > 2 else "Não informado"
    turma = estudante[3] if estudante and len(estudante) > 3 else "Não informado"
    perfil = estudante[4] if estudante and len(estudante) > 4 else "Não informado"

    if not nome_base:
        nome_base = f"Relatorio_Visual_Apoio_Docente_{codigo}_{ano_letivo}"
    nome_arquivo = f"{nome_base}.pdf".replace("/", "-").replace("\\", "-")
    caminho_pdf = caminho_relatorio_visual_docente(nome_arquivo)

    c = canvas.Canvas(str(caminho_pdf), pagesize=A4)
    W, H = A4

    def col(hex_value):
        return colors.HexColor(hex_value)

    def limpar_texto_curto(texto):
        texto = limpar_marcadores_relatorio(str(texto or ""))
        texto = texto.replace("•", " ").replace("–", "-").replace("—", "-")
        texto = re.sub(r"\s+", " ", texto).strip()
        texto = texto.replace("Demonstra demonstração", "Demonstra")
        return texto or "Informação não registrada."

    def wrap(texto, fonte, tamanho, largura):
        palavras = limpar_texto_curto(texto).split()
        linhas, atual = [], ""
        for palavra in palavras:
            teste = (atual + " " + palavra).strip()
            if stringWidth(teste, fonte, tamanho) <= largura:
                atual = teste
            else:
                if atual:
                    linhas.append(atual)
                atual = palavra
        if atual:
            linhas.append(atual)
        return linhas

    def texto_caixa(texto, x, y, largura, fonte="Helvetica", tamanho=8.2, entrelinha=10, cor="#334155", max_linhas=8):
        linhas = wrap(texto, fonte, tamanho, largura)
        if max_linhas and len(linhas) > max_linhas:
            linhas = linhas[:max_linhas]
            linhas[-1] = linhas[-1].rstrip(".,;:") + "..."
        c.setFont(fonte, tamanho)
        c.setFillColor(col(cor))
        for linha in linhas:
            c.drawString(x, y, linha)
            y -= entrelinha
        return y

    def itens(texto, fallback=None, max_itens=6, max_chars=120):
        fallback = fallback or ["Registrar observações pedagógicas relevantes para orientar a prática docente."]
        bruto = limpar_marcadores_relatorio(str(texto or ""))
        bruto = bruto.replace("\r", "\n")
        candidatos = []
        for linha in bruto.splitlines():
            linha = linha.strip()
            if not linha:
                continue
            # Se a IA devolveu tudo em uma linha com " - ", separa em itens.
            partes = re.split(r"\s+-\s+", linha) if " - " in linha else [linha]
            for parte in partes:
                parte = parte.strip().lstrip("•-–—0123456789. ").strip()
                if len(parte) > 5:
                    candidatos.append(parte)
        if not candidatos:
            partes = re.split(r"(?<=[.;])\s+", limpar_texto_curto(bruto))
            candidatos = [p.strip(" .;:-") for p in partes if len(p.strip()) > 8]
        if not candidatos:
            candidatos = fallback
        saida = []
        for item in candidatos:
            item = limpar_texto_curto(item)
            if len(item) > max_chars:
                item = item[:max_chars].rsplit(" ", 1)[0] + "..."
            if item not in saida:
                saida.append(item)
            if len(saida) >= max_itens:
                break
        return saida

    def cabecalho(titulo_pagina="RELATÓRIO VISUAL DE APOIO AO DOCENTE", pagina=1):
        c.setFillColor(col("#ffffff"))
        c.rect(0, 0, W, H, stroke=0, fill=1)
        c.setFillColor(col("#0f172a"))
        c.setFont("Helvetica-Bold", 15)
        c.drawString(1.25*cm, H-1.2*cm, titulo_pagina)
        c.setFont("Helvetica", 7.2)
        c.setFillColor(col("#64748b"))
        c.drawRightString(W-1.25*cm, H-0.75*cm, f"INCLUISRM • finalidade pedagógica • Página {pagina}")
        c.setStrokeColor(col("#e2e8f0"))
        c.line(1.25*cm, H-1.45*cm, W-1.25*cm, H-1.45*cm)

    def meta_bar(y):
        dados = [
            ("Código", codigo), ("Ano/Série", ano_serie), ("Turma", turma),
            ("Ano letivo", ano_letivo), ("Área", componente_destino), ("Data", agora_local().strftime("%d/%m/%Y")),
        ]
        x = 1.25*cm
        gap = 0.12*cm
        w = (W - 2.5*cm - 5*gap) / 6
        for label, valor in dados:
            c.setFillColor(col("#f8fafc"))
            c.setStrokeColor(col("#cbd5e1"))
            c.roundRect(x, y, w, 0.86*cm, 5, stroke=1, fill=1)
            c.setFillColor(col("#475569"))
            c.setFont("Helvetica-Bold", 6.5)
            c.drawString(x+0.12*cm, y+0.55*cm, label)
            c.setFillColor(col("#0f172a"))
            c.setFont("Helvetica", 6.8)
            texto_caixa(str(valor), x+0.12*cm, y+0.28*cm, w-0.24*cm, tamanho=6.8, entrelinha=7, max_linhas=1, cor="#0f172a")
            x += w + gap

    def card(x, y, w, h, titulo_card, texto, cor_borda, cor_fundo, max_linhas=7):
        c.setFillColor(col(cor_fundo))
        c.setStrokeColor(col(cor_borda))
        c.roundRect(x, y, w, h, 9, stroke=1, fill=1)
        c.setFillColor(col(cor_borda))
        c.roundRect(x, y+h-0.62*cm, w, 0.62*cm, 9, stroke=0, fill=1)
        c.setFillColor(col("#ffffff"))
        c.setFont("Helvetica-Bold", 8.4)
        c.drawString(x+0.28*cm, y+h-0.39*cm, titulo_card.upper()[:48])
        texto_caixa(texto, x+0.28*cm, y+h-0.92*cm, w-0.56*cm, tamanho=7.7, entrelinha=9.4, max_linhas=max_linhas)

    def bloco_lista(x, y, w, h, titulo_bloco, lista, cor_borda, cor_fundo, max_itens=6):
        c.setFillColor(col(cor_fundo))
        c.setStrokeColor(col(cor_borda))
        c.roundRect(x, y, w, h, 8, stroke=1, fill=1)
        c.setFillColor(col("#0f172a"))
        c.setFont("Helvetica-Bold", 10)
        c.drawString(x+0.3*cm, y+h-0.45*cm, titulo_bloco)
        yy = y+h-0.82*cm
        c.setFont("Helvetica", 7.8)
        for item in lista[:max_itens]:
            c.setFillColor(col(cor_borda))
            c.circle(x+0.36*cm, yy+2.2, 2.1, stroke=0, fill=1)
            yy = texto_caixa(item, x+0.55*cm, yy, w-0.85*cm, tamanho=7.55, entrelinha=8.8, max_linhas=2)
            yy -= 0.08*cm
            if yy < y + 0.22*cm:
                break

    # Dados de seções
    sec_pot = extrair_secao_relatorio_docente(conteudo_relatorio, 3) or potencialidades_interesses
    sec_bar = extrair_secao_relatorio_docente(conteudo_relatorio, 4) or barreiras_observadas
    sec_est = extrair_secao_relatorio_docente(conteudo_relatorio, 5) or estrategias_rapidas
    sec_adapt = extrair_secao_relatorio_docente(conteudo_relatorio, 6)
    sec_aval = extrair_secao_relatorio_docente(conteudo_relatorio, 7)
    sec_aee = extrair_secao_relatorio_docente(conteudo_relatorio, 8)
    sec_fechamento = extrair_secao_relatorio_docente(conteudo_relatorio, 9)

    pot = itens(sec_pot, ["Valorizar interesses e habilidades já observadas.", "Aproveitar formas de comunicação já existentes.", "Priorizar atividades práticas e visuais."], 6, 115)
    bar = itens(sec_bar, ["Comandos longos podem reduzir compreensão.", "Ambientes com excesso de estímulos podem dificultar atenção."], 6, 110)
    est = itens(sec_est, ["Usar comandos curtos e objetivos.", "Dividir tarefas em etapas.", "Oferecer tempo ampliado."], 6, 110)
    adapt = itens(sec_adapt, ["Reduzir volume sem perder o objetivo pedagógico.", "Permitir múltiplas formas de resposta."], 5, 120)
    aval = itens(sec_aval, ["Avaliar participação, engajamento e evolução individual.", "Evitar que a escrita seja o único critério."], 5, 120)
    aee = itens(sec_aee, ["Compartilhar observações com o AEE.", "Solicitar apoio para recursos visuais e materiais adaptados."], 5, 120)

    # Página 1
    cabecalho("RELATÓRIO VISUAL DE APOIO AO DOCENTE", 1)
    c.setFont("Helvetica", 8.4)
    c.setFillColor(col("#475569"))
    texto_caixa("Apoio rápido para reconhecer quem é o estudante, compreender o significado pedagógico da condição informada e valorizar habilidades, potencialidades e formas próprias de aprender.", 1.25*cm, H-1.88*cm, W-2.5*cm, tamanho=8.4, entrelinha=10, max_linhas=2, cor="#475569")
    meta_bar(H-3.0*cm)

    c.setFillColor(col("#eff6ff"))
    c.setStrokeColor(col("#93c5fd"))
    c.roundRect(1.25*cm, H-4.2*cm, W-2.5*cm, 0.88*cm, 8, stroke=1, fill=1)
    texto_caixa("Antes de adaptar a atividade, observe como o estudante compreende, responde, comunica e participa.", 1.55*cm, H-3.82*cm, W-3.1*cm, fonte="Helvetica-Bold", tamanho=8.6, entrelinha=10, max_linhas=1, cor="#1e3a8a")

    c.setFont("Helvetica-Bold", 12)
    c.setFillColor(col("#0f172a"))
    c.drawString(1.25*cm, H-4.8*cm, "Painel de percepção docente")

    left = 1.25*cm
    gap = 0.28*cm
    cw = (W - 2.5*cm - gap) / 2
    ch = 3.45*cm
    y1 = H-8.55*cm
    card(left, y1, cw, ch, "Quem é o estudante", quem_estudante, "#2563eb", "#eff6ff", 8)
    card(left+cw+gap, y1, cw, ch, "Condição em linguagem pedagógica", significado_condicao, "#7c3aed", "#f5f3ff", 8)
    y2 = y1 - ch - 0.32*cm
    card(left, y2, cw, ch, "Como aprende melhor", como_aprende, "#0891b2", "#ecfeff", 8)
    card(left+cw+gap, y2, cw, ch, "Potencialidades e interesses", potencialidades_interesses, "#16a34a", "#f0fdf4", 8)
    y3 = y2 - ch - 0.32*cm
    card(left, y3, cw, ch, "Barreiras observadas", barreiras_observadas, "#ea580c", "#fff7ed", 8)
    card(left+cw+gap, y3, cw, ch, "Estratégias e atenção", estrategias_rapidas + " " + pontos_atencao, "#dc2626", "#fff1f2", 8)

    c.showPage()

    # Página 2
    cabecalho("ORIENTAÇÕES PRÁTICAS PARA SALA REGULAR", 2)
    texto_caixa("Síntese para transformar a percepção do estudante em decisões pedagógicas: participação, comunicação, acessibilidade, atividade e avaliação.", 1.25*cm, H-1.9*cm, W-2.5*cm, tamanho=8.3, entrelinha=10, max_linhas=2, cor="#475569")
    y = H-7.0*cm
    bloco_lista(1.25*cm, y, W-2.5*cm, 4.35*cm, "Potencialidades a valorizar", pot, "#16a34a", "#f0fdf4")
    y -= 4.72*cm
    bloco_lista(1.25*cm, y, W-2.5*cm, 4.35*cm, "Barreiras que podem dificultar a participação", bar, "#ea580c", "#fff7ed")
    y -= 4.72*cm
    bloco_lista(1.25*cm, y, W-2.5*cm, 4.35*cm, "Estratégias que favorecem participação", est, "#2563eb", "#eff6ff")
    y -= 4.72*cm
    bloco_lista(1.25*cm, y, W-2.5*cm, 3.65*cm, "Adaptação das atividades", adapt, "#64748b", "#f8fafc", 5)
    c.showPage()

    # Página 3
    cabecalho("AVALIAÇÃO, ACOMPANHAMENTO E ARTICULAÇÃO COM O AEE", 3)
    texto_caixa("Pontos para registro docente, acompanhamento da evolução e alinhamento com o professor do AEE.", 1.25*cm, H-1.9*cm, W-2.5*cm, tamanho=8.3, entrelinha=10, max_linhas=2, cor="#475569")
    y = H-7.1*cm
    bloco_lista(1.25*cm, y, W-2.5*cm, 4.55*cm, "Recomendações avaliativas", aval, "#7c3aed", "#f5f3ff")
    y -= 4.95*cm
    bloco_lista(1.25*cm, y, W-2.5*cm, 4.55*cm, "Articulação com o AEE", aee, "#0891b2", "#ecfeff")

    y -= 3.0*cm
    c.setFillColor(col("#f8fafc"))
    c.setStrokeColor(col("#cbd5e1"))
    c.roundRect(1.25*cm, y, W-2.5*cm, 2.35*cm, 8, stroke=1, fill=1)
    c.setFont("Helvetica-Bold", 10)
    c.setFillColor(col("#0f172a"))
    c.drawString(1.55*cm, y+1.88*cm, "Fechamento pedagógico")
    texto_caixa(sec_fechamento or "Valorizar potencialidades, respeitar formas próprias de comunicação e promover adaptações pedagógicas são caminhos para uma inclusão efetiva.", 1.55*cm, y+1.52*cm, W-3.1*cm, tamanho=7.9, entrelinha=9.4, max_linhas=5)

    c.setFont("Helvetica", 8)
    c.setFillColor(col("#334155"))
    c.drawString(1.25*cm, 3.0*cm, "Professor(a) AEE: ___________________________________________")
    c.drawString(1.25*cm, 2.45*cm, "Coordenação/Gestão: _________________________________________")
    c.drawString(1.25*cm, 1.9*cm, "Responsável: ________________________________________________")
    texto_caixa("Fontes utilizadas pelo sistema: " + str(fontes_geradas or "Avaliação pedagógica, estudo de caso e entrevista familiar usados apenas como contexto pedagógico."), 1.25*cm, 1.18*cm, W-2.5*cm, tamanho=6.6, entrelinha=7.4, max_linhas=2, cor="#64748b")
    texto_caixa("Finalidade exclusivamente pedagógica. Não expõe dados familiares ou informações sensíveis desnecessárias.", 1.25*cm, 0.62*cm, W-2.5*cm, tamanho=6.6, entrelinha=7.4, max_linhas=1, cor="#64748b")
    c.save()
    return str(caminho_pdf)


def gerar_pdf_infografico_docente(
    estudante,
    ano_letivo,
    componente_destino,
    quem_estudante,
    significado_condicao,
    como_aprende,
    barreiras_observadas,
    potencialidades_interesses,
    pontos_atencao,
    estrategias_rapidas,
    conteudo_relatorio,
    fontes_geradas="",
    nome_base=None,
):
    """Gera painel infográfico em PDF com layout corrigido.

    V49: mantém a proposta visual, mas evita texto cortado usando resumos, cards com altura
    fixa, segunda página complementar e sem caracteres especiais problemáticos.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase.pdfmetrics import stringWidth

    codigo = estudante[1] if estudante and len(estudante) > 1 else "Não informado"
    ano_serie = estudante[2] if estudante and len(estudante) > 2 else "Não informado"
    turma = estudante[3] if estudante and len(estudante) > 3 else "Não informado"
    perfil = estudante[4] if estudante and len(estudante) > 4 else "Não informado"

    if not nome_base:
        nome_base = f"Painel_Infografico_Docente_{codigo}_{ano_letivo}"
    nome_arquivo = f"{nome_base}.pdf".replace("/", "-").replace("\\", "-")
    caminho_pdf = caminho_relatorio_visual_docente(nome_arquivo)

    c = canvas.Canvas(str(caminho_pdf), pagesize=A4)
    W, H = A4

    def col(h): return colors.HexColor(h)

    def limpar(texto):
        texto = limpar_marcadores_relatorio(str(texto or ""))
        texto = texto.replace("•", " ").replace("–", "-").replace("—", "-")
        texto = re.sub(r"\s+", " ", texto).strip()
        return texto.replace("Demonstra demonstração", "Demonstra") or "Não informado."

    def wrap(texto, fonte, tamanho, largura):
        palavras = limpar(texto).split()
        linhas, atual = [], ""
        for palavra in palavras:
            teste = (atual + " " + palavra).strip()
            if stringWidth(teste, fonte, tamanho) <= largura:
                atual = teste
            else:
                if atual: linhas.append(atual)
                atual = palavra
        if atual: linhas.append(atual)
        return linhas

    def draw_text(texto, x, y, largura, fonte="Helvetica", tamanho=7, entre=8.3, cor="#334155", max_linhas=5):
        linhas = wrap(texto, fonte, tamanho, largura)
        if max_linhas and len(linhas) > max_linhas:
            linhas = linhas[:max_linhas]
            linhas[-1] = linhas[-1].rstrip(".,;:") + "..."
        c.setFont(fonte, tamanho)
        c.setFillColor(col(cor))
        for linha in linhas:
            c.drawString(x, y, linha)
            y -= entre
        return y

    def lista(texto, fallback, max_itens=5, max_chars=82):
        bruto = limpar_marcadores_relatorio(str(texto or ""))
        candidatos = []
        for linha in bruto.splitlines():
            partes = re.split(r"\s+-\s+", linha.strip()) if " - " in linha else [linha]
            for p in partes:
                p = p.strip().lstrip("•-–—0123456789. ").strip()
                if len(p) > 5: candidatos.append(p)
        if not candidatos:
            candidatos = [p.strip(" .;:-") for p in re.split(r"(?<=[.;])\s+", limpar(bruto)) if len(p.strip()) > 8]
        if not candidatos: candidatos = fallback
        out = []
        for item in candidatos:
            item = limpar(item)
            if len(item) > max_chars:
                item = item[:max_chars].rsplit(" ", 1)[0] + "..."
            if item not in out: out.append(item)
            if len(out) >= max_itens: break
        return out

    def box(x, y, w, h, titulo, borda, fundo):
        c.setFillColor(col(fundo)); c.setStrokeColor(col(borda))
        c.roundRect(x, y, w, h, 8, stroke=1, fill=1)
        c.setFillColor(col(borda)); c.roundRect(x, y+h-0.55*cm, w, 0.55*cm, 8, stroke=0, fill=1)
        c.setFillColor(col("#ffffff")); c.setFont("Helvetica-Bold", 7.5)
        c.drawCentredString(x+w/2, y+h-0.35*cm, titulo.upper()[:42])

    def bullets(lista_it, x, y, largura, cor_bol="#2563eb", tamanho=6.35, max_linhas=2, limite_y=None):
        for item in lista_it:
            if limite_y and y < limite_y: break
            c.setFillColor(col(cor_bol)); c.circle(x+2.5, y+2.5, 2.0, stroke=0, fill=1)
            y = draw_text(item, x+8, y, largura-10, tamanho=tamanho, entre=7.1, max_linhas=max_linhas)
            y -= 2
        return y

    sec_pot = extrair_secao_relatorio_docente(conteudo_relatorio, 3) or potencialidades_interesses
    sec_bar = extrair_secao_relatorio_docente(conteudo_relatorio, 4) or barreiras_observadas
    sec_est = extrair_secao_relatorio_docente(conteudo_relatorio, 5) or estrategias_rapidas
    sec_adapt = extrair_secao_relatorio_docente(conteudo_relatorio, 6)
    sec_aval = extrair_secao_relatorio_docente(conteudo_relatorio, 7)
    sec_aee = extrair_secao_relatorio_docente(conteudo_relatorio, 8)

    pot = lista(sec_pot, ["Interesses e habilidades observadas.", "Resposta positiva a recursos visuais.", "Participação em atividades práticas."], 5, 72)
    bar = lista(sec_bar, ["Comandos longos podem dificultar a compreensão.", "Ambientes com muitos estímulos podem prejudicar atenção."], 5, 72)
    est = lista(sec_est, ["Comandos curtos e objetivos.", "Apoio visual e divisão da tarefa em etapas.", "Tempo ampliado e mediação gradual."], 5, 72)
    adapt = lista(sec_adapt, ["Reduzir volume sem perder objetivo.", "Permitir diferentes formas de resposta.", "Usar material concreto."], 5, 90)
    aval = lista(sec_aval, ["Avaliar participação e evolução individual.", "Evitar escrita como único critério."], 5, 90)
    aee = lista(sec_aee, ["Compartilhar observações com o AEE.", "Solicitar apoio para recursos visuais."], 5, 90)

    # Página 1
    c.setFillColor(col("#ffffff")); c.rect(0,0,W,H,stroke=0,fill=1)
    c.setFillColor(col("#0b3b75")); c.setFont("Helvetica-Bold", 13)
    c.drawString(1.1*cm, H-1.1*cm, "INCLUISRM")
    c.setFont("Helvetica", 6.5); c.drawString(1.1*cm, H-1.4*cm, "Sistema de Gestão do Atendimento Educacional Especializado")
    c.setFont("Helvetica-Bold", 15); c.drawCentredString(W/2, H-1.1*cm, "PAINEL INFOGRÁFICO DE APOIO AO DOCENTE")
    c.setFont("Helvetica-Oblique", 7); c.drawCentredString(W/2, H-1.55*cm, "Leitura rápida para planejamento inclusivo")

    # Meta cards
    meta = [("Código", codigo), ("Série", ano_serie), ("Turma", turma), ("Área", componente_destino), ("Data", agora_local().strftime("%d/%m/%Y"))]
    x=1.1*cm; y=H-2.65*cm; gap=0.15*cm; mw=(W-2.2*cm-4*gap)/5
    for lab,val in meta:
        c.setFillColor(col("#f8fafc")); c.setStrokeColor(col("#bfdbfe")); c.roundRect(x,y,mw,0.75*cm,5,stroke=1,fill=1)
        c.setFillColor(col("#334155")); c.setFont("Helvetica-Bold",6.2); c.drawString(x+0.12*cm,y+0.47*cm,lab)
        draw_text(val,x+0.12*cm,y+0.22*cm,mw-0.24*cm,tamanho=6.3,entre=7,max_linhas=1)
        x += mw+gap

    c.setFillColor(col("#eff6ff")); c.setStrokeColor(col("#93c5fd")); c.roundRect(1.1*cm,H-3.55*cm,W-2.2*cm,0.48*cm,7,stroke=1,fill=1)
    c.setFont("Helvetica-Bold",7.2); c.setFillColor(col("#1e3a8a")); c.drawCentredString(W/2,H-3.38*cm,"Antes de adaptar: observe como o estudante compreende, responde, comunica e participa.")

    left=1.1*cm; gap=0.25*cm; bw=(W-2.2*cm-gap)/2; bh=3.05*cm
    y1=H-6.95*cm
    box(left,y1,bw,bh,"Quem é o estudante", "#0b74b8", "#f0f9ff")
    draw_text(quem_estudante,left+0.22*cm,y1+bh-0.85*cm,bw-0.44*cm,tamanho=6.7,entre=7.7,max_linhas=8)
    box(left+bw+gap,y1,bw,bh,"Potencialidades", "#16a34a", "#f0fdf4")
    bullets(pot,left+bw+gap+0.22*cm,y1+bh-0.85*cm,bw-0.44*cm,"#16a34a",6.2,2,y1+0.25*cm)

    y2=y1-bh-0.28*cm
    box(left,y2,bw,bh,"Barreiras pedagógicas", "#dc2626", "#fff1f2")
    bullets(bar,left+0.22*cm,y2+bh-0.85*cm,bw-0.44*cm,"#dc2626",6.2,2,y2+0.25*cm)
    box(left+bw+gap,y2,bw,bh,"O que funciona melhor", "#059669", "#ecfdf5")
    bullets(est,left+bw+gap+0.22*cm,y2+bh-0.85*cm,bw-0.44*cm,"#059669",6.2,2,y2+0.25*cm)

    y3=y2-1.45*cm
    box(left,y3,W-2.2*cm,1.15*cm,"Mapa pedagógico rápido", "#0b3b75", "#f8fafc")
    c.setFont("Helvetica-Bold",7); c.setFillColor(col("#0b3b75"))
    c.drawString(left+0.4*cm,y3+0.43*cm,"COMPREENSÃO: comandos curtos + apoio visual")
    c.drawCentredString(W/2,y3+0.43*cm,"PARTICIPAÇÃO: atividade prática + resposta possível")
    c.drawRightString(W-1.5*cm,y3+0.43*cm,"AUTONOMIA: rotina + checklist")

    y4=1.2*cm; colw=(W-2.5*cm)/3
    box(left,y4,colw,3.0*cm,"Recursos sugeridos", "#f59e0b", "#fffbeb")
    recursos=["CAA ou comunicação alternativa","Recursos visuais estruturados","Materiais concretos","Atividades por etapas","Tecnologia quando significativa"]
    bullets(recursos,left+0.22*cm,y4+2.35*cm,colw-0.44*cm,"#f59e0b",6.0,1,y4+0.25*cm)
    box(left+colw+0.15*cm,y4,colw,3.0*cm,"Apoio pedagógico", "#2563eb", "#eff6ff")
    apoio=["Comunicação: apoio visual","Organização: apoio frequente","Atenção: previsibilidade","Avaliação: múltiplas respostas"]
    bullets(apoio,left+colw+0.37*cm,y4+2.35*cm,colw-0.44*cm,"#2563eb",6.0,1,y4+0.25*cm)
    box(left+2*(colw+0.15*cm),y4,colw,3.0*cm,"Foco principal", "#65a30d", "#f7fee7")
    draw_text("Promover participação, comunicação, autonomia e aprendizagem possível, valorizando potencialidades e reduzindo barreiras.",left+2*(colw+0.15*cm)+0.22*cm,y4+2.28*cm,colw-0.44*cm,fonte="Helvetica-Bold",tamanho=6.4,entre=7.5,max_linhas=7)
    c.showPage()

    # Página 2
    c.setFillColor(col("#ffffff")); c.rect(0,0,W,H,stroke=0,fill=1)
    c.setFont("Helvetica-Bold",15); c.setFillColor(col("#0b3b75")); c.drawString(1.2*cm,H-1.25*cm,"COMPLEMENTO PEDAGÓGICO DO INFOGRÁFICO")
    c.setFont("Helvetica",8); c.setFillColor(col("#475569")); c.drawString(1.2*cm,H-1.65*cm,"Avaliação, adaptação e acompanhamento docente em articulação com o AEE.")
    y=H-6.2*cm
    box(1.2*cm,y,W-2.4*cm,3.9*cm,"Adaptação das atividades", "#64748b", "#f8fafc")
    bullets(adapt,1.5*cm,y+3.15*cm,W-3.0*cm,"#64748b",7.0,2,y+0.35*cm)
    y-=4.3*cm
    box(1.2*cm,y,W-2.4*cm,3.9*cm,"Recomendações avaliativas", "#7c3aed", "#f5f3ff")
    bullets(aval,1.5*cm,y+3.15*cm,W-3.0*cm,"#7c3aed",7.0,2,y+0.35*cm)
    y-=4.3*cm
    box(1.2*cm,y,W-2.4*cm,3.9*cm,"Articulação com o AEE", "#0891b2", "#ecfeff")
    bullets(aee,1.5*cm,y+3.15*cm,W-3.0*cm,"#0891b2",7.0,2,y+0.35*cm)
    draw_text("Fontes utilizadas pelo sistema: " + str(fontes_geradas or "Avaliação pedagógica, estudo de caso e entrevista familiar usados apenas como contexto pedagógico."),1.2*cm,1.1*cm,W-2.4*cm,tamanho=6.4,entre=7,max_linhas=2,cor="#64748b")
    draw_text("Finalidade exclusivamente pedagógica. Não substitui avaliação docente, estudo de caso ou planejamento do AEE.",1.2*cm,0.55*cm,W-2.4*cm,tamanho=6.4,entre=7,max_linhas=1,cor="#64748b")
    c.save()
    return str(caminho_pdf)


def gerar_pdf_infografico_docente_dashboard(estudante, dados, ano_letivo, componente, nome_base=None):
    """Gera o Painel Inteligente de Apoio ao Docente com layout compacto.

    V53: organiza a primeira página como painel/dashboard completo e inicia o complemento
    ainda na primeira página, seguindo o modelo visual desejado pelo usuário.
    A segunda página recebe apenas a continuidade do complemento, evitando a sensação
    de relatório quebrado ou repetitivo.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase.pdfmetrics import stringWidth

    dados = normalizar_dados_infografico_docente(dados, "", estudante)

    codigo = estudante[1] if estudante and len(estudante) > 1 else "Não informado"
    ano_serie = estudante[2] if estudante and len(estudante) > 2 else "Não informado"
    turma = estudante[3] if estudante and len(estudante) > 3 else "Não informado"

    if not nome_base:
        nome_base = f"Painel_Inteligente_Apoio_Docente_{codigo}_{ano_letivo}"
    nome_arquivo = f"{nome_base}.pdf".replace("/", "-").replace("\\", "-")
    caminho_pdf = caminho_relatorio_visual_docente(nome_arquivo)

    c = canvas.Canvas(str(caminho_pdf), pagesize=A4)
    W, H = A4

    def cor(hex_value):
        return colors.HexColor(hex_value)

    def limpar_texto(texto):
        texto = limpar_marcadores_relatorio(str(texto or ""))
        texto = texto.replace("•", " ").replace("–", "-").replace("—", "-")
        texto = re.sub(r"\s+", " ", texto).strip()
        return texto or "Não informado nos registros analisados."

    def wrap(text, font, size, width):
        words = limpar_texto(text).split()
        lines, cur = [], ""
        for word in words:
            test = (cur + " " + word).strip()
            if stringWidth(test, font, size) <= width:
                cur = test
            else:
                if cur:
                    lines.append(cur)
                cur = word
        if cur:
            lines.append(cur)
        return lines

    def draw_text(text, x, y, width, font="Helvetica", size=5.7, leading=6.4,
                  max_lines=3, color="#0f172a"):
        c.setFont(font, size)
        c.setFillColor(cor(color))
        lines = wrap(text, font, size, width)
        if max_lines and len(lines) > max_lines:
            lines = lines[:max_lines]
            lines[-1] = lines[-1].rstrip(".,;:") + "..."
        for line in lines:
            c.drawString(x, y, line)
            y -= leading
        return y

    def painel(x, y, w, h, title, border, fill, title_size=6.9):
        c.setFillColor(cor(fill))
        c.setStrokeColor(cor(border))
        c.setLineWidth(0.55)
        c.roundRect(x, y, w, h, 5.5, stroke=1, fill=1)
        c.setFillColor(cor(border))
        c.roundRect(x, y+h-0.42*cm, w, 0.42*cm, 5.5, stroke=0, fill=1)
        c.setFillColor(cor("#ffffff"))
        c.setFont("Helvetica-Bold", title_size)
        c.drawCentredString(x+w/2, y+h-0.27*cm, title.upper()[:55])

    def bullet_list(items, x, y, width, dot="#2563eb", font_size=5.45, leading=6.15,
                    max_item_lines=2, bottom=None, max_items=None):
        if max_items:
            items = items[:max_items]
        for item in items:
            if bottom and y < bottom:
                break
            c.setFillColor(cor(dot))
            c.circle(x + 2.0, y + 2.05, 1.55, stroke=0, fill=1)
            y = draw_text(item, x + 6.5, y, width - 8, size=font_size,
                          leading=leading, max_lines=max_item_lines)
            y -= 0.7
        return y

    def meta_card(x, y, w, label, value):
        c.setFillColor(cor("#f8fafc"))
        c.setStrokeColor(cor("#bfdbfe"))
        c.setLineWidth(0.45)
        c.roundRect(x, y, w, 0.52*cm, 4, stroke=1, fill=1)
        c.setFont("Helvetica-Bold", 4.9)
        c.setFillColor(cor("#1e3a8a"))
        c.drawString(x + 0.07*cm, y + 0.32*cm, label)
        draw_text(value, x + 0.07*cm, y + 0.14*cm, w - 0.14*cm, size=4.9, leading=5.1, max_lines=1)

    def mini_grid(items, x, y, w, h, cols=3):
        items = normalizar_lista_infografico(items, limite=6, tamanho_item=58)
        gap = 0.06*cm
        cell_w = (w - (cols-1)*gap) / cols
        rows = 2
        cell_h = (h - gap) / rows
        for i, item in enumerate(items[:cols*rows]):
            row = i // cols
            col_idx = i % cols
            cx = x + col_idx*(cell_w+gap)
            cy = y + h - (row+1)*cell_h - row*gap
            c.setFillColor(cor("#ffffff"))
            c.setStrokeColor(cor("#fde68a"))
            c.setLineWidth(0.35)
            c.roundRect(cx, cy, cell_w, cell_h, 3.5, stroke=1, fill=1)
            draw_text(item, cx+0.06*cm, cy+cell_h-0.16*cm, cell_w-0.12*cm,
                      size=4.65, leading=5.10, max_lines=4)

    def bloco_complemento(x, y, w, h, titulo, lista, border, fill, font_size=6.0):
        painel(x, y, w, h, titulo, border, fill, title_size=6.3)
        bullet_list(lista, x+0.24*cm, y+h-0.62*cm, w-0.48*cm, dot=border,
                    font_size=font_size, leading=7.0, max_item_lines=2, bottom=y+0.18*cm)

    # ======================================================
    # PÁGINA 1 - DASHBOARD + INÍCIO DO COMPLEMENTO
    # ======================================================
    c.setFillColor(cor("#ffffff"))
    c.rect(0, 0, W, H, stroke=0, fill=1)

    # Cabeçalho compacto
    c.setFillColor(cor("#0b3b75"))
    c.setFont("Helvetica-Bold", 10.8)
    c.drawString(1.05*cm, H - 0.86*cm, "INCLUISRM")
    c.setFont("Helvetica", 4.8)
    c.drawString(1.05*cm, H - 1.10*cm, "Sistema de Gestão do Atendimento Educacional Especializado")
    c.setFont("Helvetica-Bold", 13.2)
    c.drawCentredString(W/2, H - 0.84*cm, "PAINEL INTELIGENTE DE APOIO AO DOCENTE")
    c.setFont("Helvetica", 5.8)
    c.drawCentredString(W/2, H - 1.14*cm, "Guia rápido para aprendizagem, participação e acessibilidade do estudante")

    c.setFillColor(cor("#0b3b75"))
    c.roundRect(W-4.55*cm, H-1.26*cm, 3.55*cm, 0.72*cm, 5, stroke=0, fill=1)
    c.setFont("Helvetica-Bold", 5.15)
    c.setFillColor(cor("#ffffff"))
    c.drawString(W-4.35*cm, H-0.82*cm, "Data: " + agora_local().strftime("%d/%m/%Y"))
    c.drawString(W-4.35*cm, H-1.07*cm, "Ano letivo: " + str(ano_letivo))

    # Identificação compacta
    x0 = 1.05*cm
    y_meta = H - 2.05*cm
    gap = 0.10*cm
    meta_w = (W - 2.10*cm - 4*gap) / 5
    for idx, (label, value) in enumerate([
        ("Código", codigo), ("Ano/Série", ano_serie), ("Turma", turma),
        ("Componente/área", componente), ("Destinatário", "________________")
    ]):
        meta_card(x0 + idx*(meta_w+gap), y_meta, meta_w, label, value)

    # Frase-guia
    c.setFillColor(cor("#eff6ff"))
    c.setStrokeColor(cor("#93c5fd"))
    c.setLineWidth(0.5)
    c.roundRect(1.05*cm, H-2.62*cm, W-2.10*cm, 0.38*cm, 5, stroke=1, fill=1)
    c.setFont("Helvetica-Bold", 5.95)
    c.setFillColor(cor("#1e3a8a"))
    c.drawCentredString(W/2, H-2.49*cm, "Antes de adaptar, observe como o estudante compreende, responde, comunica e participa.")

    # Cards principais em layout compacto
    left = 1.05*cm
    gutter = 0.24*cm
    col_w = (W - 2.10*cm - gutter) / 2
    top_y = H - 5.90*cm
    card_h = 2.92*cm

    painel(left, top_y, col_w, card_h, "1. Quem é este estudante?", "#0b74b8", "#f0f9ff")
    bullet_list(dados["quem_e_estudante"], left+0.18*cm, top_y+card_h-0.68*cm,
                col_w-0.36*cm, dot="#0b74b8", font_size=5.55, leading=6.45,
                max_item_lines=2, bottom=top_y+0.16*cm)

    painel(left+col_w+gutter, top_y, col_w, card_h, "2. Mapa de aprendizagem", "#55a630", "#f0fdf4")
    mid_x = left+col_w+gutter+col_w/2
    c.setFont("Helvetica-Bold", 5.15)
    c.setFillColor(cor("#15803d"))
    c.drawCentredString(left+col_w+gutter+col_w*0.25, top_y+card_h-0.62*cm, "COMO APRENDE MELHOR")
    c.setFillColor(cor("#dc2626"))
    c.drawCentredString(left+col_w+gutter+col_w*0.75, top_y+card_h-0.62*cm, "O QUE DIFICULTA")
    bullet_list(dados["como_aprende_melhor"][:4], left+col_w+gutter+0.16*cm,
                top_y+card_h-0.90*cm, col_w/2-0.28*cm, dot="#55a630",
                font_size=5.0, leading=5.80, max_item_lines=2, bottom=top_y+0.15*cm)
    bullet_list(dados["o_que_dificulta"][:4], mid_x+0.12*cm,
                top_y+card_h-0.90*cm, col_w/2-0.28*cm, dot="#dc2626",
                font_size=5.0, leading=5.80, max_item_lines=2, bottom=top_y+0.15*cm)
    c.setStrokeColor(cor("#d1d5db")); c.setLineWidth(0.4)
    c.line(mid_x, top_y+0.18*cm, mid_x, top_y+card_h-0.55*cm)

    # Segunda linha: potencialidades e barreiras
    y2 = top_y - 3.10*cm
    painel(left, y2, col_w, 2.62*cm, "3. Potencialidades a valorizar", "#f59e0b", "#fffbeb")
    mini_grid(dados["potencialidades"], left+0.14*cm, y2+0.18*cm, col_w-0.28*cm, 1.78*cm, cols=3)

    painel(left+col_w+gutter, y2, col_w, 2.62*cm, "4. Barreiras pedagógicas", "#7c3aed", "#faf5ff")
    bullet_list(dados["barreiras_pedagogicas"], left+col_w+gutter+0.18*cm, y2+1.95*cm,
                col_w-0.36*cm, dot="#7c3aed", font_size=5.45, leading=6.2,
                max_item_lines=2, bottom=y2+0.16*cm)

    # Terceira linha: 3 colunas
    y3 = y2 - 2.82*cm
    small_gap = 0.18*cm
    small_w = (W - 2.10*cm - 2*small_gap) / 3
    small_h = 2.45*cm

    painel(left, y3, small_w, small_h, "5. O que funciona melhor", "#0f9f8f", "#ecfeff")
    bullet_list(dados["o_que_funciona_em_sala"][:4], left+0.16*cm, y3+1.83*cm,
                small_w-0.32*cm, dot="#0f9f8f", font_size=5.1, leading=5.9,
                max_item_lines=2, bottom=y3+0.14*cm)

    painel(left+small_w+small_gap, y3, small_w, small_h, "6. Atenção docente", "#f97316", "#fff7ed")
    bullet_list(dados["atencao_docente"][:4], left+small_w+small_gap+0.16*cm,
                y3+1.83*cm, small_w-0.32*cm, dot="#f97316", font_size=5.1,
                leading=5.9, max_item_lines=2, bottom=y3+0.14*cm)

    painel(left+2*(small_w+small_gap), y3, small_w, small_h, "7. Avaliação flexível", "#2563eb", "#eff6ff")
    bullet_list(dados["avaliacao_flexivel"][:4], left+2*(small_w+small_gap)+0.16*cm,
                y3+1.83*cm, small_w-0.32*cm, dot="#2563eb", font_size=5.1,
                leading=5.9, max_item_lines=2, bottom=y3+0.14*cm)

    # Quarta linha: articulação e foco
    y4 = y3 - 2.62*cm
    bottom_h = 1.95*cm
    painel(left, y4, col_w, bottom_h, "8. Articulação com o AEE", "#7c3aed", "#faf5ff")
    bullet_list(dados["articulacao_aee"][:4], left+0.18*cm, y4+bottom_h-0.64*cm,
                col_w-0.36*cm, dot="#7c3aed", font_size=5.1, leading=5.9,
                max_item_lines=1, bottom=y4+0.14*cm)

    painel(left+col_w+gutter, y4, col_w, bottom_h, "9. Foco pedagógico principal", "#f59e0b", "#fffbeb")
    draw_text(dados["foco_pedagogico"], left+col_w+gutter+0.25*cm,
              y4+bottom_h-0.72*cm, col_w-0.50*cm, font="Helvetica-Bold",
              size=6.1, leading=7.2, max_lines=5, color="#111827")

    # Complemento começa ainda na página 1
    y_comp_titulo = y4 - 0.80*cm
    c.setFillColor(cor("#0b3b75"))
    c.setFont("Helvetica-Bold", 13)
    c.drawString(1.05*cm, y_comp_titulo, "COMPLEMENTO DO PAINEL INFOGRÁFICO")
    c.setFont("Helvetica", 6.2)
    c.setFillColor(cor("#475569"))
    c.drawString(1.05*cm, y_comp_titulo-0.28*cm, "Dados estruturados pela IA para orientar consulta futura e planejamento docente.")

    bloco_complemento(1.05*cm, 3.28*cm, W-2.10*cm, 2.40*cm,
                      "Recursos sugeridos", dados["recursos_sugeridos"], "#f59e0b", "#fffbeb", font_size=5.9)
    bloco_complemento(1.05*cm, 0.78*cm, W-2.10*cm, 2.25*cm,
                      "Como aprende melhor", dados["como_aprende_melhor"], "#55a630", "#f0fdf4", font_size=5.9)

    # ======================================================
    # PÁGINA 2 - CONTINUIDADE DO COMPLEMENTO
    # ======================================================
    c.showPage()
    c.setFillColor(cor("#ffffff"))
    c.rect(0, 0, W, H, stroke=0, fill=1)

    bloco_complemento(1.25*cm, H-6.15*cm, W-2.50*cm, 4.45*cm,
                      "O que dificulta", dados["o_que_dificulta"], "#dc2626", "#fff1f2", font_size=6.2)

    bloco_complemento(1.25*cm, H-12.65*cm, W-2.50*cm, 4.65*cm,
                      "Avaliação e articulação",
                      dados["avaliacao_flexivel"][:3] + dados["articulacao_aee"][:3],
                      "#2563eb", "#eff6ff", font_size=6.2)

    c.setFont("Helvetica", 6.4)
    c.setFillColor(cor("#64748b"))
    c.drawString(1.25*cm, 0.75*cm, "Gerado pelo INCLUISRM com IA a partir dos registros pedagógicos. Finalidade exclusivamente pedagógica.")
    c.drawString(1.25*cm, 0.48*cm, "As sugestões não representam diagnóstico, laudo ou definição fixa sobre o estudante.")

    c.save()
    return str(caminho_pdf)


# ======================================================
# CRUD - ATENDIMENTOS
# ======================================================
def salvar_atendimento(
    estudante_id, data_atendimento, objetivo, atividade, resposta_estudante,
    avancos, dificuldades, evolucao, qtd_atividades, nivel_resposta,
    nivel_avanco, nivel_dificuldade, nivel_engajamento, nivel_evolucao,
    encaminhamentos,
):
    """Salva atendimento e retorna o ID criado para vincular recursos pedagógicos."""
    campos = [
        "estudante_id", "data_atendimento", "objetivo", "atividade", "resposta_estudante",
        "avancos", "dificuldades", "evolucao", "qtd_atividades", "nivel_resposta",
        "nivel_avanco", "nivel_dificuldade", "nivel_engajamento", "nivel_evolucao",
        "encaminhamentos",
    ]
    valores = [
        estudante_id, data_atendimento, objetivo, atividade, resposta_estudante,
        avancos, dificuldades, evolucao, qtd_atividades, nivel_resposta,
        nivel_avanco, nivel_dificuldade, nivel_engajamento, nivel_evolucao,
        encaminhamentos,
    ]

    conn = conectar()
    cursor = conn.cursor()
    placeholders = ", ".join(["?"] * len(valores))
    colunas = ", ".join(campos)

    if USAR_POSTGRES:
        cursor.execute(
            f"INSERT INTO atendimentos ({colunas}) VALUES ({placeholders}) RETURNING id",
            valores,
        )
        atendimento_id = cursor.fetchone()[0]
    else:
        cursor.execute(
            f"INSERT INTO atendimentos ({colunas}) VALUES ({placeholders})",
            valores,
        )
        atendimento_id = cursor.lastrowid

    conn.commit()
    conn.close()
    limpar_cache_dados()
    return atendimento_id


@st.cache_data(ttl=30, show_spinner=False)
def listar_atendimentos(estudante_id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT data_atendimento, objetivo, atividade, resposta_estudante,
               avancos, dificuldades, evolucao, qtd_atividades, nivel_resposta,
               nivel_avanco, nivel_dificuldade, nivel_engajamento,
               nivel_evolucao, encaminhamentos
        FROM atendimentos
        WHERE estudante_id = ?
        ORDER BY id DESC
        """,
        (estudante_id,),
    )
    dados = cursor.fetchall()
    conn.close()
    return dados


@st.cache_data(ttl=30, show_spinner=False)
def listar_atendimentos_com_id(estudante_id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT id, data_atendimento, objetivo, atividade, resposta_estudante,
               avancos, dificuldades, evolucao, qtd_atividades, nivel_resposta,
               nivel_avanco, nivel_dificuldade, nivel_engajamento,
               nivel_evolucao, encaminhamentos
        FROM atendimentos
        WHERE estudante_id = ?
        ORDER BY id DESC
        """,
        (estudante_id,),
    )
    dados = cursor.fetchall()
    conn.close()
    return dados


def atualizar_atendimento(atendimento_id, data_atendimento, objetivo, atividade, resposta_estudante, avancos, dificuldades, evolucao, qtd_atividades, nivel_resposta, nivel_avanco, nivel_dificuldade, nivel_engajamento, nivel_evolucao, encaminhamentos):
    atualizar_registro(
        "atendimentos",
        [
            "data_atendimento", "objetivo", "atividade", "resposta_estudante", "avancos",
            "dificuldades", "evolucao", "qtd_atividades", "nivel_resposta", "nivel_avanco",
            "nivel_dificuldade", "nivel_engajamento", "nivel_evolucao", "encaminhamentos",
        ],
        [
            data_atendimento, objetivo, atividade, resposta_estudante, avancos, dificuldades,
            evolucao, qtd_atividades, nivel_resposta, nivel_avanco, nivel_dificuldade,
            nivel_engajamento, nivel_evolucao, encaminhamentos,
        ],
        atendimento_id,
    )


def excluir_atendimento(atendimento_id):
    excluir_recursos_por_atendimento(atendimento_id)
    excluir_registro("atendimentos", atendimento_id)


# ======================================================
# CRUD - RECURSOS PEDAGÓGICOS DO ATENDIMENTO
# ======================================================
def salvar_recurso_atendimento(atendimento_id, estudante_id, nome_recurso, categoria_recurso, link_recurso, observacao_uso):
    inserir_registro(
        "recursos_atendimento",
        [
            "atendimento_id", "estudante_id", "nome_recurso", "categoria_recurso",
            "link_recurso", "observacao_uso", "criado_em",
        ],
        [
            atendimento_id, estudante_id, nome_recurso, categoria_recurso,
            link_recurso, observacao_uso, hoje_str(),
        ],
    )


@st.cache_data(ttl=30, show_spinner=False)
def listar_recursos_atendimento(atendimento_id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT id, nome_recurso, categoria_recurso, link_recurso, observacao_uso, criado_em
        FROM recursos_atendimento
        WHERE atendimento_id = ?
        ORDER BY id ASC
        """,
        (atendimento_id,),
    )
    dados = cursor.fetchall()
    conn.close()
    return dados


def excluir_recursos_por_atendimento(atendimento_id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM recursos_atendimento WHERE atendimento_id=?", (atendimento_id,))
    conn.commit()
    conn.close()
    limpar_cache_dados()


def texto_recursos_atendimento(atendimento_id):
    recursos = listar_recursos_atendimento(atendimento_id)
    if not recursos:
        return "Nenhum recurso pedagógico com link registrado para este atendimento."

    linhas = []
    for i, r in enumerate(recursos, start=1):
        _, nome, categoria, link, observacao, criado_em = r
        linhas.append(
            f"{i}. {nome or 'Recurso não informado'}\n"
            f"   Categoria: {categoria or 'Não informada'}\n"
            f"   Link: {link or 'Não informado'}\n"
            f"   Observação de uso: {observacao or 'Não informada'}"
        )
    return "\n\n".join(linhas)


# ======================================================
# CRUD - AGENDA
# ======================================================
def salvar_agendamento(estudante_id, data_agendamento, dia_semana, hora_inicio, hora_fim, tipo_atendimento, observacoes):
    inserir_registro(
        "agenda",
        ["estudante_id", "data_agendamento", "dia_semana", "hora_inicio", "hora_fim", "tipo_atendimento", "observacoes", "criado_em"],
        [estudante_id, data_agendamento, dia_semana, hora_inicio, hora_fim, tipo_atendimento, observacoes, hoje_str()],
    )


@st.cache_data(ttl=30, show_spinner=False)
def listar_agenda():
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT agenda.id, estudantes.codigo, estudantes.ano_serie, estudantes.perfil,
               agenda.data_agendamento, agenda.dia_semana, agenda.hora_inicio,
               agenda.hora_fim, agenda.tipo_atendimento, agenda.observacoes
        FROM agenda
        JOIN estudantes ON estudantes.id = agenda.estudante_id
        ORDER BY agenda.data_agendamento, agenda.hora_inicio
        """
    )
    dados = cursor.fetchall()
    conn.close()
    return dados


def listar_agenda_com_id():
    return listar_agenda()


@st.cache_data(ttl=30, show_spinner=False)
def contar_registros_tabela(tabela):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(f"SELECT COUNT(*) FROM {tabela}")
    total = cursor.fetchone()[0]
    conn.close()
    return total


def excluir_agendamento(agenda_id):
    excluir_registro("agenda", agenda_id)


def data_texto_para_date_seguro(data_texto):
    """Converte data em formato dd/mm/aaaa para date. Retorna None se inválida."""
    try:
        return datetime.strptime(str(data_texto), "%d/%m/%Y").date()
    except Exception:
        return None


@st.cache_data(ttl=30, show_spinner=False)
def listar_agenda_detalhada(estudante_id=None, data_inicio=None, data_fim=None):
    """Lista agenda com status de presença calculado automaticamente.
    Status calculado:
    - Compareceu: se houver atendimento vinculado ou atendimento do estudante na mesma data.
    - Faltou: se a data já passou e não houver atendimento.
    - Agendado: se ainda não ocorreu.
    """
    conn = conectar()
    cursor = conn.cursor()
    if estudante_id:
        cursor.execute(
            """
            SELECT a.id, a.estudante_id, e.codigo, e.ano_serie, e.perfil,
                   a.data_agendamento, a.dia_semana, a.hora_inicio, a.hora_fim,
                   a.tipo_atendimento, a.observacoes,
                   COALESCE(a.status_presenca, 'Agendado') AS status_presenca,
                   a.atendimento_id,
                   COALESCE(a.preenchimento_ia, '') AS preenchimento_ia
            FROM agenda a
            JOIN estudantes e ON e.id = a.estudante_id
            WHERE a.estudante_id = ?
            ORDER BY a.data_agendamento, a.hora_inicio
            """,
            (estudante_id,),
        )
    else:
        cursor.execute(
            """
            SELECT a.id, a.estudante_id, e.codigo, e.ano_serie, e.perfil,
                   a.data_agendamento, a.dia_semana, a.hora_inicio, a.hora_fim,
                   a.tipo_atendimento, a.observacoes,
                   COALESCE(a.status_presenca, 'Agendado') AS status_presenca,
                   a.atendimento_id,
                   COALESCE(a.preenchimento_ia, '') AS preenchimento_ia
            FROM agenda a
            JOIN estudantes e ON e.id = a.estudante_id
            ORDER BY a.data_agendamento, a.hora_inicio
            """
        )
    dados = cursor.fetchall()

    # Busca atendimentos para status automático.
    cursor.execute("SELECT id, estudante_id, data_atendimento FROM atendimentos")
    atendimentos = cursor.fetchall()
    conn.close()

    atend_por_estudante_data = {}
    for atendimento_id, est_id, data_at in atendimentos:
        atend_por_estudante_data[(est_id, data_at)] = atendimento_id

    hoje = agora_local().date()
    filtrados = []
    for row in dados:
        row = list(row)
        data_ag = data_texto_para_date_seguro(row[5])
        if data_inicio and data_ag and data_ag < data_inicio:
            continue
        if data_fim and data_ag and data_ag > data_fim:
            continue

        ag_id, est_id = row[0], row[1]
        status_original = row[11] or "Agendado"
        atendimento_vinculado = row[12]
        atendimento_mesma_data = atend_por_estudante_data.get((est_id, row[5]))

        if atendimento_vinculado or atendimento_mesma_data:
            status_calc = "Compareceu"
            if not atendimento_vinculado and atendimento_mesma_data:
                row[12] = atendimento_mesma_data
        elif status_original in ["Compareceu", "Faltou", "Justificado", "Cancelado"]:
            status_calc = status_original
        elif data_ag and data_ag < hoje:
            status_calc = "Faltou"
        else:
            status_calc = "Agendado"
        row[11] = status_calc
        filtrados.append(tuple(row))
    return filtrados


def atualizar_status_agenda(agenda_id, status_presenca, atendimento_id=None, preenchimento_ia=None):
    conn = conectar()
    cursor = conn.cursor()
    if atendimento_id is not None and preenchimento_ia is not None:
        cursor.execute(
            "UPDATE agenda SET status_presenca=?, atendimento_id=?, preenchimento_ia=? WHERE id=?",
            (status_presenca, atendimento_id, preenchimento_ia, agenda_id),
        )
    elif atendimento_id is not None:
        cursor.execute(
            "UPDATE agenda SET status_presenca=?, atendimento_id=? WHERE id=?",
            (status_presenca, atendimento_id, agenda_id),
        )
    elif preenchimento_ia is not None:
        cursor.execute(
            "UPDATE agenda SET status_presenca=?, preenchimento_ia=? WHERE id=?",
            (status_presenca, preenchimento_ia, agenda_id),
        )
    else:
        cursor.execute(
            "UPDATE agenda SET status_presenca=? WHERE id=?",
            (status_presenca, agenda_id),
        )
    conn.commit()
    conn.close()
    limpar_cache_dados()


def montar_dataframe_quadro_semanal(estudante_id=None, data_inicio=None, data_fim=None):
    agenda = listar_agenda_detalhada(estudante_id=estudante_id, data_inicio=data_inicio, data_fim=data_fim)
    colunas = [
        "Agenda ID", "Estudante ID", "Código", "Ano/Série", "Perfil", "Data", "Dia", "Início", "Fim",
        "Tipo", "Observações da agenda", "Status", "Atendimento ID", "Sugestão IA"
    ]
    df_ag = pd.DataFrame(agenda, columns=colunas)
    if df_ag.empty:
        return df_ag

    # Traz dados do atendimento quando houver atendimento no mesmo dia.
    linhas = []
    for _, row in df_ag.iterrows():
        estudante = buscar_estudante(int(row["Estudante ID"]))
        atendimentos = listar_atendimentos(int(row["Estudante ID"]))
        at_data = None
        for a in atendimentos:
            if a[0] == row["Data"]:
                at_data = a
                break

        atividade = at_data[2] if at_data else row["Observações da agenda"]
        resposta = at_data[3] if at_data else ""
        dificuldades = at_data[5] if at_data else ""
        evolucao = at_data[6] if at_data else ""
        encaminhamentos = at_data[13] if at_data and len(at_data) > 13 else ""
        observacoes = " | ".join([x for x in [resposta, dificuldades, evolucao, encaminhamentos, row["Sugestão IA"]] if str(x).strip()])

        tipo = str(row["Tipo"] or "").lower()
        linhas.append({
            "Data": row["Data"],
            "Código": row["Código"],
            "Estudante": estudante[1] if estudante else row["Código"],
            "Individual": "X" if "individual" in tipo else "",
            "Coletivo": "X" if ("grupo" in tipo or "coletivo" in tipo) else "",
            "Recursos/Atividades desenvolvidas": atividade or "",
            "Observação em sala": "X" if "observ" in tipo else "",
            "Formação em serviço": "",
            "Diversos": "X" if str(row["Status"]) in ["Faltou", "Justificado", "Cancelado"] else "",
            "Presença": row["Status"],
            "Observações": observacoes or row["Observações da agenda"] or "",
            "Assinatura/Rúbrica": "",
        })
    return pd.DataFrame(linhas)


def gerar_texto_quadro_semanal_gre(professor, df_quadro, mes, ano, escola_manual="", regional_manual=""):
    escola = escola_manual or (professor[2] if professor else "") or "______________________________________________"
    regional = regional_manual or (professor[3] if professor else "") or "__________"
    professor_nome = (professor[1] if professor else "") or "______________________________________________"
    mes = mes or "________________"
    ano = ano or "20____"

    linhas = []
    if df_quadro is not None and not df_quadro.empty:
        for _, r in df_quadro.iterrows():
            linhas.append(
                f"Data: {r['Data']} | Individual: {r['Individual']} | Coletivo: {r['Coletivo']} | "
                f"Recursos/Atividades: {r['Recursos/Atividades desenvolvidas']} | "
                f"Observação em sala: {r['Observação em sala']} | Formação: {r['Formação em serviço']} | "
                f"Diversos: {r['Diversos']} | Presença: {r['Presença']} | Observações: {r['Observações']} | "
                f"Assinatura/Rúbrica: ____________________"
            )
    else:
        linhas.append("Nenhum atendimento/agendamento encontrado no período selecionado.")

    return f"""
{texto_cabecalho_gre("QUADRO DE ACOMPANHAMENTO SEMANAL - AÇÕES/PRÁTICAS DO PROFESSOR DO AEE")}

Escola: {escola}
Regional: {regional}
Professor(a): {professor_nome}
Mês: {mes}    Ano: {ano}

QUADRO DE ACOMPANHAMENTO SEMANAL - AÇÕES/PRÁTICAS DO PROFESSOR DO AEE

""".strip() + "\n\n" + "\n".join(linhas) + "\n\n*Este documento poderá ser editado ao longo do tempo para atender às necessidades e atualizações da atividade."


def gerar_docx_quadro_semanal_gre(professor, df_quadro, nome_base, mes, ano, escola_manual="", regional_manual=""):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT

    nome_arquivo = f"Quadro_Semanal_GRE_{nome_base}.docx".replace("/", "-").replace("\\", "-")
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)

    cab = doc.add_paragraph()
    cab.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = cab.add_run(
        "SECRETARIA DE EDUCAÇÃO E ESPORTES\n"
        "GERÊNCIA METROPOLITANA NORTE DE EDUCAÇÃO\n"
        "CGDE - COORDENAÇÃO GERAL DE DESENVOLVIMENTO DA EDUCAÇÃO\n"
        "NÚCLEO DE EDUCAÇÃO INCLUSIVA, DIREITOS HUMANOS E CIDADANIA"
    )
    run.bold = True
    run.font.size = Pt(8)

    escola = escola_manual or (professor[2] if professor else "") or "______________________________________________"
    regional = regional_manual or (professor[3] if professor else "") or "__________"
    professor_nome = (professor[1] if professor else "") or "______________________________________________"
    mes = mes or "________________"
    ano = ano or "20____"

    doc.add_paragraph(f"ESCOLA: {escola}    REGIONAL: {regional}")
    doc.add_paragraph(f"PROFESSOR(A): {professor_nome}    MÊS: {mes}    ANO: {ano}")

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = titulo.add_run("Quadro de Acompanhamento Semanal - Ações/Práticas do Professor do AEE")
    rt.bold = True
    rt.underline = True
    rt.font.size = Pt(10)

    headers = [
        "DATA", "INDIVIDUAL", "COLETIVO", "RECURSOS UTILIZADOS / ATIVIDADES DESENVOLVIDAS",
        "OBSERVAÇÃO EM SALA", "FORMAÇÃO EM SERVIÇO", "DIVERSOS", "PRESENÇA", "OBSERVAÇÕES", "ASSINATURA/RÚBRICA"
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(7)

    if df_quadro is not None and not df_quadro.empty:
        for _, r in df_quadro.iterrows():
            row = table.add_row().cells
            valores = [
                r.get("Data", ""), r.get("Individual", ""), r.get("Coletivo", ""),
                r.get("Recursos/Atividades desenvolvidas", ""), r.get("Observação em sala", ""),
                r.get("Formação em serviço", ""), r.get("Diversos", ""), r.get("Presença", ""),
                r.get("Observações", ""), ""
            ]
            for i, val in enumerate(valores):
                row[i].text = str(val or "")
                for p in row[i].paragraphs:
                    for rr in p.runs:
                        rr.font.size = Pt(7)
    else:
        for _ in range(5):
            table.add_row()

    rod = doc.add_paragraph("*Este documento poderá ser editado ao longo do tempo para atender às necessidades e atualizações da atividade.")
    for r in rod.runs:
        r.font.size = Pt(7)

    doc.save(nome_arquivo)
    return nome_arquivo


def gerar_sugestao_ia_quadro_semanal(df_quadro, foco=""):
    client = obter_cliente_openai()
    if client is None:
        return "IA não configurada. Configure OPENAI_API_KEY para usar esta função."

    dados = df_quadro.to_string(index=False) if df_quadro is not None and not df_quadro.empty else "Nenhum dado encontrado."
    prompt = f"""
Você é um especialista em Atendimento Educacional Especializado (AEE).
Analise os registros de agenda/atendimentos abaixo e sugira um preenchimento objetivo para o Quadro de Acompanhamento Semanal do Professor do AEE.

REGRAS:
- Não invente dados pessoais.
- Use linguagem institucional e pedagógica.
- Considere presença automática: Compareceu, Faltou, Justificado, Cancelado ou Agendado.
- Para faltas, sugira observação breve e encaminhamento adequado.
- Para comparecimento, sintetize recursos utilizados, atividades e observações pedagógicas.
- Responda em formato de tabela textual com as colunas: Data | Recursos/Atividades | Outros atendimentos | Observações.

FOCO/ORIENTAÇÃO DO PROFESSOR:
{foco}

DADOS:
{dados}
"""
    try:
        resposta = client.responses.create(model="gpt-4.1-mini", input=prompt)
        return resposta.output_text
    except Exception as e:
        return f"Não foi possível gerar sugestão com IA agora. Erro: {e}"


# ======================================================
# HISTÓRICO DE GERAÇÃO DOS DOCUMENTOS GRE
# ======================================================
def registrar_documento_gre(estudante_id, tipo_documento, nome_arquivo, observacao="Gerado sob demanda; arquivo não armazenado no banco."):
    """Registra apenas a data/tipo do documento GRE gerado.
    O conteúdo do PDF/Word não é armazenado para evitar dados sensíveis e excesso de arquivos.
    """
    inserir_registro(
        "documentos_gre_gerados",
        ["estudante_id", "tipo_documento", "nome_arquivo", "data_geracao", "observacao"],
        [estudante_id, tipo_documento, nome_arquivo, hoje_str(), observacao],
    )


@st.cache_data(ttl=30, show_spinner=False)
def listar_documentos_gre_gerados(estudante_id=None):
    conn = conectar()
    cursor = conn.cursor()
    if estudante_id:
        cursor.execute(
            """
            SELECT d.id, e.codigo, d.tipo_documento, d.nome_arquivo, d.data_geracao, d.observacao
            FROM documentos_gre_gerados d
            LEFT JOIN estudantes e ON e.id = d.estudante_id
            WHERE d.estudante_id = ?
            ORDER BY d.id DESC
            """,
            (estudante_id,),
        )
    else:
        cursor.execute(
            """
            SELECT d.id, e.codigo, d.tipo_documento, d.nome_arquivo, d.data_geracao, d.observacao
            FROM documentos_gre_gerados d
            LEFT JOIN estudantes e ON e.id = d.estudante_id
            ORDER BY d.id DESC
            LIMIT 50
            """
        )
    dados = cursor.fetchall()
    conn.close()
    return dados


def excluir_historico_documento_gre(registro_id):
    excluir_registro("documentos_gre_gerados", registro_id)


# ======================================================
# TEXTOS PARA EXPORTAÇÃO
# ======================================================
def texto_cadastro_estudante(estudante):
    return f"""
CADASTRO DO ESTUDANTE - INCLUISRM

Código interno: {estudante[1] or 'Não informado.'}
Ano/Série: {estudante[2] or 'Não informado.'}
Turma: {estudante[3] or 'Não informado.'}
Turno: {estudante[6] or 'Não informado.'}
Perfil educacional: {estudante[4] or 'Não informado.'}
Dias de atendimento: {estudante[7] or 'Não informado.'}
Horário preferencial: {estudante[8] or 'Não informado.'}

Campos sensíveis para preenchimento manual:
Nome completo: ___________________________________________
CPF/RG: _________________________________________________
Data de nascimento: ______________________________________
Responsável: ____________________________________________
Telefone: _______________________________________________
Endereço: _______________________________________________

Observações pedagógicas:
{estudante[5] or 'Não informado.'}
""".strip()



def texto_matricula_srm(estudante):
    codigo = estudante[1] or "Não informado."
    ano_serie = estudante[2] or "Não informado."
    turma = estudante[3] or "Não informado."
    perfil = estudante[4] or "Não informado."
    observacoes = estudante[5] or "Não informado."
    turno = estudante[6] or "Não informado."
    dias = estudante[7] or "Não informado."
    horario = estudante[8] or "Não informado."

    return f"""
MATRÍCULA PARA O ATENDIMENTO EDUCACIONAL ESPECIALIZADO
SALA DE RECURSOS MULTIFUNCIONAIS - SRM

1. IDENTIFICAÇÃO DO ESTUDANTE

Código interno: {codigo}
Ano/Série: {ano_serie}
Turma: {turma}
Turno: {turno}
Perfil educacional informado: {perfil}
Dias preferenciais de atendimento: {dias}
Horário preferencial: {horario}

Professor(a) AEE responsável:
{texto_professores_vinculados(estudante[0])}

2. DADOS SENSÍVEIS PARA PREENCHIMENTO MANUAL

Nome completo do estudante: ___________________________________________

Data de nascimento: ____/____/________

CPF/RG do estudante: _________________________________________________

Nome do responsável: _________________________________________________

CPF/RG do responsável: _______________________________________________

Telefone: ___________________________________________________________

Endereço: ___________________________________________________________

3. INFORMAÇÕES PEDAGÓGICAS INICIAIS

Observações registradas no sistema:
{observacoes}

4. DECLARAÇÃO

Declaro ciência da matrícula do estudante acima identificado no Atendimento Educacional Especializado (AEE), na Sala de Recursos Multifuncionais - SRM, conforme organização pedagógica da unidade escolar e orientações da rede de ensino.

Data: ____/____/________


__________________________________________
Assinatura do responsável


__________________________________________
Professor(a) do Atendimento Educacional Especializado - AEE


__________________________________________
Gestão / Coordenação Pedagógica
""".strip()

def texto_professor(p):
    return f"""
FICHA DE IDENTIFICAÇÃO DO(A) PROFESSOR(A) AEE - INCLUISRM

Referência/Nome profissional: {p[1] or 'Não informado.'}
Escola: {p[2] or 'Não informado.'}
Regional/GRE: {p[3] or 'Não informado.'}
Formação: {p[4] or 'Não informado.'}
Carga horária: {p[5] or 'Não informado.'}
Turno de atuação: {p[6] or 'Não informado.'}
Data de cadastro: {p[8] or 'Não informado.'}

PERFIL PEDAGÓGICO E TECNOLÓGICO DO(A) PROFESSOR(A):
Áreas de atuação/interesse: {p[9] if len(p) > 9 and p[9] else 'Não informado.'}
Nível de familiaridade tecnológica: {p[10] if len(p) > 10 and p[10] else 'Não informado.'}
Modo maker inclusivo ativado: {p[11] if len(p) > 11 and p[11] else 'Não informado.'}
Interesse em formação maker/tecnologias educacionais: {p[12] if len(p) > 12 and p[12] else 'Não informado.'}
Projetos/temas de interesse: {p[13] if len(p) > 13 and p[13] else 'Não informado.'}
Preferências metodológicas: {p[14] if len(p) > 14 and p[14] else 'Não informado.'}
Recursos pedagógicos/TA que o professor possui ou pode levar: {p[15] if len(p) > 15 and p[15] else 'Não informado.'}
Recursos que pretende utilizar nos atendimentos: {p[16] if len(p) > 16 and p[16] else 'Não informado.'}
Observações sobre uso dos recursos do professor: {p[17] if len(p) > 17 and p[17] else 'Não informado.'}

Observações:
{p[7] or 'Não informado.'}

Assinatura do(a) Professor(a) AEE:
________________________________________________________
""".strip()


def texto_avaliacao(estudante, a):
    dados = dict(zip(["id"] + CAMPOS_AVALIACAO, a))

    def v(campo):
        valor = dados.get(campo)
        return valor if valor not in (None, "") else "Não informado."

    tipo = dados.get("tipo_registro") or ""

    if tipo == "Avaliação Pedagógica - Extra / Documento Livre":
        return f"""
AVALIAÇÃO PEDAGÓGICA EXTRA - DOCUMENTO LIVRE - INCLUISRM

Código interno do estudante: {estudante[1]}
Data do registro: {v('data_registro')}
Ano letivo: {v('ano_letivo')}
Tipo de registro: {v('tipo_registro')}
Avaliação anterior vinculada: {v('avaliacao_anterior_id')}
Origem do documento: {v('origem_documento')}

CONTEÚDO DO DOCUMENTO / RELATÓRIO ANTERIOR:

{v('texto_documento_extra')}

Observação:
Este registro foi lançado em formato livre para preservar documentos pedagógicos anteriores, pareceres, relatórios ou registros que não se enquadram no modelo estruturado da avaliação pedagógica.
""".strip()

    return f"""
AVALIAÇÃO PEDAGÓGICA - INCLUISRM

Código interno do estudante: {estudante[1]}
Data do registro: {v('data_registro')}
Ano letivo: {v('ano_letivo')}
Tipo de registro: {v('tipo_registro')}
Avaliação anterior vinculada: {v('avaliacao_anterior_id')}

Barreiras enfrentadas:
{v('barreiras')}

Potencialidades e habilidades:
{v('potencialidades')}

Comunicação:
{v('comunicacao')}

Interação social:
{v('interacao')}

Autonomia:
{v('autonomia')}

Aprendizagem:
{v('aprendizagem')}

Resumo pedagógico do laudo, sem identificação:
{v('resumo_laudo')}

Análise comparativa gerada com IA:
{v('analise_comparativa_ia')}

Sugestão de nova avaliação pedagógica gerada com IA:
{v('sugestao_nova_avaliacao_ia')}
""".strip()

def texto_entrevista(estudante, e):
    dados = dict(zip(CAMPOS_ENTREVISTA_FAMILIA, e[1:]))
    def v(campo):
        valor = dados.get(campo)
        return valor if valor not in (None, "") else "Não informado."

    return f"""
ENTREVISTA COM A FAMÍLIA - INCLUISRM

Código interno do estudante: {estudante[1]}
Data do registro: {v('data_registro')}
Ano letivo: {v('ano_letivo')}
Tipo de registro: {v('tipo_registro')}

Campos sensíveis para preenchimento manual:
Nome do estudante: _______________________________________
Responsável: ____________________________________________
CPF do responsável: ______________________________________
Contato: ________________________________________________

1. INFORMAÇÕES DIVERSAS
Participa de programa de auxílio governamental: {v('auxilio_governamental')}
Qual(is): {v('auxilio_quais')}
Histórico familiar de doenças graves, deficiência ou transtornos: {v('historico_familiar')}
Qual(is): {v('historico_quais')}
Já repetiu de ano: {v('repetiu_ano')} | Quantas vezes: {v('repetiu_qtd')}
Trocou de escola: {v('trocou_escola')} | Quantas vezes: {v('trocou_qtd')}
Motivo da troca: {v('motivo_troca')}
Situação na escola: {v('situacao_escolar')}
Demonstra interesse em frequentar a escola: {v('interesse_escola')}
Cuida/organiza os materiais: {v('organiza_materiais')}
Apresenta resistência à escola: {v('resistencia_escola')}
Relaciona-se bem com colegas: {v('relacao_colegas')}
Relaciona-se bem com professores: {v('relacao_professores')}
Leva alimentação de casa: {v('leva_alimentacao')}
Alimenta-se da merenda escolar: {v('merenda_escolar')}
Possui alergia alimentar: {v('alergia_alimentar')}
Qual(is) alergias: {v('alergia_quais')}
Outras observações: {v('obs_diversas')}

1.1 SOBRE A ESCOLHA DA ESCOLA
Motivo da escolha da escola: {v('motivo_escolha')}
Outros motivos: {v('outros_motivos')}
O que conhece sobre o serviço do AEE: {v('conhecimento_aee')}

2. INFORMAÇÕES SOBRE A SAÚDE
Possui doença preexistente: {v('doenca_preexistente')}
Convulsões: {v('convulsoes')}
Acompanhamentos profissionais: {v('acompanhamentos')}
Outro acompanhamento: {v('acompanhamento_outro')}
Frequência dos acompanhamentos: {v('frequencia_acompanhamento')}
Outra frequência: {v('frequencia_outro')}
Alimentação saudável: {v('alimentacao_saudavel')}
Seletividade alimentar: {v('seletividade_alimentar')}
Dieta sensorial: {v('dieta_sensorial')}
Usa suplemento alimentar: {v('suplemento_alimentar')}
Qual suplemento: {v('suplemento_qual')}
Alimenta-se por sonda: {v('alimenta_sonda')}
Dorme bem: {v('dorme_bem')}
Faz uso de medicação: {v('medicacao')}
Qual(is) medicação(ões): {v('medicacao_qual')}
Tempo de uso/tratamentos realizados: {v('tempo_medicacao_tratamentos')}
Outras observações de saúde: {v('obs_saude')}

3. DESENVOLVIMENTO PSICOMOTOR
Lateralidade: {v('lateralidade')}
Apresenta estereotipias: {v('estereotipias')}
Qual(is): {v('estereotipias_quais')}
Segura objetos com as duas mãos: {v('segura_objetos_duas_maos')}
Tamanho dos objetos que segura: {v('tamanho_objetos')}
Faz a pega correta do lápis: {v('pega_lapis')}
Engatinhou: {v('engatinhou')}
Andou com que idade: {v('idade_andou')}
Usa fraldas na escola: {v('usa_fraldas')}
Usa sonda de alívio: {v('usa_sonda_alivio')}
O que consegue fazer sem ajuda: {v('autonomia_atividades')}
Outras atividades de autonomia: {v('autonomia_outros')}
Atende comandos: {v('atende_comandos')}
Gosta do toque: {v('gosta_toque')}
Outras observações psicomotoras: {v('obs_psicomotor')}

4. LINGUAGEM
É verbal: {v('verbal')}
Consegue se comunicar: {v('consegue_comunicar')}
Possui problemas na fala: {v('problemas_fala')}
Tem ecolalia: {v('ecolalia')}
Consegue dar recado: {v('da_recado')}
Usa comunicação alternativa: {v('comunicacao_alternativa')}
Qual comunicação alternativa: {v('comunicacao_alternativa_qual')}

5. SOCIALIZAÇÃO
Relaciona-se bem com o pai: {v('relacao_pai')}
Relaciona-se bem com a mãe: {v('relacao_mae')}
Relaciona-se bem com outros parentes: {v('relacao_parentes')}
Relaciona-se bem com irmãos: {v('relacao_irmaos')}
Relaciona-se bem com outros estudantes: {v('relacao_estudantes')}
Tem melhor amigo(a): {v('tem_melhor_amigo')}
Esse(a) melhor amigo(a) é: {v('tipo_melhor_amigo')}
Adapta-se facilmente ao ambiente: {v('adapta_ambiente')}
É flexível na rotina: {v('flexivel_rotina')}
Respeita regras: {v('respeita_regras')}
Chora com facilidade: {v('chora_facilidade')}
Gosta de brincar: {v('brinca_como')}
Assunto ou lazer de interesse: {v('interesses_lazer')}
O que a família mais gosta no(a) estudante: {v('familia_gosta')}
O que a família não gosta/necesita melhorar: {v('familia_nao_gosta')}
Ambiente físico em casa para estudos/brincadeiras: {v('ambiente_estudo_casa')}

6. CONTEXTO FAMILIAR
Principais habilidades:
{v('habilidades')}

Principais oportunidades de melhoria:
{v('oportunidades_melhoria')}

7. OUTRAS INFORMAÇÕES
{v('outras_info_familia')}

Assinatura do responsável:
________________________________________________________
""".strip()


def buscar_entrevista_familia_por_id(registro_id):
    """Busca uma entrevista familiar completa pelo ID, mantendo o formato (id, campos...)."""
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        f"SELECT id, {', '.join(CAMPOS_ENTREVISTA_FAMILIA)} FROM entrevistas_familia WHERE id=?",
        (registro_id,),
    )
    dado = cursor.fetchone()
    conn.close()
    return dado


def _normalizar_valor_comparacao(valor):
    if valor is None:
        return ""
    texto = str(valor).strip()
    if texto.lower() in ["", "não informado", "nao informado", "none", "null"]:
        return ""
    return texto


def _dict_entrevista_familia(registro):
    if not registro:
        return {}
    return dict(zip(["id"] + CAMPOS_ENTREVISTA_FAMILIA, registro))


def gerar_relatorio_comparativo_entrevistas_familia(estudante, entrevista_anterior=None, entrevista_atual=None):
    """Gera relatório comparativo manual/regra de entrevistas familiares, sem IA.
    Objetivo: identificar mudanças no contexto familiar que possam influenciar participação,
    autorregulação, vínculo, rotina, saúde, socialização e aprendizagem.
    """
    registros = listar_por_estudante("entrevistas_familia", CAMPOS_ENTREVISTA_FAMILIA, estudante[0])
    if not registros or len(registros) < 2:
        return f"""
RELATÓRIO COMPARATIVO DAS ENTREVISTAS COM A FAMÍLIA - INCLUISRM

Código interno do estudante: {estudante[1] or 'Não informado.'}

Ainda não há entrevistas familiares suficientes para comparação.
Registre pelo menos duas entrevistas, preferencialmente de anos letivos diferentes, para gerar uma análise comparativa.
""".strip()

    if entrevista_anterior is None or entrevista_atual is None:
        # Como listar_por_estudante retorna do mais recente para o mais antigo, usa as duas últimas registradas.
        entrevista_atual = registros[0]
        entrevista_anterior = registros[1]

    ant = _dict_entrevista_familia(entrevista_anterior)
    atu = _dict_entrevista_familia(entrevista_atual)

    rotulos = {
        "auxilio_governamental": "Participação em programa de auxílio governamental",
        "auxilio_quais": "Programas de auxílio informados",
        "historico_familiar": "Histórico familiar de doenças, deficiência ou transtornos",
        "historico_quais": "Históricos familiares informados",
        "repetiu_ano": "Repetência",
        "repetiu_qtd": "Quantidade de repetências",
        "trocou_escola": "Mudança de escola",
        "trocou_qtd": "Quantidade de mudanças de escola",
        "motivo_troca": "Motivo da troca de escola",
        "situacao_escolar": "Situação escolar",
        "interesse_escola": "Interesse em frequentar a escola",
        "organiza_materiais": "Organização dos materiais",
        "resistencia_escola": "Resistência à escola",
        "relacao_colegas": "Relação com colegas",
        "relacao_professores": "Relação com professores",
        "doenca_preexistente": "Doença preexistente",
        "convulsoes": "Convulsões",
        "acompanhamentos": "Acompanhamentos profissionais",
        "frequencia_acompanhamento": "Frequência dos acompanhamentos",
        "alimentacao_saudavel": "Alimentação saudável",
        "seletividade_alimentar": "Seletividade alimentar",
        "dieta_sensorial": "Dieta sensorial",
        "dorme_bem": "Sono",
        "medicacao": "Uso de medicação",
        "medicacao_qual": "Medicações informadas",
        "lateralidade": "Lateralidade",
        "estereotipias": "Estereotipias observadas",
        "autonomia_atividades": "Atividades de autonomia",
        "atende_comandos": "Atende comandos",
        "gosta_toque": "Aceitação do toque",
        "verbal": "Comunicação verbal",
        "consegue_comunicar": "Capacidade de comunicação",
        "problemas_fala": "Problemas na fala",
        "ecolalia": "Ecolalia",
        "comunicacao_alternativa": "Uso de comunicação alternativa",
        "relacao_pai": "Relação com o pai",
        "relacao_mae": "Relação com a mãe",
        "relacao_parentes": "Relação com outros parentes",
        "relacao_irmaos": "Relação com irmãos",
        "relacao_estudantes": "Relação com outros estudantes",
        "tem_melhor_amigo": "Presença de melhor amigo(a)",
        "adapta_ambiente": "Adaptação ao ambiente",
        "flexivel_rotina": "Flexibilidade na rotina",
        "respeita_regras": "Respeito às regras",
        "chora_facilidade": "Choro com facilidade",
        "brinca_como": "Forma de brincar",
        "interesses_lazer": "Interesses e lazer",
        "familia_gosta": "Aspectos que a família mais valoriza",
        "familia_nao_gosta": "Aspectos que a família aponta para melhorar",
        "ambiente_estudo_casa": "Ambiente físico em casa para estudo/brincadeiras",
        "habilidades": "Habilidades percebidas pela família",
        "oportunidades_melhoria": "Oportunidades de melhoria percebidas pela família",
        "outras_info_familia": "Outras informações familiares relevantes",
    }

    campos_ignorar = {"id", "data_registro", "ano_letivo", "tipo_registro"}
    campos_sensiveis_pedagogicos = {
        "relacao_pai", "relacao_mae", "relacao_parentes", "relacao_irmaos",
        "ambiente_estudo_casa", "familia_gosta", "familia_nao_gosta", "outras_info_familia",
        "interesse_escola", "resistencia_escola", "chora_facilidade", "flexivel_rotina",
        "dorme_bem", "medicacao", "medicacao_qual", "seletividade_alimentar",
        "acompanhamentos", "interesses_lazer", "adapta_ambiente"
    }

    alteracoes = []
    destaques = []
    for campo in CAMPOS_ENTREVISTA_FAMILIA:
        if campo in campos_ignorar:
            continue
        valor_ant = _normalizar_valor_comparacao(ant.get(campo))
        valor_atu = _normalizar_valor_comparacao(atu.get(campo))
        if valor_ant != valor_atu:
            rotulo = rotulos.get(campo, campo.replace("_", " ").capitalize())
            linha = f"- {rotulo}: antes: {valor_ant or 'não informado'} | atual: {valor_atu or 'não informado'}"
            alteracoes.append(linha)
            if campo in campos_sensiveis_pedagogicos:
                destaques.append(linha)

    if alteracoes:
        texto_alteracoes = "\n".join(alteracoes)
    else:
        texto_alteracoes = "Não foram identificadas alterações objetivas nos campos comparados."

    if destaques:
        texto_destaques = "\n".join(destaques)
    else:
        texto_destaques = "Não foram identificadas alterações em campos familiares de maior impacto pedagógico imediato."

    return f"""
RELATÓRIO COMPARATIVO DAS ENTREVISTAS COM A FAMÍLIA - INCLUISRM

Código interno do estudante: {estudante[1] or 'Não informado.'}
Ano/Série atual: {estudante[2] or 'Não informado.'}
Turma: {estudante[3] or 'Não informado.'}
Perfil educacional informado: {estudante[4] or 'Não informado.'}

Entrevista anterior:
- ID do registro: {ant.get('id') or 'Não informado.'}
- Data do registro: {ant.get('data_registro') or 'Não informado.'}
- Ano letivo: {ant.get('ano_letivo') or 'Não informado.'}
- Tipo de registro: {ant.get('tipo_registro') or 'Não informado.'}

Entrevista atual/comparada:
- ID do registro: {atu.get('id') or 'Não informado.'}
- Data do registro: {atu.get('data_registro') or 'Não informado.'}
- Ano letivo: {atu.get('ano_letivo') or 'Não informado.'}
- Tipo de registro: {atu.get('tipo_registro') or 'Não informado.'}

1. OBJETIVO DO RELATÓRIO
Este relatório compara entrevistas familiares registradas em momentos diferentes, com o objetivo de identificar alterações no contexto familiar, na rotina, na saúde, na socialização, na comunicação, na autonomia e nas condições de apoio ao estudante. A comparação apoia a leitura pedagógica do professor do AEE, especialmente quando mudanças familiares podem influenciar autorregulação, comportamento, participação escolar, vínculo afetivo, aprendizagem e necessidade de acolhimento.

2. ALTERAÇÕES IDENTIFICADAS ENTRE AS ENTREVISTAS
{texto_alteracoes}

3. CAMPOS DE ATENÇÃO PEDAGÓGICA E FAMILIAR
{texto_destaques}

4. LEITURA PEDAGÓGICA PARA ACOMPANHAMENTO
As alterações registradas devem ser analisadas com cuidado pelo professor do AEE, em diálogo com a família, equipe pedagógica e professores da sala comum. Mudanças como perda de familiar, alteração de moradia, novos responsáveis, conflitos familiares, mudança no sono, alimentação, medicação, humor, rotina ou vínculo podem repercutir diretamente na participação do estudante, na autorregulação emocional e na resposta às intervenções pedagógicas.

5. ENCAMINHAMENTOS SUGERIDOS PARA REGISTRO MANUAL
- Verificar se as alterações familiares coincidem com mudanças no comportamento, frequência, participação ou desempenho escolar.
- Registrar observações nos atendimentos do AEE, especialmente sinais de desregulação, isolamento, agitação, tristeza, resistência à escola ou mudança na interação social.
- Dialogar com a família para compreender melhor o contexto atual, sem exposição de dados sensíveis.
- Atualizar, se necessário, a Avaliação Pedagógica, o Estudo de Caso e o Plano AEE/PAEE.
- Acionar a equipe pedagógica ou rede de apoio da escola quando houver necessidade de acolhimento, cuidado emocional ou acompanhamento intersetorial.

6. OBSERVAÇÃO IMPORTANTE
Este relatório pode ser gerado em modo comparativo manual/regra ou com apoio de IA na etapa de análise. A entrevista familiar permanece sempre como registro manual e fiel ao que foi informado pela família; a IA não cria, altera nem preenche entrevistas. Quando utilizada, a IA apenas apoia a leitura comparativa para subsidiar a decisão pedagógica do professor do AEE.

Professor(a) AEE: _______________________________________
Coordenação/Gestão: _____________________________________
Data: ____/____/________
""".strip()


def gerar_analise_ia_comparativo_entrevistas_familia(estudante, entrevista_anterior=None, entrevista_atual=None, observacao_professor=""):
    """Gera análise comparativa com IA a partir de duas entrevistas familiares.
    A IA NÃO cria nem altera entrevistas; apenas analisa diferenças para apoiar o AEE.
    """
    client = obter_cliente_openai()
    if client is None:
        return "IA não configurada. Configure OPENAI_API_KEY para usar esta função."

    if entrevista_anterior is None or entrevista_atual is None:
        return "Selecione duas entrevistas familiares diferentes para gerar a análise com IA."

    ant = _dict_entrevista_familia(entrevista_anterior)
    atu = _dict_entrevista_familia(entrevista_atual)
    relatorio_base = gerar_relatorio_comparativo_entrevistas_familia(
        estudante,
        entrevista_anterior=entrevista_anterior,
        entrevista_atual=entrevista_atual,
    )

    prompt = f"""
Você é especialista em Atendimento Educacional Especializado (AEE), educação inclusiva e análise pedagógica contextual.

TAREFA:
Analise comparativamente duas entrevistas familiares do mesmo estudante.

REGRA CENTRAL:
A IA NÃO deve criar, alterar, complementar ou inventar dados da entrevista familiar.
A entrevista é um registro manual e fiel ao que a família informou.
Sua função é apenas analisar as diferenças entre os registros e apoiar a leitura pedagógica do professor do AEE.

IDENTIFICAÇÃO NÃO SENSÍVEL DO ESTUDANTE:
Código interno: {estudante[1] or 'Não informado'}
Ano/Série: {estudante[2] or 'Não informado'}
Turma: {estudante[3] or 'Não informado'}
Perfil educacional informado: {estudante[4] or 'Não informado'}

ENTREVISTA ANTERIOR / BASE:
{ant}

ENTREVISTA ATUAL / COMPARADA:
{atu}

RELATÓRIO COMPARATIVO BASE GERADO PELO SISTEMA:
{relatorio_base}

OBSERVAÇÃO DO PROFESSOR DO AEE PARA ORIENTAR A ANÁLISE:
{observacao_professor or 'Não informada.'}

RESPONDA EM LINGUAGEM INSTITUCIONAL, CUIDADOSA E PEDAGÓGICA, COM AS SEÇÕES:
1. Síntese comparativa geral
2. Mudanças na estrutura familiar e rede de apoio
3. Mudanças na rotina, saúde, sono, alimentação ou acompanhamento
4. Possíveis sinais de desregulação observáveis no contexto escolar
5. Possíveis avanços ou fatores de proteção
6. Pontos de atenção para o professor do AEE
7. Sugestões pedagógicas de acompanhamento, sem diagnóstico clínico
8. Cuidados éticos: dados que precisam ser confirmados com família/equipe antes de qualquer encaminhamento

LIMITES:
- Não diagnosticar.
- Não afirmar causalidade sem evidência.
- Não expor dados sensíveis desnecessários.
- Não inventar fatos.
- Usar expressões como “pode indicar”, “sugere atenção”, “recomenda-se observar”, quando houver incerteza.
"""
    try:
        resposta = client.responses.create(model="gpt-4.1-mini", input=prompt)
        return resposta.output_text
    except Exception as e:
        return f"Não foi possível gerar a análise comparativa com IA agora. Erro: {e}"


def texto_campos_sensiveis_em_branco_estudante():
    return """
DADOS SENSÍVEIS PARA PREENCHIMENTO MANUAL NO DOCUMENTO IMPRESSO/WORD
Nome completo do(a) estudante: ___________________________________________
CPF/RG do(a) estudante: _________________________________________________
Data de nascimento: ____/____/________
Nome do responsável: _________________________________________________
CPF/RG do responsável: _______________________________________________
Telefone do responsável: ______________________________________________
Endereço: ____________________________________________________________
""".strip()


def texto_estudo_caso(estudante, e):
    """Texto oficial do Estudo de Caso no formato GRE, sem armazenar dados sensíveis."""
    # Compatibilidade com estudos antigos: e = (id, data, contextualizacao, queixa, ...)
    if len(e) <= 10:
        return f"""
ESTUDO DE CASO - INCLUISRM

Código interno do estudante: {estudante[1]}
Ano/Série: {estudante[2] or 'Não informado.'}
Turma: {estudante[3] or 'Não informado.'}
Perfil educacional: {estudante[4] or 'Não informado.'}
Data do registro: {e[1]}

{texto_campos_sensiveis_em_branco_estudante()}

1. Contextualização
{e[2] or 'Não informado.'}

2. Queixa principal / motivo do acompanhamento
{e[3] or 'Não informado.'}

3. Potencialidades
{e[4] or 'Não informado.'}

4. Dificuldades observadas
{e[5] or 'Não informado.'}

5. Estratégias pedagógicas
{e[6] or 'Não informado.'}

6. Intervenções / encaminhamentos sugeridos
{e[7] or 'Não informado.'}

7. Avaliação
{e[8] or 'Não informado.'}

8. Considerações finais
{e[9] or 'Não informado.'}

Assinaturas:
Professor(a) AEE: _______________________________________
Coordenação/Gestão: _____________________________________
""".strip()

    dados = dict(zip(CAMPOS_ESTUDO_CASO, e[1:]))

    def v(campo):
        valor = dados.get(campo)
        return valor if valor not in (None, "") else "Não informado."

    return f"""
ESTUDO DE CASO E PLANO DE ATENDIMENTO EDUCACIONAL ESPECIALIZADO - INCLUISRM

O Plano de Atendimento Educacional Especializado deverá considerar os registros avaliativos do estudante público-alvo da educação especial, partindo do estudo de caso e da identificação de barreiras, recursos de acessibilidade e estratégias necessárias à promoção da autonomia e da aprendizagem.

PARTE 1 - IDENTIFICAÇÃO SEGURA DO(A) ESTUDANTE
Código interno do estudante: {estudante[1] or 'Não informado.'}
Ano/Série cadastrado no sistema: {estudante[2] or 'Não informado.'}
Turma: {estudante[3] or 'Não informado.'}
Turno: {estudante[6] or 'Não informado.'}
Perfil educacional: {estudante[4] or 'Não informado.'}
Ano letivo do estudo de caso: {v('ano_letivo')}
Tipo de registro: {v('tipo_registro')}

{texto_campos_sensiveis_em_branco_estudante()}

1.5 Etapa/modalidade da educação em que o(a) estudante está: {v('etapa_modalidade')}
Ano/etapa: {v('ano_etapa')}
1.6 Turma e turno: {estudante[3] or 'Não informado.'} / {estudante[6] or 'Não informado.'}
1.7 O(a) estudante apresenta laudo? {v('laudo')}
1.8 Apresenta deficiência? {v('deficiencia')} | CID: {v('cid')}
Síntese pedagógica/funcional do laudo: {v('observacoes_laudo')}
1.9 Altas habilidades/superdotação: {v('altas_habilidades')}
1.10 Usuário de BPC: {v('bpc')}
1.11 Escola do ensino comum: {v('escola_nome')}
1.12 Unidade educacional onde é atendido pelo AEE: {v('unidade_aee')}
1.13 Gestor(a) da escola do ensino comum e contato: {v('gestor_nome')} - {v('gestor_contato')}
1.14 Professor(a) do AEE e contato: {v('professor_nome')} - {v('professor_contato')}
1.15 Matrícula do(a) professor(a) do AEE: {v('matricula_professor')}
1.16 Especialidade do(a) professor(a) do AEE: {v('especialidade_professor')}
1.17 Período de elaboração do Plano/AEE - início: {v('periodo_inicio')}
1.18 Data final: {v('periodo_fim')}
1.19 Frequência de atendimento na SRM: {v('frequencia_atendimento')}
1.20 Tempo de atendimento por semana: {v('tempo_atendimento_semana')}
1.21 Formato do atendimento: {v('formato_atendimento')}

2. ESTUDO DE CASO / PERCURSO EDUCACIONAL DO(A) ESTUDANTE
Relato sobre o trajeto educacional do(a) estudante em turmas comuns e no Atendimento Educacional Especializado anterior, quando aplicável:
{v('percurso_educacional')}

2.1 Motivo pelo qual o(a) estudante foi encaminhado para o Atendimento Educacional Especializado:
{v('motivo_encaminhamento_aee')}

2.2 Precisa de transporte escolar inclusivo? {v('precisa_transporte_inclusivo')}
2.2.1 Recebe o serviço de transporte escolar inclusivo? {v('recebe_transporte_inclusivo')}

2.3 Precisa de profissional de apoio? {v('precisa_profissional_apoio')}
Justificativa:
{v('justificativa_apoio')}

2.3.1 É acompanhado por profissional de apoio na escola? {v('acompanhado_profissional_apoio')}
Nome do profissional de apoio: {v('nome_profissional_apoio')}

2.4 Recursos de tecnologia educacional e/ou assistiva utilizados:
{v('recursos_tecnologia_assistiva')}

2.5 Observações relevantes para o ambiente educacional, acompanhamento médico ou terapêutico:
{v('observacoes_ambiente_educacional')}

2.6 Habilidades observadas e desenvolvidas pelo(a) estudante:
{v('habilidades_observadas')}

2.7 Habilidades que precisam ser desenvolvidas pelo(a) estudante:
{v('habilidades_a_desenvolver')}

2.8 Indicadores de altas habilidades/superdotação observados:
{v('indicadores_altas_habilidades')}

2.9 Caso o estudante seja pessoa surda, recursos utilizados:
{v('recursos_surdez')}
Observações sobre uso de aparelho auditivo, implante coclear, Libras, comunicação e efetividade dos recursos:
{v('observacoes_surdez')}

SÍNTESE PEDAGÓGICA DO ESTUDO DE CASO
Contextualização:
{v('contextualizacao')}

Potencialidades:
{v('potencialidades')}

Dificuldades/barreiras observadas:
{v('dificuldades')}

Estratégias pedagógicas:
{v('estrategias')}

Intervenções/encaminhamentos sugeridos:
{v('intervencoes')}

Avaliação:
{v('avaliacao')}

Considerações finais:
{v('consideracoes')}

ANÁLISE COMPARATIVA GERADA COM APOIO DA IA
{v('analise_comparativa_ia')}

SUGESTÃO DE NOVO ESTUDO DE CASO GERADA COM APOIO DA IA
{v('sugestao_novo_estudo_ia')}

Assinaturas:
Professor(a) AEE: _______________________________________
Coordenação/Gestão: _____________________________________
Responsável: ____________________________________________
""".strip()


def texto_plano_aee(estudante, p):
    """Gera o texto do Plano AEE/PAEE manual com todos os campos do modelo GRE.
    Aceita registros antigos e novos para evitar erro em bancos já existentes.
    """
    campos_antigos = [
        "data_registro", "objetivos_gerais", "objetivos_especificos", "habilidades_prioritarias",
        "recursos_acessibilidade", "estrategias", "organizacao_atendimento", "parcerias",
        "avaliacao_acompanhamento", "observacoes",
    ]

    valores = list(p[1:]) if p and len(p) > 1 else []
    if len(valores) == len(CAMPOS_PLANO_AEE):
        dados = dict(zip(CAMPOS_PLANO_AEE, valores))
    else:
        dados = dict(zip(campos_antigos, valores))

    def v(campo):
        valor = dados.get(campo)
        return valor if valor not in (None, "") else "Não informado."

    # Compatibilidade: se o banco antigo ainda tinha organização_atendimento,
    # esse conteúdo aparece como observação complementar.
    organizacao_antiga = dados.get("organizacao_atendimento")
    observacoes_extra = v("observacoes")
    if organizacao_antiga not in (None, ""):
        observacoes_extra = f"{observacoes_extra}\n\nOrganização do atendimento:\n{organizacao_antiga}"

    return f"""
PLANO AEE / PAEE - INCLUISRM

Código interno do estudante: {estudante[1]}
Ano/Série: {estudante[2] or 'Não informado.'}
Turma: {estudante[3] or 'Não informado.'}
Perfil educacional: {estudante[4] or 'Não informado.'}
Data do registro: {v('data_registro')}

1. Habilidades específicas
1.1 Habilidades prioritárias que serão trabalhadas na SRM
{v('habilidades_prioritarias')}

1.2 Recursos de acessibilidade que serão disponibilizados ao estudante
{v('recursos_acessibilidade')}

2. Objetivos do Atendimento Educacional Especializado
2.1 Geral
{v('objetivos_gerais')}

2.2 Específicos
{v('objetivos_especificos')}

3. Metodologias e estratégias
3.1 Metodologia
{v('metodologia')}

3.2 Estratégia
{v('estrategias')}

3.3 Prazo
{v('prazo')}

4. Ações desenvolvidas no âmbito da escola
{v('acoes_escola')}

5. Barreiras identificadas na comunidade escolar
{v('barreiras_identificadas')}

6. Parcerias realizadas pelo AEE ao longo do período
{v('parcerias')}

7. Avaliação
{v('avaliacao_acompanhamento')}

8. Observações
{observacoes_extra}

Assinaturas:
Professor(a) AEE: _______________________________________
Coordenação/Gestão: _____________________________________
Responsável: ____________________________________________
""".strip()


def texto_atendimento(estudante, a):
    return f"""
REGISTRO DE ATENDIMENTO DO AEE - INCLUISRM

Código interno: {estudante[1]}
Data do atendimento: {a[1]}

Objetivo trabalhado:
{a[2] or 'Não informado.'}

Atividade realizada:
{a[3] or 'Não informado.'}

Resposta do estudante:
{a[4] or 'Não informado.'}

Avanços observados:
{a[5] or 'Não informado.'}

Dificuldades observadas:
{a[6] or 'Não informado.'}

Evolução observada:
{a[7] or 'Não informado.'}

Indicadores:
Resposta do estudante: {a[9]}/10
Avanço pedagógico: {a[10]}/10
Nível de dificuldade/barreira: {a[11]}/10
Engajamento: {a[12]}/10
Índice geral de evolução: {a[13]}/10

Encaminhamentos:
{a[14] or 'Não informado.'}

Materiais/recursos pedagógicos utilizados no atendimento:
{texto_recursos_atendimento(a[0])}
""".strip()


def texto_agenda(df):
    if df.empty:
        return "AGENDA DE ATENDIMENTOS - INCLUISRM\n\nNenhum agendamento registrado."
    return "AGENDA DE ATENDIMENTOS - INCLUISRM\n\n" + df.to_string(index=False)


# ======================================================
# PDF
# ======================================================
def gerar_pdf_documento(conteudo, codigo, tipo="documento"):
    """Gera PDF com layout mais limpo para documentos pedagógicos.

    Ajustes V38:
    - remove marcações Markdown (**, #, ---);
    - melhora cabeçalho com logo e faixa institucional;
    - evita que marcadores soltos como "• --" apareçam no relatório;
    - destaca seções numeradas em azul;
    - organiza o Perfil Pedagógico Inteligente com visual mais profissional.
    """
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer, Image, Table, TableStyle

    nomes = {
        "cadastro": ("Cadastro_Estudante", "CADASTRO DO ESTUDANTE"),
        "matricula_srm": ("Matricula_SRM", "MATRÍCULA PARA O ATENDIMENTO EDUCACIONAL ESPECIALIZADO - SRM"),
        "professor": ("Ficha_Professor_AEE", "FICHA DE IDENTIFICAÇÃO DO(A) PROFESSOR(A) AEE"),
        "entrevista": ("Entrevista_Familia", "ENTREVISTA COM A FAMÍLIA"),
        "avaliacao": ("Avaliacao_Pedagogica", "AVALIAÇÃO PEDAGÓGICA INICIAL"),
        "estudo": ("Estudo_Caso", "ESTUDO DE CASO"),
        "plano": ("Plano_AEE_PAEE", "PLANO AEE / PAEE"),
        "atendimento": ("Registro_Atendimento", "REGISTRO DE ATENDIMENTO DO AEE"),
        "agenda": ("Agenda_Atendimentos", "AGENDA DE ATENDIMENTOS"),
        "relatorio": ("Relatorio_GRE", "RELATÓRIO GRE"),
        "documento": ("Documento", "DOCUMENTO"),
    }
    prefixo, titulo_doc = nomes.get(tipo, nomes["documento"])

    conteudo_original = str(conteudo or "")
    if tipo == "plano" and ("Perfil Pedagógico" in conteudo_original[:900] or "Síntese pedagógica do estudante" in conteudo_original[:1200]):
        titulo_doc = "PERFIL PEDAGÓGICO INTELIGENTE - AEE"
        prefixo = "Perfil_Pedagogico_Inteligente"

    tema = obter_tema_documento(tipo, titulo_doc, conteudo_original)
    cor_principal = tema["cor"]
    cor_clara = tema["cor_clara"]

    nome_arquivo = f"{prefixo}_{codigo}.pdf".replace("/", "-").replace("\\", "-")

    doc = SimpleDocTemplate(
        nome_arquivo,
        pagesize=A4,
        rightMargin=1.45 * cm,
        leftMargin=1.45 * cm,
        topMargin=1.15 * cm,
        bottomMargin=1.15 * cm,
    )
    styles = getSampleStyleSheet()

    titulo_style = ParagraphStyle(
        name="TituloDocumentoV38",
        parent=styles["Title"],
        alignment=TA_CENTER,
        fontSize=16,
        leading=20,
        spaceAfter=10,
        textColor=colors.HexColor(cor_principal),
    )
    subtitulo_style = ParagraphStyle(
        name="SubtituloDocumentoV38",
        parent=styles["Normal"],
        alignment=TA_CENTER,
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#475569"),
        spaceAfter=8,
    )
    secao_style = ParagraphStyle(
        name="SecaoDocumentoV38",
        parent=styles["Heading2"],
        fontSize=12,
        leading=15,
        spaceBefore=10,
        spaceAfter=5,
        textColor=colors.HexColor(cor_principal),
    )
    normal_style = ParagraphStyle(
        name="NormalDocumentoV38",
        parent=styles["Normal"],
        fontSize=9.6,
        leading=13.5,
        spaceAfter=5,
        textColor=colors.HexColor("#111827"),
        alignment=TA_LEFT,
    )
    bullet_style = ParagraphStyle(
        name="BulletDocumentoV38",
        parent=normal_style,
        leftIndent=12,
        firstLineIndent=-8,
    )
    destaque_style = ParagraphStyle(
        name="DestaqueDocumentoV38",
        parent=normal_style,
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#0f172a"),
        backColor=colors.HexColor(cor_clara),
        borderColor=colors.HexColor(cor_principal),
        borderWidth=0.4,
        borderPadding=6,
        spaceAfter=8,
    )
    rodape_style = ParagraphStyle(
        name="RodapeDocumentoV38",
        parent=styles["Normal"],
        fontSize=8,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#64748b"),
        spaceBefore=14,
    )

    def limpar_linha_pdf(linha):
        linha = str(linha or "").strip()
        linha = linha.replace("**", "")
        linha = re.sub(r"^#{1,6}\s*", "", linha)
        linha = linha.replace("• --", "").replace("--", "").strip()
        return linha

    elementos = []

    # Cabeçalho institucional
    try:
        logo = Image(LOGO_PATH, width=4.3 * cm, height=2.1 * cm)
        logo.hAlign = "CENTER"
        elementos.append(logo)
        elementos.append(Spacer(1, 6))
    except Exception:
        elementos.append(Paragraph("<b>INCLUISRM</b>", titulo_style))

    header = Table(
        [[
            Paragraph("<b>INCLUISRM</b><br/>Sistema de Gestão do Atendimento Educacional Especializado", normal_style),
            Paragraph("<b>Documento pedagógico de apoio ao AEE</b><br/>Gerado com apoio do INCLUISRM", normal_style),
        ]],
        colWidths=[8.4 * cm, 8.4 * cm],
    )
    header.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, 0), colors.HexColor(cor_principal)),
        ("TEXTCOLOR", (0, 0), (0, 0), colors.white),
        ("BACKGROUND", (1, 0), (1, 0), colors.HexColor(cor_clara)),
        ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    elementos.append(header)
    elementos.append(Spacer(1, 10))
    elementos.append(HRFlowable(width="100%", thickness=0.7, color=colors.HexColor("#94a3b8")))
    elementos.append(Spacer(1, 10))
    elementos.append(Paragraph(f"<b>{escape(titulo_doc)}</b>", titulo_style))
    elementos.append(Paragraph("Planejamento organizado em linguagem pedagógica, com foco em acessibilidade, participação, autonomia e acompanhamento evolutivo.", subtitulo_style))

    papel_box = Table([[Paragraph("<b>Papel deste documento:</b> " + escape(tema["papel"]), normal_style)]], colWidths=[16.8 * cm])
    papel_box.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(cor_clara)),
        ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor(cor_principal)),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    elementos.append(papel_box)
    elementos.append(Spacer(1, 10))

    texto_limpo = limpar_marcadores_relatorio(conteudo_original) if 'limpar_marcadores_relatorio' in globals() else conteudo_original

    # Pequeno card com identificação inicial quando houver linhas com chave/valor
    pares = []
    for raw in texto_limpo.splitlines()[:20]:
        linha = limpar_linha_pdf(raw)
        if ":" in linha and len(linha) < 160:
            chave, valor = linha.split(":", 1)
            if chave.strip() and valor.strip() and len(chave.strip()) <= 45:
                pares.append((chave.strip(), valor.strip()))
    if pares:
        dados = [[Paragraph("<b>Campo</b>", normal_style), Paragraph("<b>Informação</b>", normal_style)]]
        for chave, valor in pares[:6]:
            dados.append([Paragraph(escape(chave), normal_style), Paragraph(escape(valor), normal_style)])
        tabela = Table(dados, colWidths=[5.0 * cm, 11.8 * cm])
        tabela.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(cor_principal)),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("BACKGROUND", (0, 1), (0, -1), colors.HexColor(cor_clara)),
            ("BACKGROUND", (1, 1), (1, -1), colors.HexColor("#f8fafc")),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#cbd5e1")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 7),
            ("RIGHTPADDING", (0, 0), (-1, -1), 7),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        elementos.append(tabela)
        elementos.append(Spacer(1, 10))

    for raw in texto_limpo.split("\n"):
        linha = limpar_linha_pdf(raw)
        if not linha:
            elementos.append(Spacer(1, 4))
            continue
        if linha in ["-", "•", "--"]:
            continue
        linha_html = escape(linha)
        if re.match(r"^\d+(\.\d+)?\s*[\.-]\s+", linha) or re.match(r"^\d+\.\s+", linha):
            elementos.append(Paragraph(f"<b>{linha_html}</b>", secao_style))
        elif linha.upper() == linha and len(linha) > 12:
            elementos.append(Paragraph(f"<b>{linha_html}</b>", secao_style))
        elif linha.startswith("-") or linha.startswith("•"):
            item = linha.lstrip("-• ").strip()
            if item:
                elementos.append(Paragraph("• " + escape(item), bullet_style))
        elif ":" in linha and len(linha) < 130 and any(k in linha.lower() for k in ["código", "ano", "turma", "turno", "perfil", "estudante"]):
            elementos.append(Paragraph(linha_html, destaque_style))
        else:
            elementos.append(Paragraph(linha_html, normal_style))

    elementos.append(Spacer(1, 14))
    rodape = Table(
        [[Paragraph(f"Gerado em {agora_local().strftime('%d/%m/%Y %H:%M')} pelo INCLUISRM • Documento pedagógico de apoio ao AEE", rodape_style)]],
        colWidths=[16.8 * cm],
    )
    rodape.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f1f5f9")),
        ("BOX", (0, 0), (-1, -1), 0.3, colors.HexColor("#cbd5e1")),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    elementos.append(rodape)
    doc.build(elementos)
    return nome_arquivo

# ======================================================
# IA + BASE DE CONHECIMENTO + MODELOS 3D
# ======================================================
def obter_api_key():
    try:
        return os.getenv("OPENAI_API_KEY") or st.secrets.get("OPENAI_API_KEY")
    except Exception:
        return os.getenv("OPENAI_API_KEY")


def obter_cliente_openai():
    api_key = obter_api_key()
    if OpenAI is None or not api_key:
        return None
    return OpenAI(api_key=api_key)


def extrair_secao_ia(texto, marcador, padrao=""):
    """Extrai uma seção marcada no texto da IA.
    Ex.: [PERCURSO_EDUCACIONAL] ... [PROXIMA_SECAO]
    """
    try:
        padrao_regex = rf"\[{re.escape(marcador)}\](.*?)(?=\n\[[A-Z0-9_]+\]|\Z)"
        m = re.search(padrao_regex, texto or "", flags=re.DOTALL)
        if not m:
            return padrao
        return m.group(1).strip()
    except Exception:
        return padrao


def gerar_novo_estudo_caso_com_ia(estudante, estudo_anterior=None, avaliacao=None, entrevista=None, atendimentos=None, ano_novo="", avaliacoes_contexto=None):
    """Gera análise comparativa e proposta de Estudo de Caso no padrão GRE.
    A resposta é apoio pedagógico: não substitui a avaliação do professor do AEE.
    """
    client = obter_cliente_openai()
    if client is None:
        sugestao_fallback = """
[PERCURSO_EDUCACIONAL]
Preencher com relato sobre o trajeto educacional do(a) estudante, situação inicial, estratégias utilizadas e progressos alcançados em turmas comuns e no AEE, quando houver.

[MOTIVO_ENCAMINHAMENTO]
Preencher com o motivo pedagógico do encaminhamento ao Atendimento Educacional Especializado.

[HABILIDADES_OBSERVADAS]
Preencher com habilidades observadas e desenvolvidas pelo(a) estudante.

[HABILIDADES_A_DESENVOLVER]
Preencher com habilidades que precisam ser desenvolvidas ou reforçadas.

[POTENCIALIDADES]
Preencher com potencialidades observadas nas avaliações pedagógicas, atendimentos e escutas docentes.

[DIFICULDADES]
Preencher com barreiras pedagógicas, comunicacionais, curriculares, atitudinais ou de participação observadas.

[ESTRATEGIAS]
Preencher com estratégias pedagógicas recomendadas para favorecer participação, autonomia e aprendizagem.

[INTERVENCOES]
Preencher com intervenções, recursos e encaminhamentos pedagógicos sugeridos.

[AVALIACAO]
Preencher com avanços, habilidades a reforçar e efetividade das ações realizadas.

[CONSIDERACOES]
Preencher com síntese final do estudo de caso, mantendo linguagem pedagógica e institucional.

[OBSERVACOES_COMPLEMENTARES]
Registrar informações relevantes que não se encaixam nos campos anteriores.
""".strip()
        return "IA não configurada. Configure OPENAI_API_KEY para gerar automaticamente. O sistema preparou um modelo GRE para preenchimento manual.", sugestao_fallback

    texto_estudo_anterior = texto_estudo_caso(estudante, estudo_anterior) if estudo_anterior else "Nenhum estudo anterior selecionado."

    # Usa a avaliação selecionada como principal e também considera todas as demais avaliações
    # pedagógicas/documentos livres existentes para construir uma memória pedagógica mais completa.
    avaliacoes_contexto = avaliacoes_contexto or []
    textos_avaliacoes = []
    ids_avaliacoes_usadas = set()

    if avaliacao:
        try:
            textos_avaliacoes.append("AVALIAÇÃO PRINCIPAL SELECIONADA:\n" + texto_avaliacao(estudante, avaliacao))
            ids_avaliacoes_usadas.add(avaliacao[0])
        except Exception:
            textos_avaliacoes.append("AVALIAÇÃO PRINCIPAL SELECIONADA:\n" + str(avaliacao))

    for av in avaliacoes_contexto[:8]:
        if av and av[0] not in ids_avaliacoes_usadas:
            try:
                textos_avaliacoes.append("AVALIAÇÃO COMPLEMENTAR / DOCUMENTO LIVRE:\n" + texto_avaliacao(estudante, av))
            except Exception:
                textos_avaliacoes.append("AVALIAÇÃO COMPLEMENTAR / DOCUMENTO LIVRE:\n" + str(av))

    texto_avaliacao_ctx = "\n\n---\n\n".join(textos_avaliacoes) if textos_avaliacoes else "Nenhuma avaliação pedagógica selecionada ou registrada."
    texto_entrevista = str(entrevista or "Nenhuma entrevista com família localizada.")
    texto_atendimentos = "\n".join([str(a) for a in (atendimentos or [])[:12]]) or "Nenhum atendimento registrado."

    prompt = f"""
Você é um especialista em Atendimento Educacional Especializado (AEE), educação inclusiva e elaboração de Estudo de Caso no padrão usado por GRE/Secretarias de Educação.

TAREFA:
Gerar uma proposta de ESTUDO DE CASO padronizado, com base nos registros disponíveis do sistema.

REGRAS OBRIGATÓRIAS:
- Não inventar dados pessoais, nomes, diagnósticos, laudos, CPF, RG, endereço ou telefone.
- Não usar linguagem médica como conclusão diagnóstica.
- Usar linguagem pedagógica, institucional, objetiva e acolhedora.
- Se alguma informação não estiver disponível, deixar o campo com orientação de preenchimento posterior.
- Usar estudo de caso anterior, se houver.
- Se não houver estudo anterior, usar avaliações pedagógicas e demais registros como base.
- Se algo não couber nos campos padronizados, colocar em OBSERVAÇÕES_COMPLEMENTARES.
- Gerar texto para revisão do professor do AEE antes de salvar.

ESTUDANTE:
Código interno: {estudante[1]}
Ano/Série atual: {estudante[2]}
Turma: {estudante[3]}
Perfil educacional cadastrado: {estudante[4]}
Ano letivo do novo estudo: {ano_novo}

ESTUDO DE CASO ANTERIOR:
{texto_estudo_anterior}

AVALIAÇÕES PEDAGÓGICAS / DOCUMENTOS LIVRES:
{texto_avaliacao_ctx}

ENTREVISTA COM A FAMÍLIA:
{texto_entrevista}

ATENDIMENTOS REGISTRADOS:
{texto_atendimentos}

Responda exatamente com as seções abaixo, mantendo os marcadores:

[ANALISE_COMPARATIVA]
Comparar registros anteriores e atuais, destacando evolução, permanências, barreiras, avanços e pontos de atenção. Se não houver estudo anterior, explicar que a proposta foi criada a partir de avaliações pedagógicas e demais registros.

[PERCURSO_EDUCACIONAL]
Texto para o campo “Estudo de caso / percurso educacional”, relatando trajetória escolar, situação inicial, estratégias utilizadas e progressos alcançados.

[MOTIVO_ENCAMINHAMENTO]
Motivo pedagógico do encaminhamento ao AEE, sem diagnóstico clínico conclusivo.

[HABILIDADES_OBSERVADAS]
Lista ou parágrafo com habilidades observadas e desenvolvidas.

[HABILIDADES_A_DESENVOLVER]
Lista ou parágrafo com habilidades que precisam ser desenvolvidas.

[POTENCIALIDADES]
Potencialidades pedagógicas, interesses, formas de participação e respostas positivas observadas.

[DIFICULDADES]
Dificuldades e barreiras pedagógicas observadas.

[ESTRATEGIAS]
Estratégias pedagógicas recomendadas.

[INTERVENCOES]
Intervenções, recursos de acessibilidade, recursos pedagógicos e encaminhamentos sugeridos.

[AVALIACAO]
Texto avaliativo sobre avanços, habilidades a reforçar e ações para eliminação de barreiras.

[CONSIDERACOES]
Síntese pedagógica final para compor o estudo de caso.

[OBSERVACOES_COMPLEMENTARES]
Informações relevantes que não se encaixaram nos campos anteriores.
"""
    try:
        resposta = client.responses.create(model="gpt-4.1-mini", input=prompt)
        texto = resposta.output_text or ""
        analise = extrair_secao_ia(texto, "ANALISE_COMPARATIVA", texto)
        return analise.strip(), texto.strip()
    except Exception as e:
        return "", f"Não foi possível gerar o novo estudo de caso com IA agora. Erro: {e}"

def gerar_nova_avaliacao_pedagogica_com_ia(estudante, avaliacao_anterior=None, estudo_caso=None, entrevista=None, atendimentos=None, ano_novo=""):
    """Gera análise comparativa e sugestão de nova Avaliação Pedagógica com IA.
    A resposta é apoio pedagógico: não substitui a avaliação do professor do AEE.
    """
    client = obter_cliente_openai()
    if client is None:
        return "", "IA não configurada. Configure OPENAI_API_KEY para usar esta função."

    texto_avaliacao_anterior = texto_avaliacao(estudante, avaliacao_anterior) if avaliacao_anterior else "Nenhuma avaliação anterior selecionada."
    texto_estudo = texto_estudo_caso(estudante, estudo_caso) if estudo_caso else "Nenhum estudo de caso localizado."
    texto_entrevista = str(entrevista or "Nenhuma entrevista com família localizada.")
    texto_atendimentos = "\n".join([str(a) for a in (atendimentos or [])[:12]]) or "Nenhum atendimento registrado."

    prompt = f"""
Você é um especialista em Atendimento Educacional Especializado (AEE) e em avaliação pedagógica inclusiva.
Analise os registros abaixo e produza uma resposta institucional, pedagógica e objetiva para apoiar o professor do AEE na construção de uma nova Avaliação Pedagógica.

REGRAS IMPORTANTES:
- Não invente dados pessoais nem diagnósticos.
- Não substitua a decisão pedagógica do professor.
- Trabalhe apenas com os dados fornecidos.
- Use linguagem adequada para documento escolar.
- Foque em barreiras, potencialidades, comunicação, interação social, autonomia, aprendizagem e recomendações pedagógicas.
- Sugira recursos acessíveis, atividades pedagógicas, recursos maker/3D quando coerente e atividades plugadas/desplugadas.
- Evite termos clínicos conclusivos; mantenha foco educacional.

ESTUDANTE:
Código interno: {estudante[1]}
Ano/Série atual: {estudante[2]}
Turma: {estudante[3]}
Perfil educacional: {estudante[4]}
Ano letivo da nova avaliação: {ano_novo}

AVALIAÇÃO PEDAGÓGICA ANTERIOR:
{texto_avaliacao_anterior}

ESTUDO DE CASO DISPONÍVEL:
{texto_estudo}

ENTREVISTA COM A FAMÍLIA:
{texto_entrevista}

ATENDIMENTOS REGISTRADOS:
{texto_atendimentos}

Responda exatamente com duas seções:

[ANALISE_COMPARATIVA]
Texto comparando a avaliação anterior com os registros atuais, destacando avanços, permanências, novas barreiras, potencialidades e pontos de atenção.

[SUGESTAO_NOVA_AVALIACAO]
Proposta inicial de nova Avaliação Pedagógica, organizada nos tópicos: barreiras enfrentadas, potencialidades e habilidades, comunicação, interação social, autonomia, aprendizagem, resumo pedagógico e recomendações para intervenção no AEE.
"""
    try:
        resposta = client.responses.create(model="gpt-4.1-mini", input=prompt)
        texto = resposta.output_text or ""
        analise = texto
        sugestao = texto
        if "[SUGESTAO_NOVA_AVALIACAO]" in texto:
            partes = texto.split("[SUGESTAO_NOVA_AVALIACAO]", 1)
            analise = partes[0].replace("[ANALISE_COMPARATIVA]", "").strip()
            sugestao = partes[1].strip()
        return analise.strip(), sugestao.strip()
    except Exception as e:
        return "", f"Não foi possível gerar a nova avaliação pedagógica com IA agora. Erro: {e}"


# ======================================================
# BASE DE CONHECIMENTO IA - PDFs + CHROMADB
# ======================================================
def listar_pdfs_base(pasta):
    """Lista arquivos PDF de uma pasta da base de conhecimento."""
    pasta = Path(pasta)
    pasta.mkdir(parents=True, exist_ok=True)
    return sorted(list(pasta.glob("*.pdf")))


def listar_pdfs_todas_bases():
    """Retorna um resumo dos PDFs cadastrados em todas as bases IA."""
    resumo = []
    for nome_base, pasta in PASTAS_BASE_IA.items():
        for pdf in listar_pdfs_base(pasta):
            resumo.append({"base": nome_base, "categoria": OPCOES_BASES_CONHECIMENTO_IA.get(nome_base, nome_base), "arquivo": pdf.name})
    return resumo


def sincronizar_pastas_base_conhecimento():
    """Garante que as pastas principais e as pastas criadas manualmente sejam reconhecidas.

    Pastas reconhecidas nesta versão:
    - cientifica
    - pedagogica
    - Lei
    - Tecnologia
    - bncc
    - bncc_computacao
    - estrategias
    - tecnologia_assistiva
    - maker_inclusivo
    - legislacao
    - projetos_interdisciplinares
    - perfis_pedagogicos
    """
    BASE_CONHECIMENTO_DIR.mkdir(parents=True, exist_ok=True)
    for pasta in PASTAS_BASE_IA.values():
        Path(pasta).mkdir(parents=True, exist_ok=True)


def salvar_pdf_base_conhecimento(uploaded_file, nome_base):
    """Salva PDF enviado pelo usuário na pasta correta da Base de Conhecimento IA."""
    if nome_base not in PASTAS_BASE_IA:
        raise ValueError(f"Base inválida: {nome_base}")
    pasta = PASTAS_BASE_IA[nome_base]
    pasta.mkdir(parents=True, exist_ok=True)
    nome_seguro = re.sub(r"[^A-Za-z0-9_.\-À-ÿ ]", "_", uploaded_file.name).strip()
    caminho = pasta / nome_seguro
    with open(caminho, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return caminho


def bases_para_componente(componente_destino):
    """Define quais bases a IA deve consultar conforme o componente/área selecionado."""
    componente = str(componente_destino or "").lower()
    if "geral" in componente:
        return []
    bases = ["pedagogica", "estrategias", "bncc"]
    if "comput" in componente or "tecnologia" in componente:
        bases.append("bncc_computacao")
    if any(x in componente for x in ["matem", "física", "fisica", "química", "quimica", "biologia", "ciências", "ciencias"]):
        bases.extend(["bncc_computacao", "maker_inclusivo", "projetos_interdisciplinares"])
    if any(x in componente for x in ["português", "portugues", "história", "historia", "geografia", "artes", "inglês", "ingles", "espanhol", "filosofia", "sociologia", "projeto de vida"]):
        bases.extend(["bncc_computacao", "estrategias"])
    # Bases complementares criadas manualmente no projeto.
    # Elas ficam disponíveis para consulta quando houver PDFs cadastrados, sem substituir as bases principais.
    if "tecnologia" in PASTAS_BASE_IA:
        bases.append("tecnologia")
    if "lei" in PASTAS_BASE_IA:
        bases.append("lei")
    # Remove duplicadas preservando ordem
    return list(dict.fromkeys([b for b in bases if b in PASTAS_BASE_IA]))


def extrair_texto_pdf(caminho_pdf):
    """Extrai texto de um PDF usando pypdf."""
    from pypdf import PdfReader

    partes = []
    reader = PdfReader(str(caminho_pdf))

    for i, pagina in enumerate(reader.pages):
        texto = pagina.extract_text() or ""
        if texto.strip():
            partes.append(f"\n--- Arquivo: {caminho_pdf.name} | Página {i + 1} ---\n{texto}")

    return "\n".join(partes).strip()


def quebrar_texto(texto, tamanho=1200, sobreposicao=200):
    """Divide textos longos em partes menores para busca semântica."""
    if not texto:
        return []

    partes = []
    inicio = 0
    while inicio < len(texto):
        fim = inicio + tamanho
        parte = texto[inicio:fim].strip()
        if parte:
            partes.append(parte)
        inicio = max(fim - sobreposicao, fim)
    return partes


def criar_embedding(textos):
    """Cria embeddings com OpenAI."""
    client = obter_cliente_openai()
    if client is None:
        raise RuntimeError("OPENAI_API_KEY não configurada ou biblioteca OpenAI indisponível.")

    resposta = client.embeddings.create(
        model="text-embedding-3-small",
        input=textos,
    )
    return [item.embedding for item in resposta.data]


def obter_chroma_collection(nome_base):
    """Abre/cria uma coleção ChromaDB para a base informada."""
    import chromadb

    cliente = chromadb.PersistentClient(path=str(CHROMA_DIR))
    return cliente.get_or_create_collection(name=f"incluisrm_{nome_base}")


def indexar_base_conhecimento(nome_base):
    """Indexa os PDFs da base de conhecimento informada."""
    if nome_base not in PASTAS_BASE_IA:
        raise ValueError(f"Base inválida: {nome_base}")

    pasta = PASTAS_BASE_IA[nome_base]
    arquivos = listar_pdfs_base(pasta)
    if not arquivos:
        return 0, f"Nenhum PDF encontrado em {pasta}."

    collection = obter_chroma_collection(nome_base)
    total_chunks = 0

    for pdf in arquivos:
        texto = extrair_texto_pdf(pdf)
        if not texto.strip():
            continue

        chunks = quebrar_texto(texto)
        if not chunks:
            continue

        ids = [f"{nome_base}_{pdf.stem}_{idx}" for idx in range(len(chunks))]
        metadados = [
            {"arquivo": pdf.name, "base": nome_base, "chunk": idx}
            for idx in range(len(chunks))
        ]
        embeddings = criar_embedding(chunks)

        collection.upsert(
            ids=ids,
            documents=chunks,
            embeddings=embeddings,
            metadatas=metadados,
        )
        total_chunks += len(chunks)

    return total_chunks, f"Base {nome_base} indexada com sucesso."


def buscar_na_base_conhecimento(pergunta, bases=None, limite=5):
    """Busca trechos relevantes nas bases indexadas."""
    if bases is None:
        bases = ["cientifica", "pedagogica"]

    try:
        embedding = criar_embedding([pergunta])[0]
    except Exception:
        return []

    resultados_finais = []

    for base in bases:
        try:
            collection = obter_chroma_collection(base)
            resultado = collection.query(
                query_embeddings=[embedding],
                n_results=limite,
            )

            docs = resultado.get("documents", [[]])[0]
            metas = resultado.get("metadatas", [[]])[0]

            for doc, meta in zip(docs, metas):
                resultados_finais.append(
                    {
                        "texto": doc,
                        "arquivo": meta.get("arquivo", "Arquivo não identificado"),
                        "base": meta.get("base", base),
                    }
                )
        except Exception:
            continue

    return resultados_finais


def montar_contexto_base(resultados):
    """Monta contexto textual para enviar à IA."""
    if not resultados:
        return "Nenhum trecho relevante encontrado nas bases científica e pedagógica."

    partes = []
    for i, item in enumerate(resultados, start=1):
        partes.append(
            f"""
[Fonte {i}]
Base: {item.get('base', 'não informada')}
Arquivo: {item.get('arquivo', 'não informado')}
Trecho:
{item.get('texto', '')}
""".strip()
        )
    return "\n\n".join(partes)


def arquivos_consultados_texto(resultados):
    arquivos = []
    for item in resultados or []:
        nome = item.get("arquivo")
        if nome and nome not in arquivos:
            arquivos.append(nome)
    return ", ".join(arquivos) if arquivos else "Nenhum arquivo consultado ou base ainda não indexada."


# ======================================================
# ATENDIMENTOS EM TEXTO
# ======================================================
def listar_atendimentos_texto(estudante_id):
    atendimentos = listar_atendimentos(estudante_id)
    if not atendimentos:
        return "Nenhum atendimento registrado."
    partes = []
    for a in atendimentos[:10]:
        partes.append(
            f"""
Data: {a[0]}
Objetivo: {a[1] or 'Não informado.'}
Atividade: {a[2] or 'Não informado.'}
Resposta: {a[3] or 'Não informado.'}
Avanços: {a[4] or 'Não informado.'}
Dificuldades: {a[5] or 'Não informado.'}
Evolução: {a[6] or 'Não informado.'}
Resposta: {a[8] if len(a) > 8 else 'Não informado.'}/10
Avanço: {a[9] if len(a) > 9 else 'Não informado.'}/10
Dificuldade: {a[10] if len(a) > 10 else 'Não informado.'}/10
Engajamento: {a[11] if len(a) > 11 else 'Não informado.'}/10
Índice: {a[12] if len(a) > 12 else 'Não informado.'}/10
"""
        )
    return "\n---\n".join(partes)


# ======================================================
# BUSCA DE MODELOS 3D - RECURSOS PEDAGÓGICOS
# ======================================================
from urllib.parse import quote_plus


def formatar_termo_busca_3d(termo):
    """Formata o termo para uso seguro em URLs de busca."""
    return quote_plus(str(termo or "").strip())


def link_busca_thingiverse(termo):
    """Gera link de busca no Thingiverse."""
    termo_formatado = formatar_termo_busca_3d(termo)
    return f"https://www.thingiverse.com/search?q={termo_formatado}&type=things"


def link_busca_printables(termo):
    """Gera link de busca no Printables."""
    termo_formatado = formatar_termo_busca_3d(termo)
    return f"https://www.printables.com/search/models?q={termo_formatado}"


def link_busca_makerworld(termo):
    """Gera link de busca no MakerWorld."""
    termo_formatado = formatar_termo_busca_3d(termo)
    return f"https://makerworld.com/pt/search/models?keyword={termo_formatado}"


def gerar_termos_3d_com_ia(conteudo_paee):
    """
    Gera exatamente 5 termos de busca para modelos 3D pedagógicos.
    Os termos são preferencialmente em inglês porque retornam mais resultados
    nas plataformas de modelos 3D.
    """
    client = obter_cliente_openai()

    termos_padrao = [
        "tactile alphabet",
        "visual schedule",
        "fine motor skills",
        "communication cards",
        "sensory shapes",
    ]

    if client is None or not conteudo_paee:
        return termos_padrao

    prompt = f"""
Analise o plano AEE abaixo e gere exatamente 5 termos curtos em inglês para busca de modelos 3D pedagógicos.

Priorize recursos:
- inclusivos;
- táteis;
- manipuláveis;
- visuais;
- para alfabetização;
- coordenação motora fina;
- CAA;
- rotina visual;
- autonomia;
- cultura maker;
- robótica educacional simples.

Não cite diagnóstico clínico.
Não use nome do estudante.
Não explique.
Retorne apenas 5 linhas.
Cada linha deve conter apenas um termo de busca.

PLANO AEE:
{conteudo_paee}
"""

    try:
        resposta = client.responses.create(
            model="gpt-4.1-mini",
            input=prompt,
        )
        texto = resposta.output_text

        termos = []
        for linha in texto.split("\n"):
            termo = linha.strip().strip("-•0123456789. )(").strip()
            if termo and termo not in termos:
                termos.append(termo)

        return termos[:5] if termos else termos_padrao

    except Exception:
        return termos_padrao


def gerar_lista_modelos_3d(conteudo_paee):
    """
    Retorna uma lista estruturada com termo e links.
    Essa lista é usada para exibir cards e botões clicáveis na tela do Streamlit.
    """
    termos = gerar_termos_3d_com_ia(conteudo_paee)

    sugestoes = []
    for termo in termos:
        sugestoes.append(
            {
                "termo": termo,
                "makerworld": link_busca_makerworld(termo),
                "printables": link_busca_printables(termo),
                "thingiverse": link_busca_thingiverse(termo),
            }
        )

    return sugestoes


def gerar_texto_modelos_3d(sugestoes):
    """
    Gera a versão textual das sugestões para exportação em TXT, PDF e Word.
    No relatório exportado ficam os links completos, enquanto na tela aparecem botões.
    """
    linhas = []
    linhas.append("SUGESTÕES DE MODELOS 3D PARA APOIO PEDAGÓGICO")
    linhas.append("")
    linhas.append(
        "Observação: os links abaixo são sugestões de busca em plataformas públicas de modelos 3D. "
        "O professor deverá avaliar a adequação pedagógica, segurança, tamanho da peça, faixa etária, "
        "necessidade de adaptação e viabilidade de impressão antes do uso com o estudante."
    )
    linhas.append("")

    for item in sugestoes:
        termo = item.get("termo", "Recurso pedagógico")
        linhas.append(f"Recurso sugerido: {termo}")
        linhas.append(f"- MakerWorld: {item.get('makerworld', '')}")
        linhas.append(f"- Printables: {item.get('printables', '')}")
        linhas.append(f"- Thingiverse: {item.get('thingiverse', '')}")
        linhas.append("")

    return "\n".join(linhas)


def gerar_links_modelos_3d(conteudo_paee):
    """
    Mantém compatibilidade com versões anteriores do sistema.
    Retorna diretamente o texto com os links completos.
    """
    return gerar_texto_modelos_3d(gerar_lista_modelos_3d(conteudo_paee))


def bloco_sugestoes_3d_streamlit(conteudo_base, estudante_id, estudante_codigo, origem="sugestao", nome_extra=""):
    """
    Bloco visual reutilizável para Streamlit.
    Na tela: exibe cards com 3 botões clicáveis por sugestão.
    Nos arquivos exportados: mantém links completos em texto.
    """
    st.markdown("### 🧩 Sugestões de modelos 3D para apoio pedagógico")
    st.caption(
        "Gera sugestões de busca em plataformas públicas de modelos 3D. "
        "Na tela, os links aparecem como botões; nos relatórios exportados, aparecem como links completos."
    )

    chave_base = f"sugestoes_3d_{origem}_{estudante_id}{nome_extra}"

    if st.button("🔎 Gerar sugestões de modelos 3D", key=f"gerar_{chave_base}"):
        st.session_state[chave_base] = gerar_lista_modelos_3d(conteudo_base)

    if chave_base in st.session_state:
        sugestoes = st.session_state[chave_base]

        st.info(
            "Os links abaixo são sugestões de busca. Avalie segurança, tamanho, faixa etária, "
            "adequação pedagógica, necessidade de adaptação e viabilidade de impressão antes do uso."
        )

        for i, item in enumerate(sugestoes, start=1):
            termo = item.get("termo", "Recurso pedagógico")

            with st.container(border=True):
                st.markdown(f"**Recurso sugerido {i}:** {termo}")

                col1, col2, col3 = st.columns(3)

                with col1:
                    st.link_button("🌐 MakerWorld", item.get("makerworld", "#"))

                with col2:
                    st.link_button("🧩 Printables", item.get("printables", "#"))

                with col3:
                    st.link_button("🔎 Thingiverse", item.get("thingiverse", "#"))

        texto_exportacao = gerar_texto_modelos_3d(sugestoes)

        with st.expander("📄 Visualizar texto que será exportado em PDF/Word/TXT"):
            st.text_area(
                "Texto com links completos",
                texto_exportacao,
                height=340,
                key=f"text_area_{chave_base}",
            )

        export_buttons(
            texto_exportacao,
            f"Links_Modelos_3D_{origem}_{estudante_codigo}{nome_extra}",
            tipo_pdf="documento",
        )

# ======================================================
# PAEE SEM IA / COM IA
# ======================================================
def gerar_paee_sem_ia(estudante, avaliacao=None, entrevista=None, estudo=None):
    codigo = estudante[1]
    ano_serie = estudante[2] or "Não informado."
    turma = estudante[3] or "Não informado."
    perfil = estudante[4] or "Não informado."
    observacoes = estudante[5] or "Não informado."

    if avaliacao:
        avaliacao_txt = texto_avaliacao(estudante, ("", *avaliacao))
    else:
        avaliacao_txt = "Nenhuma avaliação pedagógica registrada."

    if entrevista:
        entrevista_txt = texto_entrevista(estudante, ("", *entrevista))
    else:
        entrevista_txt = "Nenhuma entrevista com a família registrada."

    if estudo:
        estudo_txt = texto_estudo_caso(estudante, ("", *estudo))
    else:
        estudo_txt = "Nenhum estudo de caso GRE registrado."

    historico_txt = listar_atendimentos_texto(estudante[0])

    return f"""
PLANO AEE / PAEE - SUGESTÃO PEDAGÓGICA SEM IA

1. Identificação segura do estudante
Código interno: {codigo}
Ano/Série: {ano_serie}
Turma: {turma}
Perfil educacional informado: {perfil}

Campos sensíveis para preenchimento manual no Word/PDF:
Nome completo do(a) estudante: ___________________________________________
CPF/RG do(a) estudante: _________________________________________________
Data de nascimento: ____/____/________
Nome do responsável: _________________________________________________
CPF/RG do responsável: _______________________________________________
Telefone: ___________________________________________________________
Endereço: ___________________________________________________________

2. Caracterização pedagógica inicial
{observacoes}

3. Base documental utilizada
3.1 Avaliação pedagógica
{avaliacao_txt}

3.2 Entrevista com a família
{entrevista_txt}

3.3 Estudo de caso GRE
{estudo_txt}

4. Barreiras e necessidades educacionais específicas
- Organizar as informações da entrevista, avaliação e estudo de caso para identificar barreiras de comunicação, interação, autonomia, aprendizagem, acesso curricular e participação.
- Evitar qualquer inferência diagnóstica não registrada.
- Considerar recursos de acessibilidade, tecnologia assistiva e adaptações pedagógicas conforme resposta do estudante.

5. Objetivos gerais do AEE
- Ampliar as condições de acesso, participação e aprendizagem do estudante nas atividades escolares.
- Desenvolver estratégias que favoreçam comunicação, autonomia, interação social e organização da rotina escolar.
- Utilizar recursos pedagógicos acessíveis, materiais concretos, tecnologias educacionais inclusivas e propostas maker quando pertinentes.

6. Objetivos específicos
- Fortalecer habilidades prioritárias observadas no estudo de caso.
- Reduzir barreiras identificadas na avaliação pedagógica e nos registros de atendimento.
- Apoiar o uso de recursos visuais, táteis, manipuláveis, digitais e de comunicação alternativa quando necessário.
- Acompanhar a evolução do estudante por meio de registros sistemáticos.

7. Estratégias pedagógicas sugeridas
- Utilizar rotina visual, instruções objetivas e antecipação das atividades.
- Propor atividades com materiais concretos, jogos pedagógicos, recursos visuais, recursos táteis e tecnologias digitais.
- Articular as ações do AEE com professores do ensino comum e família.
- Registrar avanços, dificuldades, engajamento, autonomia e resposta às estratégias em cada atendimento.

8. Sugestões de tecnologias educacionais inclusivas
- Impressão 3D: letras, números, mapas táteis, peças de associação, jogos de encaixe, pranchas adaptadas e recursos manipuláveis.
- Robótica educacional: sequências lógicas, comandos simples, causa e efeito, organização espacial e resolução de problemas.
- Jogos digitais: atividades de atenção, memória, comunicação, raciocínio lógico e reforço de habilidades curriculares.
- Recursos maker: construção de materiais personalizados, placas, cartões, roletas, jogos de pareamento e recursos sensoriais.
- Comunicação alternativa e aumentativa: pranchas visuais, cartões de escolha, pictogramas e rotinas estruturadas.

9. Atividades plugadas sugeridas
- Jogos digitais educativos adaptados ao perfil do estudante.
- Uso orientado de tablet/computador com recursos visuais e auditivos.
- Atividades com Scratch, Construct ou jogos simples para trabalhar sequência, escolha, atenção e resposta.
- Robótica com comandos básicos, sensores simples e atividades de causa e efeito.

10. Atividades desplugadas sugeridas
- Sequências com cartões de rotina e organização temporal.
- Pareamento de figuras, objetos, letras, números, formas e cores.
- Trilhas pedagógicas, jogos de memória, dominós adaptados e bingo de imagens/palavras.
- Atividades sensoriais com texturas, objetos 3D, encaixes, blocos, massinha, peças táteis e materiais concretos.
- Jogos de comandos corporais, lateralidade, atenção compartilhada e imitação.
- Contação de histórias com cartões visuais, personagens concretos e recontagem orientada.

11. Organização do atendimento
- Definir frequência, duração, formato e objetivos de cada ciclo de atendimento conforme necessidade pedagógica.
- Registrar cada atendimento no sistema para análise de evolução.

12. Avaliação e acompanhamento
Histórico de atendimentos registrado:
{historico_txt}

- Revisar o plano periodicamente com base nos registros reais.
- Se os registros ainda forem limitados, ampliar o acompanhamento antes de conclusões evolutivas.

13. Responsável pelo AEE
Nome: ___________________________________________
Função: Professor(a) do Atendimento Educacional Especializado (AEE)
Assinatura: _______________________________________

14. Coordenação pedagógica
Nome: ___________________________________________
Cargo: Coordenação Pedagógica
Assinatura: _______________________________________

15. Data de elaboração: ____/____/________
""".strip()




def texto_perfil_professor_aee_para_ia():
    """Resume o perfil pedagógico/tecnológico do professor AEE atual para orientar a IA."""
    professores = listar_professores()
    if not professores:
        return (
            "Nenhum professor AEE cadastrado no sistema. As sugestões devem ser realistas, "
            "priorizando estratégias pedagógicas simples, desplugadas, acessíveis e de baixo custo."
        )

    p = professores[0]
    areas = p[9] if len(p) > 9 and p[9] else "Não informado"
    nivel = p[10] if len(p) > 10 and p[10] else "Básico"
    maker = p[11] if len(p) > 11 and p[11] else "Não"
    formacao = p[12] if len(p) > 12 and p[12] else "Não informado"
    projetos = p[13] if len(p) > 13 and p[13] else "Não informado"
    preferencias = p[14] if len(p) > 14 and p[14] else "Não informado"
    recursos_professor = p[15] if len(p) > 15 and p[15] else "Não informado"
    recursos_professor_uso = p[16] if len(p) > 16 and p[16] else "Não informado"
    recursos_professor_obs = p[17] if len(p) > 17 and p[17] else "Não informado"

    return f"""
PERFIL DO PROFESSOR AEE RESPONSÁVEL:
Nome/referência: {p[1] or 'Não informado'}
Escola: {p[2] or 'Não informado'}
Áreas de atuação/interesse: {areas}
Nível de familiaridade tecnológica: {nivel}
Modo maker inclusivo ativado: {maker}
Interesse em formação maker/tecnologias educacionais: {formacao}
Projetos/temas de interesse: {projetos}
Preferências metodológicas: {preferencias}
Recursos pedagógicos/TA que o professor possui ou pode levar: {recursos_professor}
Recursos que o professor pretende utilizar nos atendimentos: {recursos_professor_uso}
Observações sobre uso dos recursos do professor: {recursos_professor_obs}

ORIENTAÇÃO PARA A IA:
- Respeitar o perfil do professor AEE.
- Considerar SEMPRE dois contextos de recursos quando existirem: (1) recursos institucionais cadastrados pela escola e (2) recursos pedagógicos/tecnologias assistivas que o professor possui ou pretende utilizar.
- Priorizar os recursos institucionais da escola quando estiverem disponíveis e forem coerentes com o perfil do estudante.
- Complementar as sugestões com os recursos do professor quando forem adequados aos objetivos pedagógicos, às barreiras identificadas e às necessidades do estudante.
- Quando usar recursos do professor, deixar claro que são recursos de apoio disponibilizados pelo professor do AEE, e não patrimônio institucional da escola.
- Não sugerir robótica, impressão 3D, CNC, programação ou tecnologias complexas se o modo maker estiver desativado ou se o nível tecnológico for básico, salvo se o texto indicar interesse em aprender e a atividade for introdutória.
- Priorizar atividades desplugadas, recursos visuais, materiais manipuláveis, organização de rotina e estratégias de acessibilidade quando o perfil tecnológico for básico.
- Quando o modo maker estiver ativado, sugerir projetos e tecnologias de forma progressiva, segura, com etapas simples e sempre vinculadas às necessidades do estudante.
- Não inventar recursos não cadastrados nem recursos que o professor não informou possuir ou pretender utilizar.
- Se não houver recursos da escola nem recursos do professor, indicar alternativas pedagógicas simples, desplugadas, de baixo custo e possíveis de confeccionar.
""".strip()


def montar_contexto_projeto_norteador_para_ia(usar_projeto=False, projeto_norteador="", nivel_projeto="", observacoes_projeto=""):
    """Monta orientações para uso de projeto interdisciplinar no AEE sem transformar o atendimento em reforço escolar."""
    if not usar_projeto:
        return (
            "Projeto norteador não ativado neste plano. Caso sugira atividades, priorize objetivos funcionais do AEE, "
            "atividades desplugadas, recursos acessíveis e mediações pedagógicas alinhadas ao perfil do estudante."
        )

    return f"""
PROJETO NORTEADOR DO ATENDIMENTO ATIVADO:
Tema/projeto informado pelo professor: {projeto_norteador or 'Projeto não especificado'}
Nível desejado: {nivel_projeto or 'Não informado'}
Observações do professor: {observacoes_projeto or 'Não informado'}

REGRAS OBRIGATÓRIAS:
- O AEE NÃO é reforço escolar. O projeto deve ser usado como meio pedagógico para acessibilidade, participação, autonomia, organização cognitiva, comunicação, permanência, resolução de problemas e protagonismo.
- As áreas do Ensino Médio devem aparecer de forma contextualizada, interdisciplinar e significativa, sem substituir o professor da disciplina.
- Incluir atividades desplugadas e de baixo custo, especialmente quando houver comorbidades, fadiga, dificuldades motoras, sensoriais ou executivas.
- Se o projeto envolver maker/robótica/impressão 3D, dividir em etapas simples, seguras e progressivas.
- Sempre adaptar o projeto ao perfil do estudante, aos recursos da escola e ao perfil tecnológico do professor AEE.
""".strip()


def montar_contexto_plano_aee_ia(estudante, avaliacao=None, entrevista=None, estudo=None, plano_manual=None):
    """Reúne os dados pedagógicos disponíveis para o módulo Plano AEE - IA.

    Nesta versão, o Plano AEE - IA passa a considerar também:
    - Escutas docentes salvas;
    - Histórico de escutas docentes;
    - Relatórios de Apoio ao Docente gerados e salvos.

    Observação: a finalidade continua sendo exclusivamente pedagógica e educacional.
    """
    historico_txt = listar_atendimentos_texto(estudante[0])
    escutas_docentes = listar_escutas_docentes(estudante[0])[:10]
    relatorios_docentes = listar_relatorios_docente(estudante[0])[:5]

    estudante_txt = f"""
Código interno: {estudante[1]}
Ano/Série: {estudante[2]}
Turma: {estudante[3]}
Turno: {estudante[6]}
Perfil educacional informado: {estudante[4]}
Observações pedagógicas iniciais: {estudante[5]}
Dias de atendimento preferenciais: {estudante[7]}
Horário preferencial: {estudante[8]}
""".strip()

    avaliacao_txt = texto_avaliacao(estudante, ("", *avaliacao)) if avaliacao else "Nenhuma avaliação pedagógica registrada."
    entrevista_txt = texto_entrevista(estudante, ("", *entrevista)) if entrevista else "Nenhuma entrevista com a família registrada."
    estudo_txt = texto_estudo_caso(estudante, ("", *estudo)) if estudo else "Nenhum estudo de caso GRE registrado."
    plano_txt = texto_plano_aee(estudante, ("", *plano_manual)) if plano_manual else "Nenhum plano AEE manual registrado."

    # Recursos reais cadastrados pela escola/unidade.
    # A IA deve priorizar esses recursos antes de sugerir materiais externos.
    recursos_escola_txt = listar_recursos_escola_texto(limite=60)
    perfil_professor_txt = texto_perfil_professor_aee_para_ia()

    textos_escutas = []
    for esc in escutas_docentes:
        try:
            textos_escutas.append(texto_escuta_docente(estudante, esc))
        except Exception:
            textos_escutas.append(str(esc))
    escutas_txt = "\n\n---\n\n".join(textos_escutas) if textos_escutas else "Nenhuma escuta docente registrada para este estudante."

    textos_relatorios = []
    for rel in relatorios_docentes:
        try:
            textos_relatorios.append(texto_relatorio_docente(estudante, rel))
        except Exception:
            textos_relatorios.append(str(rel))
    relatorios_docente_txt = "\n\n---\n\n".join(textos_relatorios) if textos_relatorios else "Nenhum Relatório de Apoio ao Docente salvo para este estudante."

    pergunta_busca = f"""
AEE plano mensal atendimento sala de recursos multifuncionais estudante {estudante[4]} {estudante[2]}
comunicação funcional autonomia CAA tecnologia assistiva recursos visuais rotina estruturada robótica impressão 3D.
"""
    resultados_base = buscar_na_base_conhecimento(pergunta_busca, bases=["cientifica", "pedagogica"], limite=5)
    contexto_base = montar_contexto_base(resultados_base)
    arquivos_base = arquivos_consultados_texto(resultados_base)

    return {
        "estudante_txt": estudante_txt,
        "avaliacao_txt": avaliacao_txt,
        "entrevista_txt": entrevista_txt,
        "estudo_txt": estudo_txt,
        "plano_txt": plano_txt,
        "recursos_escola_txt": recursos_escola_txt,
        "perfil_professor_txt": perfil_professor_txt,
        "escutas_docentes_txt": escutas_txt,
        "relatorios_docente_txt": relatorios_docente_txt,
        "qtd_escutas_docentes": len(escutas_docentes),
        "qtd_relatorios_docente": len(relatorios_docentes),
        "historico_txt": historico_txt,
        "contexto_base": contexto_base,
        "arquivos_base": arquivos_base,
    }


def gerar_diagnostico_aee_ia(estudante, avaliacao=None, entrevista=None, estudo=None, plano_manual=None):
    """Gera diagnóstico pedagógico inicial/evolutivo para apoiar o professor do AEE."""
    ctx = montar_contexto_plano_aee_ia(estudante, avaliacao, entrevista, estudo, plano_manual)
    client = obter_cliente_openai()

    fallback = f"""
DIAGNÓSTICO PEDAGÓGICO INICIAL - PLANO AEE IA

Código interno: {estudante[1]}
Perfil educacional informado: {estudante[4]}
Ano/Série: {estudante[2]}

Síntese inicial:
O diagnóstico pedagógico deve considerar os registros já disponíveis no cadastro, entrevista familiar, avaliação pedagógica, estudo de caso, escuta docente, relatórios de apoio ao docente, plano AEE e atendimentos. Caso ainda existam poucos registros, recomenda-se utilizar este documento como roteiro de observação inicial, sem conclusões definitivas sobre evolução.

Focos de observação prioritários:
- Comunicação funcional e formas de expressão utilizadas pelo estudante.
- Autonomia na organização e realização das atividades.
- Atenção compartilhada e permanência nas propostas.
- Interação com professor, profissional de apoio e colegas.
- Resposta a recursos visuais, tecnológicos, concretos, sensoriais e maker.
- Barreiras comunicacionais, pedagógicas, atitudinais e de acessibilidade.

Encaminhamento:
Registrar os atendimentos semanalmente para que o sistema consiga gerar análises evolutivas mais consistentes.
""".strip()

    if client is None:
        return fallback

    prompt = f"""
Você é especialista em Atendimento Educacional Especializado (AEE), educação inclusiva, tecnologia assistiva, CAA, cultura maker e avaliação pedagógica funcional.

TAREFA:
Gere um DIAGNÓSTICO PEDAGÓGICO AEE com base nos dados disponíveis. O texto deve apoiar o professor do AEE na organização dos atendimentos e NÃO deve criar diagnóstico clínico.

REGRAS:
- Não usar nome real do estudante.
- Usar apenas o código interno.
- Não inventar dados.
- Se houver poucos atendimentos, informar que a análise evolutiva ainda é inicial.
- Diferenciar potencialidades, barreiras e necessidades prioritárias.
- Sugerir apenas encaminhamentos pedagógicos e funcionais.

DADOS DO ESTUDANTE:
{ctx['estudante_txt']}

ENTREVISTA FAMILIAR:
{ctx['entrevista_txt']}

AVALIAÇÃO PEDAGÓGICA:
{ctx['avaliacao_txt']}

ESTUDO DE CASO GRE:
{ctx['estudo_txt']}

PLANO AEE MANUAL:
{ctx['plano_txt']}

PERFIL DO PROFESSOR AEE:
{ctx['perfil_professor_txt']}

RECURSOS PEDAGÓGICOS E TECNOLOGIAS ASSISTIVAS (ESCOLA + PROFESSOR AEE):
{ctx['recursos_escola_txt']}

ESCUTA DOCENTE / HISTÓRICO DE ESCUTAS:
{ctx['escutas_docentes_txt']}

RELATÓRIOS DE APOIO AO DOCENTE SALVOS:
{ctx['relatorios_docente_txt']}

ATENDIMENTOS REGISTRADOS:
{ctx['historico_txt']}

BASES CONSULTADAS:
{ctx['contexto_base']}

ARQUIVOS CONSULTADOS:
{ctx['arquivos_base']}

ESTRUTURE EM:
1. Síntese pedagógica do estudante
2. Potencialidades observadas
3. Barreiras identificadas
4. Necessidades prioritárias para o AEE
5. Recursos com maior chance de resposta
6. Cuidados pedagógicos no atendimento
7. Indicadores a observar nos próximos atendimentos
8. Observação sobre suficiência dos dados
"""
    try:
        resposta = client.responses.create(model="gpt-4.1-mini", input=prompt)
        return resposta.output_text
    except Exception as e:
        return f"{fallback}\n\nObservação técnica: não foi possível gerar com IA agora. Erro: {e}"


def gerar_sugestao_geral_aee_ia(estudante, avaliacao=None, entrevista=None, estudo=None, plano_manual=None):
    """Gera uma sugestão geral de atendimento para orientar o semestre/período."""
    ctx = montar_contexto_plano_aee_ia(estudante, avaliacao, entrevista, estudo, plano_manual)
    client = obter_cliente_openai()

    fallback = f"""
SUGESTÃO GERAL DE ATENDIMENTO AEE - IA

Código interno: {estudante[1]}
Perfil educacional informado: {estudante[4]}

Objetivo geral sugerido:
Promover a participação, a autonomia, a comunicação funcional e o acesso às atividades escolares por meio de estratégias estruturadas, recursos de acessibilidade, tecnologias educacionais inclusivas e mediação pedagógica no Atendimento Educacional Especializado.

Eixos prioritários:
1. Comunicação funcional e uso progressivo de recursos visuais/CAA.
2. Organização da rotina e desenvolvimento de autonomia.
3. Atenção compartilhada e permanência nas atividades.
4. Participação social mediada.
5. Uso de recursos tecnológicos, concretos, sensoriais e maker conforme resposta do estudante.

Estratégias gerais:
- Utilizar comandos curtos, claros e objetivos.
- Trabalhar com rotina visual e previsibilidade.
- Oferecer escolhas mediadas por imagens, objetos ou recursos digitais.
- Registrar respostas, avanços, barreiras e nível de engajamento em cada atendimento.
- Iniciar com atividades de baixa complexidade e ampliar gradualmente.

Recursos possíveis:
- Pranchas de CAA, cartões visuais, tablet, Chromebook, jogos pedagógicos, materiais manipuláveis, impressão 3D, robótica educacional e atividades desplugadas.

Avaliação:
Acompanhar semanalmente a participação, comunicação, autonomia, atenção, interação e resposta aos recursos utilizados.
""".strip()

    if client is None:
        return fallback

    prompt = f"""
Você é especialista em AEE e planejamento de atendimento em Sala de Recursos Multifuncionais.

TAREFA:
Gere uma SUGESTÃO GERAL DE ATENDIMENTO AEE para orientar o professor no período letivo. Ela deve ser prática, aplicável e alimentar depois os registros de atendimento.

REGRAS:
- Não usar nome real.
- Usar código interno.
- Não inventar diagnósticos ou terapias.
- Não substituir o Plano AEE oficial; este é um apoio operacional de atendimento.
- Priorizar ações aplicáveis na SRM.

DADOS DO ESTUDANTE:
{ctx['estudante_txt']}

ENTREVISTA FAMILIAR:
{ctx['entrevista_txt']}

AVALIAÇÃO PEDAGÓGICA:
{ctx['avaliacao_txt']}

ESTUDO DE CASO GRE:
{ctx['estudo_txt']}

PLANO AEE MANUAL:
{ctx['plano_txt']}

PERFIL DO PROFESSOR AEE:
{ctx['perfil_professor_txt']}

RECURSOS PEDAGÓGICOS E TECNOLOGIAS ASSISTIVAS (ESCOLA + PROFESSOR AEE):
{ctx['recursos_escola_txt']}

ESCUTA DOCENTE / HISTÓRICO DE ESCUTAS:
{ctx['escutas_docentes_txt']}

RELATÓRIOS DE APOIO AO DOCENTE SALVOS:
{ctx['relatorios_docente_txt']}

ATENDIMENTOS REGISTRADOS:
{ctx['historico_txt']}

BASES CONSULTADAS:
{ctx['contexto_base']}

ESTRUTURE EM:
1. Objetivo geral de atendimento
2. Objetivos específicos
3. Eixos prioritários
4. Organização sugerida dos atendimentos
5. Recursos de acessibilidade e tecnologia educacional
6. Sugestões de atividades plugadas
7. Sugestões de atividades desplugadas
8. Sugestões com impressão 3D/robótica/cultura maker, quando aplicável
9. Como registrar evidências no módulo Atendimentos
10. Indicadores de acompanhamento
11. Cuidados para revisão do plano
"""
    try:
        resposta = client.responses.create(model="gpt-4.1-mini", input=prompt)
        return resposta.output_text
    except Exception as e:
        return f"{fallback}\n\nObservação técnica: não foi possível gerar com IA agora. Erro: {e}"


def gerar_plano_mensal_aee_ia(
    estudante,
    mes_referencia,
    ano_referencia,
    qtd_atendimentos_semana=1,
    avaliacao=None,
    entrevista=None,
    estudo=None,
    plano_manual=None,
    datas_atendimentos=None,
    usar_projeto_norteador=False,
    projeto_norteador="",
    nivel_projeto_norteador="",
    observacoes_projeto_norteador="",
):
    """Gera plano mensal aplicável às datas reais de atendimento do mês."""
    ctx = montar_contexto_plano_aee_ia(estudante, avaliacao, entrevista, estudo, plano_manual)
    client = obter_cliente_openai()
    qtd = max(1, min(5, int(qtd_atendimentos_semana or 1)))
    datas_atendimentos = datas_atendimentos or []
    datas_txt = datas_atendimentos_para_texto(datas_atendimentos)
    total_atendimentos = len(datas_atendimentos) if datas_atendimentos else qtd * 4
    projeto_norteador_txt = montar_contexto_projeto_norteador_para_ia(
        usar_projeto_norteador,
        projeto_norteador,
        nivel_projeto_norteador,
        observacoes_projeto_norteador,
    )

    fallback = f"""
PLANO MENSAL DE ATENDIMENTO EDUCACIONAL ESPECIALIZADO (AEE)

Código interno: {estudante[1]}
Mês de referência: {mes_referencia}/{ano_referencia}
Dias/datas previstas de atendimento:
{datas_txt}

Total previsto de atendimentos no mês: {total_atendimentos}

Objetivo do mês:
Organizar uma rotina inicial/progressiva de atendimento voltada à comunicação funcional, autonomia, atenção compartilhada, interação e participação nas atividades da SRM.

Projeto norteador / orientação interdisciplinar:
{projeto_norteador_txt}
""".strip()

    if datas_atendimentos:
        for idx, data_atual in enumerate(datas_atendimentos, start=1):
            semana_mes = ((data_atual.day - 1) // 7) + 1
            nome_dia = DIAS_SEMANA[data_atual.weekday()] if data_atual.weekday() < len(DIAS_SEMANA) else "Dia"
            fallback += f"""

SEMANA {semana_mes} - ATENDIMENTO {idx}
Data prevista: {data_atual.strftime('%d/%m/%Y')} ({nome_dia})
- Objetivo: desenvolver comunicação funcional, atenção, autonomia e participação de forma gradual.
- Atividade: proposta mediada com apoio visual, escolha dirigida, recurso tecnológico, material manipulável ou atividade maker conforme resposta do estudante.
- Recursos: prancha de CAA, rotina visual, tablet/Chromebook, material concreto, impressão 3D/robótica quando aplicável.
- Mediação: comandos curtos, demonstração prática, reforço positivo e tempo ampliado para resposta.
- Registro no sistema: resposta do estudante, engajamento, recurso utilizado, barreiras observadas, avanços e encaminhamentos.
"""
    else:
        for semana in range(1, 5):
            fallback += f"""

SEMANA {semana}
Atendimento 1:
- Objetivo: ampliar gradualmente comunicação, atenção e autonomia.
- Atividade: atividade estruturada com apoio visual, escolha mediada, sequência simples ou recurso tecnológico.
- Recursos: prancha de CAA, rotina visual, material manipulável, tablet/Chromebook ou jogo pedagógico.
- Registro no sistema: avanços, dificuldades, barreiras e encaminhamentos.
"""
            if qtd >= 2:
                fallback += f"""
Atendimento 2:
- Objetivo: generalizar a habilidade trabalhada em nova situação.
- Atividade: proposta prática com tecnologia, recurso maker, impressão 3D, robótica ou atividade desplugada.
- Recursos: materiais concretos, recurso visual e mediação individualizada.
- Registro no sistema: comparação com o atendimento anterior e necessidade de ajuste.
"""

    fallback += """

Avaliação do mês:
Ao final do mês, verificar evolução em comunicação funcional, autonomia, atenção compartilhada, interação e resposta aos recursos utilizados.
"""

    if client is None:
        return fallback

    prompt = f"""
Você é especialista em AEE e deve criar um PLANO MENSAL DE ATENDIMENTO aplicável na Sala de Recursos Multifuncionais.

TAREFA:
Crie um plano mensal considerando as DATAS REAIS de atendimento abaixo, para o mês de {mes_referencia}/{ano_referencia}.

DATAS REAIS DE ATENDIMENTO CALCULADAS PELO SISTEMA:
{datas_txt}

TOTAL REAL DE ATENDIMENTOS NO MÊS: {total_atendimentos}

PROJETO NORTEADOR DO ATENDIMENTO:
{projeto_norteador_txt}

REGRAS:
- Não usar nome real do estudante.
- Usar código interno.
- Não inventar evolução ainda não registrada.
- Gerar atividades realistas para o professor aplicar.
- O AEE não deve ser tratado como reforço escolar; o foco deve ser acessibilidade, autonomia, participação, organização cognitiva, comunicação funcional, permanência, mediação pedagógica e protagonismo.
- Se houver projeto norteador, organizar atividades interdisciplinares sem substituir o professor da disciplina.
- Incluir atividades desplugadas, manipuláveis, visuais, práticas e de baixo custo, especialmente para estudantes do Ensino Médio com comorbidades ou maior necessidade de apoio.
- Cada atendimento deve ter: objetivo, atividade, recursos, mediação, registro esperado no sistema e indicador de observação.
- Incluir progressão gradual.
- Considerar atividades plugadas e desplugadas.
- Quando fizer sentido, incluir tablet, Chromebook, CAA, recursos visuais, impressão 3D, robótica e materiais manipuláveis.
- O plano deve ser usado para alimentar o módulo Atendimentos.

DADOS DO ESTUDANTE:
{ctx['estudante_txt']}

ENTREVISTA FAMILIAR:
{ctx['entrevista_txt']}

AVALIAÇÃO PEDAGÓGICA:
{ctx['avaliacao_txt']}

ESTUDO DE CASO GRE:
{ctx['estudo_txt']}

PLANO AEE MANUAL:
{ctx['plano_txt']}

PERFIL DO PROFESSOR AEE:
{ctx['perfil_professor_txt']}

RECURSOS PEDAGÓGICOS E TECNOLOGIAS ASSISTIVAS (ESCOLA + PROFESSOR AEE):
{ctx['recursos_escola_txt']}

ESCUTA DOCENTE / HISTÓRICO DE ESCUTAS:
{ctx['escutas_docentes_txt']}

RELATÓRIOS DE APOIO AO DOCENTE SALVOS:
{ctx['relatorios_docente_txt']}

ATENDIMENTOS REGISTRADOS:
{ctx['historico_txt']}

BASES CONSULTADAS:
{ctx['contexto_base']}

FORMATO DE SAÍDA:
1. Identificação segura
2. Objetivo do mês
3. Habilidades prioritárias do mês
4. Recursos necessários
5. Projeto norteador interdisciplinar, quando ativado
6. Roteiro por data real de atendimento
   - Data e dia da semana
   - Objetivo do atendimento
   - Atividade proposta
   - Recursos
   - Mediação
   - Registro esperado no sistema
6. Como registrar cada atendimento no sistema
7. Indicadores para avaliação mensal
8. Ajustes possíveis para o mês seguinte
"""
    try:
        resposta = client.responses.create(model="gpt-4.1-mini", input=prompt)
        return resposta.output_text
    except Exception as e:
        return f"{fallback}\n\nObservação técnica: não foi possível gerar com IA agora. Erro: {e}"


def salvar_historico_plano_aee_ia(estudante_id, mes_referencia, ano_referencia, qtd_atendimentos_semana, tipo_geracao, diagnostico_ia="", sugestao_geral="", plano_mensal="", observacoes=""):
    inserir_registro(
        "plano_aee_ia",
        ["estudante_id", *CAMPOS_PLANO_AEE_IA],
        [
            estudante_id,
            hoje_str(),
            mes_referencia,
            str(ano_referencia),
            int(qtd_atendimentos_semana or 1),
            tipo_geracao,
            diagnostico_ia,
            sugestao_geral,
            "",
            "",
            "",
            plano_mensal,
            plano_mensal,
            observacoes,
        ],
    )

def gerar_paee_com_ia(estudante, avaliacao=None, entrevista=None, estudo=None):
    client = obter_cliente_openai()

    if client is None:
        return gerar_paee_sem_ia(estudante, avaliacao, entrevista, estudo)

    historico_txt = listar_atendimentos_texto(estudante[0])

    estudante_txt = f"""
Código interno: {estudante[1]}
Ano/Série: {estudante[2]}
Turma: {estudante[3]}
Turno: {estudante[6]}
Perfil educacional informado: {estudante[4]}
Observações pedagógicas iniciais: {estudante[5]}
Dias de atendimento preferenciais: {estudante[7]}
Horário preferencial: {estudante[8]}
"""

    avaliacao_txt = texto_avaliacao(estudante, ("", *avaliacao)) if avaliacao else "Nenhuma avaliação pedagógica registrada."
    entrevista_txt = texto_entrevista(estudante, ("", *entrevista)) if entrevista else "Nenhuma entrevista com a família registrada."
    estudo_txt = texto_estudo_caso(estudante, ("", *estudo)) if estudo else "Nenhum estudo de caso GRE registrado."
    recursos_escola_txt = listar_recursos_escola_texto(limite=60)
    perfil_professor_txt = texto_perfil_professor_aee_para_ia()

    pergunta_busca = f"""
Plano AEE PAEE para estudante com perfil {estudante[4]}, ano/série {estudante[2]},
barreiras, potencialidades, estratégias pedagógicas, atividades desplugadas,
tecnologia assistiva, recursos maker, robótica, impressão 3D e inclusão escolar.
"""
    resultados_base = buscar_na_base_conhecimento(pergunta_busca, bases=["cientifica", "pedagogica"], limite=5)
    contexto_base = montar_contexto_base(resultados_base)
    arquivos_base = arquivos_consultados_texto(resultados_base)

    prompt = f"""
Você é um assistente pedagógico especializado em Atendimento Educacional Especializado (AEE), Educação Inclusiva, Sala de Recursos Multifuncionais (SRM), tecnologias educacionais inclusivas e elaboração de Plano AEE/PAEE.

TAREFA:
Elabore uma sugestão de Plano AEE/PAEE com linguagem formal, técnica, objetiva e pedagógica, cruzando as informações do cadastro, entrevista com a família, avaliação pedagógica, estudo de caso GRE, escuta docente, relatórios de apoio ao docente, histórico de atendimentos e os trechos recuperados das bases científica e pedagógica.

REGRAS DE SEGURANÇA E PRIVACIDADE:
- Não usar nome real de estudante.
- Usar somente “Código interno” para identificar o estudante.
- Não solicitar CPF, RG, endereço, telefone ou dados sensíveis.
- Quando necessário, deixar campos sensíveis em branco para preenchimento manual no Word/PDF.
- Não inventar diagnóstico.
- Não criar condutas médicas.
- Não afirmar evolução não registrada.
- Não inventar citações bibliográficas.
- Quando usar a base de conhecimento, cite apenas o nome do arquivo consultado.

REGRA SOBRE TEA SEM NÍVEL:
Se o perfil educacional for TEA e o nível de suporte não estiver informado, adotar provisoriamente estratégias compatíveis com suporte moderado, sem afirmar diagnóstico clínico.

REGRA CRÍTICA SOBRE ATENDIMENTOS:
A análise da evolução deve ser baseada exclusivamente nos dados reais do histórico de atendimentos. Se os registros forem insuficientes, escrever: “Os registros de atendimento ainda são limitados para uma análise evolutiva consistente, sendo necessário ampliar o acompanhamento pedagógico.”

DADOS DO ESTUDANTE:
{estudante_txt}

ENTREVISTA COM A FAMÍLIA:
{entrevista_txt}

AVALIAÇÃO PEDAGÓGICA:
{avaliacao_txt}

ESTUDO DE CASO GRE:
{estudo_txt}

RECURSOS PEDAGÓGICOS E TECNOLOGIAS ASSISTIVAS (ESCOLA + PROFESSOR AEE):
{recursos_escola_txt}

HISTÓRICO DE ATENDIMENTOS:
{historico_txt}

TRECHOS DAS BASES CIENTÍFICA E PEDAGÓGICA:
{contexto_base}

ARQUIVOS CONSULTADOS:
{arquivos_base}

ESTRUTURE O DOCUMENTO COM:
1. Identificação segura do estudante
2. Campos sensíveis para preenchimento manual no Word/PDF
3. Caracterização pedagógica
4. Síntese da entrevista com a família
5. Síntese da avaliação pedagógica
6. Síntese do estudo de caso GRE
7. Necessidades educacionais específicas
8. Barreiras identificadas
9. Potencialidades
10. Objetivos gerais do AEE
11. Objetivos específicos do AEE
12. Estratégias pedagógicas
13. Recursos de acessibilidade e tecnologias assistivas
14. Sugestões de tecnologias educacionais inclusivas
    - impressão 3D
    - robótica educacional
    - jogos digitais
    - recursos maker
    - comunicação alternativa e aumentativa
    - materiais táteis, visuais e manipuláveis
15. Atividades plugadas sugeridas
16. Atividades desplugadas sugeridas
17. Como aplicar as atividades no atendimento da SRM
18. Organização do atendimento
19. Articulação com família, professores e gestão
20. Avaliação e acompanhamento
21. Evolução do estudante com base nos atendimentos
22. Recomendações para revisão do plano
23. Fundamentação com base nos documentos consultados
24. Arquivos utilizados como referência
25. Responsável pelo AEE:
Nome: ___________________________________________
Função: Professor(a) do Atendimento Educacional Especializado (AEE)
Assinatura: _______________________________________
26. Coordenação pedagógica:
Nome: ___________________________________________
Cargo: Coordenação Pedagógica
Assinatura: _______________________________________
27. Data de elaboração: ____/____/________

Na seção de atividades desplugadas, inclua sugestões concretas e aplicáveis sem computador, como cartões de rotina, pareamento, sequência lógica, jogos de memória, materiais táteis, objetos 3D, trilhas pedagógicas, contação de histórias com apoio visual, atividades de atenção compartilhada e recursos manipuláveis.
"""

    resposta = client.responses.create(
        model="gpt-4.1-mini",
        input=prompt,
    )
    return resposta.output_text


def gerar_relatorio_evolucao(estudante, avaliacao=None):
    api_key = obter_api_key()

    if OpenAI is None or not api_key:
        return "IA não configurada."

    historico_txt = listar_atendimentos_texto(estudante[0])

    estudante_txt = f"""
Código interno: {estudante[1]}
Ano/Série: {estudante[2]}
Turma: {estudante[3]}
Perfil educacional: {estudante[4]}
"""

    prompt = f"""
Você é um especialista em Educação Inclusiva e AEE.

Analise o histórico de atendimentos e produza um RELATÓRIO PEDAGÓGICO ANALÍTICO.

DADOS DO ESTUDANTE:
{estudante_txt}

HISTÓRICO DE ATENDIMENTOS:
{historico_txt}

REGRAS IMPORTANTES:
- NÃO inventar informações.
- Usar somente dados reais.
- Se os dados forem fracos, dizer claramente.

ESTRUTURA DO RELATÓRIO:
1. Síntese da evolução do estudante
2. Análise dos avanços
3. Análise das dificuldades
4. Qualidade dos registros pedagógicos
Classificar como: Alta, Média ou Baixa.
5. Principais problemas identificados nos registros
6. Recomendações para melhoria dos registros
7. Recomendações pedagógicas para o AEE
8. Conclusão técnica
"""

    client = OpenAI(api_key=api_key)
    resposta = client.responses.create(
        model="gpt-4.1-mini",
        input=prompt,
    )
    return resposta.output_text


# ======================================================
# DOCUMENTOS GRE - MODELOS INSTITUCIONAIS COM CAMPOS SIGILOSOS EM BRANCO
# ======================================================
def linha_preenchimento(rotulo, tamanho=60):
    return f"{rotulo}: " + "_" * tamanho


def valor_ou_linha(valor, rotulo="", tamanho=60, sigiloso=False):
    """Preenche somente dados não sigilosos já existentes no sistema."""
    if sigiloso:
        return linha_preenchimento(rotulo, tamanho) if rotulo else "_" * tamanho
    if valor not in (None, "", "Não informado"):
        return str(valor)
    return linha_preenchimento(rotulo, tamanho) if rotulo else "Não informado."


def marcar_opcao(valor, esperado):
    valor = str(valor or "").strip().lower()
    esperado = str(esperado or "").strip().lower()
    return "(X)" if valor == esperado else "( )"


def texto_campos_sensiveis_gre_estudante():
    return """
DADOS SIGILOSOS / PREENCHIMENTO MANUAL NO WORD OU IMPRESSO
Nome completo do(a) estudante: ___________________________________________
Data de nascimento: ____/____/________
CPF/RG do(a) estudante: _________________________________________________
Matrícula oficial do(a) estudante: _______________________________________
Nome do(a) responsável: _________________________________________________
CPF do(a) responsável: _________________________________________________
Contato telefônico do(a) responsável: ___________________________________
Endereço completo: ______________________________________________________
Número do NIS: __________________________________________________________
Número do Cartão SUS: __________________________________________________
Filiação: _______________________________________________________________
""".strip()


def texto_cabecalho_gre(titulo):
    return f"""
SECRETARIA DE EDUCAÇÃO E ESPORTES
GERÊNCIA METROPOLITANA NORTE DE EDUCAÇÃO
CGDE - COORDENAÇÃO GERAL DE DESENVOLVIMENTO DA EDUCAÇÃO
NÚCLEO / UNIDADE DE EDUCAÇÃO INCLUSIVA, DIREITOS HUMANOS E CIDADANIA

{titulo}
""".strip()


def texto_ficha_professor_gre(professor=None):
    if professor is None:
        professor = buscar_professor_responsavel()

    if not professor:
        return "Nenhum professor AEE cadastrado."

    return f"""
{texto_cabecalho_gre("FICHA DE IDENTIFICAÇÃO PROFESSOR(A) AEE")}

ANO LETIVO: ______________________

1. IDENTIFICAÇÃO PESSOAL
Nome: {professor[1] or 'Não informado.'}
Data de nascimento: ____/____/________
Matrícula: ___________________________________________
Endereço: ____________________________________________
Filiação: ____________________________________________
País de nacionalidade: _______________________________
CPF: _________________________________________________
UF de nascimento: ______ Município de nascimento: __________________________

Profissional escolar com deficiência, Transtorno do Espectro do Autismo e/ou Altas Habilidades/Superdotação?
( ) Não   ( ) Sim. Qual? _________________________________________________

Telefones para contato: _________________________________________________
E-mail: _________________________________________________________________

2. FORMAÇÃO ACADÊMICA
Maior nível de escolaridade concluída: {professor[4] or 'Não informado.'}

Dados sobre graduação:
Curso: _________________________________________________________________
Instituição: ____________________________________________________________
Ano de conclusão: _________________________

Dados sobre pós-graduação:
( ) Especialização   ( ) Mestrado   ( ) Doutorado   ( ) Não tem pós concluída
Curso: _________________________________________________________________
Instituição: ____________________________________________________________
Ano de conclusão: _________________________

Outras informações relevantes quanto à qualificação profissional:
{professor[7] or '________________________________________________________________________________'}

3. IDENTIFICAÇÃO PROFISSIONAL
Escola que atua na rede: {professor[2] or 'Não informado.'}
Telefone institucional: _________________________________________________
Regional: {professor[3] or 'Não informado.'}
Horário / carga horária: {professor[5] or 'Não informado.'}
Turno de atuação: {professor[6] or 'Não informado.'}
Acumulação: ___________________________________________
Trabalha em outra rede? ( ) Sim   ( ) Não
Rede que atua: _________________________________________
Horário: ____________________ Telefone: ________________________________

____________________, ____ de __________________________ de ____________

Assinatura do(a) professor(a): __________________________________________
""".strip()


def texto_matricula_srm_gre(estudante):
    return f"""
{texto_cabecalho_gre("MATRÍCULA PARA O ATENDIMENTO EDUCACIONAL ESPECIALIZADO NA SALA DE RECURSOS MULTIFUNCIONAIS - SRM")}

Escola: ________________________________________________
Estudante: _____________________________________________
Ano: {estudante[2] or 'Não informado.'}   Turma: {estudante[3] or 'Não informado.'}   Turno: {estudante[6] or 'Não informado.'}
Ano letivo: ______________________

1. IDENTIFICAÇÃO SEGURA NO SISTEMA
Código interno: {estudante[1] or 'Não informado.'}
Perfil educacional informado: {estudante[4] or 'Não informado.'}
Dias preferenciais de atendimento: {estudante[7] or 'Não informado.'}
Horário preferencial: {estudante[8] or 'Não informado.'}

{texto_campos_sensiveis_gre_estudante()}

2. REQUERIMENTO DE MATRÍCULA / TERMO DE CIÊNCIA DO SERVIÇO DE AEE
Nome da Instituição de Ensino da SRM: ___________________________________
Local e data: ______________________, ____ de __________________ de ______

Nome do(a) estudante: _________________________________________________
Data de nascimento: ____/____/________
Turma: {estudante[3] or 'Não informado.'}   Turno: {estudante[6] or 'Não informado.'}
Filiação: ______________________________________________________________
Endereço: ______________________________________________________________
Recebe auxílio? ( ) Não   ( ) Sim   Número do NIS: ______________________
Número do Cartão SUS: _________________________________________________

3. INDICAÇÃO DO PÚBLICO-ALVO DA EDUCAÇÃO ESPECIAL
( ) Deficiência Auditiva        ( ) Surdo(a)        ( ) Usa aparelho auditivo
( ) Surdocegueira
( ) Deficiência Intelectual     ( ) Síndrome de Down
( ) Altas habilidades/Superdotação
( ) Deficiência Visual          ( ) Cegueira        ( ) Baixa visão
( ) Transtorno do Espectro do Autismo - TEA
( ) Deficiência Física          ( ) Usuário(a) de cadeira de rodas
( ) Deficiência Múltipla. Quais? ________________________________________
( ) Outro: ______________________________________________________________

Perfil registrado no sistema: {estudante[4] or 'Não informado.'}

4. AUTORIZAÇÕES
Autorização para retirar o(a) estudante da escola:
( ) Deixar a escola sozinho(a)
( ) Deixar a escola apenas acompanhado(a)

Nome do acompanhante | Parentesco/relação social | Documento | Telefone
_______________________________________________________________________
_______________________________________________________________________

Permissão para uso e divulgação da imagem para fins educacionais e pedagógicos:
( ) Sim   ( ) Não

Matrícula no Atendimento Educacional Especializado - AEE, realizado na Sala de Recursos Multifuncionais desta escola,
nos dias de {estudante[7] or '_______________________'} no horário de {estudante[8] or '________ às ________'}.

5. TERMO DE CIÊNCIA
Eu, ____________________________________________, CPF nº ______________________________,
responsável pelo(a) estudante ____________________________________________, regularmente matriculado(a)
na escola ____________________________________________, no ano {estudante[2] or '_____'} turma {estudante[3] or '_____'} turno {estudante[6] or '_____'},
declaro estar ciente do Serviço do Atendimento Educacional Especializado disponibilizado pela Unidade Escolar.
Autorizo e me comprometo com a participação e frequência deste(a) estudante aos Atendimentos Educacionais
Especializados na Sala de Recursos Multifuncionais desta Unidade Educacional.

Declaro ter ciência de que o descumprimento do compromisso acima poderá resultar na perda da vaga neste serviço.

Data: ____/____/________

Assinatura do responsável: ______________________________________________
Professor(a) da SRM / AEE: ______________________________________________
Gestão / Coordenação Pedagógica: ________________________________________
""".strip()


def texto_entrevista_familia_gre(estudante, entrevista=None):
    if entrevista is None:
        entrevista = ultima_linha("entrevistas_familia", CAMPOS_ENTREVISTA_FAMILIA, estudante[0])

    dados = dict(zip(CAMPOS_ENTREVISTA_FAMILIA, entrevista or []))
    def v(campo):
        valor = dados.get(campo)
        return valor if valor not in (None, "") else "Não informado."

    return f"""
{texto_cabecalho_gre("ENTREVISTA COM A FAMÍLIA")}

Escola: ________________________________________________
Estudante: _____________________________________________
Ano letivo: {v('ano_letivo')}
Tipo de registro: {v('tipo_registro')}
Código interno no sistema: {estudante[1] or 'Não informado.'}
Ano/Série: {estudante[2] or 'Não informado.'}   Turma: {estudante[3] or 'Não informado.'}   Turno: {estudante[6] or 'Não informado.'}

{texto_campos_sensiveis_gre_estudante()}

1. INFORMAÇÕES DIVERSAS
A família participa de algum programa de auxílio governamental? {v('auxilio_governamental')}
Qual(is): {v('auxilio_quais')}
Há histórico familiar de doenças graves, deficiência ou transtornos? {v('historico_familiar')}
Qual(is): {v('historico_quais')}
Já repetiu de ano? {v('repetiu_ano')}   Quantas vezes? {v('repetiu_qtd')}
Trocou de escola? {v('trocou_escola')}   Quantas vezes? {v('trocou_qtd')}
Motivo da troca: {v('motivo_troca')}
Situação em relação à escola: {v('situacao_escolar')}
Demonstra interesse em frequentar a escola: {v('interesse_escola')}
Cuida/organiza seus materiais: {v('organiza_materiais')}
Apresenta resistência à escola: {v('resistencia_escola')}
Relaciona-se bem com colegas: {v('relacao_colegas')}
Relaciona-se bem com professores: {v('relacao_professores')}
Leva alimentação de casa: {v('leva_alimentacao')}
Alimenta-se da merenda escolar: {v('merenda_escolar')}
Possui alergia alimentar: {v('alergia_alimentar')}   Qual(is): {v('alergia_quais')}
Outras observações: {v('obs_diversas')}

2. SOBRE A ESCOLHA DA ESCOLA
Motivo da escolha: {v('motivo_escolha')}
Outros motivos: {v('outros_motivos')}
O que conhece sobre o serviço do AEE: {v('conhecimento_aee')}

3. INFORMAÇÕES SOBRE A SAÚDE
Possui doença preexistente: {v('doenca_preexistente')}
Convulsões: {v('convulsoes')}
Acompanhamentos profissionais: {v('acompanhamentos')}
Outro acompanhamento: {v('acompanhamento_outro')}
Frequência: {v('frequencia_acompanhamento')}   Outra: {v('frequencia_outro')}
Alimentação saudável: {v('alimentacao_saudavel')}
Seletividade alimentar: {v('seletividade_alimentar')}
Dieta sensorial: {v('dieta_sensorial')}
Usa suplemento alimentar: {v('suplemento_alimentar')}   Qual: {v('suplemento_qual')}
Alimenta-se por sonda: {v('alimenta_sonda')}
Dorme bem: {v('dorme_bem')}
Faz uso de medicação: {v('medicacao')}   Qual(is): {v('medicacao_qual')}
Tempo de medicação/tratamentos realizados: {v('tempo_medicacao_tratamentos')}
Outras observações de saúde: {v('obs_saude')}

4. DESENVOLVIMENTO PSICOMOTOR
Lateralidade: {v('lateralidade')}
Estereotipias: {v('estereotipias')}   Qual(is): {v('estereotipias_quais')}
Segura objetos com as duas mãos: {v('segura_objetos_duas_maos')}
Tamanho dos objetos que segura: {v('tamanho_objetos')}
Faz a pega correta do lápis: {v('pega_lapis')}
Engatinhou: {v('engatinhou')}   Andou com que idade: {v('idade_andou')}
Usa fraldas na escola: {v('usa_fraldas')}
Usa sonda de alívio: {v('usa_sonda_alivio')}
Autonomia nas atividades: {v('autonomia_atividades')}
Outras atividades de autonomia: {v('autonomia_outros')}
Atende comandos: {v('atende_comandos')}
Gosta do toque: {v('gosta_toque')}
Outras observações psicomotoras: {v('obs_psicomotor')}

5. LINGUAGEM
Verbal: {v('verbal')}
Consegue se comunicar: {v('consegue_comunicar')}
Problemas na fala: {v('problemas_fala')}
Ecolalia: {v('ecolalia')}
Consegue dar recado: {v('da_recado')}
Usa comunicação alternativa: {v('comunicacao_alternativa')}   Qual: {v('comunicacao_alternativa_qual')}

6. SOCIALIZAÇÃO
Relação com pai: {v('relacao_pai')}
Relação com mãe: {v('relacao_mae')}
Relação com parentes: {v('relacao_parentes')}
Relação com irmãos: {v('relacao_irmaos')}
Relação com outros estudantes: {v('relacao_estudantes')}
Tem melhor amigo(a): {v('tem_melhor_amigo')}   Tipo: {v('tipo_melhor_amigo')}
Adapta-se ao ambiente: {v('adapta_ambiente')}
Flexível na rotina: {v('flexivel_rotina')}
Respeita regras: {v('respeita_regras')}
Chora com facilidade: {v('chora_facilidade')}
Brinca como: {v('brinca_como')}
Interesses e lazer: {v('interesses_lazer')}
O que a família mais gosta no(a) estudante: {v('familia_gosta')}
O que a família considera necessário melhorar: {v('familia_nao_gosta')}
Ambiente físico em casa para estudos/brincadeiras: {v('ambiente_estudo_casa')}

7. CONTEXTO FAMILIAR
Principais habilidades:
{v('habilidades')}

Principais oportunidades de melhoria:
{v('oportunidades_melhoria')}

8. OUTRAS INFORMAÇÕES
{v('outras_info_familia')}

____________________, ____ de __________________________ de ____________
Assinatura do responsável: ______________________________________________
Professor(a) AEE: ______________________________________________________
""".strip()


def buscar_ultimos_planos_aee_ia(estudante_id, limite=20):
    """Retorna os últimos registros do módulo Plano AEE - IA.

    Usado para que documentos consolidados, como Estudo de Caso e Plano AEE,
    possam aproveitar também o Perfil Pedagógico Inteligente, Sugestão Geral AEE
    e Plano Mensal IA quando o Plano AEE Manual estiver vazio.
    """
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT id, data_geracao, mes_referencia, ano_referencia, qtd_atendimentos_semana,
               tipo_geracao, diagnostico_ia, sugestao_geral, objetivos_prioritarios,
               recursos_sugeridos, estrategias_recomendadas, plano_mensal,
               sugestoes_semanais, observacoes
        FROM plano_aee_ia
        WHERE estudante_id=?
        ORDER BY id DESC
        LIMIT ?
        """,
        (estudante_id, limite),
    )
    dados = cursor.fetchall()
    conn.close()
    return dados or []


def plano_aee_ia_para_dict(registro):
    if not registro:
        return {}
    campos = [
        "id", "data_geracao", "mes_referencia", "ano_referencia", "qtd_atendimentos_semana",
        "tipo_geracao", "diagnostico_ia", "sugestao_geral", "objetivos_prioritarios",
        "recursos_sugeridos", "estrategias_recomendadas", "plano_mensal",
        "sugestoes_semanais", "observacoes",
    ]
    return dict(zip(campos, registro))


def filtrar_planos_ia_por_tipo(registros, tipo_procurado):
    tipo_procurado = str(tipo_procurado or "").lower()
    for reg in registros:
        d = plano_aee_ia_para_dict(reg)
        if tipo_procurado in str(d.get("tipo_geracao", "")).lower():
            return d
    return {}


def texto_valido(valor):
    valor = str(valor or "").strip()
    if not valor:
        return ""
    if valor.lower() in ["não informado", "não informado.", "none", "null"]:
        return ""
    return valor


def extrair_secao_textual(texto, palavras_chave, max_chars=1800):
    """Extrai uma seção textual simples de relatórios gerados pela IA.

    Procura uma linha que contenha uma das palavras-chave e retorna o conteúdo até
    a próxima seção numerada/titulada. Se não encontrar, retorna vazio.
    """
    texto = str(texto or "")
    if not texto.strip():
        return ""

    linhas = texto.splitlines()
    chaves = [str(c).lower() for c in palavras_chave]
    capturando = False
    saida = []

    for linha in linhas:
        limpa = linha.strip().strip("#*-• ")
        baixa = limpa.lower()

        encontrou = any(chave in baixa for chave in chaves)
        if not capturando and encontrou:
            capturando = True
            saida.append(limpa)
            continue

        if capturando:
            # Para ao encontrar nova seção numerada/título grande.
            if saida and (
                re.match(r"^(\d+\.|\d+\)|[A-ZÁÉÍÓÚÂÊÔÃÕÇ ]{8,})", limpa)
                and not any(chave in baixa for chave in chaves)
            ):
                break
            if limpa:
                saida.append(limpa)

    resultado = "\n".join(saida).strip()
    if len(resultado) > max_chars:
        resultado = resultado[:max_chars].rstrip() + "..."
    return resultado


def montar_fontes_plano_aee_consolidado(planos_ia):
    fontes = []
    for d in planos_ia:
        tipo = texto_valido(d.get("tipo_geracao"))
        data = texto_valido(d.get("data_geracao"))
        mes = texto_valido(d.get("mes_referencia"))
        ano = texto_valido(d.get("ano_referencia"))
        if tipo:
            complemento = f" - {mes}/{ano}" if mes or ano else ""
            fontes.append(f"{tipo} ({data}{complemento})")
    return "; ".join(fontes) if fontes else "Nenhum relatório de IA localizado."


def consolidar_textos_plano_aee_ia(estudante_id):
    """Consolida os registros do módulo Plano AEE - IA para alimentar relatórios GRE.

    Prioridade de uso:
    1. Plano AEE Manual preenchido no módulo próprio.
    2. Campos estruturados salvos no Plano AEE - IA:
       - objetivos_prioritarios
       - recursos_sugeridos
       - estrategias_recomendadas
       - diagnostico_ia (mantido apenas como nome técnico de banco)
       - sugestao_geral (mantido por compatibilidade com versões antigas)
       - plano_mensal
       - sugestoes_semanais
       - observacoes
    3. Texto geral dos relatórios IA, quando não houver campo estruturado.

    Observação: mantemos os nomes antigos das colunas no banco para não quebrar
    instalações já existentes no Render/PostgreSQL.
    """
    registros_raw = buscar_ultimos_planos_aee_ia(estudante_id, limite=30)
    registros = [plano_aee_ia_para_dict(r) for r in registros_raw]

    def juntar(campo, tipos_preferidos=None, limite=6000):
        partes = []
        tipos_preferidos = [str(t).lower() for t in (tipos_preferidos or [])]

        # 1º: registros dos tipos preferidos
        for d in registros:
            tipo = str(d.get("tipo_geracao", "")).lower()
            if tipos_preferidos and not any(tp in tipo for tp in tipos_preferidos):
                continue
            valor = texto_valido(d.get(campo))
            if valor and valor not in partes:
                partes.append(valor)

        # 2º: qualquer registro, para não deixar relatório vazio
        for d in registros:
            valor = texto_valido(d.get(campo))
            if valor and valor not in partes:
                partes.append(valor)

        texto = "\n\n".join(partes).strip()
        if len(texto) > limite:
            texto = texto[:limite].rstrip() + "..."
        return texto

    perfil = juntar("diagnostico_ia", ["perfil pedagógico", "perfil pedagogico"])
    sugestao = juntar("sugestao_geral", ["perfil pedagógico", "perfil pedagogico", "sugestão", "sugestao"])
    objetivos = juntar("objetivos_prioritarios", ["perfil pedagógico", "perfil pedagogico", "plano mensal"])
    recursos = juntar("recursos_sugeridos", ["perfil pedagógico", "perfil pedagogico", "plano mensal"])
    estrategias = juntar("estrategias_recomendadas", ["perfil pedagógico", "perfil pedagogico", "plano mensal"])
    plano_mensal = juntar("plano_mensal", ["plano mensal"])
    semanas = juntar("sugestoes_semanais", ["plano mensal"])
    observacoes = juntar("observacoes")

    geral = "\n\n".join([x for x in [perfil, sugestao, objetivos, recursos, estrategias, plano_mensal, semanas, observacoes] if texto_valido(x)])

    return {
        "registros": registros,
        "perfil": perfil,
        "sugestao": sugestao,
        "objetivos": objetivos,
        "recursos": recursos,
        "estrategias": estrategias,
        "plano_mensal": plano_mensal,
        "semanas": semanas,
        "observacoes": observacoes,
        "geral": geral,
        "fontes": montar_fontes_plano_aee_consolidado(registros[:8]),
    }


def limpar_texto_gre_narrativo(texto):
    """Limpa títulos internos da IA e deixa o texto adequado ao modelo narrativo GRE.
    Não altera o sentido pedagógico; apenas remove rótulos de seções que causavam duplicidade
    dentro dos campos do formulário (ex.: "2. Objetivos específicos" dentro de 2.2). 
    """
    texto = str(texto or "").strip()
    if not texto:
        return ""

    # remove markdown simples e separadores
    texto = texto.replace("**", "")
    texto = re.sub(r"^#{1,6}\s*", "", texto, flags=re.MULTILINE)
    texto = texto.replace("---", "")

    # títulos internos comuns dos relatórios IA que NÃO devem aparecer no formulário GRE
    titulos_remover = [
        r"^\s*1\.?\s*Objetivo geral de atendimento\s*$",
        r"^\s*1\.?\s*Objetivo geral\s*$",
        r"^\s*2\.?\s*Objetivos específicos\s*$",
        r"^\s*2\.?\s*Objetivos especificos\s*$",
        r"^\s*3\.?\s*Eixos prioritários\s*$",
        r"^\s*3\.?\s*Eixos prioritarios\s*$",
        r"^\s*4\.?\s*Organização sugerida dos atendimentos\s*$",
        r"^\s*4\.?\s*Organizacao sugerida dos atendimentos\s*$",
        r"^\s*5\.?\s*Recursos de acessibilidade e tecnologia educacional\s*$",
        r"^\s*5\.?\s*Recursos de acessibilidade\s*$",
        r"^\s*Fico à disposição.*$",
        r"^\s*Fico a disposição.*$",
    ]

    linhas = []
    for linha in texto.splitlines():
        l = linha.strip()
        if not l:
            continue
        if any(re.match(p, l, flags=re.IGNORECASE) for p in titulos_remover):
            continue
        linhas.append(l)

    texto = "\n".join(linhas).strip()

    # remove repetição de espaços e linhas em excesso
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip()


def texto_lista_gre(texto, padrao="Não informado."):
    """Normaliza respostas em lista ou texto, mantendo linguagem do modelo preenchido GRE."""
    texto = limpar_texto_gre_narrativo(texto)
    if not texto or texto.lower() in ["não informado", "não informado.", "none", "null"]:
        return padrao
    return texto


def parece_objetivo_geral(texto):
    """Identifica quando um texto de estratégia veio, na verdade, como objetivo geral.

    No documento GRE, o campo 3.2 deve trazer ações operacionais em lista.
    Quando a IA ou o preenchimento manual repete frases como "Promover o desenvolvimento...",
    o sistema deve tratar isso como objetivo, não como estratégia.
    """
    t = str(texto or "").strip()
    if not t:
        return False
    tl = t.lower()
    tem_marcadores = ("•" in t) or (";" in t) or ("\n-" in t) or ("\n•" in t)
    inicia_objetivo = tl.startswith((
        "promover ", "favorecer ", "desenvolver ", "estimular ",
        "ampliar ", "proporcionar ", "garantir "
    ))
    cara_de_objetivo = (
        "por meio de" in tl or
        "visando" in tl or
        "favorecendo" in tl or
        "com foco" in tl or
        "com o objetivo" in tl
    )
    return inicia_objetivo and cara_de_objetivo and not tem_marcadores


def formatar_lista_operacional_gre(itens):
    """Formata estratégias em lista, usando o padrão narrativo aceito no documento GRE."""
    saida = []
    for item in itens:
        item = str(item or "").strip().strip("-• ").rstrip(".;")
        if item:
            saida.append(f"• {item}.")
    return "\n".join(saida)


def estrategias_operacionais_padrao_gre(perfil_estudante="", recursos_texto=""):
    """Estratégias operacionais para o campo 3.2 do documento GRE.

    O texto é global, aplicável ao Plano AEE, e não descreve projeto mensal.
    """
    itens = [
        "Utilização de comandos curtos, claros e objetivos, apresentados de forma gradual",
        "Uso de recursos visuais estruturados para apoiar a compreensão, a comunicação e a organização da rotina",
        "Mediação individualizada durante as atividades pedagógicas, com retirada progressiva de apoio quando possível",
        "Introdução ou fortalecimento de recursos de Comunicação Aumentativa e Alternativa (CAA), quando pertinente ao perfil do estudante",
        "Organização das atividades em pequenas etapas, com previsibilidade, tempo adequado e possibilidade de pausas planejadas",
        "Utilização de recursos tecnológicos acessíveis, como tablet, Chromebook ou ferramentas digitais mediadas, quando disponíveis e adequadas",
        "Uso de materiais concretos, impressos, visuais, manipuláveis ou táteis para favorecer participação e compreensão",
        "Observação contínua das respostas do estudante às estratégias utilizadas, com registro dos avanços, barreiras e necessidades de ajustes",
        "Estímulo à interação social e à participação nas atividades escolares de forma progressiva e mediada",
        "Articulação com professores da sala regular, família, gestão escolar e profissional de apoio, quando houver, para alinhar estratégias de acessibilidade"
    ]
    return formatar_lista_operacional_gre(itens)


def garantir_estrategias_operacionais_gre(texto, fallback=""):
    """Garante que o campo 3.2 Estratégia não receba objetivo geral.

    Se o texto estiver em formato operacional, preserva. Se vier como objetivo geral,
    vazio ou genérico demais, substitui por uma lista pedagógica de estratégias.
    """
    texto = texto_lista_gre(texto, "")
    if not texto or parece_objetivo_geral(texto):
        texto = texto_lista_gre(fallback, "")

    if not texto or parece_objetivo_geral(texto):
        return estrategias_operacionais_padrao_gre()

    # Se já tiver marcadores ou várias linhas, preserva a lista, apenas limpando títulos internos.
    if "•" in texto or "\n" in texto or ";" in texto:
        return limpar_texto_gre_narrativo(texto)

    # Texto único curto demais costuma ser rótulo/objetivo genérico.
    if len(texto) < 180:
        return estrategias_operacionais_padrao_gre()

    return texto


def derivar_etapa_ano_do_cadastro(estudante, dados_estudo=None):
    """Evita conflito entre ano/série do cadastro e campos do estudo.
    Quando o cadastro informa Ensino Médio/EF/EJA, ele prevalece sobre valores padrão acidentais.
    """
    dados_estudo = dados_estudo or {}
    serie = str(estudante[2] or "").strip()
    etapa_estudo = texto_valido(dados_estudo.get("etapa_modalidade"))
    ano_estudo = texto_valido(dados_estudo.get("ano_etapa"))

    serie_baixa = serie.lower()
    if serie:
        if "médio" in serie_baixa or "medio" in serie_baixa or " em" in f" {serie_baixa}" or "ensino médio" in serie_baixa:
            return "Ensino Médio", serie
        if "fundamental" in serie_baixa or " ef" in f" {serie_baixa}":
            return "Ensino Fundamental", serie
        if "eja" in serie_baixa:
            return "EJA", serie

    return etapa_estudo or "Não informado.", ano_estudo or serie or "Não informado."


def sim_nao_gre(valor, padrao="Não informado."):
    valor = str(valor or "").strip()
    if not valor:
        return padrao
    v = valor.lower()
    if v in ["sim", "s", "yes", "true", "1"]:
        return "Sim"
    if v in ["não", "nao", "n", "no", "false", "0"]:
        return "Não"
    return valor


def texto_estudo_plano_aee_gre(estudante, estudo=None, plano=None):
    """Gera o documento 'Estudo de Caso e Plano AEE – GRE' em formato narrativo institucional.

    Regras desta versão:
    - responde todos os campos do modelo GRE;
    - trata o documento como PLANO GLOBAL, não mensal;
    - não usa projeto norteador, plano mensal, sugestões semanais ou roteiro por atendimento;
    - usa primeiro os campos manuais do Estudo de Caso e do Plano AEE;
    - usa IA apenas como apoio global quando houver lacunas, com limpeza de títulos internos;
    - segue o estilo do modelo GRE preenchido: campos estruturados com narrativas pedagógicas.
    """
    if estudo is None:
        estudo = ultima_linha("estudos_caso", CAMPOS_ESTUDO_CASO, estudante[0])
    if plano is None:
        plano = ultima_linha("planos_aee", CAMPOS_PLANO_AEE, estudante[0])

    dados_ia = consolidar_textos_plano_aee_ia(estudante[0])

    # Somente registros globais. O documento GRE não deve puxar plano mensal, semanas ou projeto do mês.
    texto_perfil = dados_ia.get("perfil", "")
    texto_sugestao = dados_ia.get("sugestao", "")
    texto_objetivos_ia = dados_ia.get("objetivos", "")
    texto_recursos_ia = dados_ia.get("recursos", "")
    texto_estrategias_ia = dados_ia.get("estrategias", "")
    texto_geral_ia = "\n\n".join([texto_perfil, texto_sugestao, texto_objetivos_ia, texto_recursos_ia, texto_estrategias_ia]).strip()
    fontes_ia_txt = dados_ia.get("fontes", "Nenhum relatório de IA localizado.")

    dados_estudo = dict(zip(CAMPOS_ESTUDO_CASO, estudo or []))
    dados_plano = dict(zip(CAMPOS_PLANO_AEE, plano or [])) if plano else {}

    def e(campo, padrao="Não informado."):
        valor = texto_valido(dados_estudo.get(campo))
        return texto_lista_gre(valor, padrao) if valor else padrao

    def valor_manual(campo):
        return texto_valido(dados_plano.get(campo))

    def escolher(*opcoes, padrao="Não informado."):
        for opcao in opcoes:
            valor = texto_lista_gre(opcao, "")
            if texto_valido(valor):
                return valor
        return padrao

    def secao(texto, chaves, max_chars=1800):
        return texto_lista_gre(extrair_secao_textual(texto, chaves, max_chars=max_chars), "")

    def p(campo):
        """Busca primeiro no Plano AEE Manual; se vazio, usa IA global, nunca o Plano Mensal."""
        manual = texto_lista_gre(valor_manual(campo), "")
        if manual:
            return manual

        if campo == "habilidades_prioritarias":
            return escolher(
                secao(texto_objetivos_ia, ["habilidades prioritárias", "habilidades prioritarias", "objetivos prioritários", "objetivos prioritarios"]),
                secao(texto_sugestao, ["habilidades prioritárias", "habilidades prioritarias", "ensino do uso", "desenvolvimento de vida autônoma"]),
                padrao="Desenvolvimento de vida autônoma; ensino da informática acessível; ensino do uso da comunicação aumentativa e alternativa, quando pertinente ao perfil pedagógico do estudante."
            )

        if campo == "recursos_acessibilidade":
            return escolher(
                secao(texto_recursos_ia, ["recursos de acessibilidade", "tecnologia assistiva", "tecnologia educacional", "recursos"]),
                secao(texto_sugestao, ["recursos de acessibilidade", "recursos sugeridos", "tecnologia assistiva"]),
                padrao="Recursos visuais estruturados; materiais concretos/manipuláveis; recursos digitais acessíveis; tablet ou Chromebook quando disponíveis; pranchas ou recursos de Comunicação Aumentativa e Alternativa (CAA), quando pertinente; materiais impressos em papel; recursos de robótica educacional e impressão 3D quando disponibilizados pelo professor do AEE e adequados ao perfil pedagógico; estratégias de apoio à escrita e organização da rotina."
            )

        if campo == "objetivos_gerais":
            return escolher(
                secao(texto_objetivos_ia, ["objetivo geral", "objetivos prioritários", "objetivos prioritarios"]),
                secao(texto_sugestao, ["objetivo geral de atendimento", "objetivo geral"]),
                padrao="Promover o desenvolvimento da autonomia, participação, comunicação funcional, organização pedagógica e acesso às atividades escolares por meio de recursos de acessibilidade, mediação pedagógica e estratégias inclusivas no Atendimento Educacional Especializado."
            )

        if campo == "objetivos_especificos":
            return escolher(
                secao(texto_objetivos_ia, ["objetivos específicos", "objetivos especificos", "objetivos prioritários", "objetivos prioritarios"]),
                secao(texto_sugestao, ["objetivos específicos", "objetivos especificos"]),
                padrao="Desenvolver formas de comunicação funcional; ampliar participação nas atividades escolares; favorecer compreensão de comandos; estimular autonomia na organização das atividades; utilizar recursos acessíveis; fortalecer interação entre estudante, AEE, sala regular, família e equipe escolar."
            )

        if campo == "metodologia":
            return escolher(
                secao(texto_estrategias_ia, ["metodologia", "mediação pedagógica", "mediacao pedagogica", "rotina estruturada"]),
                secao(texto_sugestao, ["metodologia", "acompanhamento", "observação", "observacao"]),
                padrao="O acompanhamento será realizado por meio de observação pedagógica contínua, mediação individualizada, uso de recursos visuais, organização de rotina, atividades estruturadas, recursos tecnológicos acessíveis e registros sistemáticos da resposta do estudante às estratégias propostas."
            )

        if campo == "estrategias":
            manual_estrategias = valor_manual("estrategias")
            ia_estrategias = escolher(
                secao(texto_estrategias_ia, ["estratégias", "estrategias", "mediação", "mediacao", "recursos visuais"]),
                secao(texto_sugestao, ["estratégias", "estrategias", "comandos curtos", "apoio visual"]),
                padrao=""
            )
            return garantir_estrategias_operacionais_gre(manual_estrategias, fallback=ia_estrategias)

        if campo == "prazo":
            return "Durante o período letivo vigente, com acompanhamento contínuo e reavaliação das estratégias conforme a evolução pedagógica do estudante e os registros do AEE."

        if campo == "acoes_escola":
            return escolher(
                manual,
                padrao="Implemento de tecnologia assistiva para uso do(a) estudante; formação continuada de professores, profissionais de apoio, equipe gestora e famílias com a temática da educação inclusiva; articulação de horários para diálogo entre professor do AEE, professor da sala comum e profissional de apoio; articulação sobre PDDE Equidade, quando pertinente; solicitação de profissional de apoio ou transporte escolar inclusivo, quando identificado como necessário."
            )

        if campo == "barreiras_identificadas":
            return escolher(
                manual,
                e("dificuldades", ""),
                secao(texto_perfil, ["barreiras identificadas", "barreiras"]),
                secao(texto_sugestao, ["barreiras"]),
                padrao="Barreiras comunicacionais, curriculares e/ou atitudinais a serem acompanhadas pelo AEE, considerando participação, acesso às atividades escolares, organização da rotina, comunicação funcional e uso de recursos de acessibilidade."
            )

        if campo == "parcerias":
            return escolher(
                manual,
                texto_valido(dados_estudo.get("parcerias")),
                padrao="Articulação com família, equipe escolar, professor da sala comum, profissional de apoio, gestão escolar e rede intersetorial quando necessário, visando fortalecer as estratégias de inclusão, acessibilidade e participação do estudante."
            )

        if campo == "avaliacao_acompanhamento":
            return escolher(
                manual,
                e("avaliacao", ""),
                secao(texto_perfil, ["indicadores a observar", "avaliação", "avaliacao"]),
                padrao="O processo avaliativo será contínuo, considerando participação, autonomia, comunicação, interação social, uso dos recursos acessíveis, resposta às mediações e evolução pedagógica observada nos atendimentos e no contexto escolar."
            )

        if campo == "observacoes":
            return escolher(
                manual,
                padrao=f"Documento GRE/PAEE preenchido em perspectiva global, considerando estudo de caso, avaliação pedagógica, entrevista familiar e Plano AEE. Não foram utilizados projeto norteador mensal, sugestões semanais ou roteiros por atendimento. Fontes de apoio IA consideradas, quando disponíveis: {fontes_ia_txt}."
            )

        return manual or "Não informado."

    etapa_oficial, ano_oficial = derivar_etapa_ano_do_cadastro(estudante, dados_estudo)
    turno = estudante[6] if len(estudante) > 6 and estudante[6] else "Não informado."
    dias_cadastro = estudante[7] if len(estudante) > 7 and estudante[7] else ""

    # Campos de estudo: prioriza campo específico; se vazio, usa síntese antiga.
    percurso = escolher(e("percurso_educacional", ""), e("contextualizacao", ""), padrao="Não informado.")
    motivo = escolher(e("motivo_encaminhamento_aee", ""), e("queixa_principal", ""), padrao="Não informado.")
    habilidades_obs = escolher(e("habilidades_observadas", ""), e("potencialidades", ""), padrao="Não informado.")
    habilidades_dev = escolher(e("habilidades_a_desenvolver", ""), e("dificuldades", ""), padrao="Não informado.")
    indicadores_ahsd = escolher(e("indicadores_altas_habilidades", ""), padrao="Não há elementos suficientes observados até o momento para caracterizar indicadores de altas habilidades/superdotação, sendo necessário acompanhamento contínuo.")

    def texto_padrao_recursos_integrados_gre():
        """Gera observação institucional sobre recursos considerando escola + professor AEE.
        Usado quando o campo manual/IA não trouxe recursos suficientes para o item 2.4 do modelo GRE.
        """
        recursos_escola = listar_recursos_escola(apenas_disponiveis=True)
        professores = listar_professores()
        recursos_prof = ""
        uso_prof = ""
        obs_prof = ""
        if professores:
            prof = professores[0]
            recursos_prof = prof[15] if len(prof) > 15 and prof[15] else ""
            uso_prof = prof[16] if len(prof) > 16 and prof[16] else ""
            obs_prof = prof[17] if len(prof) > 17 and prof[17] else ""

        partes = []
        if recursos_escola:
            nomes_escola = []
            for r in recursos_escola[:12]:
                try:
                    nomes_escola.append(str(r[2]))
                except Exception:
                    pass
            if nomes_escola:
                partes.append(
                    "Serão considerados os recursos pedagógicos e tecnologias assistivas cadastrados no banco institucional da escola, "
                    f"tais como: {', '.join(nomes_escola)}."
                )

        if recursos_prof or uso_prof or obs_prof:
            texto_prof = "; ".join([x for x in [recursos_prof, uso_prof, obs_prof] if texto_valido(x)])
            partes.append(
                "Também serão considerados, quando adequados ao perfil pedagógico do estudante, recursos de apoio disponibilizados pelo professor do AEE "
                f"para uso mediado nos atendimentos, incluindo: {texto_prof}. Esses recursos não são registrados como patrimônio institucional da escola, "
                "mas como materiais de apoio pedagógico do professor."
            )

        if not partes:
            partes.append(
                "No momento, não há recursos pedagógicos ou tecnologias assistivas cadastradas no banco institucional da escola nem recursos do professor informados no sistema. "
                "Recomenda-se utilizar estratégias de baixo custo, recursos visuais, materiais impressos, atividades desplugadas e materiais manipuláveis confeccionados conforme a necessidade pedagógica do estudante."
            )

        partes.append(
            "A seleção e o uso dos recursos deverão observar as necessidades funcionais, comunicacionais, sensoriais e pedagógicas do estudante, "
            "priorizando acessibilidade, participação, autonomia e permanência nas atividades escolares."
        )
        return "\n".join(partes)

    recursos_gre = escolher(
        e('recursos_tecnologia_assistiva', ''),
        p('recursos_acessibilidade'),
        padrao=texto_padrao_recursos_integrados_gre(),
    )

    return f"""
{texto_cabecalho_gre("ESTUDO DE CASO E PLANO DE ATENDIMENTO EDUCACIONAL ESPECIALIZADO")}

O Plano de Atendimento Educacional Especializado deverá garantir os registros avaliativos do estudante público-alvo da educação especial, considerando o que o professor do Atendimento Educacional Especializado (AEE) observará no estudo de caso. O estudo de caso orienta a eliminação de barreiras e a busca por recursos de acessibilidade necessários à promoção da autonomia e da aprendizagem.

PARTE 1 - IDENTIFICAÇÃO DO(A) ESTUDANTE

1.1 Nome do(a) estudante atendido no AEE: ___________________________________________
1.2 Data de nascimento do(a) estudante: ____/____/________
1.3 Matrícula do(a) estudante: _______________________________________
1.4 Nome e contato telefônico do(a) responsável pelo(a) estudante: _______________________________________

1.5 Etapa/modalidade da educação em que o(a) estudante está: {etapa_oficial}
Ano/etapa: {ano_oficial}
1.6 Turma e turno: {estudante[3] or 'Não informado.'} / {turno}
1.7 O(a) estudante apresenta laudo? {sim_nao_gre(e('laudo'))}
1.8 O(a) estudante apresenta deficiência (com CID): {e('deficiencia')}  CID: {e('cid')}
1.9 Altas Habilidades e Superdotação: {e('altas_habilidades')}
1.10 Usuário de BPC? {sim_nao_gre(e('bpc'))}
1.11 Nome da escola em que o(a) estudante está matriculado no ensino comum: {e('escola_nome')}
1.12 Nome da unidade educacional onde o(a) estudante é atendido pelo AEE: {e('unidade_aee')}
1.13 Nome e contato telefônico do(a) gestor(a) da escola do ensino comum: {e('gestor_nome')} - {e('gestor_contato')}
1.14 Nome do(a) professor(a) do AEE que atende o(a) estudante e contato telefônico: {e('professor_nome')} - {e('professor_contato')}
1.15 Matrícula do(a) professor(a) do AEE que atende o(a) estudante: {e('matricula_professor')}
1.16 Especialidade do(a) professor(a) do AEE: {e('especialidade_professor')}
1.17 Período de elaboração deste Plano de Atendimento Educacional Especializado: {e('periodo_inicio')}
1.18 Data final: {e('periodo_fim')}
1.19 Frequência de atendimento na Sala de Recursos Multifuncionais – SRM: {e('frequencia_atendimento', dias_cadastro or 'Não informado.')}
1.20 Tempo de atendimento por semana: {e('tempo_atendimento_semana')}
1.21 Formato do atendimento: {e('formato_atendimento')}

2. ESTUDO DE CASO / Percurso educacional do(a) estudante
{percurso}

2.1 Motivo pelo qual o(a) estudante foi encaminhado para o Atendimento Educacional Especializado:
{motivo}

2.2 Precisa de transporte Escolar Inclusivo? {sim_nao_gre(e('precisa_transporte_inclusivo'))}
2.2.1 Recebe o Serviço de Transporte escolar Inclusivo? {sim_nao_gre(e('recebe_transporte_inclusivo'))}

2.3 O(a) estudante precisa de profissional de apoio? {sim_nao_gre(e('precisa_profissional_apoio'))}
Justifique: {e('justificativa_apoio')}

2.3.1 O(a) estudante é acompanhado por profissional de apoio na escola? {sim_nao_gre(e('acompanhado_profissional_apoio'))}
Nome do apoio: {e('nome_profissional_apoio')}

2.4 O(a) estudante utiliza qual(is) recurso(s) de tecnologia educacional e/ou assistiva?
{recursos_gre}

2.5 Registrar observações relevantes para o ambiente educacional do estudante, como acompanhamento médico ou terapêutico, caso seja necessário:
{e('observacoes_ambiente_educacional')}

2.6 Habilidades observadas e desenvolvidas pelo estudante:
{habilidades_obs}

2.7 Habilidades que precisam ser desenvolvidas pelo estudante:
{habilidades_dev}

2.8 Marque características que podem ser indicadores em altas habilidades e superdotação:
{indicadores_ahsd}

2.9 Caso o estudante seja pessoa surda, marque se faz uso de:
{e('recursos_surdez')}
Obs.: {e('observacoes_surdez')}

PARTE 2 - PLANO DE ATENDIMENTO EDUCACIONAL ESPECIALIZADO
Após a coleta dos dados para a construção do Estudo de Caso, o Plano de Atendimento Educacional Especializado deve ser consolidado com as informações que seguem:

1. Habilidades Específicas
1.1 Habilidades Prioritárias que serão trabalhadas com o estudante na SRM:
{p('habilidades_prioritarias')}

1.2 Recursos de acessibilidade que serão disponibilizados ao estudante:
{p('recursos_acessibilidade')}

2. Objetivos do Atendimento Educacional Especializado
2.1 Geral:
{p('objetivos_gerais')}

2.2 Específicos:
{p('objetivos_especificos')}

3. Metodologias e estratégias
(Registre a metodologia e estratégias desenvolvidas ao longo do período para alcançar os objetivos propostos para o(a) estudante, especificando as estratégias necessárias para desenvolvê-las):

3.1 Metodologia:
{p('metodologia')}

3.2 Estratégia:
{p('estrategias')}

3.3 Prazo:
{p('prazo')}

4. Registre as ações desenvolvidas no âmbito da escola junto à Gestão e Educador de Apoio para garantir a acessibilidade do estudante.
{p('acoes_escola')}

5. Marque as barreiras que foram identificadas na comunidade escolar que são impedimentos para a inclusão do estudante.
{p('barreiras_identificadas')}

6. Registre possíveis parcerias realizadas pelo AEE ao longo do período, especificando as ações realizadas:
{p('parcerias')}

7. AVALIAÇÃO
(Registre os avanços e as habilidades que necessitam ser reforçadas, além das ações realizadas para eliminar as barreiras encontradas no percurso escolar do estudante e se foram efetivas para garantir a inclusão.)
{p('avaliacao_acompanhamento')}

Observações complementares:
{p('observacoes')}

____________________, ____ de __________________________ de ____________
Professor(a) AEE: _______________________________________
Coordenação/Gestão: _____________________________________
Responsável: ____________________________________________
""".strip()

def resumo_indicadores_atendimentos(estudante_id):
    atendimentos = listar_atendimentos(estudante_id)
    if not atendimentos:
        return "Nenhum atendimento registrado."
    total = len(atendimentos)
    def media(indice):
        valores = []
        for a in atendimentos:
            try:
                valores.append(float(a[indice]))
            except Exception:
                pass
        return round(sum(valores)/len(valores), 1) if valores else "Não informado"
    return f"""
Total de atendimentos registrados: {total}
Média da resposta do estudante: {media(8)}/10
Média do avanço pedagógico: {media(9)}/10
Média da dificuldade/barreira observada: {media(10)}/10
Média do engajamento/participação: {media(11)}/10
Média do índice geral de evolução: {media(12)}/10
""".strip()


def gerar_relatorio_gre_texto(estudante):
    avaliacao = ultima_avaliacao(estudante[0])
    entrevista = ultima_linha("entrevistas_familia", CAMPOS_ENTREVISTA_FAMILIA, estudante[0])
    estudo = ultima_linha("estudos_caso", CAMPOS_ESTUDO_CASO, estudante[0])
    plano = ultima_linha(
        "planos_aee",
        CAMPOS_PLANO_AEE,
        estudante[0],
    )

    return f"""
{texto_cabecalho_gre("RELATÓRIO CONSOLIDADO GRE - ATENDIMENTO EDUCACIONAL ESPECIALIZADO")}

1. IDENTIFICAÇÃO SEGURA DO ESTUDANTE
Código interno: {estudante[1] or 'Não informado.'}
Ano/Série: {estudante[2] or 'Não informado.'}
Turma: {estudante[3] or 'Não informado.'}
Turno: {estudante[6] or 'Não informado.'}
Perfil educacional: {estudante[4] or 'Não informado.'}
Dias de atendimento: {estudante[7] or 'Não informado.'}
Horário preferencial: {estudante[8] or 'Não informado.'}

{texto_campos_sensiveis_gre_estudante()}

2. PROFESSOR(A) AEE RESPONSÁVEL
{texto_professores_vinculados(estudante[0])}

3. SÍNTESE DO CADASTRO PEDAGÓGICO
{estudante[5] or 'Não informado.'}

4. SÍNTESE DA ENTREVISTA COM A FAMÍLIA
{texto_entrevista_familia_gre(estudante, entrevista) if entrevista else 'Nenhuma entrevista registrada.'}

5. AVALIAÇÃO PEDAGÓGICA
{texto_avaliacao(estudante, ('', *avaliacao)) if avaliacao else 'Nenhuma avaliação registrada.'}

6. ESTUDO DE CASO E PLANO AEE
{texto_estudo_plano_aee_gre(estudante, estudo, plano) if (estudo or plano) else 'Nenhum estudo de caso ou plano registrado.'}

7. INDICADORES DOS ATENDIMENTOS
{resumo_indicadores_atendimentos(estudante[0])}

8. HISTÓRICO DOS ATENDIMENTOS
{listar_atendimentos_texto(estudante[0])}

9. ENCAMINHAMENTOS GERAIS
_______________________________________________________________________
_______________________________________________________________________
_______________________________________________________________________

10. CONSIDERAÇÕES FINAIS
_______________________________________________________________________
_______________________________________________________________________
_______________________________________________________________________

11. ASSINATURAS
Professor(a) AEE: _______________________________________
Coordenação/Gestão: _____________________________________
Responsável: ____________________________________________
""".strip()


def texto_pacote_gre_completo(estudante):
    return f"""
PACOTE GRE COMPLETO - INCLUISRM
Código interno do estudante: {estudante[1] or 'Não informado.'}

======================================================================
DOCUMENTO 1 - FICHA DE IDENTIFICAÇÃO PROFESSOR(A) AEE
======================================================================
{texto_ficha_professor_gre()}


======================================================================
DOCUMENTO 2 - MATRÍCULA SRM / TERMO DE CIÊNCIA
======================================================================
{texto_matricula_srm_gre(estudante)}


======================================================================
DOCUMENTO 3 - ENTREVISTA COM A FAMÍLIA
======================================================================
{texto_entrevista_familia_gre(estudante)}


======================================================================
DOCUMENTO 4 - ESTUDO DE CASO E PLANO AEE
======================================================================
{texto_estudo_plano_aee_gre(estudante)}


======================================================================
DOCUMENTO 5 - RELATÓRIO CONSOLIDADO GRE
======================================================================
{gerar_relatorio_gre_texto(estudante)}
""".strip()

# ======================================================
# DATAFRAMES / GRÁFICOS
# ======================================================
def montar_dataframe_evolucao(atendimentos):
    dados = []
    for atendimento in reversed(atendimentos):
        data_atendimento = atendimento[0]
        avancos = atendimento[4]
        dificuldades = atendimento[5]
        evolucao = atendimento[6]
        nivel_resposta = limitar_escala(atendimento[8] if len(atendimento) > 8 else 5)
        nivel_avanco = limitar_escala(atendimento[9] if len(atendimento) > 9 else 5)
        nivel_dificuldade = limitar_escala(atendimento[10] if len(atendimento) > 10 else 5)
        nivel_engajamento = limitar_escala(atendimento[11] if len(atendimento) > 11 else 5)
        indice = calcular_indice_geral(nivel_resposta, nivel_avanco, nivel_dificuldade, nivel_engajamento)

        try:
            data_ordenacao = datetime.strptime(data_atendimento, "%d/%m/%Y")
        except Exception:
            data_ordenacao = None

        dados.append(
            {
                "Data": data_atendimento,
                "Data ordenação": data_ordenacao,
                "Resposta do estudante": nivel_resposta,
                "Avanço pedagógico": nivel_avanco,
                "Nível de dificuldade": nivel_dificuldade,
                "Engajamento": nivel_engajamento,
                "Índice geral de evolução": indice,
                "Interpretação": interpretar_indice(indice),
                "Avanços": avancos or "Não informado.",
                "Dificuldades observadas": dificuldades or "Não informado.",
                "Evolução observada": evolucao or "Não informado.",
            }
        )
    df = pd.DataFrame(dados)
    if not df.empty and "Data ordenação" in df.columns:
        df = df.sort_values(by="Data ordenação", na_position="last")
    return df


def render_grafico_evolucao(atendimentos):
    df = montar_dataframe_evolucao(atendimentos)
    if df.empty:
        st.warning("Este estudante ainda não possui atendimentos registrados.")
        return

    df = df.copy()
    df["Ordem"] = range(1, len(df) + 1)
    contagem_datas = df.groupby("Data").cumcount() + 1
    repetidas = df["Data"].duplicated(keep=False)
    df["Atendimento"] = df["Data"]
    df.loc[repetidas, "Atendimento"] = df.loc[repetidas, "Data"] + " #" + contagem_datas.loc[repetidas].astype(str)

    indicadores = [
        "Resposta do estudante",
        "Avanço pedagógico",
        "Nível de dificuldade",
        "Engajamento",
        "Índice geral de evolução",
    ]
    for col in indicadores:
        df[col] = pd.to_numeric(df[col], errors="coerce").fillna(5).clip(1, 10)

    df_long = df.melt(
        id_vars=["Atendimento", "Data", "Interpretação"],
        value_vars=indicadores,
        var_name="Indicador",
        value_name="Pontuação",
    )

    grafico = (
        alt.Chart(df_long)
        .mark_bar()
        .encode(
            x=alt.X("Atendimento:N", title="Data / atendimento", sort=None),
            xOffset=alt.XOffset("Indicador:N"),
            y=alt.Y("Pontuação:Q", title="Pontuação", scale=alt.Scale(domain=[0, 10])),
            color=alt.Color("Indicador:N", title="Indicador"),
            tooltip=["Data", "Indicador", "Pontuação", "Interpretação"],
        )
        .properties(height=360)
    )
    st.altair_chart(grafico, use_container_width=True)

    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Atendimentos", len(df))
    col2.metric("Média do índice", f"{df['Índice geral de evolução'].mean():.1f}/10")
    col3.metric("Último índice", f"{df['Índice geral de evolução'].iloc[-1]:.1f}/10")
    col4.metric("Variação", f"{df['Índice geral de evolução'].iloc[-1] - df['Índice geral de evolução'].iloc[0]:+.1f}")

    with st.expander("Ver dados usados no gráfico"):
        st.dataframe(df, use_container_width=True, hide_index=True)


# ======================================================
# SIDEBAR
# ======================================================
with st.sidebar:
    try:
        st.markdown('<div class="sidebar-logo-card">', unsafe_allow_html=True)
        st.image(LOGO_PATH, use_container_width=True)
        st.markdown("</div>", unsafe_allow_html=True)
    except Exception:
        st.markdown('<div class="sidebar-title">INCLUISRM</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="sidebar-subtitle">{APP_SUBTITLE}</div>', unsafe_allow_html=True)

    st.markdown("---")

    menu = st.radio(
        "Navegação",
        [
            "Dashboard",
            "Cadastro do Estudante",
            "Cadastro do Professor AEE",
            "Entrevista com a Família",
            "Recursos Pedagógicos e TA da Escola",
            "Avaliação Pedagógica",
            "Articulação Pedagógica Inclusiva",
            "Estudo de Caso",
            "Plano AEE - IA",
            "Agenda de Atendimentos",
            "Atendimentos",
            "Relatórios GRE",
            "Administração",
        ],
        index=0,
        key="menu_atual"  # 🔥 ESSA LINHA É A CHAVE
    )

    st.markdown("---")
    st.caption("LabTec3DI – UFRPE")
    st.caption(APP_SUBTITLE)


render_app_header()

# ======================================================
# DASHBOARD
# ======================================================
if menu == "Dashboard":
    # Dashboard compacto para visualizar o ecossistema em uma única tela.
    st.markdown(
        """
        <style>
        .block-container {
            padding-top: 0.8rem !important;
            padding-bottom: 1.2rem !important;
            max-width: 1550px !important;
        }
        .app-hero {
            padding: 18px 24px !important;
            margin-bottom: 14px !important;
            border-radius: 18px !important;
        }
        .app-title {
            font-size: 30px !important;
        }
        .app-subtitle {
            font-size: 14px !important;
            margin-top: 4px !important;
        }
        .app-badge {
            font-size: 12px !important;
            padding: 4px 10px !important;
            margin-bottom: 6px !important;
        }
        .subtitulo {
            font-size: 21px !important;
            margin-bottom: 8px !important;
        }
        [data-testid="stMetric"] {
            padding: 12px 14px !important;
            min-height: 82px !important;
        }
        [data-testid="stMetricValue"] {
            font-size: 27px !important;
        }
        [data-testid="stMetricLabel"] {
            font-size: 13px !important;
        }
        div[data-testid="stVerticalBlock"] div[data-testid="stVerticalBlockBorderWrapper"] {
            padding-top: 8px !important;
            padding-bottom: 8px !important;
        }
        h3 {
            font-size: 21px !important;
            margin-top: 0.15rem !important;
            margin-bottom: 0.35rem !important;
        }
        hr {
            margin-top: 0.7rem !important;
            margin-bottom: 0.7rem !important;
        }
        .stDataFrame {
            margin-top: 0.2rem !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.markdown('<div class="subtitulo">📊 Painel inicial</div>', unsafe_allow_html=True)

    with st.spinner("Carregando painel..."):
        estudantes = listar_estudantes()
        total_estudantes = len(estudantes)
        total_avaliacoes = contar_registros_tabela("avaliacoes")
        total_atendimentos = contar_registros_tabela("atendimentos")
        total_agenda = contar_registros_tabela("agenda")
        total_escutas_docentes = contar_registros_tabela("escutas_docentes")
        total_relatorios_docente = contar_registros_tabela("relatorios_docente")

        try:
            ult_avaliacoes = carregar_tabela_dataframe("avaliacoes").sort_values("id", ascending=False).head(1)
        except Exception:
            ult_avaliacoes = pd.DataFrame()
        try:
            ult_escutas = carregar_tabela_dataframe("escutas_docentes").sort_values("id", ascending=False).head(1)
        except Exception:
            ult_escutas = pd.DataFrame()
        try:
            ult_relatorios = carregar_tabela_dataframe("relatorios_docente").sort_values("id", ascending=False).head(1)
        except Exception:
            ult_relatorios = pd.DataFrame()

    # Métricas principais em uma única linha para reduzir rolagem.
    col1, col2, col3, col4, col5, col6 = st.columns(6)
    col1.metric("Estudantes", total_estudantes)
    col2.metric("Avaliações", total_avaliacoes)
    col3.metric("Atendimentos", total_atendimentos)
    col4.metric("Agenda", total_agenda)
    col5.metric("Escutas", total_escutas_docentes)
    col6.metric("Relatórios", total_relatorios_docente)

    st.markdown("---")

    # Visão única do ecossistema: memória, articulação, IA e fluxo na mesma tela.
    col_esq, col_centro, col_dir = st.columns([1.15, 1.05, 1.05])

    with col_esq:
        with st.container(border=True):
            st.markdown("### 👥 Estudantes recentes")
            if estudantes:
                st.dataframe(
                    [
                        {
                            "Código": e[1],
                            "Ano/Série": e[2],
                            "Turma": e[3],
                            "Perfil": e[4],
                            "Dias": e[7],
                        }
                        for e in estudantes[:6]
                    ],
                    use_container_width=True,
                    hide_index=True,
                    height=185,
                )
            else:
                st.info("Nenhum estudante cadastrado ainda.")

        with st.container(border=True):
            st.markdown("### 🧠 IA Pedagógica")
            if OpenAI is not None and os.getenv("OPENAI_API_KEY"):
                st.success("IA configurada para gerar sínteses pedagógicas e relatórios de apoio ao docente.")
            else:
                st.warning("IA ainda não configurada. Configure OPENAI_API_KEY para ativar a geração automática.")

    with col_centro:
        with st.container(border=True):
            st.markdown("### 🧭 Fluxo pedagógico sugerido")
            st.markdown(
                """
                1. Cadastro do estudante com código interno.  
                2. Entrevista familiar e avaliação pedagógica inicial.  
                3. Documentos históricos e avaliações extras.  
                4. Estudo de caso e Plano AEE - IA.  
                5. Escuta Docente da sala regular.  
                6. Atendimentos e evolução pedagógica.  
                7. Relatório Pedagógico de Apoio ao Docente.  
                8. Relatórios GRE quando necessário.
                """
            )

    with col_dir:
        with st.container(border=True):
            st.markdown("### 🔗 Articulação Inclusiva")
            st.write(f"Escutas docentes registradas: **{total_escutas_docentes}**")
            st.write(f"Relatórios de apoio gerados: **{total_relatorios_docente}**")
            st.write(f"Avaliações pedagógicas: **{total_avaliacoes}**")
            st.write(f"Atendimentos AEE: **{total_atendimentos}**")

        with st.container(border=True):
            st.markdown("### 🧩 Últimos registros")
            if not ult_avaliacoes.empty:
                st.write(f"**Avaliação:** {ult_avaliacoes.iloc[0].get('data_registro', 'Data não informada')}")
            else:
                st.write("**Avaliação:** nenhuma")

            if not ult_escutas.empty:
                componente = ult_escutas.iloc[0].get('componente_curricular', 'Área não informada')
                codigo_docente = ult_escutas.iloc[0].get('codigo_docente', 'Docente não informado')
                st.write(f"**Escuta:** {componente} • {codigo_docente}")
            else:
                st.write("**Escuta:** nenhuma")

            if not ult_relatorios.empty:
                st.write(f"**Relatório:** {ult_relatorios.iloc[0].get('data_geracao', 'Data não informada')}")
            else:
                st.write("**Relatório:** nenhum")

# ======================================================
# CADASTRO DO ESTUDANTE
# ======================================================
elif menu == "Cadastro do Estudante":
    st.markdown('<div class="subtitulo">👤 Cadastro do estudante</div>', unsafe_allow_html=True)

    col_form, col_lista = st.columns([1, 1.2])
    with col_form:
        with st.container(border=True):
            st.markdown("### Novo estudante")
            with st.form(get_form_key("cadastro_estudante"), clear_on_submit=True):
                codigo = st.text_input("Código interno / matrícula segura", placeholder="Ex.: AEE-2026-001")
                ano_serie = st.text_input("Ano/Série", placeholder="Ex.: 8º ano")
                turma = st.text_input("Turma", placeholder="Ex.: 8º ano A")
                turno = st.selectbox("Turno", ["Não informado", "Manhã", "Tarde", "Noite", "Integral"])
                perfil = st.selectbox("Perfil educacional informado", PERFIS)
                dias = st.multiselect("Dias preferenciais de atendimento", DIAS_SEMANA)
                horario = st.text_input("Horário preferencial", placeholder="Ex.: 08:00 às 08:50")
                observacoes = st.text_area("Observações pedagógicas iniciais")
                enviar = st.form_submit_button("Cadastrar estudante")

                if enviar:
                    if not codigo.strip():
                        st.error("Informe um código interno para o estudante.")
                    else:
                        try:
                            cadastrar_estudante(codigo.strip(), ano_serie, turma, turno, perfil, observacoes, ", ".join(dias), horario)
                            st.success("Estudante cadastrado com sucesso.")
                            resetar_form("cadastro_estudante")
                            st.rerun()

                        except INTEGRITY_ERRORS:
                            st.error("Este código já existe. Use outro código interno.")

    with col_lista:
        with st.container(border=True):
            st.markdown("### 📋 Estudantes cadastrados")
            estudantes = listar_estudantes()
            if estudantes:
                st.dataframe(
                    [
                        {
                            "ID": e[0], "Código": e[1], "Ano/Série": e[2], "Turma": e[3],
                            "Turno": e[6], "Perfil": e[4], "Dias": e[7], "Horário": e[8],
                        }
                        for e in estudantes
                    ],
                    use_container_width=True,
                    hide_index=True,
                )
            else:
                st.info("Nenhum estudante cadastrado ainda.")

    with st.container(border=True):
        st.markdown("### ✏️ Editar, excluir ou exportar cadastro")
        estudantes = listar_estudantes()
        if estudantes:
            ids, mapa = opcoes_estudantes_por_id(estudantes)
            estudante_id = st.selectbox("Selecione o estudante", ids, format_func=lambda x: mapa[x], key="edit_estudante")
            estudante = buscar_estudante(estudante_id)
            with st.expander("Editar cadastro"):
                with st.form(f"form_edit_estudante_{estudante_id}"):
                    codigo_e = st.text_input("Código interno", value=estudante[1] or "")
                    ano_e = st.text_input("Ano/Série", value=estudante[2] or "")
                    turma_e = st.text_input("Turma", value=estudante[3] or "")
                    turno_e = st.selectbox("Turno", ["Não informado", "Manhã", "Tarde", "Noite", "Integral"], index=["Não informado", "Manhã", "Tarde", "Noite", "Integral"].index(estudante[6]) if estudante[6] in ["Não informado", "Manhã", "Tarde", "Noite", "Integral"] else 0)
                    perfil_e = st.selectbox("Perfil", PERFIS, index=PERFIS.index(estudante[4]) if estudante[4] in PERFIS else 0)
                    dias_e = st.text_input("Dias de atendimento", value=estudante[7] or "")
                    horario_e = st.text_input("Horário preferencial", value=estudante[8] or "")
                    obs_e = st.text_area("Observações", value=estudante[5] or "")
                    if st.form_submit_button("Salvar alterações"):
                        atualizar_estudante(estudante_id, codigo_e, ano_e, turma_e, turno_e, perfil_e, obs_e, dias_e, horario_e)
                        st.success("Cadastro atualizado.")
                        st.rerun()

            texto = texto_cadastro_estudante(estudante)
            export_buttons(texto, f"Cadastro_Estudante_{estudante[1]}", tipo_pdf="cadastro")

            confirmar = st.checkbox("Confirmar exclusão do estudante e de todos os registros vinculados", key="conf_excluir_est")
            if st.button("Excluir estudante", key="btn_excluir_est"):
                if confirmar:
                    excluir_estudante(estudante_id)
                    st.success("Estudante excluído.")
                    st.rerun()
                else:
                    st.warning("Marque a confirmação antes de excluir.")


# ======================================================
# CADASTRO DO PROFESSOR AEE
# ======================================================
elif menu == "Cadastro do Professor AEE":
    st.markdown('<div class="subtitulo">👨‍🏫 Cadastro do Professor AEE</div>', unsafe_allow_html=True)
    st.caption("Protótipo personalizado: o sistema considera o professor cadastrado mais recente como professor responsável pelos estudantes e relatórios.")

    if "prof_form_nonce" not in st.session_state:
        st.session_state["prof_form_nonce"] = 0

    nonce = st.session_state["prof_form_nonce"]

    # ===============================
    # FORMULÁRIO DE CADASTRO
    # ===============================
    with st.container(border=True):
        st.markdown("### Novo cadastro profissional")
        st.caption("Ao salvar, a página será atualizada e o formulário aparecerá limpo.")

        with st.form(f"form_professor_{nonce}", clear_on_submit=True):
            nome_ref = st.text_input(
                "Nome de referência / nome profissional",
                key=f"prof_nome_{nonce}",
            )
            escola = st.text_input(
                "Escola",
                key=f"prof_escola_{nonce}",
            )
            regional = st.text_input(
                "Regional / GRE",
                key=f"prof_regional_{nonce}",
            )
            formacao = st.text_area(
                "Formação",
                key=f"prof_formacao_{nonce}",
            )
            carga = st.text_input(
                "Carga horária",
                key=f"prof_carga_{nonce}",
            )
            turno = st.text_input(
                "Turno de atuação",
                key=f"prof_turno_{nonce}",
            )
            st.markdown("#### Perfil pedagógico e tecnológico do professor AEE")
            areas_interesse_prof = st.multiselect(
                "Áreas de atuação/interesse",
                OPCOES_AREAS_INTERESSE_DOCENTE,
                key=f"prof_areas_interesse_{nonce}",
            )
            colp1, colp2 = st.columns(2)
            with colp1:
                nivel_tecnologico_prof = st.selectbox(
                    "Nível de familiaridade tecnológica",
                    OPCOES_NIVEL_TECNOLOGICO_DOCENTE,
                    key=f"prof_nivel_tecnologico_{nonce}",
                )
                modo_maker_prof = st.toggle(
                    "🧩 Ativar modo maker inclusivo",
                    value=False,
                    key=f"prof_modo_maker_{nonce}",
                    help="Quando ativado, a IA pode sugerir cultura maker, robótica, impressão 3D, programação e projetos STEAM de forma progressiva.",
                )
            with colp2:
                interesse_formacao_prof = st.checkbox(
                    "Tenho interesse em aprender tecnologias maker e educacionais",
                    key=f"prof_interesse_formacao_{nonce}",
                )
                projetos_interesse_prof = st.text_input(
                    "Projetos/temas que gostaria de trabalhar",
                    placeholder="Ex.: carrinho solar, CAA, jogos, música, impressão 3D...",
                    key=f"prof_projetos_interesse_{nonce}",
                )
            preferencias_metodologicas_prof = st.text_area(
                "Preferências metodológicas",
                placeholder="Ex.: atividades desplugadas, materiais de baixo custo, jogos, projetos curtos, tecnologia gradual...",
                key=f"prof_preferencias_metodologicas_{nonce}",
            )

            st.markdown("#### 🧰 Recursos pedagógicos do professor")
            st.caption("Registre recursos pessoais, móveis ou produzidos pelo professor que poderão ser utilizados pedagogicamente no AEE. Esses recursos não serão tratados como patrimônio da escola.")
            recursos_professor_sel = st.multiselect(
                "Recursos que o professor possui ou pode levar",
                OPCOES_RECURSOS_PROFESSOR_AEE,
                key=f"prof_recursos_possui_{nonce}",
            )
            recursos_professor_uso = st.text_area(
                "Recursos que pretende utilizar nos atendimentos",
                placeholder="Ex.: tablet, CAA, material de robótica, peças em impressão 3D, materiais impressos em papel...",
                key=f"prof_recursos_uso_{nonce}",
            )
            recursos_professor_obs = st.text_area(
                "Observações sobre uso dos recursos do professor",
                placeholder="Ex.: recursos próprios usados de forma mediada; materiais levados conforme perfil do estudante; necessidade de agendamento/preparo prévio...",
                key=f"prof_recursos_obs_{nonce}",
            )

            obs = st.text_area(
                "Observações",
                key=f"prof_obs_{nonce}",
            )

            salvar = st.form_submit_button("Salvar cadastro do professor")

        if salvar:
            if not nome_ref.strip():
                st.error("Informe pelo menos o nome de referência do professor.")
            else:
                salvar_professor(
                    nome_ref,
                    escola,
                    regional,
                    formacao,
                    carga,
                    turno,
                    obs,
                    areas_interesse=", ".join(areas_interesse_prof),
                    nivel_tecnologico=nivel_tecnologico_prof,
                    modo_maker="Sim" if modo_maker_prof else "Não",
                    interesse_formacao_maker="Sim" if interesse_formacao_prof else "Não",
                    projetos_interesse=projetos_interesse_prof,
                    preferencias_metodologicas=preferencias_metodologicas_prof,
                    recursos_professor=", ".join(recursos_professor_sel),
                    recursos_professor_uso=recursos_professor_uso,
                    recursos_professor_observacoes=recursos_professor_obs,
                )
                st.success("Cadastro do professor salvo com sucesso.")
                st.session_state["prof_form_nonce"] += 1
                st.rerun()

    # ===============================
    # HISTÓRICO DO PROFESSOR
    # ===============================
    with st.container(border=True):
        st.markdown("### Histórico")

        professores = listar_professores()

        if professores:
            st.dataframe(
                [
                    {
                        "ID": p[0],
                        "Nome referência": p[1],
                        "Escola": p[2],
                        "Regional": p[3],
                        "Carga horária": p[5],
                        "Turno": p[6],
                        "Responsável atual": "Sim" if i == 0 else "Não",
                    }
                    for i, p in enumerate(professores)
                ],
                use_container_width=True,
                hide_index=True,
            )

            professor_atual = professores[0]
            st.success(
                f"Professor responsável atual nos documentos: {professor_atual[1] or 'Sem nome'}"
            )

            for p in professores:
                rotulo = "responsável atual" if p[0] == professor_atual[0] else "registro anterior"
                with st.expander(f"Professor(a) #{p[0]} - {p[1] or 'Sem nome'} ({rotulo})"):
                    texto = texto_professor(p)
                    st.text(texto)
                    export_buttons(texto, f"Ficha_Professor_AEE_{p[0]}", tipo_pdf="professor")

                    st.markdown("---")
                    st.markdown("### ✏️ Editar cadastro")

                    with st.form(f"edit_professor_{p[0]}"):
                        nome_e = st.text_input("Nome referência", value=p[1] or "", key=f"edit_prof_nome_{p[0]}")
                        escola_e = st.text_input("Escola", value=p[2] or "", key=f"edit_prof_escola_{p[0]}")
                        regional_e = st.text_input("Regional/GRE", value=p[3] or "", key=f"edit_prof_regional_{p[0]}")
                        formacao_e = st.text_area("Formação", value=p[4] or "", key=f"edit_prof_formacao_{p[0]}")
                        carga_e = st.text_input("Carga horária", value=p[5] or "", key=f"edit_prof_carga_{p[0]}")
                        turno_e = st.text_input("Turno", value=p[6] or "", key=f"edit_prof_turno_{p[0]}")
                        areas_e = st.multiselect(
                            "Áreas de atuação/interesse",
                            OPCOES_AREAS_INTERESSE_DOCENTE,
                            default=[x.strip() for x in (p[9] if len(p) > 9 and p[9] else "").split(",") if x.strip() in OPCOES_AREAS_INTERESSE_DOCENTE],
                            key=f"edit_prof_areas_{p[0]}",
                        )
                        nivel_e = st.selectbox(
                            "Nível de familiaridade tecnológica",
                            OPCOES_NIVEL_TECNOLOGICO_DOCENTE,
                            index=OPCOES_NIVEL_TECNOLOGICO_DOCENTE.index(p[10]) if len(p) > 10 and p[10] in OPCOES_NIVEL_TECNOLOGICO_DOCENTE else 0,
                            key=f"edit_prof_nivel_{p[0]}",
                        )
                        modo_maker_e = st.toggle(
                            "🧩 Ativar modo maker inclusivo",
                            value=(len(p) > 11 and p[11] == "Sim"),
                            key=f"edit_prof_maker_{p[0]}",
                        )
                        interesse_formacao_e = st.checkbox(
                            "Tenho interesse em aprender tecnologias maker e educacionais",
                            value=(len(p) > 12 and p[12] == "Sim"),
                            key=f"edit_prof_formacao_maker_{p[0]}",
                        )
                        projetos_e = st.text_input(
                            "Projetos/temas que gostaria de trabalhar",
                            value=p[13] if len(p) > 13 and p[13] else "",
                            key=f"edit_prof_projetos_{p[0]}",
                        )
                        preferencias_e = st.text_area(
                            "Preferências metodológicas",
                            value=p[14] if len(p) > 14 and p[14] else "",
                            key=f"edit_prof_preferencias_{p[0]}",
                        )
                        recursos_e = st.multiselect(
                            "Recursos que o professor possui ou pode levar",
                            OPCOES_RECURSOS_PROFESSOR_AEE,
                            default=[x.strip() for x in (p[15] if len(p) > 15 and p[15] else "").split(",") if x.strip() in OPCOES_RECURSOS_PROFESSOR_AEE],
                            key=f"edit_prof_recursos_{p[0]}",
                        )
                        recursos_uso_e = st.text_area(
                            "Recursos que pretende utilizar nos atendimentos",
                            value=p[16] if len(p) > 16 and p[16] else "",
                            key=f"edit_prof_recursos_uso_{p[0]}",
                        )
                        recursos_obs_e = st.text_area(
                            "Observações sobre uso dos recursos do professor",
                            value=p[17] if len(p) > 17 and p[17] else "",
                            key=f"edit_prof_recursos_obs_{p[0]}",
                        )
                        obs_e = st.text_area("Observações", value=p[7] or "", key=f"edit_prof_obs_{p[0]}")

                        if st.form_submit_button("💾 Atualizar cadastro"):
                            atualizar_professor(
                                p[0],
                                nome_e,
                                escola_e,
                                regional_e,
                                formacao_e,
                                carga_e,
                                turno_e,
                                obs_e,
                                areas_interesse=", ".join(areas_e),
                                nivel_tecnologico=nivel_e,
                                modo_maker="Sim" if modo_maker_e else "Não",
                                interesse_formacao_maker="Sim" if interesse_formacao_e else "Não",
                                projetos_interesse=projetos_e,
                                preferencias_metodologicas=preferencias_e,
                                recursos_professor=", ".join(recursos_e),
                                recursos_professor_uso=recursos_uso_e,
                                recursos_professor_observacoes=recursos_obs_e,
                            )
                            st.success("Cadastro atualizado com sucesso.")
                            st.rerun()

                    st.markdown("---")
                    st.markdown("### 🗑️ Excluir professor")
                    st.caption("Ao excluir o professor, o sistema passará a considerar o cadastro mais recente restante como responsável.")

                    confirmar_exclusao = st.checkbox(
                        "Confirmar exclusão deste professor",
                        key=f"confirmar_excluir_prof_{p[0]}",
                    )

                    if st.button("Excluir professor", key=f"excluir_professor_{p[0]}"):
                        if confirmar_exclusao:
                            excluir_professor(p[0])
                            st.warning("Professor excluído com sucesso.")
                            st.rerun()
                        else:
                            st.warning("Marque a confirmação antes de excluir.")
        else:
            st.info("Nenhum professor cadastrado.")
            st.warning("Cadastre seus dados profissionais para que eles apareçam nos relatórios e documentos.")


# ENTREVISTA COM A FAMÍLIA
# ======================================================
elif menu == "Entrevista com a Família":
    st.markdown('<div class="subtitulo">👪 Entrevista com a Família</div>', unsafe_allow_html=True)
    estudantes = listar_estudantes()

    if not estudantes:
        st.info("Cadastre um estudante primeiro.")
    else:
        ids, mapa = opcoes_estudantes_por_id(estudantes)
        estudante_id = st.selectbox(
            "Selecione o estudante",
            ids,
            format_func=lambda x: mapa[x],
            key="entrevista_estudante",
        )
        estudante = buscar_estudante(estudante_id)

        with st.container(border=True):
            st.markdown("### Nova entrevista com a família")
            st.caption("Dados organizados conforme o roteiro de entrevista familiar do AEE. Campos sensíveis continuam para preenchimento manual nos documentos impressos.")

            with st.form("form_entrevista_completa"):
                ano_padrao_entrevista = str(agora_local().year)
                anos_opcoes_entrevista = [str(a) for a in range(2020, agora_local().year + 4)]
                idx_ano_entrevista = anos_opcoes_entrevista.index(ano_padrao_entrevista) if ano_padrao_entrevista in anos_opcoes_entrevista else 0

                st.markdown("### 📅 Identificação do registro")
                col_ano_ent, col_tipo_ent = st.columns(2)
                with col_ano_ent:
                    ano_letivo_entrevista = st.selectbox(
                        "Ano letivo da entrevista familiar",
                        anos_opcoes_entrevista,
                        index=idx_ano_entrevista,
                        key="entrevista_ano_letivo",
                    )
                with col_tipo_ent:
                    tipo_registro_entrevista = st.selectbox(
                        "Tipo de registro",
                        OPCOES_TIPO_ENTREVISTA_FAMILIA,
                        key="entrevista_tipo_registro",
                    )
                st.caption(
                    "Use 'Registro histórico' para lançar entrevistas de anos anteriores e 'Entrevista familiar atual' para o ano em acompanhamento."
                )

                st.markdown("### 🧾 1. Informações diversas")
                col1, col2 = st.columns(2)
                with col1:
                    auxilio_governamental = st.radio("A família participa de programa de auxílio governamental?", ["Não", "Sim"], horizontal=True)
                    auxilio_quais = st.text_input("Qual(is) programa(s)?")
                    historico_familiar = st.radio("Há histórico de doenças graves, deficiência ou transtornos na família?", ["Não", "Sim"], horizontal=True)
                    historico_quais = st.text_input("Qual(is) histórico(s)?")
                    repetiu_ano = st.radio("O estudante já repetiu de ano?", ["Não", "Sim"], horizontal=True)
                    repetiu_qtd = st.number_input("Quantas vezes repetiu?", min_value=0, max_value=20, value=0)
                with col2:
                    trocou_escola = st.radio("O estudante trocou de escola?", ["Não", "Sim"], horizontal=True)
                    trocou_qtd = st.number_input("Quantas vezes trocou de escola?", min_value=0, max_value=20, value=0)
                    motivo_troca = st.text_area("Por qual motivo trocou de escola?")
                    situacao_escolar = st.radio("Em relação à escola, o estudante é:", ["Não informado", "Novato(a)", "Veterano(a)"], horizontal=True)

                st.markdown("### 🏫 Relação com a escola")
                col1, col2, col3 = st.columns(3)
                with col1:
                    interesse_escola = st.radio("Demonstra interesse em frequentar a escola?", ["Não", "Sim", "Não informado"], horizontal=True)
                    organiza_materiais = st.radio("Cuida/organiza seus materiais?", ["Não", "Sim", "Não informado"], horizontal=True)
                    resistencia_escola = st.radio("Apresenta resistência à escola?", ["Não", "Sim", "Não informado"], horizontal=True)
                with col2:
                    relacao_colegas = st.radio("Relaciona-se bem com colegas?", ["Não", "Sim", "Não informado"], horizontal=True)
                    relacao_professores = st.radio("Relaciona-se bem com professores?", ["Não", "Sim", "Não informado"], horizontal=True)
                with col3:
                    leva_alimentacao = st.radio("Leva alimentação de casa?", ["Não", "Sim", "Não informado"], horizontal=True)
                    merenda_escolar = st.radio("Alimenta-se da merenda escolar?", ["Não", "Sim", "Não informado"], horizontal=True)
                    alergia_alimentar = st.radio("Possui alergia alimentar?", ["Não", "Sim", "Não informado"], horizontal=True)
                    alergia_quais = st.text_input("Qual(is) alergia(s)?")
                obs_diversas = st.text_area("Outras observações relevantes sobre escola/família")

                st.markdown("### 🏫 1.1 Sobre a escolha da escola")
                motivo_escolha_lista = st.multiselect(
                    "Motivo da escolha da escola",
                    [
                        "Proximidade de casa",
                        "Tem o serviço do AEE",
                        "Irmãos matriculados na unidade educacional",
                        "Escola com boas referências",
                        "A unidade tem EJAI / estudante fora de faixa",
                        "Outro",
                    ],
                )
                motivo_escolha = ", ".join(motivo_escolha_lista)
                outros_motivos = st.text_area("Outros motivos da escolha")
                conhecimento_aee = st.text_area("O que a família conhece sobre o serviço do AEE?")

                st.markdown("### 🩺 2. Informações sobre a saúde do estudante")
                col1, col2, col3 = st.columns(3)
                with col1:
                    doenca_preexistente = st.radio("Possui doença preexistente?", ["Não", "Sim", "Não informado"], horizontal=True)
                    convulsoes = st.radio("Apresenta convulsões?", ["Não", "Sim", "Não informado"], horizontal=True)
                    acompanhamentos_lista = st.multiselect(
                        "Possui acompanhamento de",
                        ["Fonoaudiólogo", "Terapeuta Ocupacional", "Fisioterapeuta", "Psicólogo", "Outro"],
                    )
                    acompanhamentos = ", ".join(acompanhamentos_lista)
                    acompanhamento_outro = st.text_input("Outro acompanhamento")
                with col2:
                    frequencia_acompanhamento = st.selectbox("Frequência dos acompanhamentos", ["Não informado", "1x na semana", "2x na semana", "Quinzenalmente", "Outro"])
                    frequencia_outro = st.text_input("Outra frequência")
                    alimentacao_saudavel = st.radio("Tem alimentação saudável?", ["Não", "Sim", "Não informado"], horizontal=True)
                    seletividade_alimentar = st.radio("Possui seletividade alimentar?", ["Não", "Sim", "Não informado"], horizontal=True)
                    dieta_sensorial = st.radio("Possui dieta sensorial?", ["Não", "Sim", "Não informado"], horizontal=True)
                with col3:
                    suplemento_alimentar = st.radio("Usa suplemento alimentar?", ["Não", "Sim", "Não informado"], horizontal=True)
                    suplemento_qual = st.text_input("Qual suplemento?")
                    alimenta_sonda = st.radio("Alimenta-se por sonda?", ["Não", "Sim", "Não informado"], horizontal=True)
                    dorme_bem = st.radio("Dorme bem?", ["Não", "Sim", "Não informado"], horizontal=True)
                    medicacao = st.radio("Faz uso de medicação?", ["Não", "Sim", "Não informado"], horizontal=True)
                    medicacao_qual = st.text_input("Qual(is) medicação(ões)?")
                tempo_medicacao_tratamentos = st.text_area("Há quanto tempo usa medicação? Quais tratamentos realizados?")
                obs_saude = st.text_area("Outras observações relevantes sobre saúde")

                st.markdown("### 🧠 3. Desenvolvimento psicomotor")
                lateralidade = st.radio("Lateralidade", ["Não informado", "Destro(a)", "Canhoto(a)", "Ambidestro(a)"], horizontal=True)
                col1, col2 = st.columns(2)
                with col1:
                    estereotipias_lista = st.multiselect(
                        "Estereotipias observadas",
                        [
                            "Não apresenta",
                            "Balançar o corpo para frente e para trás",
                            "Balançar as mãos / flapping",
                            "Girar objetos ou girar em volta do próprio corpo",
                            "Fazer sons repetitivos ou repetir sílabas/palavras",
                            "Estalar os dedos",
                            "Pular",
                            "Correr indo e vindo sem destino claro",
                            "Andar na ponta dos pés",
                            "Movimentar os dedos na frente dos olhos",
                        ],
                    )
                    estereotipias = ", ".join(estereotipias_lista)
                    estereotipias_quais = st.text_area("Outras estereotipias / descrição")
                    segura_objetos_duas_maos = st.radio("Segura objetos com as duas mãos?", ["Não", "Sim", "Não informado"], horizontal=True)
                    tamanho_objetos_lista = st.multiselect("Segura objetos", ["Pequenos", "Médios", "Grandes"])
                    tamanho_objetos = ", ".join(tamanho_objetos_lista)
                    pega_lapis = st.radio("Faz a pega correta do lápis?", ["Não", "Sim", "Não informado"], horizontal=True)
                with col2:
                    engatinhou = st.radio("Engatinhou?", ["Não", "Sim", "Não informado"], horizontal=True)
                    idade_andou = st.text_input("Andou com que idade?")
                    usa_fraldas = st.radio("Usa fraldas na escola?", ["Não", "Sim", "Não informado"], horizontal=True)
                    usa_sonda_alivio = st.radio("Usa sonda de alívio?", ["Não", "Sim", "Não informado"], horizontal=True)
                    autonomia_lista = st.multiselect("O que consegue fazer sem ajuda?", ["Tomar banho", "Escovar os dentes", "Usar o banheiro", "Se alimentar", "Se vestir", "Outros"])
                    autonomia_atividades = ", ".join(autonomia_lista)
                    autonomia_outros = st.text_input("Outras atividades de autonomia")
                    atende_comandos = st.radio("Atende comandos?", ["Não", "Sim", "Não informado"], horizontal=True)
                    gosta_toque = st.radio("Gosta do toque?", ["Não", "Sim", "Não informado"], horizontal=True)
                obs_psicomotor = st.text_area("Outras observações relevantes sobre desenvolvimento psicomotor")

                st.markdown("### 🗣️ 4. Linguagem")
                col1, col2, col3 = st.columns(3)
                with col1:
                    verbal = st.radio("É verbal?", ["Não", "Sim", "Não informado"], horizontal=True)
                    consegue_comunicar = st.radio("Consegue se comunicar?", ["Não", "Sim", "Não informado"], horizontal=True)
                    problemas_fala = st.radio("Possui problemas na fala?", ["Não", "Sim", "Não informado"], horizontal=True)
                with col2:
                    ecolalia = st.radio("Tem ecolalia?", ["Não", "Sim", "Não informado"], horizontal=True)
                    da_recado = st.radio("Consegue dar um recado?", ["Não", "Sim", "Não informado"], horizontal=True)
                with col3:
                    comunicacao_alternativa = st.radio("Usa comunicação alternativa?", ["Não", "Sim", "Não informado"], horizontal=True)
                    comunicacao_alternativa_qual = st.text_input("Qual comunicação alternativa?")

                st.markdown("### 🤝 5. Socialização")
                col1, col2, col3 = st.columns(3)
                with col1:
                    relacao_pai = st.radio("Relaciona-se bem com o pai?", ["Não", "Sim", "Não informado"], horizontal=True)
                    relacao_mae = st.radio("Relaciona-se bem com a mãe?", ["Não", "Sim", "Não informado"], horizontal=True)
                    relacao_parentes = st.radio("Relaciona-se bem com outros parentes?", ["Não", "Sim", "Não informado"], horizontal=True)
                    relacao_irmaos = st.radio("Relaciona-se bem com irmãos?", ["Não", "Sim", "Não informado"], horizontal=True)
                with col2:
                    relacao_estudantes = st.radio("Relaciona-se bem com outros estudantes?", ["Não", "Sim", "Não informado"], horizontal=True)
                    tem_melhor_amigo = st.radio("Tem melhor amigo(a)?", ["Não", "Sim", "Não informado"], horizontal=True)
                    tipo_melhor_amigo = st.selectbox("Esse(a) melhor amigo(a) é", ["Não informado", "Parente", "Colega de escola", "Colega do local onde mora"])
                    adapta_ambiente = st.radio("Adapta-se facilmente ao ambiente?", ["Não", "Sim", "Não informado"], horizontal=True)
                with col3:
                    flexivel_rotina = st.radio("É flexível na rotina?", ["Não", "Sim", "Não informado"], horizontal=True)
                    respeita_regras = st.radio("Respeita regras?", ["Não", "Sim", "Não informado"], horizontal=True)
                    chora_facilidade = st.radio("Chora com facilidade?", ["Não", "Sim", "Não informado"], horizontal=True)
                    brinca_como = st.radio("Gosta de brincar", ["Não informado", "Sozinho(a)", "Com outros", "Das duas formas"], horizontal=True)
                interesses_lazer = st.text_area("Assunto ou lazer que interessa ao estudante")
                familia_gosta = st.text_area("O que a família mais gosta nesse(a) filho(a)?")
                familia_nao_gosta = st.text_area("O que a família não gosta/necesita melhorar nele(a)?")
                ambiente_estudo_casa = st.text_area("Em casa existe ambiente físico para atividades escolares e brincadeiras?")

                st.markdown("### 🏠 6. Contexto familiar")
                habilidades = st.text_area("Principais habilidades: o que faz bem")
                oportunidades_melhoria = st.text_area("Principais oportunidades de melhoria: o que podemos estimular mais")

                st.markdown("### 📝 7. Outras informações")
                outras_info_familia = st.text_area("Outras informações que a família considera importante registrar")

                salvar_entrevista = st.form_submit_button("Salvar entrevista com a família")

                if salvar_entrevista:
                    inserir_registro(
                        "entrevistas_familia",
                        ["estudante_id"] + CAMPOS_ENTREVISTA_FAMILIA,
                        [
                            estudante_id, hoje_str(), ano_letivo_entrevista, tipo_registro_entrevista,
                            auxilio_governamental, auxilio_quais, historico_familiar, historico_quais,
                            repetiu_ano, str(repetiu_qtd), trocou_escola, str(trocou_qtd), motivo_troca,
                            situacao_escolar, interesse_escola, organiza_materiais, resistencia_escola,
                            relacao_colegas, relacao_professores, leva_alimentacao, merenda_escolar,
                            alergia_alimentar, alergia_quais, obs_diversas,
                            motivo_escolha, outros_motivos, conhecimento_aee,
                            doenca_preexistente, convulsoes, acompanhamentos, acompanhamento_outro,
                            frequencia_acompanhamento, frequencia_outro, alimentacao_saudavel,
                            seletividade_alimentar, dieta_sensorial, suplemento_alimentar, suplemento_qual,
                            alimenta_sonda, dorme_bem, medicacao, medicacao_qual,
                            tempo_medicacao_tratamentos, obs_saude,
                            lateralidade, estereotipias, estereotipias_quais, segura_objetos_duas_maos,
                            tamanho_objetos, pega_lapis, engatinhou, idade_andou,
                            usa_fraldas, usa_sonda_alivio, autonomia_atividades, autonomia_outros,
                            atende_comandos, gosta_toque, obs_psicomotor,
                            verbal, consegue_comunicar, problemas_fala, ecolalia, da_recado,
                            comunicacao_alternativa, comunicacao_alternativa_qual,
                            relacao_pai, relacao_mae, relacao_parentes, relacao_irmaos, relacao_estudantes,
                            tem_melhor_amigo, tipo_melhor_amigo, adapta_ambiente, flexivel_rotina,
                            respeita_regras, chora_facilidade, brinca_como, interesses_lazer,
                            familia_gosta, familia_nao_gosta, ambiente_estudo_casa,
                            habilidades, oportunidades_melhoria, outras_info_familia,
                        ],
                    )
                    st.success("Entrevista com a família salva com sucesso.")
                    st.rerun()

        registros = listar_por_estudante("entrevistas_familia", CAMPOS_ENTREVISTA_FAMILIA, estudante_id)

        with st.container(border=True):
            st.markdown("### Histórico de entrevistas")
            if registros:
                for r in registros:
                    ano_reg = r[2] if len(r) > 2 and r[2] else "Ano não informado"
                    tipo_reg = r[3] if len(r) > 3 and r[3] else "Registro"
                    with st.expander(f"Entrevista em {r[1]} | {ano_reg} | {tipo_reg}"):
                        texto = texto_entrevista(estudante, r)
                        st.text(texto)
                        export_buttons(texto, f"Entrevista_Familia_{estudante[1]}_{r[0]}", tipo_pdf="entrevista")
                        if st.button("Excluir entrevista", key=f"exc_ent_{r[0]}"):
                            excluir_registro("entrevistas_familia", r[0])
                            st.success("Entrevista excluída.")
                            st.rerun()
            else:
                st.info("Nenhuma entrevista registrada.")


# ======================================================
# AVALIAÇÃO PEDAGÓGICA
# ======================================================
elif menu == "Avaliação Pedagógica":
    st.markdown('<div class="subtitulo">📝 Avaliação Pedagógica / Histórico por Ano Letivo</div>', unsafe_allow_html=True)
    estudantes = listar_estudantes()

    if not estudantes:
        st.info("Cadastre um estudante primeiro.")
    else:
        ids, mapa = opcoes_estudantes_por_id(estudantes)
        estudante_id = st.selectbox(
            "Selecione o estudante",
            ids,
            format_func=lambda x: mapa[x],
            key="avaliacao_estudante",
        )
        estudante = buscar_estudante(estudante_id)
        avaliacoes = listar_avaliacoes(estudante_id)

        st.info(
            "Este módulo permite registrar avaliações pedagógicas estruturadas, avaliações anteriores, "
            "documentos livres e anexos digitalizados. Use sempre o código interno do estudante e evite "
            "inserir dados pessoais sensíveis nos textos ou arquivos."
        )

        anos_opcoes = [str(a) for a in range(2024, agora_local().year + 4)]
        ano_padrao = str(agora_local().year)
        idx_ano = anos_opcoes.index(ano_padrao) if ano_padrao in anos_opcoes else 0

        with st.container(border=True):
            st.markdown("### Configuração do registro")
            col_ano, col_tipo = st.columns(2)

            with col_ano:
                ano_letivo = st.selectbox(
                    "Ano letivo da avaliação pedagógica",
                    anos_opcoes,
                    index=idx_ano,
                    key="avaliacao_ano_letivo",
                )

            with col_tipo:
                tipo_registro = st.selectbox(
                    "Tipo de registro",
                    OPCOES_TIPO_AVALIACAO,
                    key="avaliacao_tipo_registro",
                )

            modo_extra = tipo_registro == "Avaliação Pedagógica - Extra / Documento Livre"

            opcoes_avaliacoes = {0: "Nenhuma avaliação anterior"}
            for av in avaliacoes:
                ano_av = av[2] or "Ano não informado"
                tipo_av = av[3] or "Tipo não informado"
                opcoes_avaliacoes[av[0]] = f"ID {av[0]} • {ano_av} • {tipo_av} • {av[1]}"

            avaliacao_anterior_id = st.selectbox(
                "Avaliação anterior para comparação ou vínculo",
                list(opcoes_avaliacoes.keys()),
                format_func=lambda x: opcoes_avaliacoes[x],
                key="avaliacao_anterior_id_select",
            )

            if not modo_extra:
                foco_ia = st.text_area(
                    "Orientação para a IA (opcional)",
                    placeholder="Ex.: observar avanços em autonomia, comunicação, participação nas atividades e barreiras de aprendizagem.",
                    height=80,
                    key="avaliacao_foco_ia",
                )

                if st.button("🤖 Gerar sugestão de nova Avaliação Pedagógica com IA", key="btn_ia_avaliacao"):
                    av_base = buscar_avaliacao_por_id(avaliacao_anterior_id) if avaliacao_anterior_id else None
                    estudo_atual = ultima_linha("estudos_caso", CAMPOS_ESTUDO_CASO, estudante_id)
                    estudo_atual = ("", *estudo_atual) if estudo_atual else None
                    entrevista_atual = ultima_linha("entrevistas_familia", CAMPOS_ENTREVISTA_FAMILIA, estudante_id)
                    atendimentos_atuais = listar_atendimentos(estudante_id)
                    analise_ia, sugestao_ia = gerar_nova_avaliacao_pedagogica_com_ia(
                        estudante,
                        av_base,
                        estudo_atual,
                        entrevista_atual,
                        atendimentos_atuais,
                        ano_letivo,
                    )
                    if foco_ia.strip():
                        analise_ia = f"Orientação do professor: {foco_ia.strip()}\n\n{analise_ia}".strip()
                    st.session_state["avaliacao_analise_ia"] = analise_ia
                    st.session_state["avaliacao_sugestao_ia"] = sugestao_ia
                    st.success("Sugestão gerada. Revise e ajuste manualmente antes de salvar.")

                analise_previa = st.session_state.get("avaliacao_analise_ia", "")
                sugestao_previa = st.session_state.get("avaliacao_sugestao_ia", "")
                if analise_previa or sugestao_previa:
                    with st.expander("Prévia da análise e sugestão geradas com IA", expanded=True):
                        st.markdown("**Análise comparativa:**")
                        st.text_area("Análise comparativa IA", value=analise_previa, height=180, key="preview_analise_avaliacao_ia")
                        st.markdown("**Sugestão de nova avaliação:**")
                        st.text_area("Sugestão IA", value=sugestao_previa, height=260, key="preview_sugestao_avaliacao_ia")

        with st.container(border=True):
            if modo_extra:
                st.markdown("### Avaliação Pedagógica Extra / Documento Livre")
                st.caption(
                    "Use este espaço para registrar documentos antigos, pareceres, relatórios anteriores "
                    "ou informações pedagógicas que não cabem no formulário estruturado."
                )

                with st.form("form_avaliacao_extra"):
                    origem_documento = st.text_input(
                        "Origem do documento",
                        placeholder="Ex.: relatório antigo da escola, parecer pedagógico anterior, registro do AEE anterior...",
                    )

                    texto_documento_extra = st.text_area(
                        "Conteúdo completo do documento ou relatório",
                        height=420,
                        placeholder="Cole aqui o conteúdo do documento antigo, relatório, parecer ou histórico pedagógico anterior.",
                    )

                    arquivo_upload = st.file_uploader(
                        "Anexar documento digitalizado ou arquivo relacionado",
                        type=["pdf", "docx", "txt", "png", "jpg", "jpeg"],
                        accept_multiple_files=False,
                    )

                    observacao_upload = st.text_area(
                        "Observação sobre o arquivo anexado",
                        height=90,
                        placeholder="Ex.: documento físico arquivado em pasta protegida; arquivo sem dados sensíveis; digitalização parcial...",
                    )

                    salvar_extra = st.form_submit_button("Salvar Avaliação Pedagógica Extra")

                if salvar_extra:
                    if not texto_documento_extra.strip() and arquivo_upload is None:
                        st.warning("Digite o conteúdo do documento ou anexe um arquivo antes de salvar.")
                    else:
                        salvar_avaliacao(
                            estudante_id=estudante_id,
                            ano_letivo=ano_letivo,
                            tipo_registro=tipo_registro,
                            avaliacao_anterior_id=None if avaliacao_anterior_id == 0 else avaliacao_anterior_id,
                            origem_documento=origem_documento,
                            texto_documento_extra=texto_documento_extra,
                        )

                        avaliacoes_atualizadas = listar_avaliacoes(estudante_id)
                        avaliacao_id = avaliacoes_atualizadas[0][0] if avaliacoes_atualizadas else None

                        if arquivo_upload is not None:
                            salvar_documento_avaliacao(
                                estudante_id=estudante_id,
                                avaliacao_id=avaliacao_id,
                                ano_letivo=ano_letivo,
                                tipo_documento=tipo_registro,
                                arquivo=arquivo_upload,
                                observacao=observacao_upload,
                            )

                        st.success("Avaliação Pedagógica Extra salva com sucesso.")
                        st.rerun()

            else:
                st.markdown("### Registro manual / revisão da avaliação")

                with st.form("form_avaliacao"):
                    sugestao_base = st.session_state.get("avaliacao_sugestao_ia", "")
                    barreiras = st.text_area(
                        "Barreiras enfrentadas pelo estudante",
                        value=sugestao_base if tipo_registro == "Nova avaliação com base em avaliação anterior" else "",
                        height=120,
                    )
                    potencialidades = st.text_area("Potencialidades e habilidades já desenvolvidas", height=100)
                    comunicacao = st.text_area("Comunicação", height=90)
                    interacao = st.text_area("Interação social", height=90)
                    autonomia = st.text_area("Autonomia", height=90)
                    aprendizagem = st.text_area("Aprendizagem", height=100)
                    resumo_laudo = st.text_area("Resumo pedagógico do laudo, sem identificação", height=90)

                    analise_comparativa_ia = st.text_area(
                        "Análise comparativa IA/manual",
                        value=st.session_state.get("avaliacao_analise_ia", ""),
                        height=150,
                    )
                    sugestao_nova_avaliacao_ia = st.text_area(
                        "Sugestão IA/manual para nova avaliação",
                        value=st.session_state.get("avaliacao_sugestao_ia", ""),
                        height=180,
                    )

                    arquivo_upload = st.file_uploader(
                        "Anexar documento complementar da avaliação",
                        type=["pdf", "docx", "txt", "png", "jpg", "jpeg"],
                        accept_multiple_files=False,
                    )

                    observacao_upload = st.text_area(
                        "Observação sobre o arquivo anexado",
                        height=80,
                        placeholder="Ex.: documento complementar, imagem digitalizada ou parecer usado como apoio.",
                    )

                    salvar_estruturada = st.form_submit_button("Salvar avaliação pedagógica")

                if salvar_estruturada:
                    salvar_avaliacao(
                        estudante_id=estudante_id,
                        barreiras=barreiras,
                        potencialidades=potencialidades,
                        comunicacao=comunicacao,
                        interacao=interacao,
                        autonomia=autonomia,
                        aprendizagem=aprendizagem,
                        resumo_laudo=resumo_laudo,
                        ano_letivo=ano_letivo,
                        tipo_registro=tipo_registro,
                        avaliacao_anterior_id=None if avaliacao_anterior_id == 0 else avaliacao_anterior_id,
                        analise_comparativa_ia=analise_comparativa_ia,
                        sugestao_nova_avaliacao_ia=sugestao_nova_avaliacao_ia,
                    )

                    avaliacoes_atualizadas = listar_avaliacoes(estudante_id)
                    avaliacao_id = avaliacoes_atualizadas[0][0] if avaliacoes_atualizadas else None

                    if arquivo_upload is not None:
                        salvar_documento_avaliacao(
                            estudante_id=estudante_id,
                            avaliacao_id=avaliacao_id,
                            ano_letivo=ano_letivo,
                            tipo_documento=tipo_registro,
                            arquivo=arquivo_upload,
                            observacao=observacao_upload,
                        )

                    st.session_state.pop("avaliacao_analise_ia", None)
                    st.session_state.pop("avaliacao_sugestao_ia", None)

                    st.success("Avaliação pedagógica salva com sucesso.")
                    st.rerun()

        with st.container(border=True):
            st.markdown("### Histórico de avaliações pedagógicas")
            avaliacoes = listar_avaliacoes(estudante_id)

            if avaliacoes:
                for av in avaliacoes:
                    dados_av = dict(zip(["id"] + CAMPOS_AVALIACAO, av))
                    av_id = dados_av.get("id")
                    data_reg = dados_av.get("data_registro") or "Data não informada"
                    ano_reg = dados_av.get("ano_letivo") or "Ano não informado"
                    tipo_reg = dados_av.get("tipo_registro") or "Registro"

                    with st.expander(f"Avaliação em {data_reg} | {ano_reg} | {tipo_reg}"):
                        texto = texto_avaliacao(estudante, av)
                        st.text(texto)
                        export_buttons(texto, f"Avaliacao_Pedagogica_{estudante[1]}_{av_id}", tipo_pdf="avaliacao")

                        if st.button("Excluir avaliação", key=f"exc_avaliacao_{av_id}"):
                            excluir_registro("avaliacoes", av_id)
                            st.success("Avaliação excluída.")
                            st.rerun()
            else:
                st.info("Nenhuma avaliação pedagógica registrada.")

        with st.container(border=True):
            st.markdown("### Documentos anexados à Avaliação Pedagógica")
            documentos = listar_documentos_avaliacao(estudante_id)

            if documentos:
                for doc in documentos:
                    doc_id, avaliacao_id, ano_doc, tipo_doc, nome_original, caminho_arquivo, observacao, data_upload = doc

                    with st.expander(f"{nome_original} | {ano_doc or 'Ano não informado'} | {data_upload}"):
                        st.write(f"**Tipo:** {tipo_doc or 'Não informado'}")
                        st.write(f"**Avaliação vinculada:** {avaliacao_id or 'Não vinculada'}")
                        st.write(f"**Observação:** {observacao or 'Sem observação.'}")

                        if caminho_arquivo and Path(caminho_arquivo).exists():
                            with open(caminho_arquivo, "rb") as f:
                                st.download_button(
                                    "Baixar documento",
                                    data=f,
                                    file_name=nome_original,
                                    key=f"baixar_doc_avaliacao_{doc_id}",
                                )
                        else:
                            st.warning("Arquivo não encontrado na pasta de documentos do sistema.")

                        if st.button("Excluir documento anexado", key=f"excluir_doc_avaliacao_{doc_id}"):
                            excluir_documento_avaliacao(doc_id, caminho_arquivo)
                            st.success("Documento excluído.")
                            st.rerun()
            else:
                st.info("Nenhum documento anexado para este estudante.")


# ======================================================
# ARTICULAÇÃO PEDAGÓGICA INCLUSIVA - ESCUTA DOCENTE
# ======================================================
elif menu == "Articulação Pedagógica Inclusiva":
    st.markdown('<div class="subtitulo">🤝 Articulação Pedagógica Inclusiva – Escuta Docente</div>', unsafe_allow_html=True)
    estudantes = listar_estudantes()

    if not estudantes:
        st.info("Cadastre um estudante primeiro.")
    else:
        ids, mapa = opcoes_estudantes_por_id(estudantes)
        estudante_id = st.selectbox(
            "Selecione o estudante",
            ids,
            format_func=lambda x: mapa[x],
            key="api_estudante",
        )
        estudante = buscar_estudante(estudante_id)

        st.info(
            "Este módulo registra a escuta dos professores da sala regular e gera o "
            "Relatório Pedagógico de Apoio ao Docente por IA, cruzando avaliação pedagógica, "
            "estudo de caso, atendimentos, documentos históricos e escutas docentes."
        )

        anos_opcoes = [str(a) for a in range(2024, agora_local().year + 4)]
        ano_padrao = str(agora_local().year)
        idx_ano = anos_opcoes.index(ano_padrao) if ano_padrao in anos_opcoes else 0

        aba_escuta, aba_historico, aba_relatorio, aba_relatorios_salvos = st.tabs(
            [
                "Nova Escuta Docente",
                "Histórico de Escutas",
                "Gerar Relatório de Apoio ao Docente",
                "Relatórios Salvos",
            ]
        )

        with aba_escuta:
            with st.container(border=True):
                st.markdown("### Nova Escuta Docente da Sala Regular")
                st.caption(
                    "Use este formulário para registrar a percepção pedagógica do professor da área. "
                    "Evite dados familiares, clínicos ou informações sensíveis que não sejam necessárias ao planejamento pedagógico."
                )

                with st.form("form_escuta_docente"):
                    col1, col2 = st.columns(2)

                    with col1:
                        ano_letivo = st.selectbox(
                            "Ano letivo",
                            anos_opcoes,
                            index=idx_ano,
                            key="escuta_ano_letivo",
                        )
                        codigo_docente = st.text_input(
                            "Código interno do docente / referência pedagógica",
                            placeholder="Ex.: DOC-MAT-001",
                            help="Use um código interno ou referência pedagógica. Evite registrar nomes pessoais de docentes neste campo.",
                        )
                        componente_curricular = st.selectbox(
                            "Componente curricular / área",
                            OPCOES_AREAS_CONHECIMENTO,
                        )
                        outro_componente = ""
                        if componente_curricular == "Outras":
                            outro_componente = st.text_input(
                                "Informe o componente/área",
                                placeholder="Digite a área ou componente curricular",
                            )

                    with col2:
                        turma = st.text_input(
                            "Turma",
                            value=estudante[3] or "",
                            placeholder="Ex.: 1º A",
                        )
                        tempo_acompanhamento = st.text_input(
                            "Tempo que acompanha o estudante",
                            placeholder="Ex.: desde o início do ano letivo, há dois meses...",
                        )
                        nivel_participacao = st.slider("Nível de participação em sala", 1, 10, 5)
                        nivel_autonomia = st.slider("Nível de autonomia nas atividades", 1, 10, 5)
                        nivel_engajamento = st.slider("Nível de engajamento", 1, 10, 5)

                    participacao_sala = st.text_area(
                        "Como o estudante participa das aulas?",
                        height=110,
                        placeholder="Descreva participação, atenção, realização das atividades, necessidade de mediação etc.",
                    )
                    comunicacao = st.text_area(
                        "Comunicação em sala",
                        height=90,
                        placeholder="Como o estudante expressa dúvidas, responde às solicitações, comunica necessidades ou interage verbalmente/não verbalmente?",
                    )
                    interacao_social = st.text_area(
                        "Interação social",
                        height=90,
                        placeholder="Como interage com colegas, professor e atividades em grupo?",
                    )
                    aprendizagem = st.text_area(
                        "Aprendizagem no componente curricular",
                        height=110,
                        placeholder="O que consegue acompanhar? Em quais situações demonstra melhor compreensão?",
                    )

                    estrategias_selecionadas = st.multiselect(
                        "Estratégias que parecem funcionar",
                        OPCOES_ESTRATEGIAS_INCLUSIVAS,
                    )
                    adaptacoes_selecionadas = st.multiselect(
                        "Adaptações já utilizadas pelo docente",
                        OPCOES_ESTRATEGIAS_INCLUSIVAS,
                    )

                    st.markdown("#### Perfil de atuação do professor da sala regular")
                    st.caption("Essas informações ajudam a IA a sugerir orientações compatíveis com a realidade do docente, sem impor tecnologias que ele ainda não domina.")
                    areas_interesse_docente = st.multiselect(
                        "Áreas/estratégias que o docente gosta ou tem interesse em utilizar",
                        OPCOES_AREAS_INTERESSE_DOCENTE,
                        key="escuta_areas_interesse_docente",
                    )
                    cold1, cold2 = st.columns(2)
                    with cold1:
                        nivel_tecnologico_docente = st.selectbox(
                            "Nível de familiaridade tecnológica do docente",
                            OPCOES_NIVEL_TECNOLOGICO_DOCENTE,
                            key="escuta_nivel_tecnologico_docente",
                        )
                        modo_maker_docente = st.toggle(
                            "🧩 Ativar modo maker inclusivo para este docente",
                            value=False,
                            key="escuta_modo_maker_docente",
                        )
                    with cold2:
                        interesse_formacao_maker_docente = st.checkbox(
                            "Docente demonstra interesse em aprender tecnologias maker/educacionais",
                            key="escuta_interesse_formacao_maker_docente",
                        )
                        interesse_projetos_interdisciplinares = st.checkbox(
                            "Docente tem interesse em projetos interdisciplinares",
                            key="escuta_interesse_projetos_interdisciplinares",
                        )
                    projeto_interesse_docente = st.text_input(
                        "Projeto/tema que o docente gostaria de trabalhar",
                        placeholder="Ex.: carro seguidor de linha, carrinho solar, podcast, jogo educativo, experimento científico...",
                        key="escuta_projeto_interesse_docente",
                    )
                    preferencias_metodologicas_docente = st.text_area(
                        "Preferências metodológicas do docente",
                        height=80,
                        placeholder="Ex.: atividades desplugadas, experimentos simples, projetos curtos, recursos visuais, tecnologia gradual...",
                        key="escuta_preferencias_metodologicas_docente",
                    )

                    barreiras_percebidas = st.text_area(
                        "Barreiras percebidas",
                        height=100,
                        placeholder="Ex.: excesso de cópia, comandos longos, organização da atividade, comunicação, atenção, interação...",
                    )
                    potencialidades_observadas = st.text_area(
                        "Potencialidades observadas",
                        height=100,
                        placeholder="Ex.: criatividade, memória visual, oralidade, interesse por tecnologia, raciocínio lógico, participação em atividades práticas...",
                    )
                    recomendacoes_docente = st.text_area(
                        "Sugestões/recomendações do professor da área",
                        height=100,
                        placeholder="Registre o que o professor acredita que pode ajudar no processo de aprendizagem.",
                    )
                    observacoes = st.text_area(
                        "Observações complementares",
                        height=80,
                    )

                    salvar_escuta = st.form_submit_button("Salvar Escuta Docente")

                if salvar_escuta:
                    if not codigo_docente.strip():
                        st.warning("Informe o código interno do docente / referência pedagógica antes de salvar. Ex.: DOC-MAT-001.")
                    else:
                        salvar_escuta_docente(
                            estudante_id=estudante_id,
                            ano_letivo=ano_letivo,
                            codigo_docente=codigo_docente,
                            componente_curricular=componente_curricular,
                            outro_componente=outro_componente,
                            turma=turma,
                            tempo_acompanhamento=tempo_acompanhamento,
                            participacao_sala=participacao_sala,
                            comunicacao=comunicacao,
                            interacao_social=interacao_social,
                            aprendizagem=aprendizagem,
                            barreiras_percebidas=barreiras_percebidas,
                            potencialidades_observadas=potencialidades_observadas,
                            estrategias_funcionam=", ".join(estrategias_selecionadas),
                            adaptacoes_utilizadas=", ".join(adaptacoes_selecionadas),
                            recomendacoes_docente=recomendacoes_docente,
                            nivel_participacao=nivel_participacao,
                            nivel_autonomia=nivel_autonomia,
                            nivel_engajamento=nivel_engajamento,
                            observacoes=observacoes,
                            areas_interesse_docente=", ".join(areas_interesse_docente),
                            nivel_tecnologico_docente=nivel_tecnologico_docente,
                            modo_maker_docente="Sim" if modo_maker_docente else "Não",
                            interesse_formacao_maker_docente="Sim" if interesse_formacao_maker_docente else "Não",
                            interesse_projetos_interdisciplinares="Sim" if interesse_projetos_interdisciplinares else "Não",
                            projeto_interesse_docente=projeto_interesse_docente,
                            preferencias_metodologicas_docente=preferencias_metodologicas_docente,
                        )
                        st.success("Escuta docente salva com sucesso. Acesse a aba Histórico de Escutas para baixar o registro em Word ou PDF.")
                        st.rerun()

        with aba_historico:
            with st.container(border=True):
                st.markdown("### Histórico de Escutas Docentes")
                st.caption("Cada escuta registrada pode ser visualizada, baixada em TXT, PDF ou Word e excluída, se necessário.")
                escutas = listar_escutas_docentes(estudante_id)

                if escutas:
                    for esc in escutas:
                        dados_esc = dict(zip(["id"] + CAMPOS_ESCUTA_DOCENTE, esc))
                        esc_id = dados_esc.get("id")
                        data_reg = dados_esc.get("data_registro") or "Data não informada"
                        area = dados_esc.get("componente_curricular") or "Área não informada"
                        codigo_doc = dados_esc.get("codigo_docente") or "Código docente não informado"

                        with st.expander(f"{data_reg} | {area} | {codigo_doc}"):
                            texto = texto_escuta_docente(estudante, esc)
                            st.text(texto)
                            export_buttons(texto, f"Escuta_Docente_{estudante[1]}_{esc_id}", tipo_pdf="escuta_docente")

                            if st.button("Excluir escuta docente", key=f"excluir_escuta_{esc_id}"):
                                excluir_escuta_docente(esc_id)
                                st.success("Escuta docente excluída.")
                                st.rerun()
                else:
                    st.info("Nenhuma escuta docente registrada para este estudante.")

        with aba_relatorio:
            with st.container(border=True):
                st.markdown("### Relatório Pedagógico de Apoio ao Docente")
                st.caption(
                    "O relatório será gerado com IA a partir dos registros já existentes: avaliação pedagógica, "
                    "estudo de caso, atendimentos, documentos históricos anexados e escutas docentes."
                )

                colr1, colr2 = st.columns(2)
                with colr1:
                    ano_relatorio = st.selectbox(
                        "Ano letivo do relatório",
                        anos_opcoes,
                        index=idx_ano,
                        key="relatorio_docente_ano",
                    )
                    componente_destino = st.selectbox(
                        "Componente/área de destino",
                        OPCOES_AREAS_CONHECIMENTO,
                        key="relatorio_docente_componente",
                    )
                    if componente_destino == "Outras":
                        componente_destino = st.text_input(
                            "Informe o componente/área de destino",
                            key="relatorio_docente_componente_outro",
                        )

                with colr2:
                    professor_destino = ""
                    st.info("O destinatário ficará em branco no relatório para preenchimento manual após impressão ou edição do Word.")
                    observacoes_relatorio = st.text_area(
                        "Observações internas do AEE",
                        height=90,
                        placeholder="Opcional. Ex.: relatório para planejamento de atividades avaliativas do bimestre.",
                        key="relatorio_docente_obs",
                    )

                foco_docente = st.text_area(
                    "Foco para a IA",
                    height=100,
                    placeholder="Ex.: orientar o professor sobre participação, adaptação das atividades e formas avaliativas sem expor dados familiares.",
                    key="relatorio_docente_foco",
                )

                contexto_previo, fontes_previas = montar_contexto_relatorio_docente(estudante_id)
                with st.expander("Ver fontes que serão cruzadas pelo sistema"):
                    st.write(fontes_previas)
                    st.text_area(
                        "Contexto reunido pelo sistema",
                        value=contexto_previo,
                        height=300,
                        disabled=True,
                    )

                if st.button("🤖 Gerar Relatório Pedagógico de Apoio ao Docente", key="btn_gerar_relatorio_docente_ia"):
                    conteudo, fontes = gerar_relatorio_apoio_docente_ia(
                        estudante_id=estudante_id,
                        ano_letivo=ano_relatorio,
                        componente_destino=componente_destino,
                        professor_destino=professor_destino,
                        foco_docente=foco_docente,
                    )
                    st.session_state["relatorio_docente_conteudo"] = conteudo
                    st.session_state["relatorio_docente_fontes"] = fontes
                    st.success("Relatório gerado. Revise antes de salvar ou entregar ao docente.")

                conteudo_gerado = st.session_state.get("relatorio_docente_conteudo", "")
                fontes_geradas = st.session_state.get("relatorio_docente_fontes", fontes_previas)

                if conteudo_gerado:
                    st.markdown("### Prévia editável do relatório")
                    conteudo_editado = st.text_area(
                        "Revise o relatório antes de salvar",
                        value=conteudo_gerado,
                        height=520,
                        key="relatorio_docente_editavel",
                    )

                    titulo_relatorio = "RELATÓRIO PEDAGÓGICO DE APOIO AO DOCENTE"
                    texto_final = f"""
{titulo_relatorio}

Código interno do estudante: {estudante[1]}
Ano/Série: {estudante[2] or 'Não informado.'}
Turma: {estudante[3] or 'Não informado.'}
Ano letivo: {ano_relatorio}
Componente/área de destino: {componente_destino}
Destinatário: ________________________________________
Data de geração: {hoje_str()}

{conteudo_editado}

FONTES UTILIZADAS PELO SISTEMA:
{fontes_geradas}

Observação institucional:
Este relatório possui finalidade exclusivamente pedagógica e objetiva apoiar práticas educacionais inclusivas no contexto escolar, sem expor dados familiares ou informações sensíveis desnecessárias.
""".strip()

                    st.text(texto_final)
                    export_buttons(
                        texto_final,
                        f"Relatorio_Apoio_Docente_{estudante[1]}_{ano_relatorio}",
                        tipo_pdf="relatorio_docente",
                    )

                    st.markdown("### 🧩 Relatório Visual de Apoio ao Docente - PDF")
                    st.caption(
                        "Versão visual para professores. O foco é perceber quem é o estudante, "
                        "o que a condição informada representa pedagogicamente e quais habilidades, "
                        "potencialidades e estratégias podem favorecer sua aprendizagem."
                    )

                    dados_painel = gerar_dados_painel_visual_docente(estudante, conteudo_editado)

                    with st.expander("Editar cards do painel visual antes de gerar", expanded=False):
                        col_card_1, col_card_2 = st.columns(2)
                        with col_card_1:
                            painel_quem = st.text_area(
                                "Quem é este estudante no contexto escolar",
                                value=dados_painel["quem_estudante"],
                                height=140,
                                key="painel_quem_estudante",
                            )
                            painel_condicao = st.text_area(
                                "O que a condição informada pode significar",
                                value=dados_painel["significado_condicao"],
                                height=140,
                                key="painel_significado_condicao",
                            )
                            painel_barreiras = st.text_area(
                                "Barreiras que podem aparecer",
                                value=dados_painel["barreiras_observadas"],
                                height=140,
                                key="painel_barreiras_observadas",
                            )
                            painel_pontos = st.text_area(
                                "O que merece atenção",
                                value=dados_painel["pontos_atencao"],
                                height=140,
                                key="painel_pontos_atencao",
                            )
                        with col_card_2:
                            painel_aprende = st.text_area(
                                "Como aprende melhor",
                                value=dados_painel["como_aprende"],
                                height=140,
                                key="painel_como_aprende",
                            )
                            painel_potencialidades = st.text_area(
                                "Habilidades, potencialidades e interesses",
                                value=dados_painel["potencialidades_interesses"],
                                height=140,
                                key="painel_potencialidades_interesses",
                            )
                            painel_estrategias = st.text_area(
                                "Estratégias rápidas para a aula",
                                value=dados_painel["estrategias_rapidas"],
                                height=140,
                                key="painel_estrategias_rapidas",
                            )

                    st.info(
                        "Você pode gerar dois modelos diferentes. Ambos serão salvos automaticamente "
                        "no histórico de relatórios visuais do estudante para consulta futura."
                    )

                    dados_pdf_visual = dict(
                        estudante=estudante,
                        ano_letivo=ano_relatorio,
                        componente_destino=componente_destino,
                        quem_estudante=st.session_state.get("painel_quem_estudante", dados_painel["quem_estudante"]),
                        significado_condicao=st.session_state.get("painel_significado_condicao", dados_painel["significado_condicao"]),
                        como_aprende=st.session_state.get("painel_como_aprende", dados_painel["como_aprende"]),
                        barreiras_observadas=st.session_state.get("painel_barreiras_observadas", dados_painel["barreiras_observadas"]),
                        potencialidades_interesses=st.session_state.get("painel_potencialidades_interesses", dados_painel["potencialidades_interesses"]),
                        pontos_atencao=st.session_state.get("painel_pontos_atencao", dados_painel["pontos_atencao"]),
                        estrategias_rapidas=st.session_state.get("painel_estrategias_rapidas", dados_painel["estrategias_rapidas"]),
                        conteudo_relatorio=conteudo_editado,
                        fontes_geradas=fontes_geradas,
                    )

                    chave_pdf_limpo = f"pdf_visual_limpo_{estudante[1]}_{ano_relatorio}"
                    chave_pdf_info = f"pdf_visual_infografico_{estudante[1]}_{ano_relatorio}"

                    col_pdf_limpo, col_pdf_info = st.columns(2)

                    with col_pdf_limpo:
                        if st.button("📄 Gerar PDF pedagógico limpo", key="btn_gerar_pdf_visual_docente_limpo"):
                            arquivo_pdf_visual = gerar_pdf_relatorio_visual_docente(
                                **dados_pdf_visual,
                                nome_base=f"Relatorio_Visual_Limpo_Apoio_Docente_{estudante[1]}_{ano_relatorio}",
                            )
                            salvar_relatorio_visual_docente(
                                estudante_id=estudante_id,
                                ano_letivo=ano_relatorio,
                                componente_destino=componente_destino,
                                tipo_relatorio="PDF pedagógico limpo",
                                titulo="Relatório Visual de Apoio ao Docente - PDF pedagógico limpo",
                                caminho_arquivo=arquivo_pdf_visual,
                                fontes_utilizadas=fontes_geradas,
                                observacoes="Gerado a partir do relatório pedagógico de apoio ao docente revisado.",
                            )
                            st.session_state[chave_pdf_limpo] = arquivo_pdf_visual
                            st.success("PDF pedagógico limpo gerado e salvo no histórico.")

                        if chave_pdf_limpo in st.session_state and Path(st.session_state[chave_pdf_limpo]).exists():
                            with open(st.session_state[chave_pdf_limpo], "rb") as f:
                                st.download_button(
                                    "⬇️ Baixar PDF pedagógico limpo",
                                    data=f,
                                    file_name=Path(st.session_state[chave_pdf_limpo]).name,
                                    mime="application/pdf",
                                    key="download_pdf_visual_docente_limpo",
                                )

                    with col_pdf_info:
                        if st.button("🧩 Gerar painel infográfico docente", key="btn_gerar_pdf_infografico_docente"):
                            with st.spinner("A IA está lendo o relatório e estruturando o painel em JSON..."):
                                dados_infografico = gerar_conteudo_infografico_docente_ia(
                                    relatorio_docente_txt=conteudo_editado,
                                    estudante=estudante,
                                    ano_letivo=ano_relatorio,
                                    componente=componente_destino,
                                )

                                arquivo_pdf_infografico = gerar_pdf_infografico_docente_dashboard(
                                    estudante=estudante,
                                    dados=dados_infografico,
                                    ano_letivo=ano_relatorio,
                                    componente=componente_destino,
                                    nome_base=f"Painel_Inteligente_Apoio_Docente_{estudante[1]}_{ano_relatorio}",
                                )

                            salvar_relatorio_visual_docente(
                                estudante_id=estudante_id,
                                ano_letivo=ano_relatorio,
                                componente_destino=componente_destino,
                                tipo_relatorio="Painel infográfico docente",
                                titulo="Painel Inteligente de Apoio ao Docente",
                                caminho_arquivo=arquivo_pdf_infografico,
                                fontes_utilizadas=fontes_geradas,
                                observacoes="Gerado com IA a partir do relatório pedagógico docente e estruturado em JSON para layout visual.",
                            )
                            st.session_state[chave_pdf_info] = arquivo_pdf_infografico
                            st.session_state[f"json_infografico_{estudante[1]}_{ano_relatorio}"] = dados_infografico
                            st.success("Painel infográfico docente gerado com IA e salvo no histórico.")

                        if f"json_infografico_{estudante[1]}_{ano_relatorio}" in st.session_state:
                            with st.expander("Ver JSON estruturado pela IA", expanded=False):
                                st.json(st.session_state[f"json_infografico_{estudante[1]}_{ano_relatorio}"])

                        if chave_pdf_info in st.session_state and Path(st.session_state[chave_pdf_info]).exists():
                            with open(st.session_state[chave_pdf_info], "rb") as f:
                                st.download_button(
                                    "⬇️ Baixar painel infográfico docente",
                                    data=f,
                                    file_name=Path(st.session_state[chave_pdf_info]).name,
                                    mime="application/pdf",
                                    key="download_pdf_infografico_docente",
                                )

                    if st.button("Salvar relatório no histórico", key="btn_salvar_relatorio_docente"):
                        salvar_relatorio_docente(
                            estudante_id=estudante_id,
                            ano_letivo=ano_relatorio,
                            componente_destino=componente_destino,
                            professor_destino=professor_destino,
                            titulo=titulo_relatorio,
                            conteudo=conteudo_editado,
                            fontes_utilizadas=fontes_geradas,
                            observacoes=observacoes_relatorio,
                        )
                        st.session_state.pop("relatorio_docente_conteudo", None)
                        st.session_state.pop("relatorio_docente_fontes", None)
                        st.success("Relatório salvo no histórico.")
                        st.rerun()

        with aba_relatorios_salvos:
            with st.container(border=True):
                st.markdown("### Relatórios Pedagógicos de Apoio ao Docente salvos")
                relatorios = listar_relatorios_docente(estudante_id)

                if relatorios:
                    for rel in relatorios:
                        dados_rel = dict(zip(["id"] + CAMPOS_RELATORIO_DOCENTE, rel))
                        rel_id = dados_rel.get("id")
                        data_rel = dados_rel.get("data_geracao") or "Data não informada"
                        area_rel = dados_rel.get("componente_destino") or "Área não informada"

                        with st.expander(f"{data_rel} | {area_rel}"):
                            texto = texto_relatorio_docente(estudante, rel)
                            st.text(texto)
                            export_buttons(
                                texto,
                                f"Relatorio_Apoio_Docente_{estudante[1]}_{rel_id}",
                                tipo_pdf="relatorio_docente_salvo",
                            )

                            if st.button("Excluir relatório salvo", key=f"excluir_relatorio_docente_{rel_id}"):
                                excluir_relatorio_docente(rel_id)
                                st.success("Relatório excluído.")
                                st.rerun()
                else:
                    st.info("Nenhum relatório pedagógico de apoio ao docente salvo para este estudante.")

            with st.container(border=True):
                st.markdown("### Relatórios visuais docentes salvos")
                st.caption("Histórico dos PDFs visuais gerados para este estudante: modelo pedagógico limpo e painel infográfico.")
                relatorios_visuais = listar_relatorios_visuais_docente(estudante_id)

                if relatorios_visuais:
                    for rel_visual in relatorios_visuais:
                        dados_visual = dict(zip(["id"] + CAMPOS_RELATORIO_VISUAL_DOCENTE, rel_visual))
                        rel_visual_id = dados_visual.get("id")
                        data_visual = dados_visual.get("data_geracao") or "Data não informada"
                        tipo_visual = dados_visual.get("tipo_relatorio") or "Relatório visual"
                        caminho_visual = dados_visual.get("caminho_arquivo") or ""
                        titulo_visual = dados_visual.get("titulo") or tipo_visual

                        with st.expander(f"{data_visual} | {tipo_visual}"):
                            st.write(f"**Título:** {titulo_visual}")
                            st.write(f"**Ano letivo:** {dados_visual.get('ano_letivo') or 'Não informado'}")
                            st.write(f"**Componente/área:** {dados_visual.get('componente_destino') or 'Não informado'}")
                            if dados_visual.get("observacoes"):
                                st.write(f"**Observações:** {dados_visual.get('observacoes')}")

                            if caminho_visual and Path(caminho_visual).exists():
                                with open(caminho_visual, "rb") as f:
                                    st.download_button(
                                        "⬇️ Baixar PDF salvo",
                                        data=f,
                                        file_name=Path(caminho_visual).name,
                                        mime="application/pdf",
                                        key=f"download_relatorio_visual_salvo_{rel_visual_id}",
                                    )
                            else:
                                st.warning(
                                    "O registro existe no histórico, mas o arquivo PDF não foi encontrado no servidor. "
                                    "Se estiver usando Render, verifique se há disco persistente configurado para a pasta relatorios_visuais_docente."
                                )

                            if st.button("Excluir relatório visual salvo", key=f"excluir_relatorio_visual_{rel_visual_id}"):
                                excluir_relatorio_visual_docente(rel_visual_id, caminho_visual)
                                st.success("Relatório visual excluído.")
                                st.rerun()
                else:
                    st.info("Nenhum relatório visual docente salvo para este estudante.")


elif menu == "Estudo de Caso":
    st.markdown('<div class="subtitulo">📚 Estudo de Caso / Documento GRE</div>', unsafe_allow_html=True)
    st.caption("Os campos sensíveis como CPF, RG, endereço e nome completo ficam fora do banco. Eles aparecem em branco no Word/PDF para preenchimento manual antes da entrega à GRE.")

    estudantes = listar_estudantes()
    if not estudantes:
        st.info("Cadastre um estudante primeiro.")
    else:
        ids, mapa = opcoes_estudantes_por_id(estudantes)
        estudante_id = st.selectbox("Selecione o estudante", ids, format_func=lambda x: mapa[x], key="estudo_estudante")
        estudante = buscar_estudante(estudante_id)
        professor_resp = buscar_professor_responsavel()
        estudos_anteriores = listar_por_estudante("estudos_caso", CAMPOS_ESTUDO_CASO, estudante_id)
        avaliacoes_pedagogicas = listar_avaliacoes(estudante_id)

        with st.container(border=True):
            st.markdown("### Novo Estudo de Caso - Campos obrigatórios GRE")
            st.info("Agora o Estudo de Caso pode ser registrado por ano letivo. Você pode lançar estudos anteriores como histórico e gerar um novo estudo com apoio da IA, comparando registros antigos com dados atuais.")
            st.caption("✅ V13 GRE ativo: a IA gera o Estudo de Caso em estrutura padronizada GRE, usando estudo anterior e/ou todas as avaliações pedagógicas disponíveis como base.")

            with st.expander("🤖 Criar novo estudo de caso com apoio da IA", expanded=True):
                if estudos_anteriores or avaliacoes_pedagogicas:
                    st.caption(
                        "A IA pode gerar um novo Estudo de Caso a partir de estudo anterior e/ou avaliações pedagógicas. "
                        "Quando não houver estudo anterior, o sistema usa as avaliações pedagógicas e demais registros como base."
                    )

                    if estudos_anteriores:
                        opcoes_estudos = {0: "Nenhum estudo anterior - gerar com base nas avaliações pedagógicas e demais registros"}
                        opcoes_estudos.update({
                            e[0]: f"ID {e[0]} | Ano: {e[2] or 'não informado'} | Tipo: {e[3] or 'não informado'} | Registro: {e[1]}"
                            for e in estudos_anteriores
                        })
                        estudo_base_id_ia = st.selectbox(
                            "Estudo anterior para comparação (opcional)",
                            list(opcoes_estudos.keys()),
                            format_func=lambda x: opcoes_estudos[x],
                            key=f"estudo_base_ia_{estudante_id}",
                        )
                    else:
                        estudo_base_id_ia = 0
                        st.info("Não há estudo de caso anterior. A IA criará uma primeira versão com base nas avaliações pedagógicas e demais registros disponíveis.")

                    if avaliacoes_pedagogicas:
                        opcoes_avaliacoes_ia = {
                            a[0]: f"ID {a[0]} | Ano: {a[2] or 'não informado'} | Tipo: {a[3] or 'não informado'} | Registro: {a[1]}"
                            for a in avaliacoes_pedagogicas
                        }
                        avaliacao_base_id_ia = st.selectbox(
                            "Avaliação pedagógica principal para compor o Estudo de Caso",
                            list(opcoes_avaliacoes_ia.keys()),
                            format_func=lambda x: opcoes_avaliacoes_ia[x],
                            key=f"avaliacao_base_estudo_ia_{estudante_id}",
                        )
                        st.caption("A avaliação escolhida será a principal, mas a IA também considerará as demais avaliações pedagógicas/documentos livres registrados para este estudante.")
                    else:
                        avaliacao_base_id_ia = 0
                        st.warning("Nenhuma avaliação pedagógica localizada. A IA usará apenas o estudo anterior e demais registros disponíveis.")

                    ano_novo_ia = st.text_input("Ano letivo do novo estudo", value=str(agora_local().year), key=f"ano_novo_ia_{estudante_id}")

                    fontes_ia = []
                    if estudos_anteriores and estudo_base_id_ia != 0:
                        fontes_ia.append("estudo de caso anterior")
                    if avaliacoes_pedagogicas:
                        fontes_ia.append("avaliação pedagógica")
                    fontes_ia.extend(["entrevista familiar", "atendimentos AEE", "registros disponíveis"])
                    st.caption("Fontes previstas para a IA: " + ", ".join(fontes_ia) + ".")

                    if st.button("Gerar análise e sugestão de novo estudo", key=f"gerar_estudo_ia_{estudante_id}"):
                        estudo_base = next((e for e in estudos_anteriores if e[0] == estudo_base_id_ia), None)
                        avaliacao_atual = next((a for a in avaliacoes_pedagogicas if a[0] == avaliacao_base_id_ia), None) if avaliacao_base_id_ia else None
                        entrevista_atual = ultima_linha("entrevistas_familia", CAMPOS_ENTREVISTA_FAMILIA, estudante_id)
                        atendimentos_atuais = listar_atendimentos(estudante_id)
                        analise_ia, sugestao_ia = gerar_novo_estudo_caso_com_ia(
                            estudante,
                            estudo_base,
                            avaliacao_atual,
                            entrevista_atual,
                            atendimentos_atuais,
                            ano_novo_ia,
                            avaliacoes_contexto=avaliacoes_pedagogicas,
                        )
                        st.session_state[f"analise_estudo_ia_{estudante_id}"] = analise_ia
                        st.session_state[f"sugestao_estudo_ia_{estudante_id}"] = sugestao_ia
                        st.session_state[f"ano_estudo_ia_{estudante_id}"] = ano_novo_ia
                        st.session_state[f"estudo_anterior_id_{estudante_id}"] = None if estudo_base_id_ia == 0 else estudo_base_id_ia
                        st.session_state[f"avaliacao_base_id_estudo_ia_{estudante_id}"] = avaliacao_base_id_ia
                        st.session_state[f"fontes_estudo_ia_{estudante_id}"] = ", ".join(fontes_ia)

                        # Campos estruturados do Estudo de Caso GRE gerados pela IA.
                        st.session_state[f"ia_percurso_{estudante_id}"] = extrair_secao_ia(sugestao_ia, "PERCURSO_EDUCACIONAL")
                        st.session_state[f"ia_motivo_{estudante_id}"] = extrair_secao_ia(sugestao_ia, "MOTIVO_ENCAMINHAMENTO")
                        st.session_state[f"ia_hab_obs_{estudante_id}"] = extrair_secao_ia(sugestao_ia, "HABILIDADES_OBSERVADAS")
                        st.session_state[f"ia_hab_des_{estudante_id}"] = extrair_secao_ia(sugestao_ia, "HABILIDADES_A_DESENVOLVER")
                        st.session_state[f"ia_potencialidades_{estudante_id}"] = extrair_secao_ia(sugestao_ia, "POTENCIALIDADES")
                        st.session_state[f"ia_dificuldades_{estudante_id}"] = extrair_secao_ia(sugestao_ia, "DIFICULDADES")
                        st.session_state[f"ia_estrategias_{estudante_id}"] = extrair_secao_ia(sugestao_ia, "ESTRATEGIAS")
                        st.session_state[f"ia_intervencoes_{estudante_id}"] = extrair_secao_ia(sugestao_ia, "INTERVENCOES")
                        st.session_state[f"ia_avaliacao_{estudante_id}"] = extrair_secao_ia(sugestao_ia, "AVALIACAO")
                        st.session_state[f"ia_consideracoes_{estudante_id}"] = extrair_secao_ia(sugestao_ia, "CONSIDERACOES")
                        st.session_state[f"ia_observacoes_complementares_{estudante_id}"] = extrair_secao_ia(sugestao_ia, "OBSERVACOES_COMPLEMENTARES")

                        st.success("Sugestão GRE gerada. Os campos do formulário foram preparados para revisão e edição antes de salvar.")

                    if st.session_state.get(f"sugestao_estudo_ia_{estudante_id}"):
                        st.markdown("#### Análise IA")
                        analise_preview = st.text_area(
                            "Resultado da análise",
                            value=st.session_state.get(f"analise_estudo_ia_{estudante_id}", ""),
                            height=180,
                            key=f"preview_analise_{estudante_id}",
                        )
                        st.markdown("#### Sugestão GRE estruturada")
                        sugestao_preview = st.text_area(
                            "Sugestão gerada",
                            value=st.session_state.get(f"sugestao_estudo_ia_{estudante_id}", ""),
                            height=260,
                            key=f"preview_sugestao_{estudante_id}",
                        )

                        texto_previa_estudo = f"""
ESTUDO DE CASO GRE - SUGESTÃO GERADA POR IA

Código interno do estudante: {estudante[1]}
Ano/Série: {estudante[2] or 'Não informado.'}
Turma: {estudante[3] or 'Não informado.'}
Ano letivo: {st.session_state.get(f'ano_estudo_ia_{estudante_id}', str(agora_local().year))}

ANÁLISE PEDAGÓGICA DA IA:
{analise_preview}

SUGESTÃO ESTRUTURADA PADRÃO GRE:
{sugestao_preview}

Observação institucional:
Documento preliminar para revisão do professor do AEE antes do salvamento definitivo e/ou envio à GRE.
""".strip()

                        st.markdown("#### Download da sugestão gerada")
                        export_buttons(
                            texto_previa_estudo,
                            f"Estudo_Caso_GRE_IA_{estudante[1]}_{st.session_state.get(f'ano_estudo_ia_{estudante_id}', str(agora_local().year))}",
                            tipo_pdf="estudo_ia_previa",
                        )

                        st.markdown("#### Salvar no histórico")
                        st.caption("Esta opção salva imediatamente a sugestão gerada no histórico do estudante, mantendo a estrutura GRE. Depois você ainda pode abrir o registro salvo, baixar em Word/PDF e complementar manualmente.")
                        if st.button("💾 Salvar sugestão gerada no histórico", key=f"salvar_sugestao_estudo_ia_{estudante_id}"):
                            ano_salvar = st.session_state.get(f"ano_estudo_ia_{estudante_id}", str(agora_local().year))
                            estudo_ant_salvar = st.session_state.get(f"estudo_anterior_id_{estudante_id}")
                            analise_salvar = st.session_state.get(f"analise_estudo_ia_{estudante_id}", "")
                            sugestao_salvar = st.session_state.get(f"sugestao_estudo_ia_{estudante_id}", "")

                            dados_estudo_ia = {
                                "data_registro": hoje_str(),
                                "ano_letivo": ano_salvar,
                                "tipo_registro": "Novo estudo com base em estudo anterior ou avaliação pedagógica",
                                "estudo_anterior_id": estudo_ant_salvar,
                                "analise_comparativa_ia": analise_salvar,
                                "sugestao_novo_estudo_ia": sugestao_salvar,
                                "contextualizacao": extrair_secao_ia(sugestao_salvar, "PERCURSO_EDUCACIONAL"),
                                "queixa_principal": extrair_secao_ia(sugestao_salvar, "MOTIVO_ENCAMINHAMENTO"),
                                "potencialidades": extrair_secao_ia(sugestao_salvar, "POTENCIALIDADES") or extrair_secao_ia(sugestao_salvar, "HABILIDADES_OBSERVADAS"),
                                "dificuldades": extrair_secao_ia(sugestao_salvar, "DIFICULDADES") or extrair_secao_ia(sugestao_salvar, "HABILIDADES_A_DESENVOLVER"),
                                "estrategias": extrair_secao_ia(sugestao_salvar, "ESTRATEGIAS"),
                                "intervencoes": extrair_secao_ia(sugestao_salvar, "INTERVENCOES"),
                                "avaliacao": extrair_secao_ia(sugestao_salvar, "AVALIACAO"),
                                "consideracoes": extrair_secao_ia(sugestao_salvar, "CONSIDERACOES"),
                                "etapa_modalidade": "",
                                "ano_etapa": estudante[2] or "",
                                "laudo": "",
                                "deficiencia": estudante[4] or "",
                                "cid": "",
                                "observacoes_laudo": "",
                                "altas_habilidades": "",
                                "bpc": "",
                                "escola_nome": "",
                                "unidade_aee": "",
                                "gestor_nome": "",
                                "gestor_contato": "",
                                "professor_nome": "",
                                "professor_contato": "",
                                "matricula_professor": "",
                                "especialidade_professor": "",
                                "periodo_inicio": "",
                                "periodo_fim": "",
                                "frequencia_atendimento": "",
                                "tempo_atendimento_semana": "",
                                "formato_atendimento": "",
                                "percurso_educacional": extrair_secao_ia(sugestao_salvar, "PERCURSO_EDUCACIONAL"),
                                "motivo_encaminhamento_aee": extrair_secao_ia(sugestao_salvar, "MOTIVO_ENCAMINHAMENTO"),
                                "precisa_transporte_inclusivo": "",
                                "recebe_transporte_inclusivo": "",
                                "precisa_profissional_apoio": "",
                                "justificativa_apoio": "",
                                "acompanhado_profissional_apoio": "",
                                "nome_profissional_apoio": "",
                                "recursos_tecnologia_assistiva": extrair_secao_ia(sugestao_salvar, "INTERVENCOES"),
                                "observacoes_ambiente_educacional": extrair_secao_ia(sugestao_salvar, "OBSERVACOES_COMPLEMENTARES"),
                                "habilidades_observadas": extrair_secao_ia(sugestao_salvar, "HABILIDADES_OBSERVADAS"),
                                "habilidades_a_desenvolver": extrair_secao_ia(sugestao_salvar, "HABILIDADES_A_DESENVOLVER"),
                                "indicadores_altas_habilidades": "",
                                "recursos_surdez": "",
                                "observacoes_surdez": "",
                            }
                            inserir_registro(
                                "estudos_caso",
                                ["estudante_id"] + CAMPOS_ESTUDO_CASO,
                                [estudante_id] + [dados_estudo_ia.get(campo, "") for campo in CAMPOS_ESTUDO_CASO],
                            )
                            st.success("Estudo de caso GRE salvo no histórico do estudante.")
                            st.rerun()

                        st.info("Você também pode revisar os campos nas abas abaixo e clicar em **Salvar estudo de caso GRE** para gravar uma versão editada manualmente.")
                else:
                    st.info(
                        "Ainda não há estudo anterior ou avaliações pedagógicas registradas para este estudante. "
                        "Primeiro cadastre avaliação pedagógica e/ou o estudo histórico ou atual."
                    )

            with st.form("form_estudo_gre"):
                aba1, aba2, aba3, aba4 = st.tabs([
                    "1. Identificação",
                    "2. Percurso educacional",
                    "3. Habilidades",
                    "4. Síntese pedagógica",
                ])

                with aba1:
                    st.markdown("#### Identificação educacional")
                    col_ano_hist, col_tipo_hist = st.columns(2)
                    with col_ano_hist:
                        ano_letivo = st.text_input("Ano letivo do estudo de caso", value=st.session_state.get(f"ano_estudo_ia_{estudante_id}", str(agora_local().year)))
                    with col_tipo_hist:
                        tipo_registro = st.selectbox(
                            "Tipo de registro",
                            ["Estudo de caso atual", "Registro histórico", "Novo estudo com base em estudo anterior ou avaliação pedagógica"],
                            index=2 if st.session_state.get(f"sugestao_estudo_ia_{estudante_id}") else 0,
                        )
                    estudo_anterior_id = st.session_state.get(f"estudo_anterior_id_{estudante_id}")
                    analise_comparativa_ia = st.session_state.get(f"analise_estudo_ia_{estudante_id}", "")
                    sugestao_novo_estudo_ia = st.session_state.get(f"sugestao_estudo_ia_{estudante_id}", "")

                    etapa_modalidade = st.selectbox(
                        "Etapa/modalidade",
                        ["Ensino Fundamental", "Ensino Médio", "EJA", "Outro / não informado"],
                    )
                    ano_etapa = st.selectbox("Ano/etapa", OPCOES_ANO_ETAPA)
                    turma_turno_info = f"{estudante[3] or 'Turma não informada'} / {estudante[6] or 'Turno não informado'}"
                    st.info(f"Turma e turno cadastrados: {turma_turno_info}")

                    col_l1, col_l2, col_l3 = st.columns(3)
                    with col_l1:
                        laudo = st.radio("Possui laudo?", ["Não", "Sim"], horizontal=True)
                    with col_l2:
                        deficiencia = st.radio("Apresenta deficiência?", ["Não", "Sim"], horizontal=True)
                    with col_l3:
                        altas_habilidades = st.radio("Altas habilidades/superdotação?", ["Não", "Sim"], horizontal=True)

                    cid = st.text_input(
                        "CID, se houver",
                        placeholder="Ex.: F84.0",
                        help="Informe apenas o CID indicado no laudo, quando disponível. Não anexe laudos ou exames no sistema."
                    )
                    observacoes_laudo = st.text_area(
                        "Síntese pedagógica/funcional do laudo",
                        placeholder=(
                            "Registre apenas informações relevantes para o contexto educacional, como comunicação, "
                            "autonomia, interação social, atenção, comportamento, necessidades de apoio, "
                            "organização pedagógica e acessibilidade. Não inserir laudo completo, exames ou dados clínicos excessivamente sensíveis."
                        ),
                        height=140,
                        help="Este campo deve transformar a informação do laudo em orientação pedagógica/funcional para o AEE."
                    )
                    bpc = st.radio("Usuário de BPC?", ["Não", "Sim", "Não informado"], horizontal=True)

                    st.markdown("#### Dados institucionais")
                    escola_nome = st.text_input("Nome da escola em que o(a) estudante está matriculado no ensino comum", value="")
                    unidade_aee = st.text_input("Unidade educacional onde o(a) estudante é atendido pelo AEE", value="")
                    gestor_nome = st.text_input("Nome do(a) gestor(a) da escola do ensino comum", value="")
                    gestor_contato = st.text_input("Contato institucional do(a) gestor(a)", value="")

                    professor_nome_padrao = professor_resp[1] if professor_resp else ""
                    professor_nome = st.text_input("Nome do(a) professor(a) AEE", value=professor_nome_padrao or "")
                    professor_contato = st.text_input("Contato institucional do(a) professor(a) AEE", value="")
                    matricula_professor = st.text_input("Matrícula do(a) professor(a) AEE", value="")
                    especialidade_professor = st.multiselect("Especialidade do(a) professor(a) AEE", OPCOES_ESPECIALIDADE_AEE, default=["Professor(a) do AEE"])

                    col_p1, col_p2 = st.columns(2)
                    with col_p1:
                        periodo_inicio = st.text_input("Período de elaboração/início do atendimento", placeholder="Ex.: 01/02/2026")
                    with col_p2:
                        periodo_fim = st.text_input("Data final", placeholder="Ex.: 30/06/2026")

                    frequencia_atendimento = st.multiselect("Frequência de atendimento na SRM", DIAS_SEMANA, default=[])
                    tempo_atendimento_semana = st.text_input("Tempo de atendimento por semana", placeholder="Ex.: 2h semanais")
                    formato_atendimento = st.multiselect("Formato do atendimento", OPCOES_FORMATO_ATENDIMENTO, default=["Individual"])

                with aba2:
                    st.markdown("#### Estudo de caso / percurso educacional")
                    percurso_educacional = st.text_area(
                        "Relato sobre o trajeto educacional do(a) estudante",
                        value=st.session_state.get(f"ia_percurso_{estudante_id}", ""),
                        placeholder="Descreva situação inicial, estratégias já utilizadas e progressos alcançados em turmas comuns e/ou AEE anterior.",
                        height=180,
                    )
                    motivo_encaminhamento_aee = st.text_area(
                        "Motivo pelo qual foi encaminhado ao AEE",
                        value=st.session_state.get(f"ia_motivo_{estudante_id}", ""),
                        height=120,
                    )

                    col_t1, col_t2 = st.columns(2)
                    with col_t1:
                        precisa_transporte_inclusivo = st.radio("Precisa de transporte escolar inclusivo?", ["Não", "Sim", "Não informado"], horizontal=True)
                    with col_t2:
                        recebe_transporte_inclusivo = st.radio("Recebe transporte escolar inclusivo?", ["Não", "Sim", "Não informado"], horizontal=True)

                    precisa_profissional_apoio = st.radio("Precisa de profissional de apoio?", ["Não", "Sim", "Não informado"], horizontal=True)
                    justificativa_apoio = st.text_area("Justificativa para profissional de apoio, se necessário", height=100)
                    acompanhado_profissional_apoio = st.radio("É acompanhado por profissional de apoio na escola?", ["Não", "Sim", "Não informado"], horizontal=True)
                    nome_profissional_apoio = st.text_input("Nome do profissional de apoio, se houver")

                    recursos_tecnologia_assistiva = st.multiselect("Recursos de tecnologia educacional e/ou assistiva utilizados", OPCOES_RECURSOS_TA)
                    observacoes_ambiente_educacional = st.text_area(
                        "Observações relevantes para o ambiente educacional / acompanhamento médico ou terapêutico",
                        value=st.session_state.get(f"ia_observacoes_complementares_{estudante_id}", ""),
                        placeholder="Ex.: acompanhamento terapêutico, recomendações pedagógicas, cuidados no ambiente escolar.",
                        height=140,
                    )

                with aba3:
                    st.markdown("#### Habilidades e indicadores")
                    habilidades_observadas = st.multiselect("Habilidades observadas e desenvolvidas", OPCOES_HABILIDADES_PEDAGOGICAS)
                    habilidades_observadas_ia = st.text_area(
                        "Habilidades observadas sugeridas pela IA / complementação manual",
                        value=st.session_state.get(f"ia_hab_obs_{estudante_id}", ""),
                        height=90,
                    )
                    habilidades_a_desenvolver = st.multiselect("Habilidades que precisam ser desenvolvidas", OPCOES_HABILIDADES_A_DESENVOLVER)
                    habilidades_a_desenvolver_ia = st.text_area(
                        "Habilidades a desenvolver sugeridas pela IA / complementação manual",
                        value=st.session_state.get(f"ia_hab_des_{estudante_id}", ""),
                        height=90,
                    )
                    indicadores_altas_habilidades = st.multiselect("Indicadores de altas habilidades/superdotação", OPCOES_INDICADORES_AHSD)
                    recursos_surdez = st.multiselect("Caso seja pessoa surda, marque recursos utilizados", OPCOES_RECURSOS_SURDEZ, default=["Não se aplica"])
                    observacoes_surdez = st.text_area(
                        "Observações sobre aparelho auditivo, implante coclear, Libras, uso diário, efetividade ou incômodos",
                        height=120,
                    )

                with aba4:
                    st.markdown("#### Síntese pedagógica")
                    contextualizacao = st.text_area(
                        "Contextualização",
                        value=st.session_state.get(f"ia_percurso_{estudante_id}", ""),
                        height=130,
                    )
                    potencialidades = st.text_area("Potencialidades", value=st.session_state.get(f"ia_potencialidades_{estudante_id}", ""), height=100)
                    dificuldades = st.text_area("Dificuldades/barreiras observadas", value=st.session_state.get(f"ia_dificuldades_{estudante_id}", ""), height=100)
                    estrategias = st.text_area("Estratégias pedagógicas", value=st.session_state.get(f"ia_estrategias_{estudante_id}", ""), height=120)
                    intervencoes = st.text_area("Intervenções / encaminhamentos sugeridos", value=st.session_state.get(f"ia_intervencoes_{estudante_id}", ""), height=120)
                    avaliacao_estudo = st.text_area("Avaliação", value=st.session_state.get(f"ia_avaliacao_{estudante_id}", ""), height=100)
                    consideracoes = st.text_area("Considerações finais", value=st.session_state.get(f"ia_consideracoes_{estudante_id}", ""), height=100)

                salvar_estudo = st.form_submit_button("Salvar estudo de caso GRE")

                if salvar_estudo:
                    campos = ["estudante_id"] + CAMPOS_ESTUDO_CASO
                    valores = [
                        estudante_id,
                        hoje_str(),
                        ano_letivo,
                        tipo_registro,
                        estudo_anterior_id,
                        analise_comparativa_ia,
                        sugestao_novo_estudo_ia,
                        contextualizacao,
                        motivo_encaminhamento_aee,
                        "; ".join([x for x in [", ".join(habilidades_observadas), habilidades_observadas_ia] if x]),
                        "; ".join([x for x in [", ".join(habilidades_a_desenvolver), habilidades_a_desenvolver_ia] if x]),
                        estrategias,
                        intervencoes,
                        avaliacao_estudo,
                        consideracoes,
                        etapa_modalidade,
                        ano_etapa,
                        laudo,
                        deficiencia,
                        cid,
                        observacoes_laudo,
                        altas_habilidades,
                        bpc,
                        escola_nome,
                        unidade_aee,
                        gestor_nome,
                        gestor_contato,
                        professor_nome,
                        professor_contato,
                        matricula_professor,
                        ", ".join(especialidade_professor),
                        periodo_inicio,
                        periodo_fim,
                        ", ".join(frequencia_atendimento),
                        tempo_atendimento_semana,
                        ", ".join(formato_atendimento),
                        percurso_educacional,
                        motivo_encaminhamento_aee,
                        precisa_transporte_inclusivo,
                        recebe_transporte_inclusivo,
                        precisa_profissional_apoio,
                        justificativa_apoio,
                        acompanhado_profissional_apoio,
                        nome_profissional_apoio,
                        ", ".join(recursos_tecnologia_assistiva),
                        observacoes_ambiente_educacional,
                        "; ".join([x for x in [", ".join(habilidades_observadas), habilidades_observadas_ia] if x]),
                        "; ".join([x for x in [", ".join(habilidades_a_desenvolver), habilidades_a_desenvolver_ia] if x]),
                        ", ".join(indicadores_altas_habilidades),
                        ", ".join(recursos_surdez),
                        observacoes_surdez,
                    ]
                    inserir_registro("estudos_caso", campos, valores)
                    st.success("Estudo de caso GRE salvo.")
                    st.rerun()

        estudos = listar_por_estudante(
            "estudos_caso",
            CAMPOS_ESTUDO_CASO,
            estudante_id,
        )
        with st.container(border=True):
            st.markdown("### Histórico de estudos de caso")
            if estudos:
                for e in estudos:
                    with st.expander(f"Estudo em {e[1]} | Ano letivo: {e[2] if len(e) > 2 else 'não informado'}"):
                        texto = texto_estudo_caso(estudante, e)
                        st.text(texto)
                        export_buttons(texto, f"Estudo_Caso_GRE_{estudante[1]}_{e[0]}", tipo_pdf="estudo")
                        if st.button("Excluir estudo de caso", key=f"exc_estudo_{e[0]}"):
                            excluir_registro("estudos_caso", e[0])
                            st.success("Estudo excluído.")
                            st.rerun()
            else:
                st.info("Nenhum estudo de caso registrado.")


# ======================================================
# PLANO AEE - INTELIGÊNCIA
# ======================================================
elif menu == "Plano AEE - IA":
    st.markdown('<div class="subtitulo">🧠 Plano AEE - IA</div>', unsafe_allow_html=True)
    st.caption(
        "Módulo de planejamento pedagógico inteligente para diagnóstico, sugestão geral, plano mensal, evolução e histórico do AEE."
    )

    estudantes = listar_estudantes()
    if not estudantes:
        st.info("Cadastre um estudante primeiro.")
    else:
        ids, mapa = opcoes_estudantes_por_id(estudantes)
        estudante_id = st.selectbox(
            "Selecione o estudante",
            ids,
            format_func=lambda x: mapa[x],
            key="plano_ia_estudante_v19",
        )
        estudante = buscar_estudante(estudante_id)

        avaliacao_ia = ultima_avaliacao(estudante_id)
        entrevista_ia = ultima_linha("entrevistas_familia", CAMPOS_ENTREVISTA_FAMILIA, estudante_id)
        estudo_ia = ultima_linha("estudos_caso", CAMPOS_ESTUDO_CASO, estudante_id)
        plano_manual_ia = ultima_linha("planos_aee", CAMPOS_PLANO_AEE, estudante_id)

        pendencias = []
        if not entrevista_ia:
            pendencias.append("Entrevista com a família")
        if not avaliacao_ia:
            pendencias.append("Avaliação pedagógica")
        if not estudo_ia:
            pendencias.append("Estudo de caso GRE")

        escutas_docentes_ia = listar_escutas_docentes(estudante_id)
        relatorios_docente_ia = listar_relatorios_docente(estudante_id)

        with st.container(border=True):
            st.markdown("### ✅ Dados usados pela IA")
            c1, c2, c3, c4, c5, c6 = st.columns(6)
            c1.metric("Entrevista familiar", "OK" if entrevista_ia else "Pendente")
            c2.metric("Avaliação pedagógica", "OK" if avaliacao_ia else "Pendente")
            c3.metric("Estudo de caso GRE", "OK" if estudo_ia else "Pendente")
            c4.metric("Escuta docente", f"{len(escutas_docentes_ia)} registro(s)" if escutas_docentes_ia else "Pendente")
            c5.metric("Relatórios docentes", f"{len(relatorios_docente_ia)} salvo(s)" if relatorios_docente_ia else "Opcional")
            c6.metric("Plano AEE manual", "OK" if plano_manual_ia else "Opcional")

            fontes_ativas = []
            if entrevista_ia:
                fontes_ativas.append("entrevista familiar")
            if avaliacao_ia:
                fontes_ativas.append("avaliação pedagógica")
            if estudo_ia:
                fontes_ativas.append("estudo de caso GRE")
            if escutas_docentes_ia:
                fontes_ativas.append("escuta docente/histórico de escutas")
            if relatorios_docente_ia:
                fontes_ativas.append("relatórios de apoio ao docente")
            if plano_manual_ia:
                fontes_ativas.append("plano AEE manual")

            if pendencias:
                st.warning(
                    "Ainda faltam dados para uma sugestão mais personalizada: " + ", ".join(pendencias) + ". "
                    "Mesmo assim, é possível gerar um roteiro inicial de observação e planejamento sem presumir informações não registradas."
                )
            else:
                st.success("Dados mínimos encontrados. A IA pode gerar sugestões mais personalizadas para o estudante.")

            if fontes_ativas:
                st.caption("Fontes pedagógicas consideradas neste estudante: " + ", ".join(fontes_ativas) + ".")

        escolha_plano_ia = st.radio(
            "Escolha o que deseja fazer",
            [
                "🧩 Perfil Pedagógico Inteligente",
                "📘 Sugestão Geral AEE",
                "📅 Plano Mensal IA",
                "📈 Evolução IA",
                "📚 Bases de Conhecimento",
                "📝 Plano AEE Manual",
                "🗂️ Histórico IA",
            ],
            horizontal=True,
            key=f"radio_plano_ia_v22_{estudante_id}",
        )

        st.divider()

        if escolha_plano_ia == "🧩 Perfil Pedagógico Inteligente":
            st.markdown("### 🧩 Perfil Pedagógico Inteligente")
            st.info(
                "Organiza informações pedagógicas, funcionais e educacionais registradas no sistema para apoiar o planejamento do AEE, considerando comunicação, autonomia, participação escolar, acessibilidade, interação social e estratégias pedagógicas."
            )
            st.caption(
                "Este recurso possui finalidade exclusivamente pedagógica e educacional, não realizando diagnóstico clínico, médico ou terapêutico."
            )
            if st.button("🧩 Gerar Perfil Pedagógico Inteligente", key=f"gerar_perfil_pedagogico_v22_{estudante_id}"):
                with st.spinner("Gerando Perfil Pedagógico Inteligente com IA..."):
                    st.session_state[f"diagnostico_ia_v19_{estudante_id}"] = gerar_diagnostico_aee_ia(
                        estudante, avaliacao_ia, entrevista_ia, estudo_ia, plano_manual_ia
                    )

            if f"diagnostico_ia_v19_{estudante_id}" in st.session_state:
                diagnostico_txt = st.text_area(
                    "Perfil pedagógico gerado",
                    st.session_state[f"diagnostico_ia_v19_{estudante_id}"],
                    height=520,
                    key=f"diagnostico_txt_v19_{estudante_id}",
                )
                col_d1, col_d2 = st.columns([1, 1])
                with col_d1:
                    export_buttons(diagnostico_txt, f"Perfil_Pedagogico_Inteligente_{estudante[1]}", tipo_pdf="plano")
                with col_d2:
                    if st.button("💾 Salvar perfil pedagógico no histórico", key=f"salvar_perfil_pedagogico_v22_{estudante_id}"):
                        salvar_historico_plano_aee_ia(
                            estudante_id=estudante_id,
                            mes_referencia="",
                            ano_referencia=agora_local().year,
                            qtd_atendimentos_semana=1,
                            tipo_geracao="Perfil Pedagógico Inteligente",
                            diagnostico_ia=diagnostico_txt,
                            observacoes="Perfil Pedagógico Inteligente gerado no módulo Plano AEE - IA.",
                        )
                        st.success("Perfil pedagógico salvo no histórico IA.")
                        st.rerun()

        elif escolha_plano_ia == "📘 Sugestão Geral AEE":
            st.markdown("### 📘 Sugestão geral de atendimento AEE")
            st.info(
                "Gera o norte pedagógico do período: objetivos prioritários, recursos sugeridos, estratégias, organização dos atendimentos e indicadores de acompanhamento."
            )
            if st.button("📘 Gerar Sugestão Geral AEE - IA", key=f"gerar_sug_geral_v19_{estudante_id}"):
                with st.spinner("Gerando sugestão geral de atendimento..."):
                    st.session_state[f"sugestao_geral_ia_v19_{estudante_id}"] = gerar_sugestao_geral_aee_ia(
                        estudante, avaliacao_ia, entrevista_ia, estudo_ia, plano_manual_ia
                    )

            if f"sugestao_geral_ia_v19_{estudante_id}" in st.session_state:
                sugestao_txt = st.text_area(
                    "Sugestão geral gerada",
                    st.session_state[f"sugestao_geral_ia_v19_{estudante_id}"],
                    height=560,
                    key=f"sugestao_geral_txt_v19_{estudante_id}",
                )
                bloco_sugestoes_3d_streamlit(
                    conteudo_base=sugestao_txt,
                    estudante_id=estudante_id,
                    estudante_codigo=estudante[1],
                    origem="Sugestao_Geral_AEE",
                )

                st.markdown("### 📤 Exportar Sugestão Geral AEE")
                export_buttons(sugestao_txt, f"Sugestao_Geral_AEE_IA_{estudante[1]}", tipo_pdf="plano")

                st.markdown("### 💾 Salvar no histórico")
                if st.button("💾 Salvar sugestão geral no histórico", key=f"salvar_sug_geral_v19_{estudante_id}"):
                    salvar_historico_plano_aee_ia(
                        estudante_id=estudante_id,
                        mes_referencia="",
                        ano_referencia=agora_local().year,
                        qtd_atendimentos_semana=1,
                        tipo_geracao="Sugestão Geral AEE",
                        sugestao_geral=sugestao_txt,
                        observacoes="Sugestão geral de atendimento gerada no módulo Plano AEE - IA.",
                    )
                    st.success("Sugestão geral salva no histórico IA.")
                    st.rerun()

        elif escolha_plano_ia == "📅 Plano Mensal IA":
            st.markdown("### 📅 Plano mensal inteligente")
            st.info(
                "Organiza os atendimentos do mês por semana. Para cada atendimento, a IA sugere objetivo, atividade, recursos e registro esperado."
            )
            meses = [
                "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
                "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro",
            ]
            dias_cadastrados_txt = ""
            try:
                dias_cadastrados_txt = estudante[7] or ""
            except Exception:
                dias_cadastrados_txt = ""
            dias_padrao = [d for d in DIAS_SEMANA if d in dias_cadastrados_txt]
            if not dias_padrao:
                dias_padrao = ["Segunda-feira", "Quarta-feira"]

            col_m1, col_m2, col_m3 = st.columns([1, 1, 2])
            with col_m1:
                mes_ref = st.selectbox("Mês de referência", meses, index=agora_local().month - 1, key=f"mes_plano_ia_v19_{estudante_id}")
            with col_m2:
                ano_ref = st.text_input("Ano de referência", value=str(agora_local().year), key=f"ano_plano_ia_v19_{estudante_id}")
            with col_m3:
                dias_atendimento_ref = st.multiselect(
                    "Dias de atendimento na SRM",
                    DIAS_SEMANA,
                    default=dias_padrao,
                    key=f"dias_atendimento_plano_ia_v23_{estudante_id}",
                    help="O sistema calcula automaticamente quantos atendimentos existirão no mês com base nos dias selecionados.",
                )

            datas_calculadas = gerar_datas_atendimentos_mes(ano_ref, mes_ref, dias_atendimento_ref)
            qtd_semana = max(1, len(dias_atendimento_ref))
            total_atendimentos_calculado = len(datas_calculadas)

            col_info1, col_info2 = st.columns(2)
            with col_info1:
                st.metric("Atendimentos calculados no mês", total_atendimentos_calculado)
            with col_info2:
                st.metric("Dias por semana", qtd_semana)

            with st.expander("📆 Ver datas previstas de atendimento", expanded=True):
                st.text(datas_atendimentos_para_texto(datas_calculadas))

            st.caption(
                "O roteiro gerado usa as datas reais do mês. Após cada encontro, registre o atendimento no módulo Atendimentos para alimentar a evolução da IA."
            )

            with st.container(border=True):
                st.markdown("### 🚀 Projeto norteador interdisciplinar")
                st.caption(
                    "Opcional. Use quando desejar que o plano mensal seja organizado a partir de um projeto. "
                    "O foco continua sendo o AEE: acessibilidade, autonomia, participação, organização cognitiva, comunicação e permanência — não reforço escolar."
                )
                usar_projeto_norteador = st.checkbox(
                    "Usar projeto norteador neste plano mensal",
                    key=f"usar_projeto_norteador_{estudante_id}",
                )
                projeto_norteador = ""
                nivel_projeto_norteador = ""
                observacoes_projeto_norteador = ""
                if usar_projeto_norteador:
                    col_proj1, col_proj2 = st.columns(2)
                    with col_proj1:
                        projeto_opcao = st.selectbox(
                            "Projeto sugerido ou tema",
                            OPCOES_PROJETOS_NORTEADORES,
                            key=f"projeto_opcao_norteador_{estudante_id}",
                        )
                        if projeto_opcao == "Projeto livre / personalizado":
                            projeto_norteador = st.text_input(
                                "Descreva o projeto personalizado",
                                placeholder="Ex.: carrinho seguidor, carrinho solar, banner científico, experimento acessível...",
                                key=f"projeto_livre_norteador_{estudante_id}",
                            )
                        else:
                            projeto_norteador = projeto_opcao
                    with col_proj2:
                        nivel_projeto_norteador = st.selectbox(
                            "Nível do projeto",
                            OPCOES_NIVEL_PROJETO_NORTEADOR,
                            key=f"nivel_projeto_norteador_{estudante_id}",
                        )
                    observacoes_projeto_norteador = st.text_area(
                        "Observações sobre o projeto e o estudante",
                        height=90,
                        placeholder="Ex.: estudante do 3º ano, boa fala, dificuldade de compreensão em sala, interesse por tecnologia; trabalhar banner, texto, cálculo simples, registro da experiência...",
                        key=f"obs_projeto_norteador_{estudante_id}",
                    )

            if st.button("📅 Gerar Plano Mensal AEE - IA", key=f"gerar_plano_mensal_v19_{estudante_id}"):
                with st.spinner("Gerando plano mensal por semanas e atendimentos..."):
                    plano_base = gerar_plano_mensal_aee_ia(
                        estudante,
                        mes_ref,
                        ano_ref,
                        qtd_semana,
                        avaliacao_ia,
                        entrevista_ia,
                        estudo_ia,
                        plano_manual_ia,
                        datas_calculadas,
                        usar_projeto_norteador=usar_projeto_norteador,
                        projeto_norteador=projeto_norteador,
                        nivel_projeto_norteador=nivel_projeto_norteador,
                        observacoes_projeto_norteador=observacoes_projeto_norteador,
                    )
                    complemento = f"""

ORGANIZAÇÃO OPERACIONAL PARA REGISTRO NO SISTEMA
Mês de referência: {mes_ref}/{ano_ref}
Dias de atendimento selecionados: {", ".join(dias_atendimento_ref) if dias_atendimento_ref else "Não informado"}
Total real de atendimentos no mês: {total_atendimentos_calculado}
Projeto norteador ativado: {"Sim" if usar_projeto_norteador else "Não"}
Projeto/tema: {projeto_norteador if usar_projeto_norteador else "Não se aplica"}
Nível do projeto: {nivel_projeto_norteador if usar_projeto_norteador else "Não se aplica"}

Datas previstas:
{datas_atendimentos_para_texto(datas_calculadas)}

Modelo obrigatório para cada atendimento:
- Objetivo do atendimento:
- Atividade proposta:
- Recursos utilizados:
- Mediação necessária:
- Resposta esperada do estudante:
- Como registrar depois no módulo Atendimentos:
""".strip()
                    st.session_state[f"plano_mensal_ia_v19_{estudante_id}"] = plano_base + "\n\n" + complemento

            if f"plano_mensal_ia_v19_{estudante_id}" in st.session_state:
                plano_mensal_txt = st.text_area(
                    "Plano mensal gerado",
                    st.session_state[f"plano_mensal_ia_v19_{estudante_id}"],
                    height=680,
                    key=f"plano_mensal_txt_v19_{estudante_id}",
                )
                bloco_sugestoes_3d_streamlit(
                    conteudo_base=plano_mensal_txt,
                    estudante_id=estudante_id,
                    estudante_codigo=estudante[1],
                    origem="Plano_Mensal_AEE",
                    nome_extra=f"_{mes_ref}_{ano_ref}",
                )

                st.markdown("### 📤 Exportar Plano Mensal AEE")
                export_buttons(plano_mensal_txt, f"Plano_Mensal_AEE_IA_{estudante[1]}_{mes_ref}_{ano_ref}", tipo_pdf="plano_ia_visual")

                st.markdown("### 💾 Salvar no histórico")
                if st.button("💾 Salvar plano mensal no histórico", key=f"salvar_plano_mensal_v19_{estudante_id}"):
                    salvar_historico_plano_aee_ia(
                        estudante_id=estudante_id,
                        mes_referencia=mes_ref,
                        ano_referencia=ano_ref,
                        qtd_atendimentos_semana=int(qtd_semana),
                        tipo_geracao="Plano Mensal IA",
                        plano_mensal=plano_mensal_txt,
                        observacoes=f"Plano mensal calculado automaticamente com {total_atendimentos_calculado} atendimento(s) no mês, considerando os dias: {', '.join(dias_atendimento_ref)}.",
                    )
                    st.success("Plano mensal salvo no histórico IA.")
                    st.rerun()

        elif escolha_plano_ia == "📈 Evolução IA":
            st.markdown("### 📈 Evolução inteligente")
            st.info(
                "Usa os registros de atendimento para apoiar a análise de avanços, barreiras persistentes, recursos mais efetivos e ajustes para o próximo mês."
            )
            if st.button("📈 Gerar análise evolutiva IA", key=f"gerar_evolucao_ia_v19_{estudante_id}"):
                historico_txt = listar_atendimentos_texto(estudante_id)
                if "Nenhum atendimento registrado" in historico_txt:
                    st.warning("Ainda não há atendimentos registrados. Registre os atendimentos para gerar uma análise evolutiva consistente.")
                    st.session_state[f"evolucao_ia_v19_{estudante_id}"] = (
                        "Ainda não há atendimentos suficientes para análise evolutiva. "
                        "Registre objetivo, atividade, resposta do estudante, avanços, dificuldades, engajamento e encaminhamentos após cada atendimento."
                    )
                else:
                    with st.spinner("Gerando análise evolutiva com base nos atendimentos..."):
                        st.session_state[f"evolucao_ia_v19_{estudante_id}"] = gerar_diagnostico_aee_ia(
                            estudante, avaliacao_ia, entrevista_ia, estudo_ia, plano_manual_ia
                        )

            if f"evolucao_ia_v19_{estudante_id}" in st.session_state:
                evolucao_txt = st.text_area(
                    "Análise evolutiva gerada",
                    st.session_state[f"evolucao_ia_v19_{estudante_id}"],
                    height=540,
                    key=f"evolucao_txt_v19_{estudante_id}",
                )
                col_e1, col_e2 = st.columns([1, 1])
                with col_e1:
                    export_buttons(evolucao_txt, f"Evolucao_AEE_IA_{estudante[1]}", tipo_pdf="relatorio")
                with col_e2:
                    if st.button("💾 Salvar evolução no histórico", key=f"salvar_evolucao_ia_v19_{estudante_id}"):
                        salvar_historico_plano_aee_ia(
                            estudante_id=estudante_id,
                            mes_referencia="",
                            ano_referencia=agora_local().year,
                            qtd_atendimentos_semana=1,
                            tipo_geracao="Evolução IA",
                            diagnostico_ia=evolucao_txt,
                            observacoes="Análise evolutiva com base nos registros de atendimento.",
                        )
                        st.success("Evolução salva no histórico IA.")
                        st.rerun()

        elif escolha_plano_ia == "📚 Bases de Conhecimento":
            st.markdown("### 📚 Bases de Conhecimento IA")
            st.caption("Envie PDFs, organize por categoria e indexe para que a IA consulte a base ao gerar relatórios e planos.")

            with st.container(border=True):
                st.markdown("#### Enviar novo PDF para a base")
                col_up1, col_up2 = st.columns([2, 1])
                with col_up1:
                    arquivo_base = st.file_uploader(
                        "Selecionar PDF",
                        type=["pdf"],
                        key=f"upload_base_conhecimento_pdf_{estudante_id}",
                    )
                with col_up2:
                    base_destino = st.selectbox(
                        "Categoria da base",
                        list(OPCOES_BASES_CONHECIMENTO_IA.keys()),
                        format_func=lambda x: OPCOES_BASES_CONHECIMENTO_IA.get(x, x),
                        key=f"categoria_upload_base_{estudante_id}",
                    )
                if st.button("Salvar PDF na Base de Conhecimento", key=f"salvar_pdf_base_{estudante_id}"):
                    if arquivo_base is None:
                        st.warning("Selecione um arquivo PDF primeiro.")
                    else:
                        try:
                            caminho_salvo = salvar_pdf_base_conhecimento(arquivo_base, base_destino)
                            st.success(f"PDF salvo em: {caminho_salvo}")
                        except Exception as e:
                            st.error(f"Erro ao salvar PDF: {e}")

            st.markdown("#### Bases cadastradas")
            cols_bases = st.columns(2)
            for idx, (nome_base, rotulo_base) in enumerate(OPCOES_BASES_CONHECIMENTO_IA.items()):
                with cols_bases[idx % 2]:
                    with st.container(border=True):
                        st.markdown(f"**{rotulo_base}**")
                        pdfs = listar_pdfs_base(PASTAS_BASE_IA[nome_base])
                        st.metric("PDFs", len(pdfs))
                        if pdfs:
                            with st.expander("Ver PDFs"):
                                for pdf in pdfs:
                                    st.write(f"• {pdf.name}")
                        else:
                            st.caption("Nenhum PDF nesta categoria.")
                        if st.button("Indexar esta base", key=f"indexar_base_{nome_base}_{estudante_id}"):
                            try:
                                with st.spinner(f"Indexando {rotulo_base}..."):
                                    total, msg = indexar_base_conhecimento(nome_base)
                                st.success(f"{msg} Total de trechos indexados: {total}")
                            except Exception as e:
                                st.error(f"Erro ao indexar {rotulo_base}: {e}")

            st.markdown("#### Consultar manualmente as bases")
            pergunta_base = st.text_area(
                "Digite uma pergunta para consultar os documentos",
                placeholder="Ex: Que estratégias usar para estudante com TEA no AEE?",
                key=f"pergunta_base_conhecimento_ia_v19_{estudante_id}",
            )
            bases_escolhidas = st.multiselect(
                "Bases para consulta",
                list(OPCOES_BASES_CONHECIMENTO_IA.keys()),
                default=["cientifica", "pedagogica", "bncc"],
                format_func=lambda x: OPCOES_BASES_CONHECIMENTO_IA.get(x, x),
                key=f"bases_consulta_manual_v19_{estudante_id}",
            )
            if st.button("Consultar bases", key=f"consultar_bases_ia_v19_{estudante_id}"):
                if not pergunta_base.strip():
                    st.warning("Digite uma pergunta primeiro.")
                else:
                    with st.spinner("Buscando trechos nas bases..."):
                        resultados = buscar_na_base_conhecimento(pergunta_base, bases=bases_escolhidas, limite=6)
                        contexto = montar_contexto_base(resultados)
                        st.session_state[f"consulta_base_resultado_v19_{estudante_id}"] = contexto
            if f"consulta_base_resultado_v19_{estudante_id}" in st.session_state:
                st.text_area("Trechos encontrados", st.session_state[f"consulta_base_resultado_v19_{estudante_id}"], height=320)

        elif escolha_plano_ia == "📝 Plano AEE Manual":
            st.markdown("### 📝 Novo Plano AEE / PAEE manual")
            with st.form("form_plano_manual_ia_v19"):
                habilidades_lista = st.multiselect("1.1 Habilidades prioritárias que serão trabalhadas na SRM", OPCOES_HABILIDADES_PRIORITARIAS_SRM)
                habilidades_outros = st.text_area("Complemento de habilidades prioritárias, se necessário")
                recursos_lista = st.multiselect("1.2 Recursos de acessibilidade que serão disponibilizados ao estudante", OPCOES_RECURSOS_ACESSIBILIDADE)
                recursos_outros = st.text_area("Complemento de recursos de acessibilidade, se necessário")
                objetivos_gerais = st.text_area("2.1 Objetivo geral")
                objetivos_especificos = st.text_area("2.2 Objetivos específicos")
                metodologia = st.text_area("3.1 Metodologia")
                estrategias = st.text_area("3.2 Estratégia")
                prazo = st.text_area("3.3 Prazo")
                acoes_lista = st.multiselect("4. Ações desenvolvidas no âmbito da escola", OPCOES_ACOES_ESCOLA)
                acoes_outros = st.text_area("Descrição/complemento das ações desenvolvidas")
                barreiras_lista = st.multiselect("5. Barreiras identificadas na comunidade escolar", OPCOES_BARREIRAS)
                barreiras_outros = st.text_area("Descrição/complemento das barreiras identificadas")
                parcerias = st.text_area("6. Parcerias realizadas pelo AEE ao longo do período")
                avaliacao = st.text_area("7. Avaliação")
                observacoes = st.text_area("Observações complementares")
                if st.form_submit_button("Salvar Plano AEE / PAEE manual"):
                    habilidades = "; ".join([x for x in [", ".join(habilidades_lista), habilidades_outros] if x])
                    recursos = "; ".join([x for x in [", ".join(recursos_lista), recursos_outros] if x])
                    acoes_escola = "; ".join([x for x in [", ".join(acoes_lista), acoes_outros] if x])
                    barreiras = "; ".join([x for x in [", ".join(barreiras_lista), barreiras_outros] if x])
                    inserir_registro(
                        "planos_aee",
                        ["estudante_id", *CAMPOS_PLANO_AEE],
                        [
                            estudante_id, hoje_str(), habilidades, recursos, objetivos_gerais, objetivos_especificos,
                            metodologia, estrategias, prazo, acoes_escola, barreiras, parcerias, avaliacao, observacoes,
                        ],
                    )
                    st.success("Plano manual salvo.")
                    st.rerun()

            st.markdown("### Histórico de planos manuais")
            planos = listar_por_estudante("planos_aee", CAMPOS_PLANO_AEE, estudante_id)
            if planos:
                for p in planos:
                    with st.expander(f"Plano em {p[1]}"):
                        texto = texto_plano_aee(estudante, p)
                        st.text(texto)
                        export_buttons(texto, f"Plano_AEE_PAEE_{estudante[1]}_{p[0]}", tipo_pdf="plano")
                        if st.button("Excluir plano", key=f"exc_plano_v19_{p[0]}"):
                            excluir_registro("planos_aee", p[0])
                            st.success("Plano excluído.")
                            st.rerun()
            else:
                st.info("Nenhum plano manual registrado.")

        elif escolha_plano_ia == "🗂️ Histórico IA":
            st.markdown("### 🗂️ Histórico do Plano AEE - IA")
            campos_hist_ia = CAMPOS_PLANO_AEE_IA
            historico_ia = listar_por_estudante("plano_aee_ia", campos_hist_ia, estudante_id)
            if historico_ia:
                for item in historico_ia:
                    item_id = item[0]
                    data_geracao = item[1]
                    mes_hist = item[2]
                    ano_hist = item[3]
                    tipo_hist = item[5] or "Registro IA"
                    diagnostico_hist = item[6] or ""
                    sugestao_hist = item[7] or ""
                    plano_hist = item[11] or item[12] or ""
                    conteudo_hist = diagnostico_hist or sugestao_hist or plano_hist or item[13] or ""
                    titulo_hist = f"{tipo_hist} - {normalizar_data_historico(data_geracao)}"
                    if mes_hist or ano_hist:
                        titulo_hist += f" - {mes_hist or ''}/{ano_hist or ''}"
                    with st.expander(titulo_hist):
                        st.text_area("Conteúdo", conteudo_hist, height=440, key=f"hist_plano_ia_v19_{item_id}")
                        export_buttons(conteudo_hist, f"Plano_AEE_IA_{estudante[1]}_{item_id}", tipo_pdf=("plano_ia_visual" if "Plano Mensal" in str(tipo_hist) else "plano"))
                        if st.button("Excluir registro IA", key=f"exc_plano_ia_v19_{item_id}"):
                            excluir_registro("plano_aee_ia", item_id)
                            st.success("Registro IA excluído.")
                            st.rerun()
            else:
                st.info("Nenhum registro IA salvo ainda.")


elif menu == "Recursos Pedagógicos e TA da Escola":
    renderizar_pagina_recursos_escola()

elif menu == "Atendimentos":
    st.markdown('<div class="subtitulo">📌 Registro dos atendimentos</div>', unsafe_allow_html=True)
    estudantes = listar_estudantes()
    if not estudantes:
        st.info("Cadastre um estudante primeiro.")
    else:
        ids, mapa = opcoes_estudantes_por_id(estudantes)
        estudante_id = st.selectbox("Selecione o estudante", ids, format_func=lambda x: mapa[x], key="atendimento_estudante")
        estudante = buscar_estudante(estudante_id)

        with st.container(border=True):
            st.markdown("### Novo atendimento")
            with st.form("form_atendimento"):
                data_atendimento = st.date_input("Data do atendimento")
                objetivo = st.text_area("Objetivo trabalhado")
                atividade = st.text_area("Atividade realizada")
                resposta = st.text_area("Resposta do estudante")
                nivel_resposta = st.slider("Escala da resposta do estudante", 1, 10, 5)
                avancos = st.text_area("Avanços observados")
                nivel_avanco = st.slider("Escala do avanço pedagógico", 1, 10, 5)
                dificuldades = st.text_area("Dificuldades observadas")
                nivel_dificuldade = st.slider("Escala de dificuldade/barreira observada", 1, 10, 5)
                evolucao = st.text_area("Evolução observada")
                nivel_engajamento = st.slider("Escala de engajamento/participação", 1, 10, 5)
                indice = calcular_indice_geral(nivel_resposta, nivel_avanco, nivel_dificuldade, nivel_engajamento)
                st.info(f"Índice geral calculado automaticamente: {indice}/10")
                encaminhamentos = st.text_area("Encaminhamentos")

                st.markdown("#### 🧩 Recursos pedagógicos utilizados")
                st.caption(
                    "Registre apenas o material criado ou utilizado no atendimento. "
                    "Não insira fotos do estudante; use links do Drive, Canva, Wordwall, modelos 3D, jogos ou documentos pedagógicos."
                )

                qtd_recursos = st.number_input(
                    "Quantidade de recursos/materiais utilizados",
                    min_value=0,
                    max_value=10,
                    value=0,
                    step=1,
                    help="Use quando houver mais de um material pedagógico vinculado ao atendimento.",
                )

                categorias_recursos = [
                    "CAA / Comunicação alternativa",
                    "Rotina visual",
                    "Jogo pedagógico",
                    "Atividade digital",
                    "Wordwall",
                    "Canva",
                    "Impressão 3D",
                    "Robótica educacional",
                    "Material concreto/manipulável",
                    "Tecnologia assistiva",
                    "Outro",
                ]

                recursos_registrados = []
                for i in range(int(qtd_recursos)):
                    st.markdown(f"**Recurso {i + 1}**")
                    c1, c2 = st.columns([1.2, 1])
                    with c1:
                        nome_recurso = st.text_input(
                            "Nome do recurso/material",
                            key=f"recurso_nome_{i}",
                            placeholder="Ex.: Prancha CAA alimentar, atividade Wordwall, modelo 3D...",
                        )
                    with c2:
                        categoria_recurso = st.selectbox(
                            "Categoria",
                            categorias_recursos,
                            key=f"recurso_categoria_{i}",
                        )
                    link_recurso = st.text_input(
                        "Link do material",
                        key=f"recurso_link_{i}",
                        placeholder="Cole aqui o link do Drive, Canva, Wordwall, arquivo 3D, vídeo ou documento.",
                    )
                    observacao_uso = st.text_area(
                        "Observação sobre o uso do recurso",
                        key=f"recurso_obs_{i}",
                        placeholder="Ex.: apresentou maior engajamento; precisou de mediação; facilitou a escolha visual...",
                        height=80,
                    )
                    recursos_registrados.append((nome_recurso, categoria_recurso, link_recurso, observacao_uso))

                if st.form_submit_button("Salvar atendimento"):
                    atendimento_id = salvar_atendimento(
                        estudante_id, data_atendimento.strftime("%d/%m/%Y"), objetivo, atividade, resposta,
                        avancos, dificuldades, evolucao, 1, nivel_resposta, nivel_avanco,
                        nivel_dificuldade, nivel_engajamento, indice, encaminhamentos,
                    )

                    for nome_recurso, categoria_recurso, link_recurso, observacao_uso in recursos_registrados:
                        if nome_recurso.strip() or link_recurso.strip() or observacao_uso.strip():
                            salvar_recurso_atendimento(
                                atendimento_id,
                                estudante_id,
                                nome_recurso.strip(),
                                categoria_recurso,
                                link_recurso.strip(),
                                observacao_uso.strip(),
                            )

                    st.success("Atendimento registrado com recursos pedagógicos vinculados.")
                    st.rerun()

        atendimentos = listar_atendimentos_com_id(estudante_id)

        with st.container(border=True):
            st.markdown("### 📊 Indicadores por atendimento")
            render_grafico_evolucao(listar_atendimentos(estudante_id))

        with st.container(border=True):
            st.markdown("### Histórico de atendimentos")
            if atendimentos:
                for a in atendimentos:
                    with st.expander(f"Atendimento em {a[1]}"):
                        texto = texto_atendimento(estudante, a)
                        st.text(texto)

                        recursos_do_atendimento = listar_recursos_atendimento(a[0])
                        if recursos_do_atendimento:
                            st.markdown("#### 🧩 Recursos pedagógicos vinculados")
                            dados_recursos = []
                            for r in recursos_do_atendimento:
                                _, nome, categoria, link, observacao, criado_em = r
                                dados_recursos.append({
                                    "Recurso/material": nome or "Não informado",
                                    "Categoria": categoria or "Não informada",
                                    "Link": link or "",
                                    "Observação": observacao or "",
                                })
                            st.dataframe(pd.DataFrame(dados_recursos), use_container_width=True)
                            for r in recursos_do_atendimento:
                                _, nome, categoria, link, observacao, criado_em = r
                                if link:
                                    st.markdown(f"🔗 [{nome or 'Abrir recurso pedagógico'}]({link})")

                        export_buttons(texto, f"Registro_Atendimento_{estudante[1]}_{a[0]}", tipo_pdf="atendimento")
                        if st.button("Excluir atendimento", key=f"exc_at_{a[0]}"):
                            excluir_atendimento(a[0])
                            st.success("Atendimento excluído.")
                            st.rerun()
            else:
                st.info("Nenhum atendimento registrado.")


# ======================================================
# AGENDA DE ATENDIMENTOS
# ======================================================
elif menu == "Agenda de Atendimentos":
    st.markdown('<div class="subtitulo">📅 Agenda de Atendimentos</div>', unsafe_allow_html=True)
    estudantes = listar_estudantes()
    if not estudantes:
        st.info("Cadastre um estudante primeiro.")
    else:
        ids, mapa = opcoes_estudantes_por_id(estudantes)
        with st.container(border=True):
            st.markdown("### Novo agendamento")
            with st.form("form_agenda"):
                estudante_id = st.selectbox("Estudante", ids, format_func=lambda x: mapa[x])
                data_ag = st.date_input("Data do atendimento")
                dia_semana = st.selectbox("Dia da semana", DIAS_SEMANA)
                hora_inicio = st.time_input("Hora início", value=time(8, 0))
                hora_fim = st.time_input("Hora fim", value=time(8, 50))
                tipo = st.selectbox("Tipo de atendimento", ["Individual", "Grupo", "Acompanhamento", "Observação", "Outro"])
                obs = st.text_area("Observações")
                if st.form_submit_button("Salvar agendamento"):
                    salvar_agendamento(
                        estudante_id,
                        data_ag.strftime("%d/%m/%Y"),
                        dia_semana,
                        hora_inicio.strftime("%H:%M"),
                        hora_fim.strftime("%H:%M"),
                        tipo,
                        obs,
                    )
                    st.success("Agendamento salvo.")
                    st.rerun()

    agenda = listar_agenda()
    with st.container(border=True):
        st.markdown("### 📆 Agenda semanal")
        if agenda:
            df = pd.DataFrame(
                agenda,
                columns=["ID", "Código", "Ano/Série", "Perfil", "Data", "Dia", "Início", "Fim", "Tipo", "Observações"],
            )
            st.dataframe(df, use_container_width=True, hide_index=True)

            csv = df.to_csv(index=False).encode("utf-8")
            col_csv, col_txt, col_pdf = st.columns(3)
            with col_csv:
                st.download_button("Baixar agenda em CSV", csv, "Agenda_INCLUISRM.csv", "text/csv")
            with col_txt:
                texto = texto_agenda(df)
                st.download_button("Baixar agenda em TXT", texto, "Agenda_INCLUISRM.txt", "text/plain")
            with col_pdf:
                if st.button("Gerar PDF da agenda"):
                    arquivo = gerar_pdf_documento(texto_agenda(df), "Agenda_INCLUISRM", tipo="agenda")
                    st.session_state["pdf_agenda"] = arquivo
                if "pdf_agenda" in st.session_state:
                    with open(st.session_state["pdf_agenda"], "rb") as f:
                        st.download_button("Baixar PDF da agenda", f, "Agenda_INCLUISRM.pdf", "application/pdf")

            for _, row in df.iterrows():
                with st.expander(f"{row['Data']} - {row['Início']} - {row['Código']}"):
                    st.write(row.to_dict())
                    col_status1, col_status2, col_status3 = st.columns(3)
                    with col_status1:
                        if st.button("Marcar Compareceu", key=f"cmp_ag_{row['ID']}"):
                            atualizar_status_agenda(int(row["ID"]), "Compareceu")
                            st.success("Presença marcada como compareceu.")
                            st.rerun()
                    with col_status2:
                        if st.button("Marcar Faltou", key=f"falt_ag_{row['ID']}"):
                            atualizar_status_agenda(int(row["ID"]), "Faltou")
                            st.success("Presença marcada como falta.")
                            st.rerun()
                    with col_status3:
                        if st.button("Marcar Justificado", key=f"just_ag_{row['ID']}"):
                            atualizar_status_agenda(int(row["ID"]), "Justificado")
                            st.success("Presença marcada como justificada.")
                            st.rerun()
                    if st.button("Excluir agendamento", key=f"exc_ag_{row['ID']}"):
                        excluir_agendamento(int(row["ID"]))
                        st.success("Agendamento excluído.")
                        st.rerun()
        else:
            st.info("Nenhum agendamento registrado.")


# ======================================================
# RELATÓRIOS GRE
# ======================================================
elif menu == "Relatórios GRE":
    st.markdown('<div class="subtitulo">📄 Relatórios GRE</div>', unsafe_allow_html=True)
    st.markdown(
        """
        <div class="descricao">
        Gere os documentos solicitados pela GRE sob demanda, a partir dos dados já cadastrados no sistema.
        Os arquivos não ficam armazenados no banco; apenas a data e o tipo de documento gerado entram no histórico.
        Os campos sigilosos permanecem em branco para preenchimento manual no Word ou após impressão.
        </div>
        """,
        unsafe_allow_html=True,
    )

    estudantes = listar_estudantes()
    professores = listar_professores()

    if not estudantes:
        st.info("Cadastre um estudante primeiro.")
    else:
        ids, mapa = opcoes_estudantes_por_id(estudantes)
        estudante_id = st.selectbox("Selecione o estudante", ids, format_func=lambda x: mapa[x], key="gre_estudante")
        estudante = buscar_estudante(estudante_id)

        with st.container(border=True):
            st.markdown("### 📌 Documentos GRE disponíveis")
            tipo = st.selectbox(
                "Escolha o documento que deseja gerar",
                [
                    "Ficha de Identificação Professor(a) AEE",
                    "Matrícula SRM / Termo de Ciência",
                    "Relatório Comparativo das Entrevistas Familiares",
                    "Registro da Entrevista Familiar (última registrada)",
                    "Estudo de Caso e Plano AEE – GRE",
                    "Relatório Consolidado GRE",
                    "Quadro Semanal - Ações/Práticas do Professor AEE",
                    "Pacote GRE Completo",
                ],
            )

            st.info(
                "Os campos como nome completo, CPF, endereço, telefone pessoal, NIS, Cartão SUS e dados familiares sensíveis ficarão em branco para preenchimento manual."
            )
            if tipo == "Relatório Comparativo das Entrevistas Familiares":
                st.success("Este relatório será gerado a partir das entrevistas familiares já registradas no sistema, comparando uma entrevista anterior com uma entrevista atual. A IA atua apenas na análise comparativa, sem criar ou alterar respostas da família.")
            elif tipo == "Registro da Entrevista Familiar (última registrada)":
                st.warning("Este documento apenas imprime a última entrevista familiar registrada. Para analisar mudanças entre anos, escolha 'Relatório Comparativo das Entrevistas Familiares'.")

            # Configurações específicas do Quadro Semanal GRE
            data_inicio_qs = None
            data_fim_qs = None
            escola_qs = ""
            regional_qs = ""
            mes_qs = agora_local().strftime("%m")
            ano_qs = agora_local().strftime("%Y")
            df_quadro_qs = pd.DataFrame()

            entrevista_comp_anterior_id = None
            entrevista_comp_atual_id = None

            if tipo == "Relatório Comparativo das Entrevistas Familiares":
                st.markdown("### 👪 Configuração do relatório comparativo familiar")
                entrevistas_familia = listar_por_estudante("entrevistas_familia", CAMPOS_ENTREVISTA_FAMILIA, estudante_id)
                if len(entrevistas_familia) < 2:
                    st.warning("Registre pelo menos duas entrevistas familiares para gerar o relatório comparativo.")
                else:
                    mapa_entrevistas = {
                        e[0]: f"#{e[0]} | Data: {e[1] or 'Não informada'} | Ano: {(e[2] if len(e) > 2 else '') or 'Não informado'} | Tipo: {(e[3] if len(e) > 3 else '') or 'Não informado'}"
                        for e in entrevistas_familia
                    }
                    ids_entrevistas = [e[0] for e in entrevistas_familia]
                    col_ent_ant, col_ent_atu = st.columns(2)
                    with col_ent_ant:
                        entrevista_comp_anterior_id = st.selectbox(
                            "Entrevista anterior/base",
                            ids_entrevistas,
                            index=1 if len(ids_entrevistas) > 1 else 0,
                            format_func=lambda x: mapa_entrevistas[x],
                            key="gre_entrevista_comparativa_anterior",
                        )
                    with col_ent_atu:
                        entrevista_comp_atual_id = st.selectbox(
                            "Entrevista atual/comparada",
                            ids_entrevistas,
                            index=0,
                            format_func=lambda x: mapa_entrevistas[x],
                            key="gre_entrevista_comparativa_atual",
                        )
                    if entrevista_comp_anterior_id == entrevista_comp_atual_id:
                        st.info("Selecione duas entrevistas diferentes para uma comparação mais útil.")
                    st.caption("A entrevista familiar continua manual. A IA pode ser usada apenas para analisar a comparação entre duas entrevistas, sem criar ou alterar respostas da família.")

                    observacao_ia_entrevista = st.text_area(
                        "Orientação para a análise com IA (opcional)",
                        placeholder="Ex: observar possível impacto de luto, mudança de responsável, alteração na rotina ou sinais de desregulação.",
                        key="gre_entrevista_observacao_ia",
                    )

                    if st.button("🤖 Gerar análise comparativa das entrevistas com IA", key="gre_entrevista_comparativa_ia"):
                        entrevista_anterior_ia = buscar_entrevista_familia_por_id(entrevista_comp_anterior_id) if entrevista_comp_anterior_id else None
                        entrevista_atual_ia = buscar_entrevista_familia_por_id(entrevista_comp_atual_id) if entrevista_comp_atual_id else None
                        with st.spinner("Analisando comparação das entrevistas familiares com IA..."):
                            st.session_state["analise_ia_entrevistas_familia"] = gerar_analise_ia_comparativo_entrevistas_familia(
                                estudante,
                                entrevista_anterior=entrevista_anterior_ia,
                                entrevista_atual=entrevista_atual_ia,
                                observacao_professor=observacao_ia_entrevista,
                            )

                    if "analise_ia_entrevistas_familia" in st.session_state:
                        st.markdown("### 🤖 Análise comparativa com IA")
                        st.text_area(
                            "Análise editável antes de usar no relatório",
                            st.session_state["analise_ia_entrevistas_familia"],
                            height=360,
                            key="gre_entrevista_analise_ia_texto",
                        )

            if tipo == "Quadro Semanal - Ações/Práticas do Professor AEE":
                st.markdown("### 📋 Configuração do Quadro Semanal")
                col_periodo1, col_periodo2, col_periodo3, col_periodo4 = st.columns(4)
                with col_periodo1:
                    data_inicio_qs = st.date_input("Data inicial", value=agora_local().date().replace(day=1), key="qs_inicio")
                with col_periodo2:
                    data_fim_qs = st.date_input("Data final", value=agora_local().date(), key="qs_fim")
                with col_periodo3:
                    mes_qs = st.text_input("Mês", value=agora_local().strftime("%m"), key="qs_mes")
                with col_periodo4:
                    ano_qs = st.text_input("Ano", value=agora_local().strftime("%Y"), key="qs_ano")

                col_escola, col_regional = st.columns(2)
                with col_escola:
                    escola_qs = st.text_input("Escola", value="", placeholder="Preencha ou deixe para puxar do professor cadastrado", key="qs_escola")
                with col_regional:
                    regional_qs = st.text_input("Regional", value="", placeholder="Ex: Metropolitana Norte", key="qs_regional")

                df_quadro_qs = montar_dataframe_quadro_semanal(
                    estudante_id=estudante_id,
                    data_inicio=data_inicio_qs,
                    data_fim=data_fim_qs,
                )

                st.markdown("### ✅ Presença automática e registros do período")
                if df_quadro_qs.empty:
                    st.info("Nenhum agendamento/atendimento encontrado no período selecionado.")
                else:
                    st.dataframe(df_quadro_qs, use_container_width=True, hide_index=True)

                    col_g1, col_g2 = st.columns(2)
                    with col_g1:
                        st.markdown("#### 📈 Atendimentos por status")
                        status_df = df_quadro_qs.groupby("Presença").size().reset_index(name="Quantidade")
                        graf_status = (
                            alt.Chart(status_df)
                            .mark_bar()
                            .encode(
                                x=alt.X("Presença:N", title="Status"),
                                y=alt.Y("Quantidade:Q", title="Quantidade"),
                                tooltip=["Presença", "Quantidade"],
                            )
                            .properties(height=260)
                        )
                        st.altair_chart(graf_status, use_container_width=True)
                    with col_g2:
                        st.markdown("#### 📈 Atendimentos por data")
                        data_df = df_quadro_qs.groupby("Data").size().reset_index(name="Quantidade")
                        graf_data = (
                            alt.Chart(data_df)
                            .mark_bar()
                            .encode(
                                x=alt.X("Data:N", title="Data", sort=None),
                                y=alt.Y("Quantidade:Q", title="Quantidade"),
                                tooltip=["Data", "Quantidade"],
                            )
                            .properties(height=260)
                        )
                        st.altair_chart(graf_data, use_container_width=True)

                    foco_ia_qs = st.text_area(
                        "Orientação para a IA sugerir o preenchimento do quadro",
                        placeholder="Ex: sintetizar os recursos utilizados e as observações pedagógicas com linguagem institucional.",
                        key="qs_foco_ia",
                    )
                    if st.button("🤖 Sugerir preenchimento do Quadro Semanal com IA", key="qs_ia"):
                        with st.spinner("Gerando sugestão pedagógica para o quadro..."):
                            st.session_state["qs_sugestao_ia"] = gerar_sugestao_ia_quadro_semanal(df_quadro_qs, foco=foco_ia_qs)

                    if "qs_sugestao_ia" in st.session_state:
                        st.markdown("### 🤖 Sugestão da IA para preenchimento")
                        st.text_area("Sugestão editável", st.session_state["qs_sugestao_ia"], height=260, key="qs_sugestao_ia_texto")

            if st.button("Gerar documento GRE", key="gerar_documento_gre"):
                if tipo == "Ficha de Identificação Professor(a) AEE":
                    texto = texto_ficha_professor_gre()
                    tipo_pdf = "professor"
                    nome = "Ficha_Professor_AEE_GRE"

                elif tipo == "Matrícula SRM / Termo de Ciência":
                    texto = texto_matricula_srm_gre(estudante)
                    tipo_pdf = "matricula_srm"
                    nome = f"Matricula_SRM_Termo_Ciencia_{estudante[1]}"

                elif tipo == "Registro da Entrevista Familiar (última registrada)":
                    entrevista = ultima_linha("entrevistas_familia", CAMPOS_ENTREVISTA_FAMILIA, estudante_id)
                    texto = texto_entrevista_familia_gre(estudante, entrevista)
                    tipo_pdf = "entrevista"
                    nome = f"Registro_Entrevista_Familiar_GRE_{estudante[1]}"

                elif tipo == "Relatório Comparativo das Entrevistas Familiares":
                    entrevista_anterior = buscar_entrevista_familia_por_id(entrevista_comp_anterior_id) if entrevista_comp_anterior_id else None
                    entrevista_atual = buscar_entrevista_familia_por_id(entrevista_comp_atual_id) if entrevista_comp_atual_id else None
                    texto = gerar_relatorio_comparativo_entrevistas_familia(
                        estudante,
                        entrevista_anterior=entrevista_anterior,
                        entrevista_atual=entrevista_atual,
                    )
                    analise_ia = st.session_state.get("analise_ia_entrevistas_familia", "")
                    if analise_ia:
                        texto = texto + "\n\n7. ANÁLISE COMPARATIVA COM APOIO DE IA\n" + analise_ia + "\n\nObservação: a análise com IA é apoio à leitura pedagógica. A entrevista familiar permanece como registro manual e a decisão final cabe ao professor do AEE."
                    tipo_pdf = "relatorio"
                    nome = f"Relatorio_Comparativo_Entrevistas_Familiares_{estudante[1]}"

                elif tipo in ["Estudo de Caso e Plano AEE – GRE", "Estudo de Caso e Plano AEE – GRE"]:
                    estudo = ultima_linha("estudos_caso", CAMPOS_ESTUDO_CASO, estudante_id)
                    plano = ultima_linha(
                        "planos_aee",
                        CAMPOS_PLANO_AEE,
                        estudante_id,
                    )
                    texto = texto_estudo_plano_aee_gre(estudante, estudo, plano)
                    tipo_pdf = "estudo"
                    nome = f"Estudo_Caso_Plano_AEE_GRE_{estudante[1]}"

                elif tipo == "Relatório Consolidado GRE":
                    texto = gerar_relatorio_gre_texto(estudante)
                    tipo_pdf = "relatorio"
                    nome = f"Relatorio_Consolidado_GRE_{estudante[1]}"

                elif tipo == "Quadro Semanal - Ações/Práticas do Professor AEE":
                    professor = buscar_professor_responsavel()
                    df_quadro_qs = montar_dataframe_quadro_semanal(
                        estudante_id=estudante_id,
                        data_inicio=data_inicio_qs,
                        data_fim=data_fim_qs,
                    )
                    texto = gerar_texto_quadro_semanal_gre(
                        professor=professor,
                        df_quadro=df_quadro_qs,
                        mes=mes_qs,
                        ano=ano_qs,
                        escola_manual=escola_qs,
                        regional_manual=regional_qs,
                    )
                    tipo_pdf = "relatorio"
                    nome = f"Quadro_Semanal_GRE_{estudante[1]}_{mes_qs}_{ano_qs}"
                    try:
                        arquivo_docx_qs = gerar_docx_quadro_semanal_gre(
                            professor=professor,
                            df_quadro=df_quadro_qs,
                            nome_base=f"{estudante[1]}_{mes_qs}_{ano_qs}",
                            mes=mes_qs,
                            ano=ano_qs,
                            escola_manual=escola_qs,
                            regional_manual=regional_qs,
                        )
                        st.session_state["gre_docx_quadro"] = arquivo_docx_qs
                    except Exception as e:
                        st.warning(f"Não foi possível gerar o Word em tabela do Quadro Semanal: {e}")

                else:
                    texto = texto_pacote_gre_completo(estudante)
                    tipo_pdf = "relatorio"
                    nome = f"Pacote_GRE_Completo_{estudante[1]}"

                st.session_state["gre_texto"] = texto
                st.session_state["gre_nome"] = nome
                st.session_state["gre_tipo_pdf"] = tipo_pdf
                st.session_state["gre_tipo"] = tipo

                registrar_documento_gre(
                    estudante_id=estudante_id,
                    tipo_documento=tipo,
                    nome_arquivo=nome,
                    observacao="Documento gerado sob demanda. PDF/Word não armazenado no banco.",
                )
                st.success("Documento GRE gerado sob demanda. Apenas o registro da geração foi salvo no histórico.")

        if "gre_texto" in st.session_state:
            with st.container(border=True):
                titulo_preview = st.session_state.get('gre_tipo', '')
                if titulo_preview == "Relatório Comparativo das Entrevistas Familiares":
                    st.markdown("### Documento gerado: Relatório Comparativo das Entrevistas Familiares")
                    st.caption("Este documento é produzido a partir de entrevistas já registradas no sistema e tem como objetivo analisar possíveis mudanças familiares, emocionais, comportamentais e pedagógicas ao longo do tempo.")
                else:
                    st.markdown(f"### Documento gerado: {titulo_preview}")
                st.text_area("Pré-visualização", st.session_state["gre_texto"], height=560)
                export_buttons(st.session_state["gre_texto"], st.session_state["gre_nome"], tipo_pdf=st.session_state["gre_tipo_pdf"])
                if st.session_state.get("gre_tipo") == "Quadro Semanal - Ações/Práticas do Professor AEE" and "gre_docx_quadro" in st.session_state:
                    try:
                        with open(st.session_state["gre_docx_quadro"], "rb") as f:
                            st.download_button(
                                "Baixar Quadro Semanal em Word com tabela",
                                data=f,
                                file_name=Path(st.session_state["gre_docx_quadro"]).name,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="download_quadro_semanal_docx_tabela",
                            )
                    except Exception:
                        st.info("Gere novamente o Quadro Semanal para baixar o Word em tabela.")

        with st.container(border=True):
            st.markdown("### 🗂️ Histórico de geração GRE")
            st.caption("Este histórico guarda somente data, tipo e nome-base do documento. O arquivo e o conteúdo não ficam armazenados no banco.")
            historico_gre = listar_documentos_gre_gerados(estudante_id)
            if historico_gre:
                df_hist_gre = pd.DataFrame(
                    historico_gre,
                    columns=["ID", "Código", "Tipo de documento", "Nome-base", "Data de geração", "Observação"],
                )
                st.dataframe(df_hist_gre.drop(columns=["ID"]), use_container_width=True)

                opcoes_exclusao_gre = [0] + list(historico_gre)

                def rotulo_excluir_historico_gre(item):
                    if item == 0:
                        return "Não excluir"
                    # item = (id, codigo, tipo_documento, nome_arquivo, data_geracao, observacao)
                    try:
                        return f"ID {item[0]} • {item[1] or 'Sem código'} • {item[2] or 'Documento'} • {item[4] or 'sem data'}"
                    except Exception:
                        return str(item)

                excluir_hist = st.selectbox(
                    "Excluir item do histórico, se necessário",
                    opcoes_exclusao_gre,
                    format_func=rotulo_excluir_historico_gre,
                    key="excluir_historico_gre",
                )
                if excluir_hist != 0:
                    st.warning(f"Selecionado para exclusão: {rotulo_excluir_historico_gre(excluir_hist)}")
                    confirmar_excluir_hist = st.checkbox(
                        "Confirmo que desejo excluir este registro do histórico GRE",
                        key="confirmar_excluir_historico_gre",
                    )
                    if st.button("Excluir registro do histórico GRE"):
                        if confirmar_excluir_hist:
                            excluir_historico_documento_gre(excluir_hist[0])
                            st.success("Registro removido do histórico.")
                            st.rerun()
                        else:
                            st.warning("Marque a confirmação antes de excluir.")
            else:
                st.info("Nenhum documento GRE gerado para este estudante nesta base de dados.")

        with st.container(border=True):
            st.markdown("### 🧾 Conferência dos dados usados")
            st.write(f"**Código interno:** {estudante[1]}")
            st.write(f"**Ano/Série:** {estudante[2] or 'Não informado'}")
            st.write(f"**Turma:** {estudante[3] or 'Não informado'}")
            st.write(f"**Turno:** {estudante[6] or 'Não informado'}")
            st.write(f"**Perfil educacional:** {estudante[4] or 'Não informado'}")
            st.write(f"**Professor(a) AEE:** {texto_professores_vinculados(estudante_id)}")
            st.write("**Campos sigilosos:** permanecerão em branco no documento gerado.")


# ======================================================
# ADMINISTRAÇÃO
# ======================================================
elif menu == "Administração":
    st.markdown('<div class="subtitulo">⚙️ Administração e backup</div>', unsafe_allow_html=True)

    with st.container(border=True):
        st.markdown("### Backup geral do banco")
        tabelas = listar_tabelas_banco()
        backup = {}
        for tabela in tabelas:
            backup[tabela] = carregar_tabela_dataframe(tabela)

        st.write("Tabelas encontradas:", ", ".join(tabelas))

        for tabela, df in backup.items():
            with st.expander(f"Tabela: {tabela} ({len(df)} registros)"):
                st.dataframe(df, use_container_width=True, hide_index=True)
                st.download_button(
                    f"Baixar {tabela}.csv",
                    df.to_csv(index=False).encode("utf-8"),
                    f"{tabela}_incluisrm.csv",
                    "text/csv",
                    key=f"backup_{tabela}",
                )

        # Backup único em JSON
        json_texto = "{\n" + ",\n".join(
            [f'"{tabela}": {df.to_json(orient="records", force_ascii=False)}' for tabela, df in backup.items()]
        ) + "\n}"
        st.download_button(
            "Baixar backup completo em JSON",
            json_texto,
            f"backup_completo_incluisrm_{agora_local().strftime('%Y%m%d_%H%M')}.json",
            "application/json",
        )
